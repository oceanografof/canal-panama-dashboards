"""
╔══════════════════════════════════════════════════════════════════════╗
║   AGUA Y SOSTENIBILIDAD — Generador de Informes                     ║
║   Autoridad del Canal de Panamá · HIMH - Hidrología                ║
║   v3.0 — Informes Mensual · Trimestral · Semestral · Anual          ║
╚══════════════════════════════════════════════════════════════════════╝
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import io, zipfile, re, base64, mimetypes, calendar
import xml.etree.ElementTree as ET
import openpyxl
from datetime import datetime
from pathlib import Path

# ── Intento importar python-docx (opcional) ────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ══════════════════════════════════════════════════════════════════════════════
#  BRAND / ASSETS
# ══════════════════════════════════════════════════════════════════════════════
BASE_DIR = Path(__file__).resolve().parent if "__file__" in globals() else Path.cwd()

def _find_local_file(patterns):
    """Busca archivos locales en la carpeta del app o en /mnt/data."""
    search_dirs = [BASE_DIR, Path("/mnt/data")]
    seen = set()
    for d in search_dirs:
        if not d.exists():
            continue
        for pat in patterns:
            for p in sorted(d.glob(pat)):
                key = str(p.resolve())
                if p.is_file() and key not in seen:
                    seen.add(key)
                    return p
    return None

def img_to_data_uri(filename: str) -> str:
    """Retorna la imagen embebida como data URI para usarla en HTML."""
    stem = Path(filename).stem
    suffix = Path(filename).suffix or ".jpg"
    patterns = [
        filename,
        f"{stem}(1){suffix}",
        f"{stem}(2){suffix}",
        f"{stem}*{suffix}",
    ]
    p = _find_local_file(patterns)
    if p and p.exists():
        mime = mimetypes.guess_type(p.name)[0] or "image/jpeg"
        return f"data:{mime};base64," + base64.b64encode(p.read_bytes()).decode()
    return ""

ACP_LOGO_URI = img_to_data_uri("CP_RGB_p_Ver.jpg")
HIMH_LOGO_URI = img_to_data_uri("LOGO_HIMH.jpg")

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Agua y Sostenibilidad · ACP",
    page_icon="♻️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
#  CSS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
/* ════════════════════════════════════════════════════
   TEMA ADAPTATIVO — Light · Dark · System
   Usa variables nativas de Streamlit + media queries
   ════════════════════════════════════════════════════ */

/* ── Tokens de color base (modo claro) ─────────────── */
:root {
    --c-bg:          #f0f5fb;
    --c-surface:     #ffffff;
    --c-surface2:    #eef3fa;
    --c-border:      rgba(0,114,184,.22);
    --c-border-soft: rgba(148,163,184,.28);
    --c-text:        #0f172a;
    --c-text-muted:  #475569;
    --c-accent:      #0072b8;
    --c-accent-dark: #002a4d;
    --c-accent-glow: rgba(0,114,184,.12);
    --c-up:          #16a34a;
    --c-down:        #dc2626;
    --c-warn-bg:     rgba(245,158,11,.13);
    --c-warn-bdr:    rgba(245,158,11,.32);
    --c-ok-bg:       rgba(16,185,129,.13);
    --c-ok-bdr:      rgba(16,185,129,.32);
    --c-info-bg:     rgba(0,114,184,.10);
    --c-info-bdr:    rgba(0,114,184,.26);
}

/* ── Modo oscuro automático (respeta configuración del OS / Streamlit) ── */
@media (prefers-color-scheme: dark) {
    :root {
        --c-bg:          #000000;
        --c-surface:     #0f1115;
        --c-surface2:    #11151b;
        --c-border:      rgba(0,160,220,.28);
        --c-border-soft: rgba(100,130,160,.28);
        --c-text:        #e6edf3;
        --c-text-muted:  #8b9ab5;
        --c-accent:      #38b6f0;
        --c-accent-dark: #0d2a46;
        --c-accent-glow: rgba(56,182,240,.14);
        --c-up:          #22c55e;
        --c-down:        #f87171;
        --c-warn-bg:     rgba(251,191,36,.11);
        --c-warn-bdr:    rgba(251,191,36,.30);
        --c-ok-bg:       rgba(34,197,94,.11);
        --c-ok-bdr:      rgba(34,197,94,.28);
        --c-info-bg:     rgba(56,182,240,.11);
        --c-info-bdr:    rgba(56,182,240,.26);
    }
}

/* ── Streamlit también expone data-theme="dark" en el html root ── */
[data-theme="dark"] {
    --c-bg:          #000000;
    --c-surface:     #0f1115;
    --c-surface2:    #11151b;
    --c-border:      rgba(0,160,220,.28);
    --c-border-soft: rgba(100,130,160,.28);
    --c-text:        #e6edf3;
    --c-text-muted:  #8b9ab5;
    --c-accent:      #38b6f0;
    --c-accent-dark: #0d2a46;
    --c-accent-glow: rgba(56,182,240,.14);
    --c-up:          #22c55e;
    --c-down:        #f87171;
    --c-warn-bg:     rgba(251,191,36,.11);
    --c-warn-bdr:    rgba(251,191,36,.30);
    --c-ok-bg:       rgba(34,197,94,.11);
    --c-ok-bdr:      rgba(34,197,94,.28);
    --c-info-bg:     rgba(56,182,240,.11);
    --c-info-bdr:    rgba(56,182,240,.26);
}

/* ── Base ────────────────────────────────────────── */
html, body, .stApp {
    background-color: var(--c-bg) !important;
    color: var(--c-text) !important;
    font-family: 'Segoe UI', system-ui, sans-serif;
}
.block-container { padding-top: 1.8rem !important; }
.main .block-container { max-width: 1400px; }

/* ── Sidebar ─────────────────────────────────────── */
[data-testid="stSidebar"] {
    background: linear-gradient(170deg,#002a4d 0%,#004a80 50%,#006ab0 100%) !important;
    box-shadow: 2px 0 16px rgba(0,0,0,.35);
}
[data-testid="stSidebar"],
[data-testid="stSidebar"] * { color: #dceeff !important; }
[data-testid="stSidebar"] hr { border-color: rgba(255,255,255,.18) !important; }
[data-testid="stSidebar"] input,
[data-testid="stSidebar"] [data-baseweb="select"] > div,
[data-testid="stSidebar"] .stSelectbox > div > div {
    background: rgba(0,50,100,.55) !important;
    border-color: rgba(0,160,220,.45) !important;
    color: #dceeff !important;
}
[data-testid="stSidebar"] label { color: #a8d4f0 !important; }
.sidebar-brand { text-align:center; margin-bottom:.55rem; }
.sidebar-brand-logos {
    display:flex; justify-content:center; align-items:center;
    gap:10px; margin-bottom:.4rem;
}
.sidebar-brand-logos img {
    background: rgba(255,255,255,.94);
    border-radius:12px; padding:5px;
    box-shadow: 0 4px 14px rgba(0,0,0,.22);
}
.sidebar-brand-note { font-size:.71rem; color:#b8d9f4 !important; }

/* ── Header encabezado ───────────────────────────── */
.brand-header {
    display: grid;
    grid-template-columns: auto 1fr auto;
    align-items: center;
    gap: 14px;
    padding: 10px 20px 12px;
    background: linear-gradient(135deg, #002a4d 0%, #005a96 100%);
    border-radius: 14px;
    margin-bottom: 14px;
    box-shadow: 0 4px 18px rgba(0,42,77,.30);
    width: 100%;
    box-sizing: border-box;
    overflow: visible;
    min-height: 86px;
    position: relative;
    z-index: 1;
}
/* Logos a la izquierda */
.brand-left {
    display: flex;
    align-items: center;
    gap: 10px;
    flex-shrink: 0;
}
.brand-left img {
    background: rgba(255,255,255,.97);
    border-radius: 10px;
    padding: 4px 6px;
    box-shadow: 0 2px 8px rgba(0,0,0,.20);
    height: 60px;
    width: auto;
    display: block;
}
/* Título central */
.brand-center {
    display: flex;
    align-items: center;
    gap: 10px;
    min-width: 0;               /* permite que el texto se trunque */
    overflow: hidden;
}
.brand-emoji-big {
    font-size: 2rem;
    flex-shrink: 0;
    line-height: 1;
    filter: drop-shadow(0 2px 4px rgba(0,0,0,.25));
}
.brand-text-col {
    min-width: 0;
    overflow: hidden;
}
.brand-title {
    font-size: clamp(1.1rem, 2.2vw, 1.65rem);
    font-weight: 800;
    color: #ffffff !important;
    line-height: 1.15;
    text-shadow: 0 1px 6px rgba(0,0,0,.3);
    white-space: normal;
    overflow: visible;
    text-overflow: clip;
}
.brand-subtitle {
    font-size: clamp(.72rem, 1.3vw, .93rem);
    color: #9dd6ff !important;
    font-weight: 600;
    margin-top: 2px;
    white-space: normal;
    overflow: visible;
    text-overflow: clip;
}
/* Pastilla derecha */
.brand-right {
    flex-shrink: 0;
    max-width: 300px;
    background: rgba(255,255,255,.13);
    border: 1px solid rgba(255,255,255,.22);
    border-radius: 9px;
    padding: 10px 13px;
    color: #e0f2ff !important;
    font-size: .76rem;
    line-height: 1.45;
    text-align: right;
    white-space: normal;
    display:flex;
    flex-direction:column;
    justify-content:center;
    min-height:64px;
}
.brand-right b  { color: #7ddcff !important; }
.brand-right .brand-creator { color: #a8d8ff !important; font-weight:700; }
/* Responsive: en pantallas pequeñas colapsa a columna */
@media (max-width: 720px) {
    .brand-header {
        grid-template-columns: 1fr;
        text-align: center;
    }
    .brand-left   { justify-content: center; }
    .brand-center { justify-content: center; }
    .brand-right  { text-align: center; }
    .brand-title, .brand-subtitle { white-space: normal; }
}

/* ── KPI bullet rows ─────────────────────────────── */
.kpi {
    padding: 3px 0 3px 10px;
    border-left: 3px solid var(--c-accent);
    margin: 2px 0;
    background: transparent;
    border-radius: 0;
    box-shadow: none;
    border-top: none; border-right: none; border-bottom: none;
}
.kpi:hover { transform: none; box-shadow: none; }
.kpi-label {
    font-size: .68rem; color: var(--c-text-muted);
    font-weight: 600; text-transform: uppercase;
    letter-spacing: .04em; margin-bottom: 1px;
}
.kpi-val { font-size: 1.05rem; font-weight: 700; color: var(--c-text); line-height: 1.2; }
.kpi-sub { font-size: .67rem; color: var(--c-text-muted); margin-top: 1px; }
.kpi-up   { color: var(--c-up);   font-size: .68rem; font-weight: 700; }
.kpi-down { color: var(--c-down); font-size: .68rem; font-weight: 700; }

/* ── Section header ──────────────────────────────── */
.sec-hdr {
    background: linear-gradient(90deg, #002a4d, #0070b8);
    color: white !important;
    padding: 5px 14px; border-radius: 6px;
    font-size: .85rem; font-weight: 700;
    margin: 10px 0 5px;
    display: flex; align-items: center; gap: 6px;
    box-shadow: 0 1px 4px rgba(0,42,77,.18);
}

/* ── Tabs ────────────────────────────────────────── */
.stTabs [data-baseweb="tab-list"] { gap: 4px; background: transparent; }
/* Compact metrics inside tabs 6 & 7 (Meteo & Salinidad) */
[data-testid="stMetricValue"] { font-size: 1.05rem !important; }
[data-testid="stMetricLabel"] { font-size: .72rem !important; }
[data-testid="stMetricDelta"] { font-size: .70rem !important; }
.stTabs [data-baseweb="tab"] {
    background: var(--c-surface);
    border-radius: 10px 10px 0 0;
    border: 1px solid var(--c-border-soft);
    padding: 8px 15px;
    font-size: .81rem; font-weight: 600;
    color: var(--c-text);
}
.stTabs [aria-selected="true"] {
    background: #002a4d !important;
    color: white !important;
    border-color: #002a4d !important;
}
/* ── Tab content panel ── */
.stTabs [data-baseweb="tab-panel"] {
    background: var(--c-surface);
    border: 1px solid var(--c-border-soft);
    border-top: none;
    border-radius: 0 10px 10px 10px;
    padding: 14px 16px;
}

/* ── Buttons ─────────────────────────────────────── */
.stButton > button {
    background: linear-gradient(135deg, #003a6e, #0072b8) !important;
    color: white !important; border: none !important;
    border-radius: 9px !important; padding: 10px 26px !important;
    font-weight: 700 !important; font-size: .87rem !important;
}
.stButton > button:hover { filter: brightness(1.1) !important; }
.stDownloadButton > button {
    background: linear-gradient(135deg, #065f46, #059669) !important;
    color: white !important; border: none !important;
    border-radius: 9px !important; font-weight: 700 !important;
}

/* ── Info / warn / success boxes ────────────────── */
.info-box {
    background: var(--c-info-bg); border: 1px solid var(--c-info-bdr);
    border-radius: 10px; padding: 11px 15px;
    font-size: .84rem; color: var(--c-text); margin: 7px 0;
}
.warn-box {
    background: var(--c-warn-bg); border: 1px solid var(--c-warn-bdr);
    border-radius: 10px; padding: 11px 15px;
    font-size: .84rem; color: var(--c-text); margin: 7px 0;
}
.success-box {
    background: var(--c-ok-bg); border: 1px solid var(--c-ok-bdr);
    border-radius: 10px; padding: 11px 15px;
    font-size: .84rem; color: var(--c-text); margin: 7px 0;
}
.ok-box {
    background: var(--c-ok-bg); border: 1px solid var(--c-ok-bdr);
    border-radius: 10px; padding: 11px 15px;
    font-size: .84rem; color: var(--c-text); margin: 7px 0;
}

/* ── Data tables ─────────────────────────────────── */
.styled-table { width: 100%; border-collapse: collapse; font-size: .82rem; }
.styled-table th {
    background: #002a4d; color: white;
    padding: 8px 11px; text-align: center;
}
.styled-table td {
    padding: 6px 11px;
    border-bottom: 1px solid var(--c-border-soft);
    color: var(--c-text);
}
.styled-table tr:nth-child(even) { background: var(--c-accent-glow); }
.styled-table tr.subtotal {
    background: rgba(0,114,184,.13); font-weight: 700; color: var(--c-text);
}
.styled-table tr.total { background: #002a4d; color: white; font-weight: 700; }

/* ── Inputs / widgets (main area) ───────────────── */
.stNumberInput input, .stTextInput input, .stSelectbox [data-baseweb="select"] > div {
    background: var(--c-surface2) !important;
    border-color: var(--c-border) !important;
    color: var(--c-text) !important;
    border-radius: 8px !important;
}

/* ── Metric ──────────────────────────────────────── */
[data-testid="stMetric"] {
    background: var(--c-surface);
    border-radius: 12px;
    padding: 10px 14px;
    border: 1px solid var(--c-border);
}
[data-testid="stMetricValue"] { color: var(--c-text) !important; }
[data-testid="stMetricLabel"] { color: var(--c-text-muted) !important; }

/* ── Dataframe ───────────────────────────────────── */
[data-testid="stDataFrame"] {
    border: 1px solid var(--c-border-soft) !important;
    border-radius: 10px !important;
    overflow: hidden;
}

/* ── Divider ─────────────────────────────────────── */
hr { border-color: var(--c-border-soft) !important; }

/* ── Footer ──────────────────────────────────────── */
.footer {
    text-align: center;
    color: var(--c-text-muted);
    font-size: .71rem;
    padding: 16px 0 6px;
    border-top: 1px solid var(--c-border-soft);
    margin-top: 26px;
}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  PLOTLY LAYOUT — adaptativo por tema
# ══════════════════════════════════════════════════════════════════════════════
import plotly.io as pio

def plotly_layout(**extra):
    """Retorna un dict de layout Plotly adaptativo para light/dark."""
    base = dict(
        plot_bgcolor  = "rgba(0,0,0,0)",   # transparente → hereda el fondo
        paper_bgcolor = "rgba(0,0,0,0)",
        font          = dict(family="Segoe UI, system-ui, sans-serif", size=12),
        margin        = dict(t=28, b=28, l=18, r=18),
        xaxis = dict(
            gridcolor = "rgba(148,163,184,.18)",
            zerolinecolor = "rgba(148,163,184,.30)",
            linecolor = "rgba(148,163,184,.30)",
            tickfont  = dict(size=11),
        ),
        yaxis = dict(
            gridcolor = "rgba(148,163,184,.18)",
            zerolinecolor = "rgba(148,163,184,.30)",
            linecolor = "rgba(148,163,184,.30)",
            tickfont  = dict(size=11),
        ),
        legend = dict(
            bgcolor      = "rgba(0,0,0,0)",
            bordercolor  = "rgba(148,163,184,.22)",
            borderwidth  = 1,
            font         = dict(size=11),
        ),
    )
    base.update(extra)
    return base

# Paleta de colores institucional ACP
PALETTE_MAIN  = ["#003a6e","#0072b8","#38b6f0","#7dd3fc","#bae6fd","#e0f2fe"]
PALETTE_SPLIT = ["#003a6e","#0072b8"]
PALETTE_KPI   = ["#003a6e","#0072b8","#15803d","#f59e0b","#ef4444","#8b5cf6"]

# ══════════════════════════════════════════════════════════════════════════════
#  CONSTANTS
# ══════════════════════════════════════════════════════════════════════════════
MESES_ES  = ["Octubre","Noviembre","Diciembre","Enero","Febrero","Marzo",
             "Abril","Mayo","Junio","Julio","Agosto","Septiembre"]
MESES_NUM = {m: n for m, n in zip(MESES_ES, [10,11,12,1,2,3,4,5,6,7,8,9])}
MESES_COL = {m: c for m, c in zip(MESES_ES,
             ["OCT","NOV","DIC","ENE","FEB","MAR","ABR","MAY","JUN","JUL","AGO","SEP"])}
CFS_TO_M3S = 0.028316847
HM3_PER_M3S_DAY = 0.0864          # 1 m3/s · día = 0.0864 hm3
CHCP_AREA_KM2 = 3306.0
EED_FLOW_M3S = (55 * 3.785 * 1000) / (24 * 3600)   # criterio de Hidroestadística
EE_HM3 = 0.2081976                                   # hm³ por esclusaje equivalente para ahorros W/X del DAILY

def hm3_to_m3s_period(hm3, dias):
    return (float(hm3) / (float(dias) * HM3_PER_M3S_DAY)) if dias else 0.0


def sig3(v, fallback="—"):
    """Formatea v con 3 cifras significativas; elimina .0 final."""
    try:
        import math as _m
        v = float(v)
        if v == 0:
            return "0"
        exp = _m.floor(_m.log10(abs(v)))
        dec = 2 - exp
        rounded = round(v, dec)
        if dec > 0:
            return f"{rounded:.{dec}f}"
        elif dec == 0:
            return f"{rounded:.0f}"
        else:
            return str(int(rounded))
    except Exception:
        return str(fallback)

def hm3_to_cfs_period(hm3, dias):
    m3s = hm3_to_m3s_period(hm3, dias)
    return (m3s / CFS_TO_M3S) if CFS_TO_M3S else 0.0

def m3s_to_cfs(m3s):
    return (float(m3s) / CFS_TO_M3S) if CFS_TO_M3S else 0.0

def m3s_to_eed(m3s):
    return (float(m3s) / EED_FLOW_M3S) if EED_FLOW_M3S else 0.0

def hm3_to_eed_period(hm3, dias):
    return m3s_to_eed(hm3_to_m3s_period(hm3, dias))

def _sync_balance_fields(datos, dias):
    """Recalcula campos derivados usando hm³ como base.
    ZZ Flush es manual; GATSPILL se interpreta como vertidos Gatún.
    """
    if not isinstance(datos, dict):
        return datos

    def g(key, default=0.0):
        try:
            v = datos.get(key, default)
            if v is None or v == "":
                return float(default)
            fv = float(v)
            import math
            return float(default) if math.isnan(fv) or math.isinf(fv) else fv
        except Exception:
            return float(default)

    pan = g("panamax_hm3")
    neo = g("neopanamax_hm3")
    # Solo recalcular esclusaje_hm3 si no existe un valor oficial del DAILY
    if pan + neo > 0 and g("esclusaje_hm3") < 0.001:
        datos["esclusaje_hm3"] = round(pan + neo, 3)
    datos["esclusaje_m3s"] = round(hm3_to_m3s_period(datos.get("esclusaje_hm3", 0.0), dias), 3) if dias else 0.0
    datos["esclusaje_cfs"] = round(hm3_to_cfs_period(datos.get("esclusaje_hm3", 0.0), dias), 3) if dias else 0.0
    datos["panamax_m3s"] = round(hm3_to_m3s_period(pan, dias), 3) if dias else 0.0
    datos["panamax_cfs"] = round(hm3_to_cfs_period(pan, dias), 3) if dias else 0.0
    datos["neopanamax_m3s"] = round(hm3_to_m3s_period(neo, dias), 3) if dias else 0.0
    datos["neopanamax_cfs"] = round(hm3_to_cfs_period(neo, dias), 3) if dias else 0.0

    pot_g = g("potabilizacion_gat_hm3")
    pot_a = g("potabilizacion_alh_hm3")
    datos["potabilizacion_hm3"] = round(pot_g + pot_a, 3)
    datos["potabilizacion_m3s"] = round(hm3_to_m3s_period(datos["potabilizacion_hm3"], dias), 3) if dias else 0.0
    datos["potabilizacion_cfs"] = round(hm3_to_cfs_period(datos["potabilizacion_hm3"], dias), 3) if dias else 0.0
    datos["potabilizacion_gat_m3s"] = round(hm3_to_m3s_period(pot_g, dias), 3) if dias else 0.0
    datos["potabilizacion_gat_cfs"] = round(hm3_to_cfs_period(pot_g, dias), 3) if dias else 0.0
    datos["potabilizacion_alh_m3s"] = round(hm3_to_m3s_period(pot_a, dias), 3) if dias else 0.0
    datos["potabilizacion_alh_cfs"] = round(hm3_to_cfs_period(pot_a, dias), 3) if dias else 0.0

    con_g = g("concesiones_gat_hm3")
    con_a = g("concesiones_alh_hm3")
    # Solo LEAK GAT sale del sistema; LEAK MAD es trasvase interno Alh→Gat
    datos["concesiones_hm3"]     = round(con_g, 3)           # solo lo que sale del sistema
    datos["concesiones_gat_hm3"] = round(con_g, 3)           # LEAK GAT → balance
    datos["concesiones_m3s"] = round(hm3_to_m3s_period(datos["concesiones_hm3"], dias), 3) if dias else 0.0
    datos["concesiones_cfs"] = round(hm3_to_cfs_period(datos["concesiones_hm3"], dias), 3) if dias else 0.0
    datos["concesiones_gat_m3s"] = round(hm3_to_m3s_period(con_g, dias), 3) if dias else 0.0
    datos["concesiones_gat_cfs"] = round(hm3_to_cfs_period(con_g, dias), 3) if dias else 0.0
    datos["concesiones_alh_m3s"] = round(hm3_to_m3s_period(con_a, dias), 3) if dias else 0.0
    datos["concesiones_alh_cfs"] = round(hm3_to_cfs_period(con_a, dias), 3) if dias else 0.0

    vgat = g("vertidos_gat_hm3")
    datos["vertidos_gat_m3s"] = round(hm3_to_m3s_period(vgat, dias), 3) if dias else 0.0
    datos["vertidos_gat_cfs"] = round(hm3_to_cfs_period(vgat, dias), 3) if dias else 0.0

    vmad = g("vertidos_mad_ops_hm3")
    datos["vertidos_mad_ops_m3s"] = round(hm3_to_m3s_period(vmad, dias), 3) if dias else 0.0
    datos["vertidos_mad_ops_cfs"] = round(hm3_to_cfs_period(vmad, dias), 3) if dias else 0.0

    zz = g("zzflush_auto_hm3")
    datos["zzflush_auto_m3s"] = round(hm3_to_m3s_period(zz, dias), 3) if dias else 0.0
    datos["zzflush_auto_cfs"] = round(hm3_to_cfs_period(zz, dias), 3) if dias else 0.0

    datos["aporte_total_m3s"] = round(hm3_to_m3s_period(g("aporte_total_hm3"), dias), 3) if dias else g("aporte_total_m3s")
    datos["aporte_neto_hm3"] = round(g("aporte_total_hm3") - g("evaporacion_hm3"), 3)
    datos["aporte_neto_m3s"] = round(hm3_to_m3s_period(datos["aporte_neto_hm3"], dias), 3) if dias else 0.0
    datos["aporte_total_cfs"] = round(hm3_to_cfs_period(g("aporte_total_hm3"), dias), 3) if dias else 0.0
    datos["aporte_neto_cfs"] = round(hm3_to_cfs_period(datos["aporte_neto_hm3"], dias), 3) if dias else 0.0
    datos["evaporacion_cfs"] = round(hm3_to_cfs_period(g("evaporacion_hm3"), dias), 3) if dias else 0.0
    datos["hidro_gatun_cfs"] = round(hm3_to_cfs_period(g("hidro_gatun_hm3"), dias), 3) if dias else 0.0
    datos["trasvase_cfs"] = round(hm3_to_cfs_period(g("trasvase_hm3"), dias), 3) if dias else 0.0

    total_sal = (
        g("esclusaje_hm3") + g("potabilizacion_hm3") + g("concesiones_gat_hm3") +
        g("evaporacion_hm3") + g("hidro_gatun_hm3") + g("zzflush_auto_hm3") +
        g("vertidos_gat_hm3")
        # concesiones_gat solo (Gatún) — concesiones_alh/LEAK MAD va en trasvase
    )
    datos["total_salidas_hm3"] = round(total_sal, 3)
    datos["total_salidas_m3s"] = round(hm3_to_m3s_period(total_sal, dias), 3) if dias else 0.0
    datos["total_salidas_cfs"] = round(hm3_to_cfs_period(total_sal, dias), 3) if dias else 0.0
    datos["excedente_hm3"] = round(g("aporte_total_hm3") - total_sal, 3)  # balance: entradas_totales − salidas_totales

    # Evaporación diaria (mm/día) y por embalse, siguiendo la Hidroestadística
    area_g = g("sup_prom_gatun_km2")
    area_a = g("sup_prom_alh_km2")
    evap_g = g("evap_gatun_hm3")
    evap_a = g("evap_alh_hm3")
    if dias and area_g > 0:
        datos["evap_gatun_hm3_dia"] = round(evap_g / dias, 3)
        datos["evap_gatun_mm_dia"] = round(evap_g * 1000 / (area_g * dias), 3)
    if dias and area_a > 0:
        datos["evap_alh_hm3_dia"] = round(evap_a / dias, 3)
        datos["evap_alh_mm_dia"] = round(evap_a * 1000 / (area_a * dias), 3)
    area_tot = area_g + area_a
    if dias and area_tot > 0:
        datos["evaporacion_diaria_mm"] = round(g("evaporacion_hm3") * 1000 / (area_tot * dias), 3)

    return datos

# ══════════════════════════════════════════════════════════════════════════════
#  DEFINICIÓN DE PERÍODOS (meses y días)
# ══════════════════════════════════════════════════════════════════════════════

# Meses y sus días promedio (para series históricas)
MES_DIAS = {
    "OCT":31,"NOV":30,"DIC":31,
    "ENE":31,"FEB":28.25,"MAR":31,
    "ABR":30,"MAY":31,"JUN":30,
    "JUL":31,"AGO":31,"SEP":30,
}
# Mapeo número de mes → columna abreviada
MES_N_TO_COL = {10:"OCT",11:"NOV",12:"DIC",1:"ENE",2:"FEB",3:"MAR",
                4:"ABR",5:"MAY",6:"JUN",7:"JUL",8:"AGO",9:"SEP"}
# Mapeo columna → offset de año respecto al Año Fiscal (AF):
#   -1 significa: está en el año calendario AF-1
COL_ANIO_OFF = {
    "OCT":-1,"NOV":-1,"DIC":-1,
    "ENE":0,"FEB":0,"MAR":0,"ABR":0,"MAY":0,"JUN":0,"JUL":0,"AGO":0,"SEP":0,
}

# Lista de (mes_num, año_offset) para cada tipo de período
PERIODO_MESES = {
    "Q1":  [(10,-1),(11,-1),(12,-1)],
    "Q2":  [(1,0),(2,0),(3,0)],
    "Q3":  [(4,0),(5,0),(6,0)],
    "Q4":  [(7,0),(8,0),(9,0)],
    "S1":  [(10,-1),(11,-1),(12,-1),(1,0),(2,0),(3,0)],
    "S2":  [(4,0),(5,0),(6,0),(7,0),(8,0),(9,0)],
    "Anual":[(10,-1),(11,-1),(12,-1),(1,0),(2,0),(3,0),
             (4,0),(5,0),(6,0),(7,0),(8,0),(9,0)],
}

# Días promedio del período
PERIODO_DIAS = {
    "Q1":92,"Q2":90.25,"Q3":91,"Q4":92,
    "S1":182.25,"S2":183,"Anual":365.25,
}

def _periodo_key(tipo, trimestre=None, semestre=None):
    """Retorna clave interna ('Q1','Q2','S1','Anual', etc.)."""
    if tipo == "Mensual":
        return "Mensual"
    if tipo == "Trimestral":
        return (trimestre or "Q1").split()[0]
    if tipo == "Semestral":
        return "S1" if "Primer" in (semestre or "Primer") else "S2"
    return "Anual"

def _label_periodo(tipo, mes_sel=None, anio_sel=None, trimestre=None, semestre=None):
    """Genera la etiqueta legible del período."""
    if tipo == "Mensual":
        return f"{mes_sel} {anio_sel}"
    if tipo == "Trimestral":
        q = (trimestre or "Q1").split()[0]
        nombres_q = {"Q1":"I Trimestre (Oct–Dic)","Q2":"II Trimestre (Ene–Mar)",
                     "Q3":"III Trimestre (Abr–Jun)","Q4":"IV Trimestre (Jul–Sep)"}
        return f"{nombres_q.get(q, q)} AF{anio_sel}"
    if tipo == "Semestral":
        sem = "Primer" if "Primer" in (semestre or "Primer") else "Segundo"
        return f"{sem} Semestre AF{anio_sel}"
    return f"Año Fiscal {anio_sel}"

def _badge_periodo(tipo):
    badges = {"Mensual":"Informe Mensual","Trimestral":"Informe Trimestral",
              "Semestral":"Informe Semestral","Anual":"Informe Anual"}
    return badges.get(tipo, tipo)

def _get_meses_periodo(tipo, anio_sel, trimestre=None, semestre=None):
    """Retorna lista de (mes_n, año_real) para el período."""
    pk = _periodo_key(tipo, trimestre, semestre)
    meses_off = PERIODO_MESES.get(pk, [])
    return [(mn, anio_sel + off) for mn, off in meses_off]

def _fuentes_sheet_target(tipo, anio_sel, mes_sel=None, trimestre=None, semestre=None):
    """Genera el nombre objetivo de hoja en Fuentes_Agua."""
    pk = _periodo_key(tipo, trimestre, semestre)
    ay = str(anio_sel)
    ap = str(anio_sel-1)
    y2 = ay[2:]; y1 = ap[2:]
    if tipo == "Mensual":
        return f"{mes_sel[:3].capitalize()}{y2}"
    targets = {
        "Q1": f"Oct{y1}-Dic{y1}",
        "Q2": f"Ene{y2}-Mar{y2}",
        "Q3": f"Abr{y2}-Jun{y2}",
        "Q4": f"Jul{y2}-Sep{y2}",
        "S1": f"Oct{y1}-Mar{y2}",
        "S2": f"Abr{y2}-Sep{y2}",
        "Anual": f"Oct{y1}-Sep{y2}",
    }
    return targets.get(pk, f"Oct{y1}-Sep{y2}")

def _plural_mes(mes):
    """Pluraliza nombre de mes en español: marzo→marzos, abril→abriles."""
    mes = mes.strip().lower()
    if mes.endswith(("a", "e", "i", "o", "u")):
        return mes + "s"
    return mes + "es"

def _meteo_rank_label_default(tipo, mes_sel=None):
    """Etiqueta para ranking de precipitación adaptable al período activo."""
    if tipo == "Mensual" and mes_sel:
        return f"{_plural_mes(mes_sel)} más húmedos"
    if tipo == "Trimestral":
        return "trimestres más húmedos"
    if tipo == "Semestral":
        return "semestres más húmedos"
    if tipo == "Anual":
        return "años fiscales más húmedos"
    return "períodos más húmedos"

def _periodo_registro_n_anios(periodo_txt):
    """Calcula años inclusivos desde texto tipo 1950-2026; retorna 0 si no aplica."""
    try:
        nums = re.findall(r"\d{4}", str(periodo_txt or ""))
        if len(nums) >= 2:
            a, b = int(nums[0]), int(nums[-1])
            return max(0, b - a + 1)
    except Exception:
        pass
    return 0

def _anios_registro_incluyendo_actual(periodo_txt, anio_actual=None, fallback=0):
    """Años de registro para el relato histórico; incluye el año evaluado si el período histórico termina antes."""
    try:
        nums = re.findall(r"\d{4}", str(periodo_txt or ""))
        if len(nums) >= 2:
            a, b = int(nums[0]), int(nums[-1])
            y = int(anio_actual) if anio_actual else b
            return max(int(fallback or 0), max(b, y) - a + 1)
    except Exception:
        pass
    try:
        return int(fallback or 0)
    except Exception:
        return 0

def _posicion_historica_texto(pos_txt, tipo=None, mes_sel=None):
    """Convierte rankings tipo '1 de 128' en una frase más natural para el informe.
    Solo convierte posición 1 a "el X más húmedo". Para otras posiciones retorna
    "N de M" limpio (ej: "4 de 128").
    """
    s = str(pos_txt or "").strip()
    if not s:
        return ""
    m = re.match(r"^\s*(\d+)\s+de\s+(\d+)", s, flags=re.IGNORECASE)
    if m:
        pos, total = int(m.group(1)), int(m.group(2))
        if pos == 1 and tipo == "Mensual" and mes_sel:
            return f"el {str(mes_sel).strip().lower()} más húmedo"
        return f"{pos} de {total}"
    return s

def _safe_label_file(tipo, mes_sel, anio_sel, trimestre=None, semestre=None):
    """Genera sufijo seguro para nombre de archivo exportado."""
    pk = _periodo_key(tipo, trimestre, semestre)
    if tipo == "Mensual":
        return f"{mes_sel}_{anio_sel}"
    suffixes = {
        "Q1":f"Q1_Oct-Dic_{anio_sel-1}","Q2":f"Q2_Ene-Mar_{anio_sel}",
        "Q3":f"Q3_Abr-Jun_{anio_sel}","Q4":f"Q4_Jul-Sep_{anio_sel}",
        "S1":f"PrimerSemestre_AF{anio_sel}","S2":f"SegundoSemestre_AF{anio_sel}",
        "Anual":f"AnioFiscal_{anio_sel}",
    }
    return suffixes.get(pk, f"AF{anio_sel}")




# ══════════════════════════════════════════════════════════════════════════════
#  HELPER — cargar datos
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_data(show_spinner=False)
def load_daily(file_bytes: bytes) -> pd.DataFrame:
    """Carga la hoja DAILYINPUT.
    Acepta bytes en lugar de un objeto-file para evitar el bug de Streamlit
    donde @st.cache_data agota el puntero del UploadedFile antes de leerlo.
    Patrón idéntico al de load_hidro_audit."""
    df = pd.read_excel(io.BytesIO(file_bytes), sheet_name="DAILYINPUT")
    df.columns = [str(c).strip() for c in df.columns]  # normalizar encabezados
    df["ACTDATE"] = pd.to_datetime(df["ACTDATE"], errors="coerce")
    # Pre-coerce all numeric-expected columns to avoid silent object dtype issues
    _numeric_cols = ["ACTGATEL","ACTMADEL","MADMWH","GATMWH","MADMW","GATMW","MADMCF","GATMCF",
                     "MADSPILL","GATSPILL","LEAK MAD","LEAK GAT","MUNIC GAT","MUNIC MAD",
                     "PMLOCKMCF","GATLOCKMCF","ACLOCKMCF","CCLLOCKMCF",
                     "NUMLOCKPM","NUMLOCKGAT","NUMLOCKACL","NUMLOCKCCL",
                     "ESCLUSAJES, hm3","TOFCHCP (hm3)","TOFGL (hm3)","TOFMD (hm3)"]
    for _nc in _numeric_cols:
        if _nc in df.columns:
            df[_nc] = pd.to_numeric(df[_nc], errors="coerce")
    return df

@st.cache_data(show_spinner=False)
def load_hidro(f) -> dict:
    """Retorna dict con DataFrames de cada hoja de Hidroestadística."""
    wb = {}
    for sh in ["Hidroestadística 1","Hidroestadística 2"]:
        try:
            wb[sh] = pd.read_excel(f, sheet_name=sh, header=None)
        except Exception:
            wb[sh] = pd.DataFrame()
    return wb

@st.cache_data(show_spinner=False)
def load_hidro_audit(file_bytes: bytes) -> dict:
    """Parsea dinámicamente las hojas de Hidroestadística para cualquier período cargado."""
    import unicodedata
    bio1 = io.BytesIO(file_bytes)
    bio2 = io.BytesIO(file_bytes)
    wb_vals = openpyxl.load_workbook(bio1, data_only=True, read_only=True)
    wb_form = openpyxl.load_workbook(bio2, data_only=False, read_only=True)

    def norm(s):
        s = "" if s is None else str(s)
        s = unicodedata.normalize("NFD", s)
        s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
        s = re.sub(r"\s+", " ", s).strip().lower()
        return s

    out = {"periodo_grupo": None, "periodo_detalle": None, "sheet1_rows": [], "sheet2_rows": []}

    if "Hidroestadística 1" in wb_vals.sheetnames:
        wsv = wb_vals["Hidroestadística 1"]
        wsf = wb_form["Hidroestadística 1"]
        out["periodo_grupo"] = wsv["C2"].value
        out["periodo_detalle"] = wsv["C3"].value
        for r in range(4, wsv.max_row + 1):
            label = wsv.cell(r, 2).value
            vals = [wsv.cell(r, c).value for c in range(3, 9)]
            fmls = [wsf.cell(r, c).value for c in range(3, 9)]
            if label is None and not any(v is not None for v in vals):
                continue
            out["sheet1_rows"].append({
                "row": r, "label": label, "label_norm": norm(label),
                "value1": vals[0], "unit1": vals[1],
                "value2": vals[2], "unit2": vals[3],
                "value3": vals[4], "unit3": vals[5],
                "formula1": fmls[0] if isinstance(fmls[0], str) and fmls[0].startswith("=") else None,
                "formula2": fmls[2] if isinstance(fmls[2], str) and fmls[2].startswith("=") else None,
                "formula3": fmls[4] if isinstance(fmls[4], str) and fmls[4].startswith("=") else None,
            })

    if "Hidroestadística 2" in wb_vals.sheetnames:
        wsv = wb_vals["Hidroestadística 2"]
        wsf = wb_form["Hidroestadística 2"]
        for r in range(4, wsv.max_row + 1):
            label = wsv.cell(r, 2).value
            val = wsv.cell(r, 3).value
            unit = wsv.cell(r, 4).value
            pct = wsv.cell(r, 5).value
            if label is None and val is None and pct is None:
                continue
            out["sheet2_rows"].append({
                "row": r, "label": label, "label_norm": norm(label),
                "value": val, "unit": unit, "pct": pct,
                "formula_value": wsf.cell(r, 3).value if isinstance(wsf.cell(r, 3).value, str) and str(wsf.cell(r, 3).value).startswith("=") else None,
                "formula_pct": wsf.cell(r, 5).value if isinstance(wsf.cell(r, 5).value, str) and str(wsf.cell(r, 5).value).startswith("=") else None,
            })
    return out

def hidro_row(rows, label, unit_hint=None, row_min=None):
    import unicodedata
    def norm(s):
        s = "" if s is None else str(s)
        s = unicodedata.normalize("NFD", s)
        s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
        s = re.sub(r"\s+", " ", s).strip().lower()
        return s
    label_n = norm(label)
    cands = [r for r in rows if r.get("label_norm") == label_n]
    if unit_hint:
        unit_n = norm(unit_hint)
        cands = [r for r in cands if norm(r.get("unit1") or r.get("unit") or "") == unit_n]
    if row_min is not None:
        cands = [r for r in cands if r.get("row", 0) >= row_min]
    return cands[0] if cands else {}

def hidro_num(rows, label, which="value1", unit_hint=None, row_min=None, default=None):
    r = hidro_row(rows, label, unit_hint=unit_hint, row_min=row_min)
    if not r:
        return default
    try:
        v = r.get(which)
        return float(v) if v is not None else default
    except Exception:
        return default

@st.cache_data(show_spinner=False)
def load_hist(f_alh, f_gat, f_chcp, f_netfl=None):
    """Lee archivos históricos y normaliza todas las series a m³/s.
    - 03_QAVGMCHCP: aportes totales históricos CHCP
    - 04_GATNETFL:  aportes netos históricos (normalmente en CFS)
    Devuelve dataframes con columnas: AÑO, OCT..SEP, ANUAL (si existe), en m³/s.
    """
    import unicodedata as _ud
    import importlib, subprocess, sys, tempfile, shutil, os

    def _ensure(pkg_import, pkg_install=None):
        try:
            return importlib.import_module(pkg_import)
        except Exception:
            if pkg_install:
                try:
                    subprocess.run([sys.executable, "-m", "pip", "install", pkg_install, "-q"],
                                   capture_output=True, timeout=90)
                    return importlib.import_module(pkg_import)
                except Exception:
                    return None
            return None

    _xlrd = _ensure("xlrd", "xlrd>=2.0.1")
    _mso = _ensure("msoffcrypto", "msoffcrypto-tool")

    def _to_bytes(f):
        if f is None:
            return None
        if isinstance(f, (bytes, bytearray)):
            return bytes(f)
        if hasattr(f, "getvalue"):
            try:
                return f.getvalue()
            except Exception:
                pass
        if hasattr(f, "read"):
            try:
                if hasattr(f, "seek"):
                    f.seek(0)
                return f.read()
            except Exception:
                pass
        try:
            return Path(str(f)).read_bytes()
        except Exception:
            return None

    def _src_name(f):
        if f is None:
            return ""
        name = getattr(f, "name", "")
        if name:
            return str(name)
        return str(f)

    def _norm(s):
        s = _ud.normalize("NFD", str(s).strip())
        s = "".join(c for c in s if _ud.category(c) != "Mn")
        return s.upper().replace(" ", "")

    _MES_MAP = {
        "ENE": "ENE", "ENERO": "ENE", "JAN": "ENE", "JANUARY": "ENE",
        "FEB": "FEB", "FEBRERO": "FEB", "FEBRUARY": "FEB",
        "MAR": "MAR", "MARZO": "MAR", "MARCH": "MAR",
        "ABR": "ABR", "ABRIL": "ABR", "APR": "ABR", "APRIL": "ABR",
        "MAY": "MAY", "MAYO": "MAY",
        "JUN": "JUN", "JUNIO": "JUN", "JUNE": "JUN",
        "JUL": "JUL", "JULIO": "JUL", "JULY": "JUL",
        "AGO": "AGO", "AGOSTO": "AGO", "AUG": "AGO", "AUGUST": "AGO",
        "SEP": "SEP", "SEPTIEMBRE": "SEP", "SEPTEMBER": "SEP",
        "OCT": "OCT", "OCTUBRE": "OCT", "OCTOBER": "OCT",
        "NOV": "NOV", "NOVIEMBRE": "NOV", "NOVEMBER": "NOV",
        "DIC": "DIC", "DICIEMBRE": "DIC", "DEC": "DIC", "DECEMBER": "DIC",
    }
    month_cols = ["ENE","FEB","MAR","ABR","MAY","JUN","JUL","AGO","SEP","OCT","NOV","DIC"]

    def _norm_month(c):
        s = _norm(c).replace(".", "")
        return _MES_MAP.get(s)

    def _decrypt_if_needed(raw):
        if not raw:
            return raw, None
        if raw[:2] != b'\xD0\xCF':
            return raw, None
        if _mso is None:
            return raw, "Workbook protegido y msoffcrypto no disponible"
        try:
            office = _mso.OfficeFile(io.BytesIO(raw))
            office.load_key(password='VelvetSweatshop')
            out = io.BytesIO()
            office.decrypt(out)
            return out.getvalue(), None
        except Exception as e:
            return raw, str(e)

    def _read_xls_via_xlrd(raw, sheet_name, header_row):
        if _xlrd is None:
            raise RuntimeError("xlrd no disponible")
        book = _xlrd.open_workbook(file_contents=raw)
        sh = book.sheet_by_name(sheet_name) if isinstance(sheet_name, str) else book.sheet_by_index(int(sheet_name))
        rows = [sh.row_values(r) for r in range(sh.nrows)]
        header = [str(x).strip() for x in rows[header_row]]
        data = rows[header_row + 1:]
        import pandas as pd
        return pd.DataFrame(data, columns=header)

    def _read_xlsx_with_pandas(raw, sheet_name, header_row):
        return pd.read_excel(io.BytesIO(raw), sheet_name=sheet_name, header=header_row, engine="openpyxl")

    def _convert_xls_to_xlsx_bytes(raw, src_name):
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            return None
        with tempfile.TemporaryDirectory() as td:
            inp = Path(td) / (Path(src_name or 'tmp').stem + '.xls')
            outdir = Path(td) / 'out'
            outdir.mkdir(exist_ok=True)
            inp.write_bytes(raw)
            try:
                subprocess.run([soffice, '--headless', '--convert-to', 'xlsx', '--outdir', str(outdir), str(inp)],
                               capture_output=True, timeout=120)
                outs = list(outdir.glob('*.xlsx'))
                if outs:
                    return outs[0].read_bytes()
            except Exception:
                return None
        return None

    def _standardize(df, src_name, sheet_name, role='generic'):
        if df is None or df.empty:
            return pd.DataFrame()
        df = df.copy()
        df.columns = [str(c).strip() for c in df.columns]

        # localizar columna de año
        year_col = None
        for c in df.columns:
            cn = _norm(c)
            if cn in {"AÑO", "ANO", "YEAR", "YEARS", "YRS", "YR", "Y"} or cn.startswith(("AÑO", "ANO", "YEAR", "YRS")):
                year_col = c
                break
        if year_col is None:
            return pd.DataFrame()
        df = df.rename(columns={year_col: 'AÑO'})
        df['AÑO'] = pd.to_numeric(df['AÑO'], errors='coerce')
        df = df.dropna(subset=['AÑO'])
        df = df[(df['AÑO'] >= 1800) & (df['AÑO'] <= 2100)].reset_index(drop=True)

        # Regla específica para 03_QAVGMCHCP:
        # el usuario requiere usar el período 1898–2025 del archivo histórico CHCP.
        src_lower_tmp = (str(src_name) + ' ' + str(sheet_name)).lower()
        if ('qavgmchcp' in src_lower_tmp) or ('chcp' in src_lower_tmp and 'qavgcdlrev' in src_lower_tmp):
            df = df[(df['AÑO'] >= 1898) & (df['AÑO'] <= 2025)].reset_index(drop=True)

        if df.empty:
            return pd.DataFrame()

        rename = {}
        for c in df.columns:
            m = _norm_month(c)
            if m:
                rename[c] = m
            elif _norm(c) in {"ANUAL", "ANNUAL", "AVERAGE", "PROMEDIO"}:
                rename[c] = 'ANUAL'
        df = df.rename(columns=rename)

        # En GATNETFL/MCM pueden venir columnas duplicadas auxiliares (hm3, %, ranking).
        # Para la serie histórica neta solo necesitamos AÑO + meses (+ anual si existe),
        # evitando columnas duplicadas que rompan la estandarización.
        if role == 'netfl' and any(tag in src_lower_tmp for tag in ['mcm', 'hm3']):
            extra_cols = []
        else:
            extra_cols = [c for c in ['hm3', 'Dif%', '%ACUM.', 'Clas. Aporte', 'Umbral inf', 'Umbral sup'] if c in df.columns]
        keep = ['AÑO'] + [c for c in month_cols + ['ANUAL'] if c in df.columns] + extra_cols
        if len([c for c in keep if c not in ['AÑO'] + extra_cols]) < 2:
            return pd.DataFrame()
        df = df[keep].copy()
        for c in keep:
            if c == 'AÑO':
                continue
            if str(c).strip() == 'Clas. Aporte':
                df[c] = df[c].astype(str).replace({'nan': '', 'None': ''}).str.strip().str.rstrip(',')
            else:
                df[c] = pd.to_numeric(df[c], errors='coerce')
        df = df.dropna(subset=['AÑO'])

        src_lower = (str(src_name) + ' ' + str(sheet_name)).lower()
        is_cfs = any(tag in src_lower for tag in ['cfs', 'pie3/s'])
        is_m3s = any(tag in src_lower for tag in ['mcs', 'm3s', 'm3/s'])
        is_hm3 = any(tag in src_lower for tag in ['mcm', 'hm3'])

        # Regla especial para GATNETFL:
        # para aportes netos históricos se prefiere la hoja de volúmenes (MCM/hm3),
        # porque el usuario audita el promedio histórico neto en hm³ y luego lo
        # convierte a m³/s usando los días reales del período seleccionado.
        if role == 'netfl' and is_hm3:
            out_df = df.reset_index(drop=True)
            out_df.attrs['value_unit'] = 'hm3'
            return out_df

        if ('gatnetfl' in src_lower) and not is_m3s and not is_hm3:
            # Si no hay hoja MCM/hm3, usar CFS como respaldo.
            is_cfs = True
        if is_cfs:
            for c in df.columns:
                if c != 'AÑO':
                    df[c] = pd.to_numeric(df[c], errors='coerce') * CFS_TO_M3S
        out_df = df.reset_index(drop=True)
        out_df.attrs['value_unit'] = 'm3s'
        return out_df

    def _read_one(f, role='generic'):
        raw = _to_bytes(f)
        src_name = _src_name(f)
        if not raw:
            return pd.DataFrame(), "Archivo vacío o no legible"

        raw2, dec_err = _decrypt_if_needed(raw)
        raw = raw2 or raw

        is_zip_xlsx = raw[:2] == b'PK'
        is_cfb = raw[:2] == b'\xD0\xCF'

        preferred = []
        if role == 'netfl' or 'gatnetfl' in src_name.lower():
            preferred = ['MCM', 'CFS', 'MCF', 'G-CFS', 'CFS ', 'G-Current', 0]
        elif role == 'chcp' or 'qavgmchcp' in src_name.lower() or 'chcp' in src_name.lower():
            preferred = ['QAVGCDLREV (mcs)', 'QAVGCDLREV (m3s)', 'QAVGCDLREV (cfs)', 'MCS', 'm3s', 'CFS', 0]
        else:
            preferred = ['QAVGCDLREV (mcs)', 'QAVGCDLREV (m3s)', 'QAVGCDLREV (cfs)', 'MCS', 'm3s', 'CFS', 0]

        sheet_names = []
        xlsx_bytes = None
        if is_zip_xlsx:
            try:
                wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
                sheet_names = wb.sheetnames
                wb.close()
            except Exception:
                sheet_names = []
        elif is_cfb:
            if _xlrd is not None:
                try:
                    book = _xlrd.open_workbook(file_contents=raw)
                    sheet_names = book.sheet_names()
                except Exception:
                    sheet_names = []
            if not sheet_names:
                xlsx_bytes = _convert_xls_to_xlsx_bytes(raw, src_name)
                if xlsx_bytes:
                    try:
                        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), read_only=True, data_only=True)
                        sheet_names = wb.sheetnames
                        wb.close()
                    except Exception:
                        sheet_names = []
        else:
            xlsx_bytes = raw

        if not sheet_names:
            return pd.DataFrame(), f"No se pudieron enumerar hojas. {dec_err or ''}".strip()

        order = []
        for s in preferred:
            if s == 0:
                if sheet_names:
                    order.append(sheet_names[0])
            elif s in sheet_names:
                order.append(s)
        for s in sheet_names:
            if s not in order:
                order.append(s)

        last_diag = []
        fallback_df = None
        fallback_err = None
        for sh in order:
            for header_row in [2, 0, 1, 3, 4, 5]:
                try:
                    if is_zip_xlsx or xlsx_bytes:
                        df0 = _read_xlsx_with_pandas(xlsx_bytes or raw, sh, header_row)
                    else:
                        df0 = _read_xls_via_xlrd(raw, sh, header_row)
                    df = _standardize(df0, src_name, sh, role=role)
                    if not df.empty and sum(1 for c in df.columns if c in month_cols) >= 6:
                        if role == 'netfl':
                            if str(getattr(df, 'attrs', {}).get('value_unit', 'm3s')).lower() == 'hm3':
                                return df, None
                            if fallback_df is None:
                                fallback_df = df
                        else:
                            return df, None
                    else:
                        last_diag.append(f"{sh}@h{header_row}: sin columnas estándar")
                except Exception as e:
                    last_diag.append(f"{sh}@h{header_row}: {e}")
        if fallback_df is not None:
            return fallback_df, fallback_err
        return pd.DataFrame(), " | ".join(last_diag[:8])

    alh, err_alh = _read_one(f_alh, 'alh') if f_alh is not None else (pd.DataFrame(), None)
    gat, err_gat = _read_one(f_gat, 'gat') if f_gat is not None else (pd.DataFrame(), None)
    chcp, err_chcp = _read_one(f_chcp, 'chcp') if f_chcp is not None else (pd.DataFrame(), None)
    netfl, err_netfl = _read_one(f_netfl, 'netfl') if f_netfl is not None else (pd.DataFrame(), None)
    errors = [e for e in [err_alh, err_gat, err_chcp, err_netfl] if e]
    return alh, gat, chcp, netfl, errors

@st.cache_data(show_spinner=False)
def load_fuentes(f):
    """Lee las 3 hojas de Fuentes de Agua."""
    out = {}
    for sh in ["Mar26","Ene26-Mar26","Oct25-Mar26"]:
        try:
            out[sh] = pd.read_excel(f, sheet_name=sh, header=None)
        except Exception:
            out[sh] = pd.DataFrame()
    return out

# ──────────────────────────────────────────────────────────────────────────────
def _v(df, label, col=2):
    """Extrae valor de Hidroestadística por etiqueta."""
    try:
        mask = df.iloc[:, 1].astype(str).str.strip().str.lower() == label.lower()
        row = df[mask]
        if row.empty:
            return None
        return float(row.iloc[0, col])
    except Exception:
        return None

def mes_col(mes_es):
    return MESES_COL.get(mes_es, "MAR")

# ══════════════════════════════════════════════════════════════════════════════
#  AUTO-CÁLCULO desde DAILY
# ══════════════════════════════════════════════════════════════════════════════

def _empty_period_result(dias: int):
    keys_zero = [
        "aporte_total_hm3","aporte_total_m3s","aporte_total_cfs","evaporacion_hm3","evaporacion_m3s","evaporacion_cfs","evap_gatun_hm3","evap_alh_hm3",
        "evaporacion_diaria_mm","sup_prom_gatun_km2","sup_prom_alh_km2","evap_gatun_hm3_dia","evap_alh_hm3_dia","evap_gatun_mm_dia","evap_alh_mm_dia",
        "aporte_neto_hm3","aporte_neto_m3s","aporte_neto_cfs","aporte_gl_hm3","aporte_md_hm3","nivel_gatun_pies","nivel_gatun_m",
        "nivel_alh_pies","esclusaje_hm3","esclusaje_m3s","esclusaje_cfs","panamax_hm3","panamax_m3s","panamax_cfs","neopanamax_hm3","neopanamax_m3s","neopanamax_cfs","panamax_pct","neopanamax_pct",
        "uso_prom_diario_hm3","uso_prom_diario_m3s","agua_panamax_trans","agua_neo_trans","potabilizacion_hm3","potabilizacion_m3s","potabilizacion_cfs",
        "potabilizacion_gat_hm3","potabilizacion_gat_m3s","potabilizacion_gat_cfs","potabilizacion_alh_hm3","potabilizacion_alh_m3s","potabilizacion_alh_cfs",
        "concesiones_hm3","concesiones_m3s","concesiones_cfs","concesiones_gat_hm3","concesiones_gat_m3s","concesiones_gat_cfs","concesiones_alh_hm3","concesiones_alh_m3s","concesiones_alh_cfs",
        "trasvase_hm3","trasvase_m3s","trasvase_cfs","trasvase_pct_esc","trasvase_fugas_hm3","trasvase_total_hm3","trasvase_total_m3s","energia_madden_mw","energia_gatun_mw","energia_madden_mwh_dia",
        "energia_gatun_mwh_dia","hidro_gatun_hm3","hidro_gatun_cfs","hidro_madden_hm3","ahorro_panamax_hm3","ahorro_neopanamax_hm3",
        "ahorro_total_hm3","ahorro_lamina_pies","zzflush_auto_hm3","zzflush_auto_m3s","zzflush_auto_cfs","vertidos_mad_ops_hm3","vertidos_mad_ops_m3s","vertidos_mad_ops_cfs","vertidos_gat_hm3","vertidos_gat_m3s","vertidos_gat_cfs",
        "vertidos_mad_hm3","alhajuela_hm3","alhajuela_pct","gatun_hm3","gatun_pct","total_salidas_hm3","total_salidas_m3s","total_salidas_cfs","movimientos_operativos_hm3","excedente_hm3",
        "salinidad_spv","salinidad_dec_pct","precipitacion_mm","precipitacion_hist_mm","precip_rank","precip_rank_total",
        "hist_total_hm3","hist_total_m3s","hist_rank","hist_rank_total",
        "agua_disp_alh_mm3","agua_disp_gat_mm3","agua_disp_total_mm3"
    ]
    r = {k: 0.0 for k in keys_zero}
    r["transitos_panamax"] = 0
    r["transitos_neopanamax"] = 0
    r["salinidad_spc"] = 0.0
    r["salinidad_source"] = "manual"
    r["sal_series"] = pd.Series(dtype=float)
    r["_missing_period_data"] = True
    r["_dias_periodo"] = dias
    return r

def calcular_mes(df_daily, mes_es, anio):
    mes_n = MESES_NUM[mes_es]
    yr = anio - 1 if mes_n in [10, 11, 12] else anio
    dias = calendar.monthrange(int(yr), int(mes_n))[1]
    m = df_daily[(df_daily["ACTDATE"].dt.year == yr) &
                 (df_daily["ACTDATE"].dt.month == mes_n)].copy()
    if m.empty:
        return _empty_period_result(dias), dias

    r = _empty_period_result(dias)
    r["_missing_period_data"] = False

    def _col_by_letter(df, letter, agg="sum"):
        """Acceso robusto a columna del DAILY por letra Excel (AS, R, etc.)."""
        try:
            idx = openpyxl.utils.column_index_from_string(letter) - 1
            if idx < 0 or idx >= len(df.columns):
                return 0.0
            s = pd.to_numeric(df.iloc[:, idx], errors="coerce")
            if agg == "mean":
                v = s.mean()
            else:
                v = s.sum(min_count=1)
            return 0.0 if (v is None or (isinstance(v, float) and pd.isna(v))) else float(v)
        except Exception:
            return 0.0

    # ── Aportes ──────────────────────────────────────────────────────────────
    # NOTA: TOFCHCP (col AY del DAILY) NO es el aporte hidrológico real.
    # Es un balance operacional (salidas + evaporación). Los aportes reales
    # (caudales de tributarios medidos) provienen EXCLUSIVAMENTE de Fuentes_Agua.
    # Se conservan _tofchcp_ops, _tofgl_ops y _tofmd_ops como referencia interna
    # (p.ej. para la distribución Alhajuela/Gatún cuando Fuentes_Agua no está cargada),
    # pero NO se asignan a los campos de aporte del resultado.
    _tofchcp_ops = float(m["TOFCHCP (hm3)"].sum(min_count=1) or 0)
    _tofgl_ops   = float(m["TOFGL (hm3)"].sum(min_count=1) or 0)
    _tofmd_ops   = float(m["TOFMD (hm3)"].sum(min_count=1) or 0)
    # aporte_total_hm3, aporte_total_m3s, aporte_neto_*, aporte_gl_*, aporte_md_*
    # se dejan en 0.0 (valor por defecto de _empty_period_result) hasta que
    # Fuentes_Agua los sobreescriba más adelante en el flujo de cálculo.

    # Evaporación — nombre completo primero, fallback a posición col BZ/CA
    def _evap_col(named_col, fallback_letter):
        if named_col in m.columns:
            v = pd.to_numeric(m[named_col], errors="coerce").sum(min_count=1)
            if not pd.isna(v): return float(v)
        return _col_by_letter(m, fallback_letter)
    evap_gat = _evap_col("Volumen evaporado embalse Gatún, hm3", "BZ")
    evap_alh = _evap_col("Volumen evaporado embalse Alhajuela, hm3", "CA")
    evap_tot = evap_gat + evap_alh
    r["evaporacion_hm3"] = round(evap_tot, 3)
    r["evaporacion_m3s"] = round(evap_tot / (dias * HM3_PER_M3S_DAY), 3) if dias else 0
    r["evap_gatun_hm3"]  = round(evap_gat, 3)
    r["evap_alh_hm3"]    = round(evap_alh, 3)
    # aporte_neto = aporte_total - evap: se recalcula en la sección de Fuentes_Agua

    # ── Niveles embalses ─────────────────────────────────────────────────────
    def _safe_col_mean(col):
        if col not in m.columns: return 0.0
        v = pd.to_numeric(m[col], errors="coerce").mean()
        return 0.0 if pd.isna(v) else float(v)
    nivel_gat_pies = _safe_col_mean("ACTGATEL")
    nivel_alh_pies = _safe_col_mean("ACTMADEL")
    r["nivel_gatun_pies"]  = round(nivel_gat_pies, 3)
    r["nivel_gatun_m"]     = round(nivel_gat_pies * 0.3048, 3)
    r["nivel_alh_pies"]    = round(nivel_alh_pies, 3)

    # ── Superficie promedio y evaporación diaria (criterio Hidroestadística) ─
    area_gat = float(m["Área espejo embalse Gatún, Km2"].mean() or 0) if "Área espejo embalse Gatún, Km2" in m.columns else 0.0
    area_alh = float(m["Área espejo embalse Alhajuela, Km2"].mean() or 0) if "Área espejo embalse Alhajuela, Km2" in m.columns else 0.0
    r["sup_prom_gatun_km2"] = round(area_gat, 3)
    r["sup_prom_alh_km2"]   = round(area_alh, 3)
    r["evap_gatun_hm3_dia"] = round(evap_gat / dias, 3) if dias else 0
    r["evap_alh_hm3_dia"]   = round(evap_alh / dias, 3) if dias else 0
    r["evap_gatun_mm_dia"]  = round(evap_gat * 1000 / (area_gat * dias), 3) if dias and area_gat > 0 else 0
    r["evap_alh_mm_dia"]    = round(evap_alh * 1000 / (area_alh * dias), 3) if dias and area_alh > 0 else 0
    r["evaporacion_diaria_mm"] = round(evap_tot * 1000 / ((area_gat + area_alh) * dias), 3) if dias and (area_gat + area_alh) > 0 else 0

    # ── Esclusajes ──────────────────────────────────────────────────────────
    def _ss(col):  # safe sum — returns 0.0 on missing col or all-NaN
        if col not in m.columns: return 0.0
        v = pd.to_numeric(m[col], errors="coerce").sum(min_count=1)
        return 0.0 if pd.isna(v) else float(v)
    esc_direct = _ss("ESCLUSAJES, hm3")
    pm_r  = _ss("PMLOCKMCF")
    gat_r = _ss("GATLOCKMCF")
    acl_r = _ss("ACLOCKMCF")
    ccl_r = _ss("CCLLOCKMCF")
    # esc_total: use ESCLUSAJES column if available, otherwise derive from lock volumes
    esc_total = esc_direct if esc_direct > 0 else (pm_r+gat_r+acl_r+ccl_r)*0.028317
    # Mapeo verificado en DAILYINPUT:
    #   Panamax    = columnas J + L = GATLOCKMCF + PMLOCKMCF
    #   Neopanamax = columnas N + P = ACLOCKMCF + CCLLOCKMCF
    pan_raw    = gat_r + pm_r
    neo_raw    = acl_r + ccl_r

    MMCF_HM3 = 0.028317
    panamax_hm3    = pan_raw * MMCF_HM3
    neopanamax_hm3 = neo_raw * MMCF_HM3
    esc_total_calc = panamax_hm3 + neopanamax_hm3
    if esc_total_calc > 0:
        pan_pct = panamax_hm3 / esc_total_calc * 100
        neo_pct = neopanamax_hm3 / esc_total_calc * 100
    elif esc_total > 0:
        pan_pct = neo_pct = 50.0
    else:
        pan_pct = neo_pct = 0.0

    # Tránsitos — mapeo consistente con volúmenes:
    #   Panamax    = NUMLOCKGAT + NUMLOCKPM
    #   Neopanamax = NUMLOCKACL + NUMLOCKCCL
    r["transitos_panamax"]    = round((((m["NUMLOCKGAT"].sum(min_count=1) or 0) +
                                    (m["NUMLOCKPM"].sum(min_count=1) or 0)) / 2.0), 1)
    r["transitos_neopanamax"] = round((((m["NUMLOCKACL"].sum(min_count=1) or 0) +
                                    (m["NUMLOCKCCL"].sum(min_count=1) or 0)) / 2.0), 1)
    r["esclusaje_hm3"]       = round(esc_total, 3)
    r["esclusaje_m3s"]       = round(esc_total / (dias * HM3_PER_M3S_DAY), 3) if dias else 0
    r["panamax_hm3"]         = round(panamax_hm3, 3)
    r["neopanamax_hm3"]      = round(neopanamax_hm3, 3)
    r["panamax_pct"]         = round(pan_pct, 2)
    r["neopanamax_pct"]      = round(neo_pct, 2)
    r["uso_prom_diario_hm3"] = round(esc_total / dias, 4) if dias else 0
    r["uso_prom_diario_m3s"] = round(esc_total / (dias * HM3_PER_M3S_DAY), 3) if dias else 0
    r["agua_panamax_trans"]  = round(panamax_hm3 / r["transitos_panamax"], 4) if (r["transitos_panamax"] or 0) > 0 else 0
    r["agua_neo_trans"]      = round(neopanamax_hm3 / r["transitos_neopanamax"], 4) if (r["transitos_neopanamax"] or 0) > 0 else 0

    # ── Potabilización ──────────────────────────────────────────────────────
    # MUNIC GAT (col T) y MUNIC MAD (col S) están en MCF (millones de pies cúbicos).
    # × 0.028317 → hm³. Esto da ~31.9 hm³ para Feb-2026.
    # Los 44 hm³ del informe incluyen APSA (fuente externa al DAILY).
    # HIMH debe ingresar la diferencia manualmente en la pestaña de Exportar.
    # NOTA AUDITORÍA: DAILY da ~31.9 hm³; ingresar manualmente el total oficial.
    MMCF_TO_HM3 = 0.028317
    pot_alh = _ss("MUNIC MAD") * MMCF_TO_HM3
    pot_gat = _ss("MUNIC GAT") * MMCF_TO_HM3
    r["potabilizacion_gat_hm3"] = round(pot_gat, 3)
    r["potabilizacion_alh_hm3"] = round(pot_alh, 3)

    # ── Concesiones / Fugas ─────────────────────────────────────────────────
    # DAILY (col U = LEAK MAD fugas Alhajuela/Madden, col V = LEAK GAT fugas Gatún)
    # LEAK MAD también participa en trasvase_fugas_hm3 (es un movimiento de Alh→Gat).
    # Se incluye en concesiones_alh para que aparezca en la tabla de usos del informe.
    # El balance total_salidas usa SOLO concesiones_gat (Gatún) para no duplicar
    # el volumen de trasvase que ya sale por separado en el balance hídrico.
    leak_gat = _ss("LEAK GAT") * MMCF_TO_HM3
    leak_alh = _ss("LEAK MAD") * MMCF_TO_HM3
    r["concesiones_alh_hm3"] = round(leak_alh, 3)   # LEAK MAD — solo para display/informe
    r["concesiones_gat_hm3"] = round(leak_gat, 3)   # LEAK GAT — entra al balance Gatún

    # ── Trasvase Alhajuela → Gatún ───────────────────────────────────────────
    # Verificación DAILYINPUT:
    #   Columna G = MADMCF    (hidrogeneración Madden, en millones de pies cúbicos)
    #   Columna R = MADSPILL  (vertidos Madden, en millones de pies cúbicos)
    #   Columna U = LEAK MAD  (fugas Madden/Alhajuela, en millones de pies cúbicos)
    # Conversión auditada: 1 millón de pies cúbicos = 0.028316846592 hm³.
    # La columna BB "Trasvase ALA _GAT (hm3)" coincide con G+R+U convertido a hm³.
    MILLON_PIE3_A_HM3_TRASVASE = 0.028316846592
    # NOTA AUDITORÍA Feb-2026: DAILY (G+R+U) = 203.8 hm³; informe = 199 hm³.
    # Col BB "Trasvase ALA_GAT (hm3)" también da 203.8. Diferencia ~4.8 hm³ es
    # de fuente (HIMH ajusta manualmente vertidos). No es un bug del app.
    trasvase_hidro_hm3 = _ss("MADMCF") * MILLON_PIE3_A_HM3_TRASVASE
    trasvase_vert_hm3  = (_ss("MADSPILL") if "MADSPILL" in m.columns
                          else _col_by_letter(m, "R")) * MILLON_PIE3_A_HM3_TRASVASE
    trasvase_fugas_hm3 = _ss("LEAK MAD") * MILLON_PIE3_A_HM3_TRASVASE
    trasvase_calc = trasvase_hidro_hm3 + trasvase_vert_hm3 + trasvase_fugas_hm3
    trasvase_excel = float(m["Trasvase ALA _GAT (hm3)"].sum(min_count=1) or 0) if "Trasvase ALA _GAT (hm3)" in m.columns else 0.0
    trasvase = trasvase_calc if trasvase_calc > 0 else trasvase_excel
    r["trasvase_hidro_hm3"] = round(trasvase_hidro_hm3, 3)
    r["trasvase_vertidos_hm3"] = round(trasvase_vert_hm3, 3)
    r["trasvase_fugas_hm3"] = round(trasvase_fugas_hm3, 3)
    r["trasvase_hm3"] = round(trasvase, 3)
    r["trasvase_m3s"] = round(trasvase / (dias * HM3_PER_M3S_DAY), 3) if dias else 0
    r["trasvase_cfs"] = round(hm3_to_cfs_period(trasvase, dias), 3) if dias else 0
    r["trasvase_pct_esc"] = round(trasvase / esc_total * 100, 2) if esc_total > 0 else 0

    # ── Energía / Hidrogeneración ────────────────────────────────────────────
    # NOTA: float(x or 0) NO protege contra NaN — NaN es truthy en Python.
    # Usar pd.isna() explícitamente para devolver 0.0 cuando no hay datos válidos.
    def _safe_mean(s):
        v = pd.to_numeric(s, errors="coerce").mean()
        return 0.0 if pd.isna(v) else float(v)
    def _safe_sum(s):
        v = pd.to_numeric(s, errors="coerce").sum(min_count=1)
        return 0.0 if pd.isna(v) else float(v)

    # ── Energía desde columnas MADMWH (col E) y GATMWH (col F) ──────────
    # Estas columnas almacenan MWh producidos cada día.
    # MWh/día = sum / días = mean (promedio diario)
    # MW promedio = MWh/día / 24  (conversión energía → potencia)
    def _energy_from_col(named_col, fallback_letter):
        """Suma MWh del período y devuelve (MWh/día promedio, MW promedio)."""
        if named_col in m.columns:
            s = _safe_sum(m[named_col])
            if s > 0:
                mwh_dia = s / dias if dias else 0.0
                return round(mwh_dia, 2), round(mwh_dia / 24, 4)
        # Fallback a letra Excel
        s = _col_by_letter(m, fallback_letter, "sum")
        mwh_dia = s / dias if (s > 0 and dias) else 0.0
        return round(mwh_dia, 2), round(mwh_dia / 24, 4)

    _mad_mwh_dia, _mad_mw = _energy_from_col("MADMWH", "E")
    _gat_mwh_dia, _gat_mw = _energy_from_col("GATMWH", "F")

    # MWh/día promedio = total_MWh / días  (variable local _mad_mwh_dia / _gat_mwh_dia)
    # MW promedio       = MWh/día / 24    (variable local _mad_mw / _gat_mw)
    r["energia_madden_mw"]      = _mad_mw
    r["energia_madden_mwh_dia"] = _mad_mwh_dia   # MWh/día promedio = total_MADMWH / días
    r["energia_gatun_mw"]       = _gat_mw
    r["energia_gatun_mwh_dia"]  = _gat_mwh_dia   # MWh/día promedio = total_GATMWH / días
    MMCF_TO_HM3_E = 0.028317
    r["hidro_gatun_hm3"]  = round(float(m["GATMCF"].sum(min_count=1) or 0) * MMCF_TO_HM3_E, 4)
    r["hidro_madden_hm3"] = round(trasvase_hidro_hm3, 3)

    # ── Salinidad ────────────────────────────────────────────────────────────
    # La salinidad SPC/SPV se define manualmente en la pestaña de Salinidad
    r["salinidad_spc"] = 0.0
    r["salinidad_source"] = "manual"
    r["sal_series"] = pd.Series(dtype=float)

    # ── Ahorros ─────────────────────────────────────────────────────────────
    # Fuente auditada solicitada:
    #   Panamax    = columna W del DAILY, en esclusajes equivalentes (EE)
    #   NeoPanamax = columna X del DAILY, en esclusajes equivalentes (EE)
    # Conversión: EE × 0.2081976 = hm³.
    # Importante: se usa la POSICIÓN de la columna para evitar que encabezados
    # similares de CCA, cambio de dirección u otras métricas sobrescriban el cálculo.
    def _sum_excel_col(df, col_letter):
        try:
            idx = openpyxl.utils.column_index_from_string(col_letter) - 1
            if idx < 0 or idx >= len(df.columns):
                return 0.0, ""
            col_name = df.columns[idx]
            vals = pd.to_numeric(df.iloc[:, idx], errors="coerce")
            v = vals.sum(min_count=1)
            return (0.0 if (v is None or pd.isna(v)) else float(v)), str(col_name)
        except Exception:
            return 0.0, ""

    ahorro_pan_ee, ahorro_pan_col = _sum_excel_col(m, "W")
    ahorro_neo_ee, ahorro_neo_col = _sum_excel_col(m, "X")
    ahorro_pan = ahorro_pan_ee * EE_HM3
    ahorro_neo = ahorro_neo_ee * EE_HM3

    r["ahorro_panamax_ee"]     = round(ahorro_pan_ee, 3)
    r["ahorro_neopanamax_ee"]  = round(ahorro_neo_ee, 3)
    r["ahorro_panamax_col"]    = ahorro_pan_col
    r["ahorro_neopanamax_col"] = ahorro_neo_col
    r["ahorro_panamax_hm3"]    = round(ahorro_pan, 3)
    r["ahorro_neopanamax_hm3"] = round(ahorro_neo, 3)
    r["ahorro_total_hm3"]      = round(ahorro_pan + ahorro_neo, 3)
    r["ahorro_lamina_pies"]    = round((ahorro_pan + ahorro_neo) / 148.0, 3)

    # ── Agua disponible en embalses (del DAILY) ─────────────────────────────
    col_disp_alh = "Agua disponible embalse Alhajuela (NAM 205 pies PLD)"
    col_disp_gat = "Agua disponible  embalse Gatún (NAM 78.80 pies PLD)"
    col_disp_tot = "Agua disponible  embalses Alhajuela +  Gatún "
    if col_disp_alh in m.columns:
        r["agua_disp_alh_mm3"] = round(float(m[col_disp_alh].iloc[-1] or 0), 3)
    if col_disp_gat in m.columns:
        r["agua_disp_gat_mm3"] = round(float(m[col_disp_gat].iloc[-1] or 0), 3)
    if col_disp_tot in m.columns:
        r["agua_disp_total_mm3"] = round(float(m[col_disp_tot].iloc[-1] or 0), 3)

    # ── ZZ Flush / Vertidos ─────────────────────────────────────────────────
    # ZZFlush/CCA: calcular automáticamente desde GATSPILL si está disponible.
    # El usuario puede sobreescribir manualmente en la pestaña de Usos.
    # NOTA AUDITORÍA: DAILY GATSPILL (col Q) da ~25.2 hm³; el informe reporta ~22.8 hm³.
    # Diferencia de ~2.4 hm³ se debe a ajuste HIMH / fuente SCADA. No es bug del app.
    # ZZFlush (CCA) se ingresa manualmente en Tab Salinidad y sobreescribe este valor.
    _gatspill_hm3 = round(_ss("GATSPILL") * MMCF_TO_HM3, 3)
    # GATSPILL = Vertidos Gatún + ZZFlush/CCA (ambos salen por las esclusas neopanamax).
    # Se guarda el bruto para que el override de CCA pueda derivar vertidos = GATSPILL − CCA.
    r["gatspill_raw_hm3"]  = _gatspill_hm3          # bruto col Q — NO editar
    r["zzflush_auto_hm3"]  = _gatspill_hm3           # initial; overwritten by CCA in Tab Salinidad
    r["zzflush_auto_avail"]= _gatspill_hm3 > 0
    r["vertidos_gat_hm3"]  = _gatspill_hm3           # initial; recalculated after CCA override
    r["vertidos_mad_ops_hm3"] = round(trasvase_vert_hm3, 3)
    r["vertidos_mad_hm3"] = 0.0

    # ── Distribución por subcuenca ──────────────────────────────────────────
    # Los aportes por subcuenca (Alhajuela / Gatún) provienen de Fuentes_Agua.

    # ── Totales derivados ───────────────────────────────────────────────────
    _sync_balance_fields(r, dias)
    # Trasvase total Alhajuela→Gatún = hidrogeneración Madden + vertidos Madden + fugas Madden.
    # No sumar nuevamente la columna BB, porque ya representa el total G+R+U.
    r["trasvase_hidro_hm3"]     = r["hidro_madden_hm3"]
    r["trasvase_vertidos_hm3"]  = r["vertidos_mad_ops_hm3"]
    r["trasvase_total_hm3"]     = round(r["trasvase_hm3"], 3)
    r["trasvase_total_m3s"]     = round(hm3_to_m3s_period(r["trasvase_total_hm3"], dias), 3) if dias else 0
    r["movimientos_operativos_hm3"] = r["trasvase_total_hm3"]
    return r, dias


# ──────────────────────────────────────────────────────────────────────────────

def calcular_semestral(df_daily, anio_fiscal):
    """Primer Semestre del Año Fiscal (Oct año-1 → Mar año). Wrapper de calcular_periodo_generico."""
    meses = _get_meses_periodo("Semestral", anio_fiscal, semestre="Primer")
    return calcular_periodo_generico(df_daily, meses)


def calcular_periodo_generico(df_daily, meses_lista):
    """
    Calcula acumulados para cualquier período dado una lista de (mes_n, año_real).
    Reutiliza calcular_mes() para cada mes individual.
    """
    inv = {v:k for k,v in MESES_NUM.items()}
    resultados = []
    meses_con_datos = []
    meses_sin_datos = []
    total_dias = 0

    for mes_n, yr in meses_lista:
        # calcular_mes espera el año FISCAL (el año en que termina el semestre)
        # pero internamente ajusta yr según mes_n. Para pasar el año correcto
        # necesitamos invertir la lógica: si mes_n en [10,11,12] yr_fiscal = yr+1 else yr_fiscal = yr
        if mes_n in [10, 11, 12]:
            anio_fiscal_equiv = yr + 1
        else:
            anio_fiscal_equiv = yr
        mes_es = inv.get(mes_n, str(mes_n))
        r_mes, dias_mes = calcular_mes(df_daily, mes_es, anio_fiscal_equiv)
        resultados.append(r_mes)
        total_dias += dias_mes
        tag = f"{mes_es[:3]} {yr}"
        if r_mes.get("_missing_period_data"):
            meses_sin_datos.append(tag)
        else:
            meses_con_datos.append(tag)

    if not resultados:
        return None, 0

    def suma(k): return sum(float(r.get(k, 0) or 0) for r in resultados)
    def prom(k):
        vals = [float(r.get(k, 0) or 0) for r in resultados]
        return sum(vals) / len(vals) if vals else 0
    def prom_pond_dias(k):
        peso_total = sum(float(r.get("_dias_periodo", 0) or 0) for r in resultados)
        if peso_total <= 0:
            return prom(k)
        return sum(float(r.get(k, 0) or 0) * float(r.get("_dias_periodo", 0) or 0) for r in resultados) / peso_total

    s = _empty_period_result(total_dias)
    s["_meses_disponibles"] = meses_con_datos
    s["_meses_faltantes"] = meses_sin_datos
    s["_n_meses"] = len(resultados)
    s["_es_parcial"] = len(meses_sin_datos) > 0
    s["_missing_period_data"] = len(meses_con_datos) == 0

    for k in [
        "aporte_total_hm3","aporte_neto_hm3","evaporacion_hm3","evap_gatun_hm3","evap_alh_hm3",
        "esclusaje_hm3","panamax_hm3","neopanamax_hm3",
        "potabilizacion_hm3","potabilizacion_gat_hm3","potabilizacion_alh_hm3","concesiones_hm3","concesiones_gat_hm3","concesiones_alh_hm3",
        "trasvase_hm3","trasvase_fugas_hm3","trasvase_total_hm3","hidro_gatun_hm3","hidro_madden_hm3","zzflush_auto_hm3",
        "ahorro_panamax_hm3","ahorro_neopanamax_hm3","ahorro_total_hm3",
        "vertidos_gat_hm3","vertidos_mad_hm3","vertidos_mad_ops_hm3",
        "alhajuela_hm3","gatun_hm3","total_salidas_hm3","movimientos_operativos_hm3","excedente_hm3",
    ]:
        s[k] = round(suma(k), 3)

    s["aporte_total_m3s"]   = round(s["aporte_total_hm3"]   / (total_dias * HM3_PER_M3S_DAY), 3) if total_dias else 0
    s["aporte_neto_m3s"]    = round(s["aporte_neto_hm3"]    / (total_dias * HM3_PER_M3S_DAY), 3) if total_dias else 0
    s["evaporacion_m3s"]    = round(s["evaporacion_hm3"]    / (total_dias * HM3_PER_M3S_DAY), 3) if total_dias else 0
    s["esclusaje_m3s"]      = round(s["esclusaje_hm3"]      / (total_dias * HM3_PER_M3S_DAY), 3) if total_dias else 0
    s["potabilizacion_m3s"] = round(s["potabilizacion_hm3"] / (total_dias * HM3_PER_M3S_DAY), 3) if total_dias else 0
    s["concesiones_m3s"]    = round(s["concesiones_hm3"]    / (total_dias * HM3_PER_M3S_DAY), 3) if total_dias else 0
    s["trasvase_m3s"]       = round(s["trasvase_hm3"]       / (total_dias * HM3_PER_M3S_DAY), 3) if total_dias else 0

    s["nivel_gatun_pies"]       = round(prom("nivel_gatun_pies"), 3)
    s["nivel_gatun_m"]          = round(prom("nivel_gatun_m"), 3)
    s["nivel_alh_pies"]         = round(prom("nivel_alh_pies"), 3)
    s["sup_prom_gatun_km2"]     = round(prom("sup_prom_gatun_km2"), 3)
    s["sup_prom_alh_km2"]       = round(prom("sup_prom_alh_km2"), 3)
    # Promedio ponderado por días de MW; MWh/día = MW_prom × 24.
    _peso_e = sum(float(r.get("_dias_periodo", 0) or 0) for r in resultados)
    _e_mad_mw_pond = (sum(float(r.get("energia_madden_mw", 0) or 0) * float(r.get("_dias_periodo", 0) or 0) for r in resultados) / _peso_e) if _peso_e else 0.0
    _e_gat_mw_pond = (sum(float(r.get("energia_gatun_mw",  0) or 0) * float(r.get("_dias_periodo", 0) or 0) for r in resultados) / _peso_e) if _peso_e else 0.0
    s["energia_madden_mw"]      = round(_e_mad_mw_pond, 3)
    s["energia_gatun_mw"]       = round(_e_gat_mw_pond, 4)
    s["energia_madden_mwh_dia"] = round(_e_mad_mw_pond * 24, 2)   # MWh/día = MW × 24
    s["energia_gatun_mwh_dia"]  = round(_e_gat_mw_pond * 24, 2)   # MWh/día = MW × 24

    s["transitos_panamax"]    = round(sum(float(r.get("transitos_panamax", 0) or 0) for r in resultados), 1)
    s["transitos_neopanamax"] = round(sum(float(r.get("transitos_neopanamax", 0) or 0) for r in resultados), 1)
    s["panamax_pct"]  = round(s["panamax_hm3"]  / s["esclusaje_hm3"] * 100, 2) if s["esclusaje_hm3"] > 0 else 0
    s["neopanamax_pct"] = round(s["neopanamax_hm3"] / s["esclusaje_hm3"] * 100, 2) if s["esclusaje_hm3"] > 0 else 0
    s["uso_prom_diario_hm3"] = round(s["esclusaje_hm3"] / total_dias, 4) if total_dias else 0
    s["uso_prom_diario_m3s"] = round(s["esclusaje_hm3"] / (total_dias * HM3_PER_M3S_DAY), 3) if total_dias else 0
    s["evap_gatun_hm3_dia"]  = round(s["evap_gatun_hm3"] / total_dias, 3) if total_dias else 0
    s["evap_alh_hm3_dia"]    = round(s["evap_alh_hm3"] / total_dias, 3) if total_dias else 0
    s["evap_gatun_mm_dia"]   = round(s["evap_gatun_hm3"] * 1000 / (s["sup_prom_gatun_km2"] * total_dias), 3) if total_dias and s["sup_prom_gatun_km2"] > 0 else 0
    s["evap_alh_mm_dia"]     = round(s["evap_alh_hm3"] * 1000 / (s["sup_prom_alh_km2"] * total_dias), 3) if total_dias and s["sup_prom_alh_km2"] > 0 else 0
    _area_tot = s["sup_prom_gatun_km2"] + s["sup_prom_alh_km2"]
    s["evaporacion_diaria_mm"] = round(s["evaporacion_hm3"] * 1000 / (_area_tot * total_dias), 3) if total_dias and _area_tot > 0 else 0
    s["agua_panamax_trans"]  = round(s["panamax_hm3"] / s["transitos_panamax"], 4) if s["transitos_panamax"] > 0 else 0
    s["agua_neo_trans"]      = round(s["neopanamax_hm3"] / s["transitos_neopanamax"], 4) if s["transitos_neopanamax"] > 0 else 0
    s["trasvase_pct_esc"]    = round(s["trasvase_hm3"] / s["esclusaje_hm3"] * 100, 2) if s["esclusaje_hm3"] > 0 else 0
    s["ahorro_lamina_pies"]  = round(s["ahorro_total_hm3"] / 148.0, 3) if s["ahorro_total_hm3"] > 0 else 0
    s["alhajuela_pct"] = round(s["alhajuela_hm3"] / s["aporte_total_hm3"] * 100, 2) if s["aporte_total_hm3"] > 0 else 0
    s["gatun_pct"]     = round(s["gatun_hm3"]     / s["aporte_total_hm3"] * 100, 2) if s["aporte_total_hm3"] > 0 else 0

    s["salinidad_spc"]    = 0.0
    s["salinidad_source"] = "manual"

    s["_meses"] = resultados
    s["_total_dias"] = total_dias
    _sync_balance_fields(s, total_dias)
    # Reafirmar trasvase auditado para períodos acumulados:
    # total = G MADMCF + R MADSPILL + U LEAK MAD, ya sumado mes a mes en trasvase_hm3.
    s["trasvase_hidro_hm3"] = round(float(s.get("hidro_madden_hm3", 0) or 0), 3)
    s["trasvase_vertidos_hm3"] = round(float(s.get("vertidos_mad_ops_hm3", 0) or 0), 3)
    s["trasvase_total_hm3"] = round(float(s.get("trasvase_hm3", 0) or 0), 3)
    s["trasvase_total_m3s"] = round(hm3_to_m3s_period(s["trasvase_total_hm3"], total_dias), 3) if total_dias else 0.0
    s["movimientos_operativos_hm3"] = s["trasvase_total_hm3"]
    return s, total_dias


# ──────────────────────────────────────────────────────────────────────────────

def _hist_period_days(tipo, mes_es=None):
    if tipo == "Mensual" and mes_es:
        # usar año bisiesto genérico para febrero
        dias = {"Enero":31,"Febrero":28.25,"Marzo":31,"Abril":30,"Mayo":31,"Junio":30,"Julio":31,"Agosto":31,"Septiembre":30,"Octubre":31,"Noviembre":30,"Diciembre":31}
        return dias.get(mes_es, 30)
    return 182


def _default_hist_range_end_year(tipo, anio_sel, mes_es=None, trimestre=None, semestre=None):
    """Rango histórico por defecto: hasta el año ANTERIOR al evaluado.
    La serie histórica excluye el año en curso para evitar que el año
    se cuente a sí mismo como 'más húmedo' por diferencias de redondeo.
    El +1 al total en el lambda ya incluye el año actual en el denominador.
    """
    try:
        return int(anio_sel) - 1
    except Exception:
        return 2025

def _default_hist_range_start_year(kind="chcp"):
    """Año inicial por defecto para cada familia histórica."""
    return 1914 if str(kind).lower() == "netfl" else 1898

def _filter_hist_range_df(df, year_start=None, year_end=None):
    if df is None or not isinstance(df, pd.DataFrame) or df.empty or "AÑO" not in df.columns:
        return df
    tmp = df.copy()
    tmp["AÑO"] = pd.to_numeric(tmp["AÑO"], errors="coerce")
    if year_start is not None:
        tmp = tmp[tmp["AÑO"] >= int(year_start)]
    if year_end is not None:
        tmp = tmp[tmp["AÑO"] <= int(year_end)]
    return tmp.dropna(subset=["AÑO"]).sort_values("AÑO").reset_index(drop=True)

def _classify_chcp_diff_pct(diff_pct):
    """Clasificación compatible con la métrica observada en la columna R de 03_QAVGMCHCP."""
    try:
        x = float(diff_pct)
    except Exception:
        return ""
    if x <= -30:
        return "Muy Seco"
    if x < -10:
        return "Seco"
    if x <= 10:
        return "Promedio"
    if x < 30:
        return "Húmedo"
    return "Muy Húmedo"

def _fiscal_sem_series(df):
    """Wrapper S1 para compatibilidad. Usa _fiscal_period_series genérica."""
    return _fiscal_period_series(df, "S1")


def _fiscal_period_series(df, periodo_key, value_unit="m3s"):
    """
    Genera series históricas anuales para cualquier tipo de período.
    periodo_key: 'Q1','Q2','Q3','Q4','S1','S2','Anual'
    value_unit: 'm3s' para series de caudal promedio, 'hm3' para series de volumen.
    Retorna DataFrame con columnas [AÑO, valor_m3s, valor_hm3].
    """
    rows = []
    if df is None or df.empty:
        return pd.DataFrame(columns=["AÑO","valor_m3s","valor_hm3"])
    meses_off = PERIODO_MESES.get(periodo_key, [])
    total_dias_prom = PERIODO_DIAS.get(periodo_key, 182)
    years = sorted(pd.to_numeric(df["AÑO"], errors="coerce").dropna().astype(int).unique())
    unit = str(value_unit or "m3s").lower()
    for y in years:
        hm3_total = 0.0
        ok = True
        for mes_n, off in meses_off:
            col = MES_N_TO_COL[mes_n]
            yr = y + off   # año calendario de esta columna
            row = df[df["AÑO"] == yr]
            if row.empty or col not in row.columns:
                ok = False; break
            try:
                v = float(row[col].iloc[0])
            except Exception:
                ok = False; break
            if unit == "hm3":
                hm3_total += v
            else:
                hm3_total += v * MES_DIAS[col] * HM3_PER_M3S_DAY
        if ok:
            m3s = hm3_total / (total_dias_prom * HM3_PER_M3S_DAY) if total_dias_prom else 0
            rows.append({"AÑO": y, "valor_m3s": m3s, "valor_hm3": hm3_total})
    return pd.DataFrame(rows)

def calcular_historicos(df_alh, df_gat, df_chcp, mes_es, tipo="Mensual", anio_fiscal=None, df_netfl=None,
                        year_start=None, year_end=None, year_start_netfl=None, year_end_netfl=None):
    """Calcula promedios históricos de forma tolerante.
    Puede operar solo con CHCP y opcionalmente con GATNETFL.
    df_netfl = GATNETFL (aportes netos CHCP, opcional).
    """
    try:
        out = {}
        if year_start_netfl is None:
            year_start_netfl = year_start
        if year_end_netfl is None:
            year_end_netfl = year_end
        _netfl_unit = str(getattr(df_netfl, 'attrs', {}).get('value_unit', 'm3s')).lower() if df_netfl is not None else 'm3s'


        def _safe_monthly(df, col):
            if df is None or not isinstance(df, pd.DataFrame) or df.empty:
                return pd.DataFrame(columns=["AÑO", "valor"])
            cols = set(map(str, df.columns))
            if "AÑO" not in cols or col not in cols:
                return pd.DataFrame(columns=["AÑO", "valor"])
            tmp = df[["AÑO", col]].copy()
            tmp["AÑO"] = pd.to_numeric(tmp["AÑO"], errors="coerce")
            tmp["valor"] = pd.to_numeric(tmp[col], errors="coerce")
            return (tmp[["AÑO", "valor"]]
                    .dropna(subset=["AÑO", "valor"])
                    .sort_values("AÑO")
                    .reset_index(drop=True))

        def _safe_fiscal(df, pk, value_unit="m3s"):
            if df is None or not isinstance(df, pd.DataFrame) or df.empty:
                return pd.DataFrame(columns=["AÑO", "valor_m3s", "valor_hm3"])
            if "AÑO" not in map(str, df.columns):
                return pd.DataFrame(columns=["AÑO", "valor_m3s", "valor_hm3"])
            try:
                return _fiscal_period_series(df, pk, value_unit=value_unit)
            except Exception:
                return pd.DataFrame(columns=["AÑO", "valor_m3s", "valor_hm3"])

        if tipo == "Mensual":
            col = mes_col(mes_es)
            try:
                mes_n = MESES_NUM.get(mes_es, None)
                if mes_n is not None and anio_fiscal is not None:
                    yr_cal = int(anio_fiscal) - 1 if mes_n in [10, 11, 12] else int(anio_fiscal)
                    days = calendar.monthrange(int(yr_cal), int(mes_n))[1]
                else:
                    days = _hist_period_days("Mensual", mes_es)
            except Exception:
                days = _hist_period_days("Mensual", mes_es)
            last_alh = _filter_hist_range_df(_safe_monthly(df_alh, col), year_start, year_end)
            last_gat = _filter_hist_range_df(_safe_monthly(df_gat, col), year_start, year_end)
            last_chcp = _filter_hist_range_df(_safe_monthly(df_chcp, col), year_start, year_end)
            chcp_raw = _filter_hist_range_df(df_chcp, year_start, year_end)

            if _netfl_unit == 'hm3':
                last_netfl_hm3 = _filter_hist_range_df(_safe_monthly(df_netfl, col), year_start_netfl, year_end_netfl)
                h_netfl_hm3 = float(last_netfl_hm3["valor"].mean()) if not last_netfl_hm3.empty else 0
                h_netfl = (h_netfl_hm3 / (days * HM3_PER_M3S_DAY)) if (days and h_netfl_hm3) else 0
                series_netfl = last_netfl_hm3.copy()
                if not series_netfl.empty:
                    series_netfl["valor_hm3"] = pd.to_numeric(series_netfl["valor"], errors="coerce")
                    series_netfl["valor"] = series_netfl["valor_hm3"] / (days * HM3_PER_M3S_DAY) if days else 0
                last_netfl = series_netfl
            else:
                last_netfl = _filter_hist_range_df(_safe_monthly(df_netfl, col), year_start_netfl, year_end_netfl)
                h_netfl = float(last_netfl["valor"].mean()) if not last_netfl.empty else 0
                h_netfl_hm3 = h_netfl * days * HM3_PER_M3S_DAY

            h_alh = float(last_alh["valor"].mean()) if not last_alh.empty else 0
            h_gat = float(last_gat["valor"].mean()) if not last_gat.empty else 0
            h_chcp = float(last_chcp["valor"].mean()) if not last_chcp.empty else 0

            out.update({
                "hist_alh_m3s": round(h_alh, 3),
                "hist_gat_m3s": round(h_gat, 3),
                "hist_chcp_m3s": round(h_chcp, 3),
                "hist_neto_m3s": round(h_netfl, 3),
                "hist_alh_hm3": round(h_alh * days * HM3_PER_M3S_DAY, 3),
                "hist_gat_hm3": round(h_gat * days * HM3_PER_M3S_DAY, 3),
                "hist_chcp_hm3": round(h_chcp * days * HM3_PER_M3S_DAY, 3),
                "hist_neto_hm3": round(h_netfl_hm3, 3),
                "n_years_chcp": len(last_chcp),
                "n_years_netfl": len(last_netfl),
                "years_chcp": (int(last_chcp["AÑO"].min()), int(last_chcp["AÑO"].max())) if not last_chcp.empty else None,
                "years_netfl": (int(last_netfl["AÑO"].min()), int(last_netfl["AÑO"].max())) if not last_netfl.empty else None,
                "series_chcp": last_chcp,
                "series_alh": last_alh,
                "series_gat": last_gat,
                "series_netfl": last_netfl,
                "series_chcp_raw": chcp_raw if isinstance(chcp_raw, pd.DataFrame) else pd.DataFrame(),
            })
        else:
            pk = _periodo_key(tipo, getattr(calcular_historicos, "_trimestre", None),
                              getattr(calcular_historicos, "_semestre", None))
            s_alh = _filter_hist_range_df(_safe_fiscal(df_alh, pk), year_start, year_end)
            s_gat = _filter_hist_range_df(_safe_fiscal(df_gat, pk), year_start, year_end)
            s_chcp = _filter_hist_range_df(_safe_fiscal(df_chcp, pk), year_start, year_end)
            s_netfl = _filter_hist_range_df(_safe_fiscal(df_netfl, pk, value_unit=('hm3' if _netfl_unit == 'hm3' else 'm3s')), year_start_netfl, year_end_netfl)
            chcp_raw = _filter_hist_range_df(df_chcp, year_start, year_end)

            h_alh = float(s_alh["valor_m3s"].mean()) if not s_alh.empty else 0
            h_gat = float(s_gat["valor_m3s"].mean()) if not s_gat.empty else 0
            h_chcp = float(s_chcp["valor_m3s"].mean()) if not s_chcp.empty else 0
            h_netfl = float(s_netfl["valor_m3s"].mean()) if not s_netfl.empty else 0
            h_netfl_hm3 = float(s_netfl["valor_hm3"].mean()) if not s_netfl.empty else 0

            out.update({
                "hist_alh_m3s": round(h_alh, 3),
                "hist_gat_m3s": round(h_gat, 3),
                "hist_chcp_m3s": round(h_chcp, 3),
                "hist_neto_m3s": round(h_netfl, 3),
                "hist_alh_hm3": round(float(s_alh["valor_hm3"].mean()) if not s_alh.empty else 0, 3),
                "hist_gat_hm3": round(float(s_gat["valor_hm3"].mean()) if not s_gat.empty else 0, 3),
                "hist_chcp_hm3": round(float(s_chcp["valor_hm3"].mean()) if not s_chcp.empty else 0, 3),
                "hist_neto_hm3": round(h_netfl_hm3, 3),
                "n_years_chcp": len(s_chcp),
                "n_years_netfl": len(s_netfl),
                "years_chcp": (int(s_chcp["AÑO"].min()), int(s_chcp["AÑO"].max())) if not s_chcp.empty else None,
                "years_netfl": (int(s_netfl["AÑO"].min()), int(s_netfl["AÑO"].max())) if not s_netfl.empty else None,
                "series_chcp": s_chcp.rename(columns={"valor_m3s": "valor"}),
                "series_alh": s_alh.rename(columns={"valor_m3s": "valor"}),
                "series_gat": s_gat.rename(columns={"valor_m3s": "valor"}),
                "series_netfl": s_netfl.rename(columns={"valor_m3s": "valor"}) if not s_netfl.empty else pd.DataFrame(),
                "series_chcp_raw": chcp_raw if isinstance(chcp_raw, pd.DataFrame) else pd.DataFrame(),
            })

        # Basta con que exista CHCP o NETFL para considerar que hubo cálculo histórico útil.
        if (out.get("n_years_chcp", 0) == 0) and (out.get("n_years_netfl", 0) == 0):
            return {}
        return out
    except Exception as _e_hist:
        import traceback as _tb
        import streamlit as _st
        _st.session_state["_hist_calc_error"] = f"{_e_hist}\n{_tb.format_exc()}"
        return {}


# ──────────────────────────────────────────────────────────────────────────────
def _clean_text(v):
    return str(v).strip() if pd.notna(v) else ""

def _clean_norm(v):
    s = _clean_text(v).lower()
    return re.sub(r"\s+", " ", s)

def _pick_first_text(row, idxs):
    for idx in idxs:
        if idx < len(row):
            s = _clean_text(row.iloc[idx])
            if s and s.lower() not in {"nan", "none"}:
                return s
    return ""

def _pick_first_float(row, idxs):
    for idx in idxs:
        if idx < len(row):
            try:
                if pd.notna(row.iloc[idx]):
                    return float(row.iloc[idx])
            except Exception:
                pass
    return None

def parse_fuentes(df_raw):
    """Parsea la hoja mensual/semestral de Fuentes_Agua con tolerancia a columnas desplazadas."""
    rios = []
    subtotales = {}
    for _, row in df_raw.iterrows():
        nombre = _pick_first_text(row, [1, 0, 8, 7])
        nombre_n = _clean_norm(nombre)
        if not nombre or nombre_n in {"no.", "fuente (tributario)", "fuente (tributario) "}:
            continue
        hm3 = _pick_first_float(row, [3, 8, 2, 9])
        pct = _pick_first_float(row, [4, 5])
        m3s = _pick_first_float(row, [2, 9, 3])

        if "subcuenca" in nombre_n:
            cuenca = "Alhajuela" if "alhajuela" in nombre_n else "Gatún"
            if hm3 is not None or m3s is not None:
                subtotales[cuenca] = {"hm3": hm3 if hm3 is not None else 0.0,
                                      "m3s": m3s if m3s is not None else 0.0,
                                      "pct": pct if pct is not None else 0.0}
            continue

        if "chcp" in nombre_n or "aporte total" in nombre_n:
            continue

        if hm3 is None and m3s is None:
            continue

        cuenca = "Alhajuela" if any(x in nombre_n for x in ["chagres", "pequ", "boquer", "indio", "madden", "alhajuela"]) else "Gatún"
        rios.append({
            "nombre": nombre,
            "m3s": m3s,
            "hm3": hm3,
            "pct": pct,
            "cuenca": cuenca,
        })
    return rios, subtotales


def _enrich_fuentes_items(items, total_hm3, dias):
    """Completa m3/s, hm3 y % para tributarios a partir de lo que venga del Excel."""
    rows = []
    for r in items or []:
        if not isinstance(r, dict):
            continue
        nombre = str(r.get("nombre", "") or "").strip()
        if not nombre:
            continue
        hm3 = r.get("hm3", None)
        m3s = r.get("m3s", None)
        pct = r.get("pct", None)
        hm3 = float(hm3) if hm3 not in [None, ""] and pd.notna(hm3) else None
        m3s = float(m3s) if m3s not in [None, ""] and pd.notna(m3s) else None
        pct = float(pct) if pct not in [None, ""] and pd.notna(pct) else None

        if hm3 is None and m3s is not None and dias:
            hm3 = float(m3s) * float(dias) * HM3_PER_M3S_DAY
        if (m3s is None or abs(m3s) < 1e-12) and hm3 is not None and dias:
            m3s = hm3_to_m3s_period(hm3, dias)
        if pct is None and hm3 is not None and total_hm3 > 0:
            pct = hm3 / total_hm3 * 100.0

        rows.append({
            "Fuente": nombre,
            "m³/s": float(m3s or 0.0),
            "hm³": float(hm3 or 0.0),
            "%": float(pct or 0.0),
        })
    return pd.DataFrame(rows, columns=["Fuente", "m³/s", "hm³", "%"])


def _build_fuentes_excel_section(rios_alh, rios_gat, fuentes_sub, aporte_total_hm3, dias):
    """Arma una tabla tipo Excel de Fuentes_Agua con subtotales y total CHCP."""
    aporte_total_hm3 = float(aporte_total_hm3 or 0)
    df_alh = _enrich_fuentes_items(rios_alh, aporte_total_hm3, dias)
    df_gat = _enrich_fuentes_items(rios_gat, aporte_total_hm3, dias)

    alh_hm3 = float(df_alh["hm³"].sum()) if not df_alh.empty else 0.0
    gat_hm3 = float(df_gat["hm³"].sum()) if not df_gat.empty else 0.0

    sub_alh = fuentes_sub.get("Alhajuela", {}) if isinstance(fuentes_sub, dict) else {}
    sub_gat = fuentes_sub.get("Gatún", {}) if isinstance(fuentes_sub, dict) else {}
    sub_alh_hm3 = float(sub_alh.get("hm3", alh_hm3) or 0.0)
    sub_gat_hm3 = float(sub_gat.get("hm3", gat_hm3) or 0.0)

    if aporte_total_hm3 <= 0:
        aporte_total_hm3 = sub_alh_hm3 + sub_gat_hm3
    aporte_total_m3s = hm3_to_m3s_period(aporte_total_hm3, dias) if dias and aporte_total_hm3 else 0.0
    sub_alh_m3s = hm3_to_m3s_period(sub_alh_hm3, dias) if dias and sub_alh_hm3 else 0.0
    sub_gat_m3s = hm3_to_m3s_period(sub_gat_hm3, dias) if dias and sub_gat_hm3 else 0.0

    rows = []
    for _, row in df_alh.iterrows():
        rows.append({**row.to_dict(), "Grupo": "Alhajuela", "Tipo": "tributario"})
    rows.append({"Fuente": "Subcuenca embalse Alhajuela", "m³/s": sub_alh_m3s, "hm³": sub_alh_hm3,
                 "%": (sub_alh_hm3 / aporte_total_hm3 * 100 if aporte_total_hm3 > 0 else 0.0),
                 "Grupo": "Alhajuela", "Tipo": "subtotal"})
    for _, row in df_gat.iterrows():
        rows.append({**row.to_dict(), "Grupo": "Gatún", "Tipo": "tributario"})
    rows.append({"Fuente": "Subcuenca embalse Gatún", "m³/s": sub_gat_m3s, "hm³": sub_gat_hm3,
                 "%": (sub_gat_hm3 / aporte_total_hm3 * 100 if aporte_total_hm3 > 0 else 0.0),
                 "Grupo": "Gatún", "Tipo": "subtotal"})
    rows.append({"Fuente": "Aporte total a la CHCP", "m³/s": aporte_total_m3s, "hm³": aporte_total_hm3,
                 "%": 100.0 if aporte_total_hm3 > 0 else 0.0, "Grupo": "CHCP", "Tipo": "total"})

    return pd.DataFrame(rows, columns=["Fuente", "m³/s", "hm³", "%", "Grupo", "Tipo"]), df_alh, df_gat


def _period_tag(tipo, mes_sel, anio_sel, periodo_sem=None):
    if tipo == "Mensual":
        return f"{mes_sel}{anio_sel}".replace(" ", "").lower()
    return str(periodo_sem or f"AF{anio_sel}").replace(" ", "").lower()

def _matches_period_name(name, tipo, mes_sel, anio_sel, periodo_sem=None):
    n = _norm_text(name).replace(" ", "")
    if tipo == "Mensual":
        return _norm_text(mes_sel).replace(" ", "") in n and str(anio_sel) in n
    return _norm_text(periodo_sem or f"AF{anio_sel}").replace(" ", "") in n

def _hist_rank(series_df, value_col="valor", current=None, descending=True):
    if series_df is None or series_df.empty or current is None:
        return 0, 0
    ser = pd.to_numeric(series_df[value_col], errors="coerce").dropna()
    if ser.empty:
        return 0, 0
    rank = int((ser > current).sum() + 1) if descending else int((ser < current).sum() + 1)
    return rank, int(len(ser))

def _ordinal_es(n: int) -> str:
    try:
        n = int(n)
    except Exception:
        return ""
    if n <= 0:
        return ""
    return f"{n}.º"

def _rank_scale_label(rank: int, total: int) -> str:
    try:
        rank = int(rank); total = int(total)
    except Exception:
        return ""
    if rank <= 0 or total <= 0:
        return ""
    frac = rank / total  # 1 = más húmedo
    if frac <= 0.10:
        return "Muy Húmedo"
    if frac <= 0.30:
        return "Húmedo"
    if frac <= 0.70:
        return "Promedio"
    if frac <= 0.90:
        return "Seco"
    return "Muy Seco"

def _annual_class_from_raw(raw_df, current_hm3):
    if raw_df is None or not isinstance(raw_df, pd.DataFrame) or raw_df.empty:
        return ""
    if "hm3" not in raw_df.columns or "Clas. Aporte" not in raw_df.columns:
        return ""
    tmp = raw_df[["hm3", "Clas. Aporte"]].copy()
    tmp["hm3"] = pd.to_numeric(tmp["hm3"], errors="coerce")
    tmp["Clas. Aporte"] = tmp["Clas. Aporte"].astype(str).str.strip().str.rstrip(',')
    tmp = tmp.dropna(subset=["hm3"])
    tmp = tmp[tmp["Clas. Aporte"].str.len() > 0]
    if tmp.empty:
        return ""
    order = ["Muy Seco", "Seco", "Promedio", "Húmedo", "Muy Húmedo"]
    bounds = []
    for cls in order:
        s = tmp.loc[tmp["Clas. Aporte"] == cls, "hm3"]
        if not s.empty:
            bounds.append((cls, float(s.min()), float(s.max())))
    if not bounds:
        return ""
    try:
        v = float(current_hm3)
    except Exception:
        return ""
    for cls, lo, hi in bounds:
        if lo <= v <= hi:
            return cls
    if v < bounds[0][1]:
        return bounds[0][0]
    if v > bounds[-1][2]:
        return bounds[-1][0]
    # si cae en un hueco entre clases, tomar la clase más cercana por centro
    centers = [(cls, (lo + hi) / 2.0) for cls, lo, hi in bounds]
    return min(centers, key=lambda x: abs(v - x[1]))[0]

def _hist_scale_auto(series_df, current=None, value_col="valor", descending=True, raw_df=None, current_hm3=None, prefer_raw_class=False):
    rank, total = _hist_rank(series_df, value_col=value_col, current=current, descending=descending)
    cls = ""
    if prefer_raw_class:
        cls = _annual_class_from_raw(raw_df, current_hm3)
    if not cls:
        cls = _rank_scale_label(rank, total)
    return {
        "rank": rank,
        "total": total,
        "ordinal": _ordinal_es(rank),
        "position_text": (f"{rank} de {total}" if rank and total else ""),
        "class_label": cls,
    }

def _pct_rel_text(diff_pct, positive_text="por encima al", negative_text="por debajo del", zero_text="igual al", decimals=0):
    try:
        x = float(diff_pct)
    except Exception:
        return "0", zero_text
    mag = abs(x)
    fmt = f"{{:.{int(decimals)}f}}"
    if x > 0:
        return fmt.format(mag), positive_text
    if x < 0:
        return fmt.format(mag), negative_text
    return fmt.format(0), zero_text

def _build_internal_hydro(tipo, mes_sel, anio_sel, periodo_sem, datos, n_dias, hist, fuentes_sub,
                          trimestre=None, semestre=None):
    period_label = _label_periodo(tipo, mes_sel, anio_sel, trimestre, semestre)
    days = n_dias if n_dias else (_hist_period_days("Mensual", mes_sel) if tipo == "Mensual" else
                                  PERIODO_DIAS.get(_periodo_key(tipo, trimestre, semestre), 182))

    def d(key, default=0.0):
        try:
            v = datos.get(key, default)
            return float(v) if v is not None else float(default)
        except Exception:
            return float(default)

    aporte_hist_m3s = float(hist.get("hist_chcp_m3s", 0) or 0)
    # Para informes del período seleccionado, el hm³ histórico debe salir del m³/s histórico
    # usando los días reales del período actual (por ejemplo FEB-2026 = 28 días),
    # de modo que valide con la conversión hm³ = m³/s * días * 0.0864.
    aporte_hist_hm3 = float((aporte_hist_m3s * days * HM3_PER_M3S_DAY) if aporte_hist_m3s else (hist.get("hist_chcp_hm3", 0) or 0))
    sub_hist_alh_hm3 = float(hist.get("hist_alh_hm3", 0) or 0)
    sub_hist_gat_hm3 = float(hist.get("hist_gat_hm3", 0) or 0)
    chcp_series = hist.get("series_chcp")
    current_hist_val = d("aporte_total_m3s") if tipo == "Mensual" else d("aporte_total_hm3")
    _raw_cls_df = hist.get("series_chcp_raw") if isinstance(hist, dict) else None
    _scale_info = _hist_scale_auto(
        chcp_series,
        current=current_hist_val,
        value_col=("valor" if tipo == "Mensual" else "valor_hm3"),
        descending=True,
        raw_df=_raw_cls_df,
        current_hm3=d("aporte_total_hm3"),
        prefer_raw_class=(tipo == "Anual")
    )
    rank = int(_scale_info.get("rank", 0) or 0)
    n_rank = int(_scale_info.get("total", 0) or 0)
    scale_pos_text = _scale_info.get("position_text", "") or (str(rank) if rank else "")
    scale_class = _scale_info.get("class_label", "") or ""
    datos["hist_total_hm3"] = aporte_hist_hm3
    datos["hist_total_m3s"] = aporte_hist_m3s
    datos["hist_rank"] = rank
    datos["hist_rank_total"] = n_rank
    datos["hist_scale_pos"] = scale_pos_text
    datos["hist_scale_class"] = scale_class

    netfl_series = hist.get("series_netfl") if isinstance(hist, dict) else None
    current_net_hist_val = d("aporte_neto_m3s") if tipo == "Mensual" else d("aporte_neto_hm3")
    _scale_info_net = _hist_scale_auto(
        netfl_series,
        current=current_net_hist_val,
        value_col=("valor" if tipo == "Mensual" else "valor_hm3"),
        descending=True,
        raw_df=None,
        current_hm3=d("aporte_neto_hm3"),
        prefer_raw_class=False
    )
    datos["hist_neto_rank"] = int(_scale_info_net.get("rank", 0) or 0)
    datos["hist_neto_rank_total"] = int(_scale_info_net.get("total", 0) or 0)
    datos["hist_neto_scale_pos"] = _scale_info_net.get("position_text", "") or (str(datos["hist_neto_rank"]) if datos["hist_neto_rank"] else "")
    datos["hist_neto_scale_class"] = _scale_info_net.get("class_label", "") or ""
    if "Alhajuela" in fuentes_sub:
        datos["alhajuela_hm3"] = float(fuentes_sub["Alhajuela"].get("hm3", datos.get("alhajuela_hm3", 0)) or 0)
    if "Gatún" in fuentes_sub:
        datos["gatun_hm3"] = float(fuentes_sub["Gatún"].get("hm3", datos.get("gatun_hm3", 0)) or 0)
    total_ap = d("aporte_total_hm3")
    if total_ap > 0:
        datos["alhajuela_pct"] = round((datos.get("alhajuela_hm3", 0) or 0) / total_ap * 100, 2)
        datos["gatun_pct"] = round((datos.get("gatun_hm3", 0) or 0) / total_ap * 100, 2)

    esc_mm = d("aporte_total_hm3") * 1000 / CHCP_AREA_KM2 if CHCP_AREA_KM2 else 0
    precip = float(datos.get("precipitacion_mm", 0) or 0)
    coef = esc_mm / precip if precip > 0 else 0
    rendimiento = d("aporte_total_m3s") * 1000 / CHCP_AREA_KM2 if CHCP_AREA_KM2 else 0
    evap_diaria = float(datos.get("evaporacion_diaria_mm", 0) or 0)

    rows1 = []
    def add1(row, label, v1=None, u1=None, v2=None, u2=None, v3=None, u3=None):
        rows1.append({
            "row": row, "label": label, "label_norm": _norm_text(label),
            "value1": v1, "unit1": u1, "value2": v2, "unit2": u2, "value3": v3, "unit3": u3,
            "formula1": None, "formula2": None, "formula3": None,
        })

    add1(3, "Precipitación", precip, "mm")
    add1(4, "Aporte Total", d("aporte_total_hm3"), "hm3", d("aporte_total_m3s"), "m3/s", hm3_to_eed_period(d("aporte_total_hm3"), days), "EED")
    add1(5, "Escorrentía", esc_mm, "mm")
    add1(6, "Coef. De escorrentía", coef, "")
    add1(7, "Rendimiento", rendimiento, "l/s/Km2")
    add1(8, "Evaporación mensual", d("evaporacion_hm3"), "hm3", d("evaporacion_m3s"), "m3/s", hm3_to_eed_period(d("evaporacion_hm3"), days), "EED")
    add1(9, "Evaporación diaria", evap_diaria, "mm")
    add1(10, "Aporte Neto", d("aporte_neto_hm3"), "hm3", d("aporte_neto_m3s"), "m3/s", hm3_to_eed_period(d("aporte_neto_hm3"), days), "EED")
    add1(15, "Evaporación Gatún", d("evap_gatun_hm3"), "hm3", hm3_to_m3s_period(d("evap_gatun_hm3"), days), "m3/s", hm3_to_eed_period(d("evap_gatun_hm3"), days), "EED")
    add1(16, "Evaporación Alhajuela", d("evap_alh_hm3"), "hm3", hm3_to_m3s_period(d("evap_alh_hm3"), days), "m3/s", hm3_to_eed_period(d("evap_alh_hm3"), days), "EED")
    add1(17, "Elev. Prom. Gatún", d("nivel_gatun_pies"), "pies PLD")
    add1(18, "Elev. Prom. Alhajuela", d("nivel_alh_pies"), "pies PLD")
    add1(19, "Sup. Promedio Gatún", d("sup_prom_gatun_km2"), "Km2")
    add1(20, "Sup. Promedio Alhajuela", d("sup_prom_alh_km2"), "Km2")
    add1(21, "Evaporación Gatún", d("evap_gatun_hm3_dia"), "hm3/día")
    add1(22, "Evaporación Gatún", d("evap_gatun_mm_dia"), "mm/día")
    add1(23, "Evaporación Alhajuela", d("evap_alh_hm3_dia"), "hm3/día")
    add1(24, "Evaporación Alhajuela", d("evap_alh_mm_dia"), "mm/día")
    add1(25, "Esclusajes PNX", d("panamax_hm3"), "hm3", hm3_to_m3s_period(d("panamax_hm3"), days), "m3/s", hm3_to_eed_period(d("panamax_hm3"), days), "EED")
    add1(29, "Esclusajes NPX", d("neopanamax_hm3"), "hm3", hm3_to_m3s_period(d("neopanamax_hm3"), days), "m3/s", hm3_to_eed_period(d("neopanamax_hm3"), days), "EED")
    add1(26, "Tránsitos Panamax", float(datos.get("transitos_panamax", 0) or 0), "tránsitos")
    add1(30, "Tránsitos NeoPanamax", float(datos.get("transitos_neopanamax", 0) or 0), "tránsitos")
    add1(31, "Potabilización Gatún", d("potabilizacion_gat_hm3"), "hm3", hm3_to_m3s_period(d("potabilizacion_gat_hm3"), days), "m3/s", hm3_to_eed_period(d("potabilizacion_gat_hm3"), days), "EED")
    add1(32, "Potabilización Alhajuela", d("potabilizacion_alh_hm3"), "hm3", hm3_to_m3s_period(d("potabilizacion_alh_hm3"), days), "m3/s", hm3_to_eed_period(d("potabilizacion_alh_hm3"), days), "EED")
    add1(33, "Concesiones y Misceláneos", d("concesiones_hm3"), "hm3", d("concesiones_m3s"), "m3/s", hm3_to_eed_period(d("concesiones_hm3"), days), "EED")
    add1(34, "Hidrogeneración Gatún", d("hidro_gatun_hm3"), "hm3", hm3_to_m3s_period(d("hidro_gatun_hm3"), days), "m3/s", hm3_to_eed_period(d("hidro_gatun_hm3"), days), "EED")
    add1(35, "Vertidos ZZ Flush", d("zzflush_auto_hm3"), "hm3", hm3_to_m3s_period(d("zzflush_auto_hm3"), days), "m3/s", hm3_to_eed_period(d("zzflush_auto_hm3"), days), "EED")
    add1(36, "Usos Totales", d("total_salidas_hm3"), "hm3", hm3_to_m3s_period(d("total_salidas_hm3"), days), "m3/s", hm3_to_eed_period(d("total_salidas_hm3"), days), "EED")
    add1(37, "Hidrogeneración Gatún", d("energia_gatun_mw"), "MW", d("energia_gatun_mwh_dia"), "MWh/día")
    add1(38, "Deficit", d("aporte_neto_hm3") - d("total_salidas_hm3"), "hm3", hm3_to_m3s_period(d("aporte_neto_hm3") - d("total_salidas_hm3"), days), "m3/s")
    add1(39, "Trasvase Madden", d("trasvase_hm3"), "hm3", d("trasvase_m3s"), "m3/s", hm3_to_eed_period(d("trasvase_hm3"), days), "EED")
    add1(45, "Hidrogeneración Madden", d("energia_madden_mw"), "MW", d("energia_madden_mwh_dia"), "MWh/día")
    add1(46, "Movimientos operativos internos", d("movimientos_operativos_hm3"), "hm3", hm3_to_m3s_period(d("movimientos_operativos_hm3"), days), "m3/s", hm3_to_eed_period(d("movimientos_operativos_hm3"), days), "EED")
    add1(54, "Panamax", d("ahorro_panamax_hm3"), "hm3")
    add1(55, "NeoPanamax Total", d("ahorro_neopanamax_hm3"), "hm3")
    add1(56, "Panamax + NeoPanamax", d("ahorro_total_hm3"), "hm3")
    add1(57, "Lámina Gatún", d("ahorro_lamina_pies"), "pies")

    rows2 = []
    def add2(row, label, val=None, unit=None, pct=None):
        rows2.append({
            "row": row, "label": label, "label_norm": _norm_text(label),
            "value": val, "unit": unit, "pct": pct, "formula_value": None, "formula_pct": None
        })

    total_sal = d("total_salidas_hm3")
    add2(4, "Evaporación Gatún", d("evap_gatun_hm3"), "hm3", d("evap_gatun_hm3")/total_sal if total_sal else 0)
    add2(5, "Evaporación Alhajuela", d("evap_alh_hm3"), "hm3", d("evap_alh_hm3")/total_sal if total_sal else 0)
    add2(6, "Esclusajes PNX", d("panamax_hm3"), "hm3", d("panamax_hm3")/total_sal if total_sal else 0)
    add2(7, "Esclusajes NPX", d("neopanamax_hm3"), "hm3", d("neopanamax_hm3")/total_sal if total_sal else 0)
    add2(8, "Potabilización Gatún", d("potabilizacion_gat_hm3"), "hm3", d("potabilizacion_gat_hm3")/total_sal if total_sal else 0)
    add2(9, "Potabilización Alhajuela", d("potabilizacion_alh_hm3"), "hm3", d("potabilizacion_alh_hm3")/total_sal if total_sal else 0)
    add2(10, "Concesiones y Misceláneos", d("concesiones_hm3"), "hm3", d("concesiones_hm3")/total_sal if total_sal else 0)
    add2(11, "Hidrogeneración Gatún", d("hidro_gatun_hm3"), "hm3", d("hidro_gatun_hm3")/total_sal if total_sal else 0)
    add2(12, "Mitigación de salinidad en esclusas neopanamax (ZZFlush)", d("zzflush_auto_hm3"), "hm3", d("zzflush_auto_hm3")/total_sal if total_sal else 0)
    add2(13, "Vertidos Gatún", d("vertidos_gat_hm3"), "hm3", d("vertidos_gat_hm3")/total_sal if total_sal else 0)
    add2(14, "Usos Totales", total_sal, "hm3", 1 if total_sal else 0)

    return {
        "periodo_grupo": tipo,
        "periodo_detalle": period_label,
        "sheet1_rows": rows1,
        "sheet2_rows": rows2,
        "file_matched": False,
        "source": "internal",
        "hist_total_hm3": aporte_hist_hm3,
        "hist_total_m3s": aporte_hist_m3s,
        "hist_rank": rank,
        "hist_rank_total": n_rank,
        "hist_scale_pos": scale_pos_text,
        "hist_scale_class": scale_class,
        "hist_sub_alh_hm3": sub_hist_alh_hm3,
        "hist_sub_gat_hm3": sub_hist_gat_hm3,
    }



# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
RESET_NUM_KEYS = [
    "zzflush_hm3", "zz_input",
    "rep_aporte_total", "rep_aporte_neto", "rep_alh_hm3", "rep_gat_hm3",
    "rep_esclusaje_hm3", "rep_pan_hm3", "rep_neo_hm3", "rep_trasvase_hm3",
    "rep_tpan", "rep_tneo", "rep_ahorro_pan", "rep_ahorro_neo",
    "rep_cca_hm3", "rep_zzflush_hm3",
    "x_prec", "x_phist", "x_ppos", "x_pct",
    "x_hist_hm3", "x_hist_m3s", "x_cpan", "x_cneo",
    "x_sspc", "x_sspv", "x_sdec", "x_demanda_acp",
]
# Estos keys pertenecen a st.text_input → deben resetearse a "" no a 0.0
RESET_STR_KEYS = ["x_prec_dias", "x_pper", "x_pesc", "met_texto_libre", "met_sal_texto_libre", "rocc_texto"]
RESET_ZERO_KEYS = RESET_NUM_KEYS + RESET_STR_KEYS  # compatibilidad

def _reset_manual_inputs_to_zero():
    for k in RESET_NUM_KEYS:
        st.session_state[k] = 0.0
    st.session_state["x_demanda_acp"] = 19.0
    for k in RESET_STR_KEYS:
        st.session_state[k] = ""

def _fix_session_state_types():
    """Corrige tipos incorrectos en session_state que causan TypeError con protobuf.
    Llama una vez al inicio del app antes de renderizar widgets."""
    for k in RESET_STR_KEYS:
        v = st.session_state.get(k)
        if v is not None and not isinstance(v, str):
            st.session_state[k] = "" if v == 0 or v == 0.0 else str(v)
    for k in RESET_NUM_KEYS:
        v = st.session_state.get(k)
        if v is not None and isinstance(v, str):
            try:
                st.session_state[k] = float(v)
            except (ValueError, TypeError):
                st.session_state[k] = 0.0


def _ensure_numeric_state(key, default_val=0.0, *, replace_if_zero=False):
    """Inicializa un key numérico ANTES de crear el widget asociado.
    Evita StreamlitAPIException por modificar session_state después de instanciar el widget."""
    try:
        default_num = float(default_val or 0.0)
    except Exception:
        default_num = 0.0

    if key not in st.session_state:
        st.session_state[key] = default_num
    elif replace_if_zero:
        try:
            cur = float(st.session_state.get(key, 0.0) or 0.0)
            if abs(cur) < 0.0001 and abs(default_num) > 0.0001:
                st.session_state[key] = default_num
        except Exception:
            st.session_state[key] = default_num
    try:
        return float(st.session_state.get(key, default_num) or 0.0)
    except Exception:
        return default_num

def _ensure_text_state(key, default_val="", *, replace_if_blank=False):
    """Inicializa un key de texto ANTES de crear el widget asociado."""
    default_txt = "" if default_val is None else str(default_val)
    if key not in st.session_state:
        st.session_state[key] = default_txt
    elif replace_if_blank:
        cur = st.session_state.get(key, "")
        if (cur is None or str(cur).strip() == "") and default_txt.strip():
            st.session_state[key] = default_txt
    cur = st.session_state.get(key, default_txt)
    return "" if cur is None else str(cur)

_fix_session_state_types()  # ejecutar inmediatamente al cargar el app

with st.sidebar:
    st.markdown(f"""
    <div class='sidebar-brand'>
        <div class='sidebar-brand-logos'>
            <img src='{ACP_LOGO_URI}' alt='Logo ACP' style='height:64px;'>
            <img src='{HIMH_LOGO_URI}' alt='Logo HIMH' style='height:64px;'>
        </div>
        <div class='sidebar-brand-note'>Autoridad del Canal de Panamá · HIMH</div>
        <div class='sidebar-brand-note'><b>Creado por JFRodriguez</b></div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("## ♻️ Agua y Sostenibilidad")
    st.markdown("**Generador de Informes · ACP**")
    st.markdown("---")

    tipo = st.selectbox("📋 Tipo de informe",
                         ["Mensual", "Trimestral", "Semestral", "Anual"],
                         help="Selecciona la granularidad del informe")
    anio_sel = st.number_input("Año fiscal", 2020, 2045, 2026, step=1,
                               help="Año en que termina el Año Fiscal (p.ej. 2026 = Oct 2025 – Sep 2026)")
    trimestre = None
    semestre  = None
    periodo_sem = f"AF{int(anio_sel)}"

    if tipo == "Mensual":
        mes_sel = st.selectbox("📅 Mes", MESES_ES, index=5)
    elif tipo == "Trimestral":
        trimestre = st.selectbox("📅 Trimestre",
            ["Q1 — I Trim (Oct–Dic)", "Q2 — II Trim (Ene–Mar)",
             "Q3 — III Trim (Abr–Jun)", "Q4 — IV Trim (Jul–Sep)"])
        mes_sel = {"Q1":"Diciembre","Q2":"Marzo","Q3":"Junio","Q4":"Septiembre"}.get(
            trimestre.split()[0], "Marzo")
    elif tipo == "Semestral":
        semestre = st.selectbox("📅 Semestre",
            ["Primer Semestre (Oct–Mar)", "Segundo Semestre (Abr–Sep)"])
        mes_sel = "Marzo" if "Primer" in semestre else "Septiembre"
    else:  # Anual
        mes_sel = "Septiembre"

    st.markdown("---")
    st.markdown("### 📂 Archivos de datos")

    f_daily  = st.file_uploader("📊 DAILY_AND_ACUMMEN (xlsx)",  type=["xlsx","xls"], key="f_daily")
    f_fuentes= st.file_uploader("🏞️ Fuentes_Agua (xlsx)",       type=["xlsx","xls"], key="f_fuentes")
    f_hist_alh = None
    f_hist_gat = None
    f_hist_chcp = st.file_uploader("📈 03_QAVGMCHCP (xls/xlsx)", type=["xlsx","xls"], key="f_hist_chcp")
    f_hist_netfl = st.file_uploader("📉 04_GATNETFL (xls/xlsx)", type=["xlsx","xls"], key="f_netfl")
    st.caption("ℹ️ Se mantienen visibles: DAILY_AND_ACUMMEN, Fuentes_Agua, 03_QAVGMCHCP y 04_GATNETFL. Los históricos 01_QAVGMALHA y 02_QAVGMGAT solo se buscarán automáticamente si activas la detección de archivos locales.")
    st.caption("⚠️ Fuentes_Agua es OBLIGATORIO para los aportes hidrológicos (tributarios por río y subcuenca). 03_QAVGMCHCP y 04_GATNETFL son OBLIGATORIOS para las estadísticas históricas. Sin ellos esas secciones quedarán en cero.")
    auto_local_files = st.checkbox(
        "Usar archivos locales detectados automáticamente",
        value=False,
        key="use_local_auto_files",
        help="Desactivado por defecto. Si no cargas archivos manualmente, el app iniciará en 0. Actívalo solo si quieres que busque archivos en la carpeta local o en /mnt/data."
    )
    st.markdown("### 📄 Documento oficial")
    f_doc_oficial = st.file_uploader("DOCX oficial del período", type=["docx"], key="f_doc_oficial")
    f_sal = st.file_uploader("🌊 Salinidad extra (xlsx/csv)", type=["xlsx","xls","csv"], key="f_sal",
                              help="Solo si la salinidad no está en el DAILY. El modo Manual/Archivo se configura en la pestaña Meteo & Salinidad.")

    st.markdown("---")
    st.markdown("""
    <div style='font-size:.72rem;color:rgba(200,230,255,.92);line-height:1.8'>
    <b>Flujo recomendado:</b><br>
    1️⃣ Selecciona el tipo de informe y año fiscal<br>
    2️⃣ Carga DAILY_AND_ACUMMEN + Fuentes_Agua<br>
    3️⃣ Carga Fuentes_Agua (obligatorio para aportes) y los históricos 03_QAVGMCHCP / 04_GATNETFL (obligatorios para estadísticas históricas)<br>
    4️⃣ La Hydro Interna se genera automáticamente<br>
    5️⃣ Ajusta salinidad o valores manuales si necesario<br>
    6️⃣ Coteja con el DOCX oficial del período<br>
    7️⃣ Exporta Informe DOCX / datos Excel
    </div>
    <div style='font-size:.68rem;color:rgba(180,210,255,.80);margin-top:6px'>
    <b>Tipos de informe disponibles:</b><br>
    📅 Mensual · Trimestral (Q1–Q4)<br>
    📆 Semestral (S1/S2) · Anual completo
    </div>""", unsafe_allow_html=True)
    if st.button("🧹 Resetear ajustes manuales a 0", use_container_width=True):
        _reset_manual_inputs_to_zero()
        st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
#  CARGA Y CÁLCULO DE DATOS
# ══════════════════════════════════════════════════════════════════════════════
def _pick_source(uploaded, patterns=None, allow_auto=False):
    """Devuelve el archivo subido; opcionalmente busca un archivo local detectado automáticamente."""
    if uploaded is not None:
        return uploaded
    if allow_auto and patterns:
        return _find_local_file(patterns)
    return None

def _source_bytes(src):
    if src is None:
        return None
    if hasattr(src, "getvalue"):
        return src.getvalue()
    return Path(src).read_bytes()

def _source_token(src):
    if src is None:
        return "NONE"
    try:
        if hasattr(src, "name"):
            return f"upload::{getattr(src, 'name', 'file')}::{getattr(src, 'size', 'na')}"
        p = Path(src)
        stat = p.stat()
        return f"local::{p.name}::{int(stat.st_mtime)}::{stat.st_size}"
    except Exception:
        return str(getattr(src, "name", src))

AUTO_RESEED_KEYS = list(dict.fromkeys(RESET_ZERO_KEYS + [
    "rep_aporte_total", "rep_aporte_neto", "rep_alh_hm3", "rep_gat_hm3",
    "rep_esclusaje_hm3", "rep_pan_hm3", "rep_neo_hm3", "rep_trasvase_hm3",
    "rep_tpan", "rep_tneo", "rep_ahorro_pan", "rep_ahorro_neo",
    "rep_cca_hm3", "rep_zzflush_hm3",
    "met_precip_mm", "met_precip_hist", "met_periodo", "met_n_anios",
    "met_ranking_txt", "met_dias_lluvia",
    "fn_hist_chcp_hm3", "fn_hist_chcp_m3s", "fn_calado_pan", "fn_calado_neo",
    "fn_pos_hist_txt", "fn_periodo_hist_txt",
    "h_precip_hist_tab7", "h_pos_ranking_tab7"
]))

def _reset_auto_seeded_widget_state():
    for k in AUTO_RESEED_KEYS:
        st.session_state.pop(k, None)

def _norm_text(s):
    import unicodedata
    s = "" if s is None else str(s)
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    s = re.sub(r"\s+", " ", s).strip().lower()
    return s

def _to_float(val):
    if val is None:
        return None
    if isinstance(val, (int, float, np.integer, np.floating)):
        try:
            return float(val)
        except Exception:
            return None
    s = str(val).strip()
    if not s:
        return None
    s = s.replace("\xa0", " ").replace(",", ".")
    m = re.search(r"[-+]?\d+(?:\.\d+)?", s)
    if not m:
        return None
    try:
        return float(m.group(0))
    except Exception:
        return None

def _read_docx_content(src):
    """
    Lee párrafos y tablas de un DOCX sin depender obligatoriamente de python-docx.
    Retorna {"paragraphs": [...], "tables": [[...], ...]}.
    """
    if src is None:
        return {"paragraphs": [], "tables": []}
    raw = _source_bytes(src)
    if not raw:
        return {"paragraphs": [], "tables": []}

    # Preferir python-docx si está disponible
    if DOCX_OK:
        try:
            bio = io.BytesIO(raw)
            doc = Document(bio)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            tables = []
            for tbl in doc.tables:
                rows = []
                for row in tbl.rows:
                    vals = [cell.text.strip() for cell in row.cells]
                    if any(v for v in vals):
                        rows.append(vals)
                if rows:
                    tables.append(rows)
            return {"paragraphs": paragraphs, "tables": tables}
        except Exception:
            pass

    # Fallback XML nativo
    try:
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            xml_bytes = zf.read("word/document.xml")
        root = ET.fromstring(xml_bytes)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        body = root.find("w:body", ns)
        paragraphs, tables = [], []
        if body is None:
            return {"paragraphs": [], "tables": []}
        for child in list(body):
            tag = child.tag.split("}")[-1]
            if tag == "p":
                txt = "".join(t.text or "" for t in child.findall(".//w:t", ns)).strip()
                if txt:
                    paragraphs.append(txt)
            elif tag == "tbl":
                rows = []
                for tr in child.findall("w:tr", ns):
                    vals = []
                    for tc in tr.findall("w:tc", ns):
                        txt = "".join(t.text or "" for t in tc.findall(".//w:t", ns)).strip()
                        vals.append(txt)
                    if any(v for v in vals):
                        rows.append(vals)
                if rows:
                    tables.append(rows)
        return {"paragraphs": paragraphs, "tables": tables}
    except Exception:
        return {"paragraphs": [], "tables": []}

def _parse_official_table_rows(table_rows, kind="fuentes"):
    out = []
    if not table_rows:
        return out
    for row in table_rows[1:]:
        vals = [str(c).strip() for c in row]
        vals = vals + [""] * (3 - len(vals))
        name, val, pct = vals[0], vals[1], vals[2]
        if not name:
            continue
        n = _norm_text(name)
        if ("fuente" in n and "tributario" in n) or ("tipo de uso" in n):
            continue
        if n.startswith("*embalse"):
            continue
        is_total = ("subcuenca" in n) or ("aporte total" in n) or ("total de salidas" in n) or (n == "total, de salidas") or (n == "total de salidas")
        num = _to_float(val)
        out.append((name, round(num, 3) if num is not None else val, pct, is_total))
    return out

def _parse_official_doc_metrics(text, tables=None):
    """
    Extrae métricas clave de la narrativa/tablas oficiales.
    """
    tables = tables or []
    t = re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()
    metrics = {}

    def grab(pattern, key, idx=1, flags=re.I):
        m = re.search(pattern, t, flags)
        if m:
            metrics[key] = _to_float(m.group(idx))

    # Narrativa principal
    grab(r"precipitación total .*? fue de ([\d.,]+)\s*mm", "precipitacion_mm")
    grab(r"promedio histórico de ([\d.,]+)\s*mm", "precipitacion_hist_mm")
    grab(r"aportes totales .*? fueron ([\d.,]+)\s*hm3", "aporte_total_hm3")
    m = re.search(r"desglosados en ([\d.,]+)\s*hm3 .*? y ([\d.,]+)\s*hm3 .*?evapor", t, re.I)
    if m:
        metrics["aporte_neto_hm3"] = _to_float(m.group(1))
        metrics["evaporacion_hm3"] = _to_float(m.group(2))
    m = re.search(r"Alhajuela el ([\d.,]+)%\s*\(([\d.,]+)\s*hm3\).*?Gatún el ([\d.,]+)%\s*\(([\d.,]+)\s*hm3\)", t, re.I)
    if m:
        metrics["alhajuela_pct"] = _to_float(m.group(1))
        metrics["alhajuela_hm3"] = _to_float(m.group(2))
        metrics["gatun_pct"] = _to_float(m.group(3))
        metrics["gatun_hm3"] = _to_float(m.group(4))
    m = re.search(r"uso de agua en las esclusas .*? ha sido de ([\d.,]+)\s*hm[³3].*?([\d.,]+)%\s*\(([\d.,]+)\s*hm3\).*?([\d.,]+)%\s*\(([\d.,]+)\s*hm3\)", t, re.I)
    if m:
        metrics["esclusaje_hm3"] = _to_float(m.group(1))
        metrics["panamax_pct"] = _to_float(m.group(2))
        metrics["panamax_hm3"] = _to_float(m.group(3))
        metrics["neopanamax_pct"] = _to_float(m.group(4))
        metrics["neopanamax_hm3"] = _to_float(m.group(5))
    grab(r"uso promedio diario de agua en las esclusas fue de ([\d.,]+)\s*hm3", "uso_prom_diario_hm3")
    m = re.search(r"NeoPanamax.*?volumen de ([\d.,]+)\s*hm3 .*? Panamax.*? volumen de ([\d.,]+)\s*hm3, para un total de ([\d.,]+)\s*hm3", t, re.I)
    if m:
        metrics["ahorro_neopanamax_hm3"] = _to_float(m.group(1))
        metrics["ahorro_panamax_hm3"] = _to_float(m.group(2))
        metrics["ahorro_total_hm3"] = _to_float(m.group(3))
    else:
        m = re.search(r"panamax, se ahorraron ([\d.,]+)\s*hm3.*?neopanamax .*? ([\d.,]+)\s*hm3", t, re.I)
        if m:
            metrics["ahorro_panamax_hm3"] = _to_float(m.group(1))
            metrics["ahorro_neopanamax_hm3"] = _to_float(m.group(2))
            metrics["ahorro_total_hm3"] = (metrics.get("ahorro_panamax_hm3") or 0) + (metrics.get("ahorro_neopanamax_hm3") or 0)
    grab(r"uso de agua promedio por tránsito de un buque panamax fue de ([\d.,]+)\s*hm3", "agua_panamax_trans_hm3")
    m = re.search(r"CCA.*? fue de ([\d.,]+)\s*hm[³3].*?para ([\d.,]+)\s*tránsitos.*?equivale a .*?([\d.,]+)\s*hm3", t, re.I)
    if m:
        metrics["cca_hm3"] = _to_float(m.group(1))
        metrics["cca_transitos"] = _to_float(m.group(2))
        metrics["cca_unit_hm3"] = _to_float(m.group(3))
    m = re.search(r"Madden fue de ([\d.,]+)\s*MW .*? Gatún fue de ([\d.,]+)\s*MW", t, re.I)
    if m:
        metrics["energia_madden_mw"] = _to_float(m.group(1))
        metrics["energia_gatun_mw"] = _to_float(m.group(2))
    grab(r"trasvasaron ([\d.,]+)\s*hm[³3]", "trasvase_hm3")
    grab(r"Los usos o salidas de agua.*? fueron de? ([\d.,]+)\s*hm3", "total_salidas_hm3")
    grab(r"Mitigación de salinidad .*?\(ZZFlush\)\s*.*?([\d.,]+)\s*hm3", "zzflush_hm3")
    grab(r"salinidad promedio diaria .*? fue .*? ([\d.,]+)\s*unidades prácticas", "salinidad_spc")
    grab(r"SPV\) se registró en ([\d.,]+)\s*ups", "salinidad_spv")
    grab(r"decremento del ([\d.,]+)%", "salinidad_dec_pct")

    # Tablas oficiales
    usage_map = {}
    for table in tables:
        header = _norm_text(" ".join(table[0]))
        rows = _parse_official_table_rows(table, "usos" if "tipo de uso" in header else "fuentes")
        if "tipo de uso" in header:
            for name, val, pct, _ in rows:
                usage_map[_norm_text(name)] = {"value": _to_float(val), "pct": _to_float(pct)}
        elif "fuente" in header and "tributario" in header:
            for name, val, pct, _ in rows:
                key = _norm_text(name)
                usage_map[f"fuente::{key}"] = {"value": _to_float(val), "pct": _to_float(pct)}

    alias = {
        "esclusaje": "esclusaje_hm3",
        "potabilizacion": "potabilizacion_hm3",
        "concesiones fugas filtraciones miscelaneos": "concesiones_hm3",
        "evaporacion": "evaporacion_hm3",
        "hidrogeneracion gatun": "hidro_gatun_hm3",
        "mitigacion de salinidad en esclusas neopanamax (zzflush)": "zzflush_hm3",
        "cca": "cca_hm3",
        "vertidos preventivos": "vertidos_hm3",
        "vertidos": "vertidos_hm3",
        "total de salidas": "total_salidas_hm3",
        "total, de salidas": "total_salidas_hm3",
    }
    for name_n, key in alias.items():
        if key not in metrics and name_n in usage_map and usage_map[name_n]["value"] is not None:
            metrics[key] = usage_map[name_n]["value"]

    return metrics

def _parse_official_report(src):
    if src is None:
        return {}
    payload = _read_docx_content(src)
    paras = [p.strip() for p in payload.get("paragraphs", []) if p and str(p).strip()]
    tables = payload.get("tables", [])
    if not paras and not tables:
        return {}
    title = paras[0] if paras else "Agua y Sostenibilidad"
    subtitle = paras[1] if len(paras) > 1 else ""
    bullets = []
    in_body = False
    for p in paras[2:]:
        pn = _norm_text(p)
        if pn.startswith("balance hidrico"):
            break
        if p:
            bullets.append(p)
    fuentes_rows = []
    usos_rows = []
    for table in tables:
        if not table:
            continue
        header = _norm_text(" ".join(table[0]))
        if "fuente" in header and "tributario" in header:
            fuentes_rows = _parse_official_table_rows(table, "fuentes")
        elif "tipo de uso" in header:
            usos_rows = _parse_official_table_rows(table, "usos")
    text = "\n".join(paras)
    return {
        "title": title,
        "subtitle": subtitle,
        "paragraphs": paras,
        "bullets": bullets,
        "tables": tables,
        "fuentes_rows": fuentes_rows,
        "usos_rows": usos_rows,
        "text": text,
        "metrics": _parse_official_doc_metrics(text, tables=tables),
    }

def _official_doc_patterns(tipo, mes_sel, anio_sel, periodo_sem=None, trimestre=None, semestre=None):
    af = periodo_sem or f"AF{anio_sel}"
    if tipo == "Mensual":
        return [
            f"*Agua*y*sostenibilidad*{mes_sel}*{anio_sel}*.docx",
            f"*agua*y*sostenibilidad*{mes_sel}*{anio_sel}*.docx",
        ]
    if tipo == "Trimestral":
        q = (trimestre or "Q1").split()[0]
        return [
            f"*Agua*y*sostenibilidad*{q}*{anio_sel}*.docx",
            f"*agua*y*sostenibilidad*{q}*{anio_sel}*.docx",
        ]
    if tipo == "Semestral":
        if semestre and "Segundo" in semestre:
            return [
                f"*Agua*y*sostenibilidad*Segundo*Semestre*{af}*.docx",
                f"*agua*y*sostenibilidad*Segundo*Semestre*{af}*.docx",
            ]
        return [
            f"*Agua*y*sostenibilidad*Primer*Semestre*{af}*.docx",
            f"*agua*y*sostenibilidad*Primer*Semestre*{af}*.docx",
        ]
    return [
        f"*Agua*y*sostenibilidad*Anual*{af}*.docx",
        f"*agua*y*sostenibilidad*Anual*{af}*.docx",
        f"*Agua*y*sostenibilidad*{af}*.docx",
    ]

def _hidro_patterns(tipo, mes_sel, anio_sel, periodo_sem=None, trimestre=None, semestre=None):
    af = periodo_sem or f"AF{anio_sel}"
    if tipo == "Mensual":
        return [
            f"*Hidroestadistica*{mes_sel}*{anio_sel}*.xlsx",
            f"*Hidroestadística*{mes_sel}*{anio_sel}*.xlsx",
        ]
    if tipo == "Trimestral":
        q = (trimestre or "Q1").split()[0]
        return [
            f"*Hidroestadistica*{q}*{anio_sel}*.xlsx",
            f"*Hidroestadística*{q}*{anio_sel}*.xlsx",
            f"*Hidroestadistica*{af}*{q}*.xlsx",
        ]
    if tipo == "Semestral":
        if semestre and "Segundo" in semestre:
            return [
                f"*Hidroestadistica*Segundo*Semestre*{anio_sel}*.xlsx",
                f"*Hidroestadística*Segundo*Semestre*{anio_sel}*.xlsx",
                f"*Hidroestadistica*{af}*S2*.xlsx",
            ]
        return [
            f"*Hidroestadistica*Primer*Semestre*{anio_sel}*.xlsx",
            f"*Hidroestadística*Primer*Semestre*{anio_sel}*.xlsx",
            f"*Hidroestadistica*{af}*.xlsx",
        ]
    # Anual
    return [
        f"*Hidroestadistica*Anual*{anio_sel}*.xlsx",
        f"*Hidroestadística*Anual*{anio_sel}*.xlsx",
        f"*Hidroestadistica*{af}*Anual*.xlsx",
        f"*Hidroestadistica*{af}*.xlsx",
    ]

def _build_doc_compare_df(app_vals, ref_vals):
    rows = []
    labels = [
        ("Aporte total", "aporte_total_hm3"),
        ("Aporte neto", "aporte_neto_hm3"),
        ("Evaporación", "evaporacion_hm3"),
        ("Subtotal Alhajuela", "alhajuela_hm3"),
        ("Subtotal Gatún", "gatun_hm3"),
        ("Esclusajes total", "esclusaje_hm3"),
        ("Panamax", "panamax_hm3"),
        ("NeoPanamax", "neopanamax_hm3"),
        ("Potabilización", "potabilizacion_hm3"),
        ("ZZ Flush / CCA", "zzflush_hm3"),
        ("Concesiones", "concesiones_hm3"),
        ("Hidrogeneración Gatún", "hidro_gatun_hm3"),
        ("Total salidas", "total_salidas_hm3"),
        ("Trasvase", "trasvase_hm3"),
    ]
    for lbl, key in labels:
        av = _to_float(app_vals.get(key))
        rv = _to_float(ref_vals.get(key))
        if av is None and rv is None:
            continue
        delta = None if (av is None or rv is None) else av - rv
        estado = "Sin referencia" if rv is None else ("OK" if delta is not None and abs(delta) <= 0.05 else "Revisar")
        rows.append({
            "Variable": lbl,
            "App": av,
            "Documento oficial": rv,
            "Δ": delta,
            "Estado": estado,
        })
    return pd.DataFrame(rows)


# Inicializar sal_mode desde session_state (se actualiza en Tab Meteo & Salinidad)
sal_mode = st.session_state.get("_sal_mode_select", "Manual / Calculado del Daily")

f_daily_src = _pick_source(f_daily, ["*DAILY*ACUMMEN*.xlsx", "*DAILY*ACUMMEN*.xls", "*DAILY AND ACUMMEN*.xlsx", "*DAILY_AND_ACUMMEN*.xlsx"], allow_auto=auto_local_files)
f_hidro_src = _pick_source(None, _hidro_patterns(tipo, mes_sel, anio_sel, periodo_sem, trimestre, semestre), allow_auto=auto_local_files)
f_fuentes_src = _pick_source(f_fuentes, ["*Fuentes_Agua*.xlsx", "*Fuentes*Agua*.xlsx"], allow_auto=auto_local_files)
f_hist_alh_src = _pick_source(f_hist_alh, ["01_QAVGMALHA*.xlsx"], allow_auto=auto_local_files)
f_hist_gat_src = _pick_source(None, ["02_QAVGMGAT*.xlsx"], allow_auto=auto_local_files)
f_hist_chcp_src  = _pick_source(f_hist_chcp,  ["03_QAVGMCHCP*.xlsx"],  allow_auto=auto_local_files)
f_hist_netfl_src = _pick_source(f_hist_netfl, ["04_GATNETFL*.xls","GATNETFL*.xls","GATNETFL*.xlsx"], allow_auto=auto_local_files)
f_doc_oficial_src = _pick_source(f_doc_oficial, _official_doc_patterns(tipo, mes_sel, anio_sel, periodo_sem, trimestre, semestre), allow_auto=auto_local_files)

# ── Rangos históricos configurables ───────────────────────────────────────
_hist_range_seed_key = f"{tipo}|{anio_sel}|{mes_sel}|{trimestre or ''}|{semestre or ''}"
_hist_default_start = _default_hist_range_start_year("chcp")
_hist_default_end = _default_hist_range_end_year(tipo, anio_sel, mes_sel, trimestre, semestre)
_hist_netfl_default_start = _default_hist_range_start_year("netfl")
_hist_netfl_default_end = _default_hist_range_end_year(tipo, anio_sel, mes_sel, trimestre, semestre)
if st.session_state.get("_hist_range_seed_key") != _hist_range_seed_key:
    st.session_state["hist_year_start"] = int(_hist_default_start)
    st.session_state["hist_year_end"] = int(_hist_default_end)
    st.session_state["hist_neto_year_start"] = int(_hist_netfl_default_start)
    st.session_state["hist_neto_year_end"] = int(_hist_netfl_default_end)
    st.session_state["_hist_range_seed_key"] = _hist_range_seed_key

def _reset_chcp():
    st.session_state["hist_year_start"] = int(_hist_default_start)
    st.session_state["hist_year_end"] = int(_hist_default_end)

with st.sidebar.expander("📈 Rango histórico CHCP total", expanded=True):
    st.caption("Escoge manualmente desde qué año hasta qué año quieres calcular el histórico total CHCP.")
    st.caption("Valor por defecto: 1898 hasta el año del cálculo seleccionado (inclusive).")
    hist_year_start = int(st.number_input("Desde (año)", min_value=1800, max_value=2100,
                                          value=int(st.session_state.get("hist_year_start", _hist_default_start)),
                                          step=1, key="hist_year_start"))
    hist_year_end = int(st.number_input("Hasta (año)", min_value=1800, max_value=2100,
                                        value=int(st.session_state.get("hist_year_end", _hist_default_end)),
                                        step=1, key="hist_year_end"))
    _c1, _c2 = st.columns(2)
    with _c1:
        st.button("Usar rango por defecto · CHCP", key="btn_hist_reset_chcp", on_click=_reset_chcp)
    with _c2:
        st.markdown(f"**Rango activo:** {hist_year_start}-{hist_year_end}")
    if hist_year_end < hist_year_start:
        st.warning("El año final es menor que el inicial. Se ajustará automáticamente.")
        hist_year_end = hist_year_start

def _reset_netfl():
    st.session_state["hist_neto_year_start"] = int(_hist_netfl_default_start)
    st.session_state["hist_neto_year_end"] = int(_hist_netfl_default_end)

with st.sidebar.expander("📉 Rango histórico GATNETFL (netos)", expanded=True):
    st.caption("Escoge manualmente desde qué año hasta qué año quieres calcular el histórico neto.")
    st.caption("Valor por defecto: 1914 hasta el año anterior del cálculo seleccionado.")
    hist_neto_year_start = int(st.number_input("Desde (año) · netos", min_value=1800, max_value=2100,
                                               value=int(st.session_state.get("hist_neto_year_start", _hist_netfl_default_start)),
                                               step=1, key="hist_neto_year_start"))
    hist_neto_year_end = int(st.number_input("Hasta (año) · netos", min_value=1800, max_value=2100,
                                             value=int(st.session_state.get("hist_neto_year_end", _hist_netfl_default_end)),
                                             step=1, key="hist_neto_year_end"))
    _c3, _c4 = st.columns(2)
    with _c3:
        st.button("Usar rango por defecto · netos", key="btn_hist_reset_netfl", on_click=_reset_netfl)
    with _c4:
        st.markdown(f"**Rango activo:** {hist_neto_year_start}-{hist_neto_year_end}")
    if hist_neto_year_end < hist_neto_year_start:
        st.warning("El año final de netos es menor que el inicial. Se ajustará automáticamente.")
        hist_neto_year_end = hist_neto_year_start

current_source_signature = "|".join([
    str(tipo), str(anio_sel), str(mes_sel), str(trimestre or ""), str(semestre or ""),
    f"hist_range={hist_year_start}-{hist_year_end}",
    f"hist_range_netfl={hist_neto_year_start}-{hist_neto_year_end}",
    f"auto={int(bool(auto_local_files))}",
    _source_token(f_daily_src), _source_token(f_fuentes_src), _source_token(f_hist_alh_src),
    _source_token(f_hist_gat_src), _source_token(f_hist_chcp_src),
    _source_token(f_hist_netfl_src), _source_token(f_hidro_src),
    _source_token(f_doc_oficial_src), _source_token(f_sal),
])
if st.session_state.get("_last_source_signature") != current_source_signature:
    _reset_auto_seeded_widget_state()
    st.session_state["_last_source_signature"] = current_source_signature

datos = {}        # resultados calculados
hist  = {}        # datos históricos
fuentes_rios = [] # tributarios por río
fuentes_sub  = {} # subtotales subcuenca
hidro_audit  = {}
hidro_rows1  = []
hidro_rows2  = []
hidro_file_audit = {}
hidro_file_rows1 = []
hidro_file_rows2 = []
oficial_doc  = {}
oficial_metrics = {}
data_ready   = False
n_dias       = 0

if f_daily_src:
    with st.spinner("⚙️ Procesando datos diarios…"):
        try:
            _daily_bytes = _source_bytes(f_daily_src)
            df_daily = load_daily(_daily_bytes)

            # ── Detectar rango de fechas disponible en el DAILY ──────────────────
            _daily_dates = df_daily["ACTDATE"].dropna()
            if not _daily_dates.empty:
                _min_date = _daily_dates.min()
                _max_date = _daily_dates.max()
                st.sidebar.info(
                    f"📅 **DAILY cargado:** {_min_date.strftime('%d/%m/%Y')} – "
                    f"{_max_date.strftime('%d/%m/%Y')}"
                )
                # Verificar que el período seleccionado esté cubierto
                if tipo == "Mensual":
                    _mes_n = MESES_NUM.get(mes_sel, 1)
                    _yr_check = int(anio_sel) - 1 if _mes_n in [10, 11, 12] else int(anio_sel)
                    _has_data = not df_daily[
                        (df_daily["ACTDATE"].dt.year == _yr_check) &
                        (df_daily["ACTDATE"].dt.month == _mes_n)
                    ].empty
                    if not _has_data:
                        _avail_years = sorted(_daily_dates.dt.year.unique().astype(int).tolist())
                        st.sidebar.warning(
                            f"⚠️ No se encontraron datos para **{mes_sel} {anio_sel}** "
                            f"en el DAILY cargado.\n\n"
                            f"Años disponibles: {', '.join(map(str, _avail_years))}.\n"
                            f"Cambia el **Año fiscal** en el selector."
                        )

            if tipo == "Mensual":
                datos, n_dias = calcular_mes(df_daily, mes_sel, anio_sel)
            else:
                meses_calc = _get_meses_periodo(tipo, int(anio_sel), trimestre, semestre)
                datos, n_dias = calcular_periodo_generico(df_daily, meses_calc)
            if datos:
                data_ready = True
        except Exception as e:
            st.error(f"Error al leer DAILY_AND_ACUMMEN: {e}")

# Históricos
if f_hist_chcp_src or f_hist_netfl_src or f_hist_alh_src or f_hist_gat_src:
    with st.spinner("Calculando estadísticas históricas…"):
        try:
            df_alh, df_gat, df_chcp, df_netfl, hist_load_errors = load_hist(
                f_hist_alh_src, f_hist_gat_src, f_hist_chcp_src, f_hist_netfl_src)

            # Mostrar solo errores realmente relevantes para las fuentes visibles o detectadas
            relevant_errors = []
            if f_hist_chcp_src and df_chcp.empty:
                relevant_errors.extend([e for e in hist_load_errors if e])
            elif f_hist_netfl_src and df_netfl.empty:
                relevant_errors.extend([e for e in hist_load_errors if e and "Archivo vacío" not in e])

            if relevant_errors:
                st.warning(f"⚠️ Error al leer históricos: {'; '.join(relevant_errors)}")

            if f_hist_chcp_src and not df_chcp.empty:
                st.sidebar.success(f"✅ 03_QAVGMCHCP: {len(df_chcp)} años")
            elif f_hist_chcp_src and df_chcp.empty:
                st.sidebar.warning("⚠️ 03_QAVGMCHCP cargado pero no se pudo interpretar.")

            if f_hist_netfl_src and df_netfl.empty:
                st.sidebar.info("ℹ️ 04_GATNETFL cargado, pero no se pudo usar para netos históricos. El promedio histórico CHCP sí puede calcularse con 03_QAVGMCHCP.")
            elif f_hist_netfl_src and not df_netfl.empty:
                _netfl_unit_msg = str(getattr(df_netfl, 'attrs', {}).get('value_unit', 'm3s')).lower()
                st.sidebar.success(f"✅ 04_GATNETFL: {len(df_netfl)} años de aportes netos ({_netfl_unit_msg})")

            if not df_alh.empty or not df_gat.empty or not df_chcp.empty or not df_netfl.empty:
                calcular_historicos._trimestre = trimestre
                calcular_historicos._semestre = semestre
                hist = calcular_historicos(
                    df_alh, df_gat, df_chcp,
                    mes_sel if tipo == "Mensual" else "Marzo",
                    tipo=tipo, anio_fiscal=int(anio_sel),
                    df_netfl=df_netfl if not df_netfl.empty else None,
                    year_start=hist_year_start, year_end=hist_year_end,
                    year_start_netfl=hist_neto_year_start, year_end_netfl=hist_neto_year_end
                )
                if not hist and f_hist_chcp_src:
                    st.sidebar.warning("⚠️ Se cargó 03_QAVGMCHCP pero no se pudo calcular el promedio histórico para el período seleccionado.")
        except Exception as e:
            st.warning(f"Error históricos: {e}")

# Hidroestadística local detectada automáticamente = referencia / auditoría.
fuentes_rios_raw = []
if f_hidro_src:
    try:
        hidro_data = load_hidro(f_hidro_src)
        hidro_file_audit = load_hidro_audit(_source_bytes(f_hidro_src))
        _hname = getattr(f_hidro_src, "name", None) or Path(str(f_hidro_src)).name
        if _matches_period_name(_hname, tipo, mes_sel, anio_sel, periodo_sem):
            hidro_file_rows1 = hidro_file_audit.get("sheet1_rows", [])
            hidro_file_rows2 = hidro_file_audit.get("sheet2_rows", [])
            h1 = hidro_data.get("Hidroestadística 1", pd.DataFrame())
            if not h1.empty and datos:
                prec = _v(h1, "Precipitación", 2)
                if prec is not None:
                    datos["precipitacion_mm"] = float(prec)
                evap_d = _v(h1, "Evaporación diaria", 2)
                if evap_d is not None and float(datos.get("evaporacion_diaria_mm", 0) or 0) == 0:
                    datos["evaporacion_diaria_mm"] = float(evap_d)
        else:
            st.sidebar.warning("⚠️ La Hidroestadística local detectada no coincide con el período seleccionado. Se usará solo la Hidroestadística interna generada desde DAILY + históricos.")
    except Exception as e:
        st.warning(f"Error Hidroestadística: {e}")

# Fuentes por tributario — desde archivo Fuentes_Agua
fuentes_rios_alh = []
fuentes_rios_gat = []
if f_fuentes_src:
    with st.spinner("⚙️ Cargando fuentes de agua…"):
        try:
            # Leer bytes UNA sola vez → dos BytesIO independientes
            # (pd.ExcelFile agota el stream; openpyxl necesita su propia copia)
            _fa_raw_bytes = _source_bytes(f_fuentes_src)
            _fa_bio1 = io.BytesIO(_fa_raw_bytes)
            _fa_bio2 = io.BytesIO(_fa_raw_bytes)

            xls_fa = pd.ExcelFile(_fa_bio1)
            target = _fuentes_sheet_target(tipo, int(anio_sel), mes_sel, trimestre, semestre)
            ay_full = str(anio_sel); ay_short = ay_full[2:]
            target_alt = target.replace(ay_short, ay_full) if ay_short in target else None

            def _norm(s): return s.replace(" ", "").replace("-", "").lower()
            sh_name = next((s for s in xls_fa.sheet_names
                            if _norm(s) == _norm(target)
                            or (target_alt and _norm(s) == _norm(target_alt))), None)
            if sh_name is None:
                sh_name = next((s for s in xls_fa.sheet_names
                                if _norm(target) in _norm(s)
                                or (target_alt and _norm(target_alt) in _norm(s))), None)
            if sh_name is None:
                st.sidebar.warning("⚠️ Fuentes_Agua no contiene la hoja del período. Hojas disponibles: " + ", ".join(xls_fa.sheet_names))
            else:
                # Usar la segunda copia para openpyxl data_only=True
                # (lee valores cacheados de fórmulas en lugar de strings "=C5*0.0864*...")
                try:
                    import openpyxl as _opx
                    _wb_d = _opx.load_workbook(_fa_bio2, data_only=True, read_only=True)
                    _ws_d = _wb_d[sh_name]
                    _rows_d = [[cell.value for cell in row] for row in _ws_d.iter_rows()]
                    _wb_d.close()
                    df_fa_raw = pd.DataFrame(_rows_d)
                except Exception as _e_opx:
                    # Fallback a pd.read_excel desde bytes frescos
                    df_fa_raw = pd.read_excel(io.BytesIO(_fa_raw_bytes), sheet_name=sh_name, header=None)
                _rios, _subtot = parse_fuentes(df_fa_raw)
                fuentes_sub.update(_subtot)
                for entry in _rios:
                    if entry.get("cuenca") == "Alhajuela":
                        fuentes_rios_alh.append(entry)
                    else:
                        fuentes_rios_gat.append(entry)
                st.session_state["_fa_sheet_used"] = sh_name
                st.sidebar.success(f"✅ Fuentes: {sh_name}")
        except Exception as e:
            st.sidebar.warning(f"⚠️ Fuentes_Agua: {e}")

if not fuentes_rios_alh:
    fuentes_rios_alh = []
if not fuentes_rios_gat:
    fuentes_rios_gat = []

# ── Asignar aportes CHCP desde Fuentes de Agua (ÚNICA fuente autoritativa) ────
# Los campos aporte_total_hm3/m3s, aporte_neto_*, alhajuela_*, gatun_* y aporte_gl/md
# se dejan en 0.0 por calcular_mes(). Solo se populan aquí, desde Fuentes_Agua.
# TOFCHCP / TOFGL / TOFMD del DAILY son balances operacionales (salidas + evaporación)
# y NO representan aportes hidrológicos reales.
if data_ready and fuentes_sub:
    _fa_alh = float(fuentes_sub.get("Alhajuela", {}).get("hm3", 0) or 0)
    _fa_gat = float(fuentes_sub.get("Gatún", {}).get("hm3", 0) or 0)
    if _fa_alh > 0 or _fa_gat > 0:
        _fa_total = round(_fa_alh + _fa_gat, 3)
        _fa_total_m3s = round(_fa_total / (n_dias * HM3_PER_M3S_DAY), 3) if n_dias else 0
        datos["aporte_total_hm3"] = _fa_total
        datos["aporte_total_m3s"] = _fa_total_m3s
        datos["alhajuela_hm3"]    = round(_fa_alh, 3)
        datos["gatun_hm3"]        = round(_fa_gat, 3)
        if _fa_total > 0:
            datos["alhajuela_pct"] = round(_fa_alh / _fa_total * 100, 2)
            datos["gatun_pct"]     = round(_fa_gat / _fa_total * 100, 2)
        _evap = float(datos.get("evaporacion_hm3", 0) or 0)
        datos["aporte_neto_hm3"] = round(_fa_total - _evap, 3)
        datos["aporte_neto_m3s"] = round((_fa_total - _evap) / (n_dias * HM3_PER_M3S_DAY), 3) if n_dias else 0
        datos["excedente_hm3"]   = round(datos["aporte_total_hm3"] - float(datos.get("total_salidas_hm3", 0) or 0), 3)  # balance correcto
        datos["_aporte_source"]  = "fuentes_agua"
        st.sidebar.success(f"✅ Aportes CHCP: {datos['aporte_total_hm3']:.1f} hm³ "
                           f"(Alh: {datos['alhajuela_hm3']:.1f} | Gat: {datos['gatun_hm3']:.1f})")
    else:
        datos["_aporte_source"] = "missing"  # fuentes_sub vacío — aportes = 0
        if data_ready:
            st.sidebar.warning(f"⚠️ Fuentes_Agua cargada pero subtotales = 0. "
                               f"Verifica que la hoja tenga filas 'Subcuenca embalse Alhajuela/Gatún'.")
else:
    if data_ready:
        datos["_aporte_source"] = "missing"  # Fuentes_Agua no cargada — aportes = 0


# Hidroestadística interna generada desde DAILY + históricos.
if data_ready:
    hidro_audit = _build_internal_hydro(tipo, mes_sel, anio_sel, periodo_sem,
                                        datos, n_dias, hist, fuentes_sub,
                                        trimestre=trimestre, semestre=semestre)
    hidro_rows1 = hidro_audit.get("sheet1_rows", [])
    hidro_rows2 = hidro_audit.get("sheet2_rows", [])
    if not f_hidro_src:
        st.sidebar.info("ℹ️ No se encontró un Excel local de Hidroestadística para este período. El informe usará únicamente la Hydro Interna calculada por el app.")

# Salinidad / ZZ manuales
sal_manual = None
sal_series_ext = None
if sal_mode == "Subir archivo Excel" and f_sal:
    try:
        if f_sal.name.endswith(".csv"):
            df_sal_raw = pd.read_csv(f_sal)
        else:
            df_sal_raw = pd.read_excel(f_sal)
        sal_col_found = next(
            (c for c in df_sal_raw.columns if "salin" in c.lower() or "ups" in c.lower()), None)
        if sal_col_found:
            sal_vals_ext = df_sal_raw[sal_col_found].dropna().astype(float)
            sal_manual   = float(sal_vals_ext.mean())
            sal_series_ext = sal_vals_ext.reset_index(drop=True)
    except Exception as e:
        st.warning(f"No se pudo leer salinidad del archivo: {e}")

if data_ready:
    # ZZFlush = CCA manual (Tab Salinidad) → sobreescribe el valor inicial (GATSPILL)
    _cca_override = round(float(st.session_state.get("rep_cca_hm3", 0.0) or 0.0), 4)
    datos["zzflush_auto_hm3"]  = _cca_override
    datos["zzflush_auto_avail"]= False
    # Vertidos Gatún = GATSPILL bruto − CCA (evitar doble conteo en total_salidas)
    _gatspill_raw = float(datos.get("gatspill_raw_hm3", 0.0) or 0.0)
    datos["vertidos_gat_hm3"]  = round(max(0.0, _gatspill_raw - _cca_override), 3)
    if sal_mode == "Subir archivo Excel" and sal_manual is not None:
        datos["salinidad_spc"] = round(float(sal_manual), 4)
        datos["salinidad_source"] = "archivo"
        datos["sal_series"] = sal_series_ext if sal_series_ext is not None else pd.Series(dtype=float)
        datos["salinidad_spv"] = float(st.session_state.get("_sal_spv_file", datos.get("salinidad_spv", 0.0)) or 0.0)
        datos["salinidad_dec_pct"] = float(st.session_state.get("_sal_dec_file", datos.get("salinidad_dec_pct", 0.0)) or 0.0)
    else:
        datos["salinidad_spc"] = float(st.session_state.get("x_sspc", datos.get("salinidad_spc", 0.0)) or 0.0)
        datos["salinidad_spv"] = float(st.session_state.get("x_sspv", datos.get("salinidad_spv", 0.0)) or 0.0)
        datos["salinidad_dec_pct"] = float(st.session_state.get("x_sdec", datos.get("salinidad_dec_pct", 0.0)) or 0.0)
        datos["salinidad_source"] = "manual"
        datos["sal_series"] = pd.Series(dtype=float)
    _sync_balance_fields(datos, n_dias)
# Documento oficial del período (narrativa/tablas de referencia)
if f_doc_oficial_src:
    try:
        oficial_doc = _parse_official_report(f_doc_oficial_src)
        oficial_metrics = oficial_doc.get("metrics", {}) if oficial_doc else {}
        _doc_name = getattr(f_doc_oficial_src, "name", None) or Path(str(f_doc_oficial_src)).name
        st.sidebar.success(f"✅ DOCX oficial: {_doc_name}")
    except Exception as e:
        st.sidebar.warning(f"⚠️ DOCX oficial: {e}")
        oficial_doc = {}
        oficial_metrics = {}

# ══════════════════════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════════════════════
label_per = _label_periodo(tipo, mes_sel, int(anio_sel), trimestre, semestre)

# ── Banner de cobertura de datos — muestra exactamente qué rango cubre cada fuente ──
def _banner_rango():
    items = []

    # ── DAILY ──────────────────────────────────────────────────────────────────
    if f_daily_src and data_ready and "df_daily" in globals():
        try:
            _dts = df_daily["ACTDATE"].dropna()
            if not _dts.empty:
                _d0 = _dts.min(); _d1 = _dts.max()
                # Exact rows for selected period
                if tipo == "Mensual":
                    _mn = MESES_NUM.get(mes_sel,1)
                    _yr = int(anio_sel)-1 if _mn in [10,11,12] else int(anio_sel)
                    _per_rows = df_daily[(df_daily["ACTDATE"].dt.year==_yr)&
                                        (df_daily["ACTDATE"].dt.month==_mn)]
                else:
                    _meses_p = _get_meses_periodo(tipo, int(anio_sel), trimestre, semestre)
                    _mask = pd.Series(False, index=df_daily.index)
                    for _mn2,_yr2 in _meses_p:
                        _mask |= ((df_daily["ACTDATE"].dt.year==_yr2)&
                                  (df_daily["ACTDATE"].dt.month==_mn2))
                    _per_rows = df_daily[_mask]
                if not _per_rows.empty:
                    _p0 = _per_rows["ACTDATE"].min(); _p1 = _per_rows["ACTDATE"].max()
                    items.append(
                        f"<span class='_br-item _br-ok'>"
                        f"<b>📋 DAILY</b>&nbsp; "
                        f"{_p0.strftime('%d %b %Y')} → {_p1.strftime('%d %b %Y')}"
                        f"&nbsp;<span class='_br-days'>({len(_per_rows)} días)</span>"
                        f"&nbsp;<span class='_br-total'>· total: "
                        f"{_d0.strftime('%d/%m/%Y')} – {_d1.strftime('%d/%m/%Y')}"
                        f"</span></span>"
                    )
                else:
                    items.append(
                        f"<span class='_br-item _br-warn'>"
                        f"<b>📋 DAILY</b>&nbsp; ⚠️ Sin datos para <b>{label_per}</b>"
                        f"</span>"
                    )
        except Exception:
            pass
    elif f_daily_src:
        items.append("<span class='_br-item _br-warn'><b>📋 DAILY</b>&nbsp; ⚠️ Cargando…</span>")
    else:
        items.append("<span class='_br-item _br-miss'><b>📋 DAILY</b>&nbsp; No cargado</span>")

    # ── Fuentes de Agua ─────────────────────────────────────────────────────────
    if f_fuentes_src:
        _fa_name = getattr(f_fuentes_src,"name",None) or "Fuentes_Agua"
        _fa_ok   = bool(fuentes_rios_alh or fuentes_rios_gat)
        if _fa_ok:
            _sh = st.session_state.get("_fa_sheet_used","—")
            items.append(
                f"<span class='_br-item _br-ok'>"
                f"<b>🏞️ Fuentes_Agua</b>&nbsp; ✅ hoja <b>{_sh}</b>"
                f"</span>"
            )
        else:
            items.append(
                f"<span class='_br-item _br-warn'>"
                f"<b>🏞️ Fuentes_Agua</b>&nbsp; ⚠️ sin hoja para <b>{label_per}</b>"
                f"</span>"
            )
    else:
        items.append("<span class='_br-item _br-miss'><b>🏞️ Fuentes_Agua</b>&nbsp; No cargado → Aportes = 0</span>")

    # ── Salinidad ───────────────────────────────────────────────────────────────
    sal_src_label = str(datos.get("salinidad_source","—") if data_ready else "—").upper()
    sal_val = datos.get("salinidad_spc",0.0) if data_ready else 0.0
    if sal_val and float(sal_val)>0:
        items.append(
            f"<span class='_br-item _br-ok'>"
            f"<b>🌊 Salinidad</b>&nbsp; {float(sal_val):.4f} ups [{sal_src_label}]"
            f"</span>"
        )
    else:
        items.append("<span class='_br-item _br-miss'><b>🌊 Salinidad SPC</b>&nbsp; ✏️ Sin dato — ingresar en Tab Salinidad</span>")

    # ── Históricos QAVGMCHCP ─────────────────────────────────────────────────────
    if f_hist_chcp_src:
        _n_yr = hist.get("n_years_chcp",0) if hist else 0
        _yrs  = hist.get("years_chcp")     if hist else None
        _rng  = f"{_yrs[0]}–{_yrs[1]}" if _yrs else "—"
        items.append(
            f"<span class='_br-item {'_br-ok' if _n_yr else '_br-warn'}'>"
            f"<b>📈 QAVGMCHCP</b>&nbsp; {_rng} ({_n_yr} años)"
            f"</span>"
        )
    else:
        items.append("<span class='_br-item _br-miss'><b>📈 Hist.CHCP</b>&nbsp; No cargado</span>")

    # ── Históricos GATNETFL ───────────────────────────────────────────────────────
    if f_hist_netfl_src:
        _n_netfl = hist.get("n_years_netfl",0) if hist else 0
        _yrs_n   = hist.get("years_netfl")     if hist else None
        _rng_n   = f"{_yrs_n[0]}–{_yrs_n[1]}" if _yrs_n else "—"
        items.append(
            f"<span class='_br-item {'_br-ok' if _n_netfl else '_br-warn'}'>"
            f"<b>🌊 GATNETFL</b>&nbsp; {_rng_n} ({_n_netfl} años)"
            f"</span>"
        )
    else:
        items.append("<span class='_br-item _br-miss'><b>🌊 GATNETFL</b>&nbsp; No cargado</span>")

    # ── Datos meteorológicos ──────────────────────────────────────────────────────
    _prec_mm   = float(st.session_state.get("x_prec",  0.0) or 0.0)
    _prec_hist = float(st.session_state.get("x_phist", 0.0) or 0.0)
    _prec_pos  = int(st.session_state.get("x_ppos",    0)   or 0)
    _prec_per  = str(st.session_state.get("x_pper",    "")  or "")
    if _prec_mm > 0:
        _prec_pct = float(st.session_state.get("x_met_pct_informe",
            round(abs(_prec_mm-_prec_hist)/_prec_hist*100,1) if _prec_hist>0 else 0.0))
        _dir = "▲" if _prec_mm >= _prec_hist else "▼"
        items.append(
            f"<span class='_br-item _br-ok'>"
            f"<b>🌧️ Meteorología</b>&nbsp; {_prec_mm:.0f} mm"
            f"&nbsp;{_dir}&nbsp;{_prec_pct:.1f}% vs hist."
            f"{f' · pos {_prec_pos}' if _prec_pos else ''}"
            f"{f' · {_prec_per}' if _prec_per else ''}"
            f"</span>"
        )
    else:
        items.append("<span class='_br-item _br-miss'><b>🌧️ Meteorología</b>&nbsp; ✏️ Sin dato — Tab Meteo</span>")

    # ── Documento oficial ───────────────────────────────────────────────────────
    if f_doc_oficial_src:
        _dname = getattr(f_doc_oficial_src,"name",None) or "DOCX oficial"
        items.append(f"<span class='_br-item _br-ok'><b>📄 DOCX oficial</b>&nbsp; {_dname}</span>")

    # ── Render ──────────────────────────────────────────────────────────────────
    if not items:
        return
    sep = "<span class='_br-sep'>│</span>"
    st.markdown(f"""
<style>
._br-banner{{display:flex;flex-wrap:wrap;align-items:center;gap:4px 8px;
  background:var(--c-surface,#1e2a3a);border:1px solid var(--c-border,#334);
  border-radius:8px;padding:7px 14px;margin:6px 0 10px;font-size:.75rem;}}
._br-item{{display:inline-flex;align-items:center;gap:4px;padding:2px 8px;
  border-radius:5px;white-space:nowrap;}}
._br-ok  {{background:rgba(22,163,74,.15);color:#6ee7b7;border:1px solid rgba(22,163,74,.3);}}
._br-warn{{background:rgba(234,179,8,.12);color:#fde68a;border:1px solid rgba(234,179,8,.3);}}
._br-miss{{background:rgba(100,116,139,.12);color:#94a3b8;border:1px solid rgba(100,116,139,.25);}}
._br-days{{font-weight:800;color:#60a5fa;}}
._br-total{{opacity:.65;font-size:.7rem;}}
._br-sep{{color:var(--c-border,#334);font-size:.9rem;}}
</style>
<div class='_br-banner'>
  <span style='font-size:.72rem;font-weight:700;color:var(--c-text-muted,#94a3b8);
               margin-right:4px;white-space:nowrap'>📌 Cobertura&nbsp;{label_per}:</span>
  {f"&nbsp;{sep}&nbsp;".join(items)}
</div>""", unsafe_allow_html=True)

_banner_rango()

_logo_html = ""
if ACP_LOGO_URI:
    _logo_html += f"<img src='{ACP_LOGO_URI}' alt='ACP'>"
if HIMH_LOGO_URI:
    _logo_html += f"<img src='{HIMH_LOGO_URI}' alt='HIMH'>"

_logos_block = f"<div class='brand-left'>{_logo_html}</div>" if _logo_html else "<div class='brand-left'></div>"
_tipo_badge  = _badge_periodo(tipo)

st.markdown(f"""
<div class='brand-header'>
  {_logos_block}
  <div class='brand-center'>
    <span class='brand-emoji-big'>♻️</span>
    <div class='brand-text-col'>
      <div class='brand-title'>Agua y Sostenibilidad</div>
      <div class='brand-subtitle'>
        Cuenca Hidrográfica del Canal de Panamá &nbsp;·&nbsp; <b>{label_per}</b>
      </div>
    </div>
  </div>
  <div class='brand-right'>
    <div>Autoridad del Canal de Panamá</div>
    <div>HIMH - Hidrología</div>
    <div class='brand-creator'>👤 JFRodriguez</div>
    <div style='color:#6dcfff!important;font-size:.71rem'>📅 {_tipo_badge}</div>
  </div>
</div>
""", unsafe_allow_html=True)

# JS: sincroniza data-theme en el root para que el CSS dark-mode funcione
st.markdown("""
<script>
(function(){
  function applyTheme(){
    var dark = window.matchMedia('(prefers-color-scheme: dark)').matches;
    // Streamlit also sets class on the body
    var hasDark = document.body.classList.contains('dark') ||
                  document.documentElement.getAttribute('data-theme') === 'dark';
    var isDark = dark || hasDark ||
                 document.body.style.backgroundColor === 'rgb(14,17,23)' ||
                 document.body.style.backgroundColor === 'rgb(13,17,23)';
    document.documentElement.setAttribute('data-theme', isDark ? 'dark' : 'light');
  }
  applyTheme();
  window.matchMedia('(prefers-color-scheme: dark)').addEventListener('change', applyTheme);
  var obs = new MutationObserver(applyTheme);
  obs.observe(document.body, {attributes:true, attributeFilter:['class','style']});
})();
</script>
""", unsafe_allow_html=True)

if not data_ready:
    st.markdown(f"""
    <div class='info-box'>
    ℹ️ <b>Para comenzar</b>, carga el archivo <code>DAILY_AND_ACUMMEN.xlsx</code> en el panel izquierdo.<br>
    Período seleccionado: <b>{label_per}</b> · Tipo: <b>{tipo}</b><br>
    Los demás archivos son <b>obligatorios</b> para el análisis completo: <b>Fuentes_Agua</b> para los aportes hidrológicos por río y subcuenca, y <b>03_QAVGMCHCP / 04_GATNETFL</b> para las estadísticas históricas. Sin ellos esas secciones quedarán en cero.
    </div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  TABS
# ══════════════════════════════════════════════════════════════════════════════
TABS = ["📊 Dashboard","🏞️ Fuentes de Agua","🌡️ Meteorología","🌊 Salinidad","📋 ROCC",
        "♻️ Balance Hídrico","🚢 Esclusajes","📤 Usos del Agua","⚡ Energía",
        "💧 Evaporación","📈 Comparativa Histórica","📝 Exportar"]
tabs = st.tabs(TABS)

# Valores con fallback
def D(key, default=0.0):
    return datos.get(key, default) if data_ready else default

def H(label, which="value1", unit_hint=None, row_min=None, default=None):
    return hidro_num(hidro_rows1, label, which=which, unit_hint=unit_hint, row_min=row_min, default=default)

def H2(label, which="value", default=None):
    try:
        r = hidro_row(hidro_rows2, label)
        v = r.get(which) if r else None
        return float(v) if v is not None else default
    except Exception:
        return default

def fmt_num(v, nd=2):
    try:
        return f"{float(v):,.{nd}f}"
    except Exception:
        return "—"

def audit_compare_df():
    if not hidro_rows1:
        return pd.DataFrame()
    rows = []
    def add(item, app_hm3=None, hydro_label=None, hydro_hm3=None, app_m3s=None, hydro_m3s=None):
        h_hm3 = hydro_hm3 if hydro_hm3 is not None else H(hydro_label, "value1", "hm3")
        h_m3s = hydro_m3s if hydro_m3s is not None else H(hydro_label, "value2", "m3/s")
        a_hm3 = app_hm3
        a_m3s = app_m3s if app_m3s is not None else (hm3_to_m3s_period(app_hm3, n_dias) if app_hm3 is not None else None)
        rows.append({
            "Concepto": item,
            "App hm³": a_hm3,
            "Hidro hm³": h_hm3,
            "Δ hm³": (a_hm3 - h_hm3) if (a_hm3 is not None and h_hm3 is not None) else None,
            "App m³/s": a_m3s,
            "Hidro m³/s": h_m3s,
            "App EED": m3s_to_eed(a_m3s) if a_m3s is not None else None,
            "Hidro EED": m3s_to_eed(h_m3s) if h_m3s is not None else None,
        })
    add("Aporte Total", D("aporte_total_hm3"), "Aporte Total", app_m3s=D("aporte_total_m3s"))
    add("Evaporación mensual", D("evaporacion_hm3"), "Evaporación mensual", app_m3s=D("evaporacion_m3s"))
    add("Aporte Neto", D("aporte_neto_hm3"), "Aporte Neto", app_m3s=D("aporte_neto_m3s"))
    add("Esclusajes PNX", D("panamax_hm3"), "Esclusajes PNX")
    add("Esclusajes NPX", D("neopanamax_hm3"), "Esclusajes NPX")
    add("Esclusajes Totales", D("esclusaje_hm3"),
        hydro_hm3=(H("Esclusajes PNX","value1","hm3", default=0) or 0) + (H("Esclusajes NPX","value1","hm3", default=0) or 0),
        app_m3s=D("esclusaje_m3s"),
        hydro_m3s=(H("Esclusajes PNX","value2","m3/s", default=0) or 0) + (H("Esclusajes NPX","value2","m3/s", default=0) or 0))
    add("Potabilización Total", D("potabilizacion_hm3"),
        hydro_hm3=(H("Potabilización Gatún","value1","hm3", default=0) or 0) + (H("Potabilización Alhajuela","value1","hm3", default=0) or 0),
        app_m3s=D("potabilizacion_m3s"),
        hydro_m3s=(H("Potabilización Gatún","value2","m3/s", default=0) or 0) + (H("Potabilización Alhajuela","value2","m3/s", default=0) or 0))
    add("Concesiones y Misceláneos", D("concesiones_hm3"), "Concesiones y Misceláneos", app_m3s=D("concesiones_m3s"))
    add("Hidrogeneración Gatún", D("hidro_gatun_hm3"), "Hidrogeneración Gatún")
    # ZZFlush: app usa valor auto de GATSPILL; hidro usa hoja Hidroestadística 2
    add("ZZ Flush (manual)",
        D("zzflush_auto_hm3"),
        "Vertidos ZZ Flush",
        hydro_hm3=H("Vertidos ZZ Flush","value1","hm3"),
        hydro_m3s=H("Vertidos ZZ Flush","value2","m3/s"))
    add("Ahorro Panamax", D("ahorro_panamax_hm3"), "Panamax", hydro_hm3=H("Panamax","value1","hm3", row_min=54))
    add("Ahorro NeoPanamax", D("ahorro_neopanamax_hm3"), "NeoPanamax Total")
    add("Ahorro Total", D("ahorro_total_hm3"), "Panamax + NeoPanamax")
    add("Usos Totales", D("total_salidas_hm3"), "Usos Totales",
        app_m3s=hm3_to_m3s_period(D("total_salidas_hm3"), n_dias))
    add("Trasvase Madden", D("trasvase_hm3"), "Trasvase Madden", app_m3s=D("trasvase_m3s"))
    return pd.DataFrame(rows)

# ═══════════════════════════════════════════════════════════════
#  TAB 0 — DASHBOARD
# ═══════════════════════════════════════════════════════════════
with tabs[0]:
    if not data_ready:
        st.info("Carga el archivo DAILY_AND_ACUMMEN para ver el dashboard.")
    else:
        # Banner para semestral con meses disponibles
        if tipo != "Mensual" and datos.get("_es_parcial"):
            disp = ", ".join(datos.get("_meses_disponibles", []))
            falt = ", ".join(datos.get("_meses_faltantes", []))
            st.markdown(f"""
            <div class='warn-box'>
            ⚠️ <b>Datos parciales:</b> El archivo solo contiene datos para <b>{disp}</b>.<br>
            Meses faltantes: {falt}. 
            Para el período completo asegúrate de que el archivo DAILY contenga todos los meses requeridos.
            </div>""", unsafe_allow_html=True)
        # ── KPI helper ──────────────────────────────────────────────────────
        def kpi(label, val, sub="", delta=None):
            delta_html = ""
            if delta is not None:
                cls = "kpi-up" if delta >= 0 else "kpi-down"
                arrow = "▲" if delta >= 0 else "▼"
                delta_html = f'<div class="{cls}">{arrow} {abs(delta):.1f}% vs histórico</div>'
            return f"""
            <div class='kpi'>
                <div class='kpi-label'>{label}</div>
                <div class='kpi-val'>{val}</div>
                <div class='kpi-sub'>{sub}</div>
                {delta_html}
            </div>"""

        hist_chcp_hm3 = (hist.get("hist_chcp_m3s", 0) * n_dias * HM3_PER_M3S_DAY) \
                        if hist and n_dias > 0 else None
        pct_vs_hist   = ((D("aporte_total_hm3") - hist_chcp_hm3) / hist_chcp_hm3 * 100) \
                        if hist_chcp_hm3 and hist_chcp_hm3 > 0 else None
        _aporte_src_label = "📂 Fuentes Agua" if datos.get("_aporte_source") == "fuentes_agua" else "⚠️ Sin Fuentes_Agua"

        # ── SEC 1 — Aportes a la CHCP ─────────────────────────────────────────
        st.markdown('<div class="sec-hdr">💧 Aportes a la Cuenca Hidrográfica del Canal (CHCP)</div>', unsafe_allow_html=True)
        a1, a2, a3, a4 = st.columns(4)
        a1.markdown(kpi("Aportes Totales CHCP",
            f"{D('aporte_total_hm3'):.1f} hm³",
            f"{D('aporte_total_m3s'):.1f} m³/s · {m3s_to_cfs(D('aporte_total_m3s')):.0f} pie³/s · {_aporte_src_label}",
            pct_vs_hist), unsafe_allow_html=True)
        _hist_neto_hm3  = (hist.get("hist_neto_m3s", 0) * n_dias * HM3_PER_M3S_DAY)\
                           if hist and n_dias > 0 else None
        _neto_rank_pos  = str(datos.get("hist_neto_scale_pos", "") or "")
        _neto_rank_pct  = ((D("aporte_neto_hm3") - _hist_neto_hm3) / _hist_neto_hm3 * 100)\
                          if _hist_neto_hm3 and _hist_neto_hm3 > 0 else None
        _neto_sub       = (f"{D('aporte_neto_m3s'):.1f} m³/s · {m3s_to_cfs(D('aporte_neto_m3s')):.0f} pie³/s"
                           + (f" · pos. {_neto_rank_pos}" if _neto_rank_pos else ""))
        a2.markdown(kpi("Aportes Netos CHCP",
            f"{D('aporte_neto_hm3'):.1f} hm³",
            _neto_sub,
            _neto_rank_pct), unsafe_allow_html=True)
        a3.markdown(kpi("Distribución — Cuenca Gatún",
            f"{D('gatun_hm3'):.1f} hm³" if D('gatun_hm3') else "⚠️ Sin Fuentes_Agua",
            f"{D('gatun_pct'):.1f}% del total · Subcuenca Gatún"), unsafe_allow_html=True)
        a4.markdown(kpi("Distribución — Cuenca Madden",
            f"{D('alhajuela_hm3'):.1f} hm³" if D('alhajuela_hm3') else "⚠️ Sin Fuentes_Agua",
            f"{D('alhajuela_pct'):.1f}% del total · Subcuenca Alhajuela"), unsafe_allow_html=True)

        # ── SEC 2 — Niveles de Embalses ───────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="sec-hdr">🏞️ Niveles de los Embalses Gatún y Alhajuela</div>', unsafe_allow_html=True)
        n1, n2, n3, n4 = st.columns(4)
        n1.markdown(kpi("Nivel Gatún",
            f"{D('nivel_gatun_pies'):.2f} pies",
            f"{D('nivel_gatun_m'):.2f} m PLD"), unsafe_allow_html=True)
        n2.markdown(kpi("Nivel Alhajuela",
            f"{D('nivel_alh_pies'):.2f} pies",
            f"{D('nivel_alh_m'):.2f} m PLD"), unsafe_allow_html=True)
        n3.markdown(kpi("Calado Panamax",
            f"{float(st.session_state.get('x_cpan', 39.50)):.2f} pies",
            "Máx. permisible"), unsafe_allow_html=True)
        n4.markdown(kpi("Calado Neopanamax",
            f"{float(st.session_state.get('x_cneo', 50.0)):.1f} pies",
            "Máx. permisible"), unsafe_allow_html=True)

        # ── SEC 3 — Uso de Agua en Esclusas ──────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="sec-hdr">🔁 Uso de Agua en las Esclusas del Canal de Panamá</div>', unsafe_allow_html=True)
        e1, e2, e3, e4, e5, e6, e7 = st.columns(7)
        e1.markdown(kpi("Uso Esclusas Total",
            f"{D('esclusaje_hm3'):.1f} hm³",
            f"{D('esclusaje_m3s'):.1f} m³/s"), unsafe_allow_html=True)
        e2.markdown(kpi("Uso Panamax",
            f"{D('panamax_hm3'):.1f} hm³",
            f"{D('panamax_pct'):.1f}% del esclusaje"), unsafe_allow_html=True)
        e3.markdown(kpi("Uso Neopanamax",
            f"{D('neopanamax_hm3'):.1f} hm³",
            f"{D('neopanamax_pct'):.1f}% del esclusaje"), unsafe_allow_html=True)
        e4.markdown(kpi("Uso Prom. Diario Esclusas",
            f"{D('uso_prom_diario_hm3'):.2f} hm³/día",
            f"{D('uso_prom_diario_m3s'):.1f} m³/s"), unsafe_allow_html=True)
        e5.markdown(kpi("Ahorros Implementados",
            f"{D('ahorro_total_hm3'):.2f} hm³",
            f"{D('ahorro_lamina_pies'):.3f} pies Gatún"), unsafe_allow_html=True)
        e6.markdown(kpi("Agua/Tránsito PNX",
            f"{D('agua_panamax_trans'):.4f} hm³",
            f"{D('panamax_pct'):.1f}% del esclusaje"), unsafe_allow_html=True)
        e7.markdown(kpi("Agua/Tránsito NPX",
            f"{D('agua_neo_trans'):.4f} hm³",
            f"{D('neopanamax_pct'):.1f}% del esclusaje"), unsafe_allow_html=True)

        # ── SEC 4 — Salinidad & CCA ───────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        _cca_d = float(st.session_state.get("rep_cca_hm3", 0.0) or 0.0)
        _spc_d = D("salinidad_spc", 0.0) or 0.0

        def _dash_alert(spc):
            if   spc < 0.44: return "🟢 Normal",                "#16a34a"
            elif spc < 0.52: return "🟡 Vigilancia",            "#d97706"
            elif spc < 0.54: return "🟠 Pre-alerta",            "#ea580c"
            elif spc < 0.56: return "🔴 Preparar lavado",       "#dc2626"
            elif spc < 0.58: return "🔴 Alerta lavado urgente", "#b91c1c"
            else:            return "🚨 Crítico",                "#7f1d1d"

        _alrt_lbl_d, _alrt_clr_d = _dash_alert(_spc_d)
        st.markdown('<div class="sec-hdr">🌊 Salinidad y Conservación de Calidad del Agua (CCA)</div>', unsafe_allow_html=True)
        s1, s2, s3, s4 = st.columns(4)
        s1.markdown(kpi("Salinidad SPC",
            f"{_spc_d:.4f} ups" if _spc_d else "Sin datos",
            datos.get("salinidad_source","—").upper()), unsafe_allow_html=True)
        s2.markdown(f"""
        <div class='kpi' style='border-left-color:{_alrt_clr_d}'>
            <div class='kpi-label'>ALERTA SALINIDAD</div>
            <div class='kpi-val' style='font-size:1rem;color:{_alrt_clr_d}'>{_alrt_lbl_d}</div>
            <div class='kpi-sub'>Umbral alerta: 0.52 ups</div>
        </div>""", unsafe_allow_html=True)
        s3.markdown(kpi("CCA (hm³)",
            f"{_cca_d:.2f}" if _cca_d else "⚠️ Sin dato",
            "Conservación Calidad Agua"), unsafe_allow_html=True)
        s4.markdown(kpi("ZZFlush (hm³)",
            f"{D('zzflush_auto_hm3'):.2f}",
            "Auto desde GATSPILL"), unsafe_allow_html=True)

        # ── SEC 5 — Energía Hidroeléctrica ────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="sec-hdr">⚡ Producción de Energía Hidroeléctrica — Madden y Gatún</div>', unsafe_allow_html=True)
        eg1, eg2, eg3, eg4 = st.columns(4)
        eg1.markdown(kpi("Madden — Promedio",
            f"{D('energia_madden_mw'):.2f} MW",
            f"{D('energia_madden_mwh_dia'):.0f} MWh/día"), unsafe_allow_html=True)
        eg2.markdown(kpi("Madden — Hidrogen. (hm³)",
            f"{D('hidro_madden_hm3'):.2f} hm³",
            "Columna G · MADMCF"), unsafe_allow_html=True)
        eg3.markdown(kpi("Gatún — Promedio",
            f"{D('energia_gatun_mw'):.4f} MW",
            f"{D('energia_gatun_mwh_dia'):.2f} MWh/día"), unsafe_allow_html=True)
        eg4.markdown(kpi("Gatún — Hidrogen. (hm³)",
            f"{D('hidro_gatun_hm3'):.2f} hm³",
            "Columna AS · GATMCF"), unsafe_allow_html=True)

        # ── SEC 6 — Trasvase Alhajuela → Gatún ───────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="sec-hdr">🔄 Trasvase Alhajuela → Gatún</div>', unsafe_allow_html=True)
        t1, t2, t3, t4 = st.columns(4)
        t1.markdown(kpi("Trasvase Total",
            f"{D('trasvase_hm3'):.2f} hm³",
            f"{D('trasvase_m3s'):.2f} m³/s"), unsafe_allow_html=True)
        t2.markdown(kpi("Por Hidrogeneración",
            f"{D('trasvase_hidro_hm3'):.2f} hm³",
            "Col G · MADMCF"), unsafe_allow_html=True)
        t3.markdown(kpi("Por Vertidos",
            f"{D('trasvase_vertidos_hm3'):.2f} hm³",
            "Col R · MADSPILL"), unsafe_allow_html=True)
        t4.markdown(kpi("Por Fugas",
            f"{D('trasvase_fugas_hm3'):.2f} hm³",
            "Col U · LEAK MAD"), unsafe_allow_html=True)

        # ── SEC 7 — Meteorología ──────────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        _prec_mm_d    = float(st.session_state.get("x_prec",  0.0) or 0.0)
        _prec_hist_d  = float(st.session_state.get("x_phist", 0.0) or 0.0)
        _prec_sup_d   = round(_prec_mm_d - _prec_hist_d, 1)
        _prec_pct_d   = float(st.session_state.get("x_met_pct_informe",
                            round(abs(_prec_sup_d / _prec_hist_d * 100), 1)
                            if _prec_hist_d > 0 else 0.0))
        _prec_pos_d   = int(st.session_state.get("x_ppos", 0) or 0)
        _prec_rank_d  = str(st.session_state.get("x_prank_label", "") or "")
        _prec_per_d   = str(st.session_state.get("x_pper", "") or "")
        _prec_anios_d = int(st.session_state.get("met_n_anios", 0) or 0)
        _prec_dir_d   = "por encima" if _prec_sup_d >= 0 else "por debajo"
        st.markdown('<div class="sec-hdr">🌧️ Meteorología — Precipitación CHCP</div>', unsafe_allow_html=True)
        pm1, pm2, pm3, pm4, pm5, pm6 = st.columns(6)
        pm1.markdown(kpi("Precipitación período",
            f"{_prec_mm_d:.0f} mm" if _prec_mm_d else "⚠️ Sin dato",
            "Total del período"), unsafe_allow_html=True)
        pm2.markdown(kpi("Prom. histórico",
            f"{_prec_hist_d:.0f} mm" if _prec_hist_d else "⚠️ Sin dato",
            "Referencia histórica"), unsafe_allow_html=True)
        _sup_lbl = "▲ Superávit" if _prec_sup_d >= 0 else "▼ Déficit"
        pm3.markdown(kpi("Superávit / Déficit",
            f"{_prec_sup_d:+.1f} mm", _sup_lbl), unsafe_allow_html=True)
        pm4.markdown(kpi("% vs hist.",
            f"{_prec_pct_d:.1f}%" if _prec_pct_d else "⚠️ Sin dato",
            _prec_dir_d.capitalize()), unsafe_allow_html=True)
        pm5.markdown(kpi("Posición ranking",
            f"# {_prec_pos_d}" if _prec_pos_d else "⚠️ Sin dato",
            _prec_rank_d or "—  (configura en Tab Meteo)"), unsafe_allow_html=True)
        pm6.markdown(kpi("Años de registro",
            f"{_prec_anios_d} años" if _prec_anios_d else "⚠️ Sin dato",
            _prec_per_d or "—  (configura en Tab Meteo)"), unsafe_allow_html=True)

        # ── SEC 8 — Otros Usos del Agua ───────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="sec-hdr">🚰 Otros Usos del Agua</div>', unsafe_allow_html=True)
        o1, o2, o3, o4, o5, o6 = st.columns(6)
        o1.markdown(kpi("Potabilización Total",
            f"{D('potabilizacion_hm3'):.2f} hm³",
            f"{D('potabilizacion_m3s'):.2f} m³/s"), unsafe_allow_html=True)
        o2.markdown(kpi("Potabiliz. Gatún",
            f"{D('potabilizacion_gat_hm3'):.2f} hm³",
            "Agua potable"), unsafe_allow_html=True)
        o3.markdown(kpi("Potabiliz. Alhajuela",
            f"{D('potabilizacion_alh_hm3'):.2f} hm³",
            "Agua potable"), unsafe_allow_html=True)
        o4.markdown(kpi("Evaporación Total",
            f"{D('evaporacion_hm3'):.2f} hm³",
            f"{D('evaporacion_m3s'):.2f} m³/s"), unsafe_allow_html=True)
        o5.markdown(kpi("Hidrogen. Gatún (hm³)",
            f"{D('hidro_gatun_hm3'):.2f} hm³",
            f"{D('hidro_gatun_m3s'):.2f} m³/s"), unsafe_allow_html=True)
        o6.markdown(kpi("Concesiones Gatún",
            f"{D('concesiones_gat_hm3'):.2f} hm³",
            "LEAK GAT · balance Gatún"), unsafe_allow_html=True)

        # ── Fila adicional: concesiones Alhajuela ─────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        ca1, ca2, ca3, ca4, ca5, ca6 = st.columns(6)
        ca1.markdown(kpi("Concesiones/Fugas Alh.",
            f"{D('concesiones_alh_hm3'):.2f} hm³",
            "LEAK MAD col U · solo display"), unsafe_allow_html=True)
        ca2.markdown(kpi("Concesiones Total",
            f"{D('concesiones_hm3'):.2f} hm³",
            "Gatún + Alhajuela"), unsafe_allow_html=True)

        # ── Gráficos ───────────────────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        g1, g2 = st.columns(2)

        with g1:
            st.markdown("#### Balance Hídrico")
            cats  = ["Aporte Total","Aporte Neto","Total Salidas","Trasvase"]
            vals  = [D("aporte_total_hm3"), D("aporte_neto_hm3"),
                     D("total_salidas_hm3"), D("trasvase_hm3")]
            clrs  = ["#0072b8","#00a0dc","#e55c00","#8b5cf6"]
            fig_b = go.Figure(go.Bar(x=cats, y=vals, marker_color=clrs,
                text=[f"{v:.1f}" for v in vals], textposition="auto",
                textfont=dict(size=11, color="white")))
            fig_b.update_layout(height=300, plot_bgcolor="rgba(0,0,0,0)",
                paper_bgcolor="rgba(0,0,0,0)", margin=dict(t=10,b=10,l=10,r=10),
                xaxis=dict(tickfont=dict(size=11)),
                yaxis=dict(title="hm³", gridcolor="rgba(148,163,184,.2)"))
            st.plotly_chart(fig_b, use_container_width=True)

        with g2:
            st.markdown("#### Usos del Agua")
            u_names = ["Esclusajes","Potabilización","Evaporación",
                       "Concesiones/Fugas","Hidrogen. Gatún","ZZFlush"]
            u_vals  = [D("esclusaje_hm3"), D("potabilizacion_hm3"),
                       D("evaporacion_hm3"), D("concesiones_hm3"),
                       D("hidro_gatun_hm3"), D("zzflush_auto_hm3")]
            fig_p = px.pie(names=u_names, values=u_vals,
                color_discrete_sequence=["#003a6e","#0072b8","#fbbf24","#38bdf8","#f97316","#0ea5e9"])
            fig_p.update_traces(textinfo="percent+label", textfont_size=10)
            fig_p.update_layout(plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                height=300, margin=dict(t=10,b=10,l=10,r=10), showlegend=False)
            st.plotly_chart(fig_p, use_container_width=True)

        g3, g4 = st.columns(2)
        with g3:
            st.markdown("#### Panamax vs Neopanamax")
            fig_esc = go.Figure()
            fig_esc.add_trace(go.Bar(
                x=["Volumen (hm³)","Tránsitos (÷10)"],
                y=[D("panamax_hm3"), D("transitos_panamax")/10],
                name="Panamax", marker_color="#003a6e",
                text=[f"{D('panamax_hm3'):.1f}",f"{D('transitos_panamax'):.0f}"],
                textposition="auto"))
            fig_esc.add_trace(go.Bar(
                x=["Volumen (hm³)","Tránsitos (÷10)"],
                y=[D("neopanamax_hm3"), D("transitos_neopanamax")/10],
                name="Neopanamax", marker_color="#0072b8",
                text=[f"{D('neopanamax_hm3'):.1f}",f"{D('transitos_neopanamax'):.0f}"],
                textposition="auto"))
            fig_esc.update_layout(barmode="group", height=290,
                plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                margin=dict(t=10,b=10,l=10,r=10), legend=dict(orientation="h",y=1.1))
            st.plotly_chart(fig_esc, use_container_width=True)

        with g4:
            st.markdown("#### Ahorros en Esclusas")
            fig_aho = go.Figure(go.Bar(
                x=["Panamax","Neopanamax","Total"],
                y=[D("ahorro_panamax_hm3"), D("ahorro_neopanamax_hm3"), D("ahorro_total_hm3")],
                marker_color=["#003a6e","#0072b8","#15803d"],
                text=[f"{D('ahorro_panamax_hm3'):.2f}",
                      f"{D('ahorro_neopanamax_hm3'):.2f}",
                      f"{D('ahorro_total_hm3'):.2f}"],
                textposition="auto", textfont=dict(color="white", size=11)))
            fig_aho.update_layout(height=290, plot_bgcolor="rgba(0,0,0,0)",
                paper_bgcolor="rgba(0,0,0,0)", margin=dict(t=10,b=10,l=10,r=10),
                yaxis=dict(title="hm³", gridcolor="rgba(148,163,184,.2)"))
            st.plotly_chart(fig_aho, use_container_width=True)

        # ── Tabla resumen por mes (para períodos multi-mes) ─────────────────────
        if tipo != "Mensual" and datos.get("_meses") and len(datos.get("_meses",[])) > 1:
            st.markdown("---")
            st.markdown(f"#### 📋 Detalle por mes — {label_per}")
            _m_data_det = datos["_meses"]
            _m_disp_det = datos.get("_meses_disponibles", [])
            if _m_disp_det and _m_data_det:
                det_rows = []
                for mi, mes_tag in enumerate(_m_disp_det):
                    r = _m_data_det[mi] if mi < len(_m_data_det) else {}
                    if r.get("_missing_period_data"):
                        continue
                    det_rows.append({
                        "Mes": mes_tag,
                        "Aporte (hm³)": round(float(r.get("aporte_total_hm3",0) or 0), 2),
                        "Evap. (hm³)":  round(float(r.get("evaporacion_hm3",0) or 0), 2),
                        "Esclusajes (hm³)": round(float(r.get("esclusaje_hm3",0) or 0), 2),
                        "Panamax (hm³)":    round(float(r.get("panamax_hm3",0) or 0), 2),
                        "NeoPanamax (hm³)": round(float(r.get("neopanamax_hm3",0) or 0), 2),
                        "Potabiliz. (hm³)": round(float(r.get("potabilizacion_hm3",0) or 0), 2),
                        "Tránsitos PM":     int(r.get("transitos_panamax",0) or 0),
                        "Tránsitos NEO":    int(r.get("transitos_neopanamax",0) or 0),
                        "Nivel Gatún (pies)": round(float(r.get("nivel_gatun_pies",0) or 0), 2),
                        "Total Salidas (hm³)": round(float(r.get("total_salidas_hm3",0) or 0), 2),
                    })
                if det_rows:
                    df_det = pd.DataFrame(det_rows)
                    st.dataframe(df_det, use_container_width=True, hide_index=True)
                    # Mini spark charts
                    spark1, spark2 = st.columns(2)
                    with spark1:
                        fig_spark = go.Figure()
                        fig_spark.add_trace(go.Scatter(
                            x=[r["Mes"] for r in det_rows],
                            y=[r["Aporte (hm³)"] for r in det_rows],
                            mode="lines+markers+text",
                            name="Aporte Total",
                            line=dict(color="#0072b8", width=2.5),
                            text=[f"{r['Aporte (hm³)']:.1f}" for r in det_rows],
                            textposition="top center"))
                        fig_spark.add_trace(go.Scatter(
                            x=[r["Mes"] for r in det_rows],
                            y=[r["Total Salidas (hm³)"] for r in det_rows],
                            mode="lines+markers+text",
                            name="Total Salidas",
                            line=dict(color="#e55c00", width=2, dash="dot"),
                            text=[f"{r['Total Salidas (hm³)']:.1f}" for r in det_rows],
                            textposition="bottom center"))
                        fig_spark.update_layout(
                            height=260, title="Aportes vs Salidas (hm³)",
                            plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                            margin=dict(t=35,b=10,l=10,r=10),
                            yaxis=dict(gridcolor="rgba(148,163,184,.2)"),
                            legend=dict(orientation="h", y=1.15))
                        st.plotly_chart(fig_spark, use_container_width=True)
                    with spark2:
                        fig_esc2 = go.Figure()
                        fig_esc2.add_trace(go.Bar(
                            x=[r["Mes"] for r in det_rows],
                            y=[r["Esclusajes (hm³)"] for r in det_rows],
                            name="Esclusajes", marker_color="#003a6e",
                            text=[f"{r['Esclusajes (hm³)']:.1f}" for r in det_rows],
                            textposition="auto"))
                        fig_esc2.update_layout(
                            height=260, title="Esclusajes mensuales (hm³)",
                            plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                            margin=dict(t=35,b=10,l=10,r=10),
                            yaxis=dict(title="hm³", gridcolor="rgba(148,163,184,.2)"))
                        st.plotly_chart(fig_esc2, use_container_width=True)

        if hidro_rows1:
            st.markdown("---")
            st.markdown("#### Resumen auditable desde Hidroestadística")
            p1, p2, p3, p4 = st.columns(4)
            p1.markdown(kpi("Usos Totales (Hidro)",
                            f"{fmt_num(H('Usos Totales','value1','hm3'),1)} hm³",
                            f"{fmt_num(H('Usos Totales','value2','m3/s'),1)} m³/s · {fmt_num(H('Usos Totales','value3','EED'),1)} EED"), unsafe_allow_html=True)
            p2.markdown(kpi("Trasvase Madden",
                            f"{fmt_num(H('Trasvase Madden','value1','hm3'),1)} hm³",
                            f"{fmt_num(H('Trasvase Madden','value2','m3/s'),1)} m³/s · {fmt_num(H('Trasvase Madden','value3','EED'),1)} EED"), unsafe_allow_html=True)
            p3.markdown(kpi("Déficit",
                            f"{fmt_num(H('Deficit','value1','hm3'),1)} hm³",
                            f"{fmt_num(H('Deficit','value2','m3/s'),1)} m³/s · {fmt_num(H('Deficit','value3','EED'),1)} EED"), unsafe_allow_html=True)
            p4.markdown(kpi("Ahorro Total",
                            f"{fmt_num(H('Panamax + NeoPanamax','value1','hm3'),2)} hm³",
                            f"{fmt_num(H('Panamax + NeoPanamax','value2','m3/s'),2)} m³/s · {fmt_num(H('Panamax + NeoPanamax','value3','EED'),2)} EED"), unsafe_allow_html=True)
            df_audit = audit_compare_df()
            if not df_audit.empty:
                st.dataframe(df_audit.style.format({
                    "App hm³":"{:.2f}","Hidro hm³":"{:.2f}","Δ hm³":"{:+.2f}",
                    "App m³/s":"{:.2f}","Hidro m³/s":"{:.2f}",
                    "App EED":"{:.2f}","Hidro EED":"{:.2f}"
                }), use_container_width=True, hide_index=True)
                st.caption(f"Hidro cargada: {hidro_audit.get('periodo_grupo','—')} · {hidro_audit.get('periodo_detalle','—')}")
        # ── Semestral: desglose por mes ──────────────────────────────────────
        if tipo != "Mensual" and datos.get("_meses") and len(datos["_meses"]) > 1:
            st.markdown("---")
            st.markdown(f"#### Evolución mensual del período — {label_per}")
            meses_disp = datos.get("_meses_disponibles", [])
            _m_data = datos["_meses"]
            col_sm1, col_sm2 = st.columns(2)
            with col_sm1:
                fig_sem_esc = go.Figure()
                fig_sem_esc.add_trace(go.Bar(
                    name="Esclusajes", x=meses_disp,
                    y=[r.get("esclusaje_hm3",0) for r in _m_data],
                    marker_color="#003a6e", text=[f"{r.get('esclusaje_hm3',0):.1f}" for r in _m_data],
                    textposition="auto"))
                fig_sem_esc.add_trace(go.Bar(
                    name="Potabilización", x=meses_disp,
                    y=[r.get("potabilizacion_hm3",0) for r in _m_data],
                    marker_color="#0072b8", text=[f"{r.get('potabilizacion_hm3',0):.1f}" for r in _m_data],
                    textposition="auto"))
                fig_sem_esc.update_layout(barmode="group", height=280, title="Usos por mes",
                    plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                    margin=dict(t=40,b=10,l=10,r=10), legend=dict(orientation="h",y=1.15),
                    yaxis=dict(title="hm³", gridcolor="rgba(148,163,184,.2)"))
                st.plotly_chart(fig_sem_esc, use_container_width=True)
            with col_sm2:
                fig_sem_ap = go.Figure()
                fig_sem_ap.add_trace(go.Bar(
                    name="Aporte Total", x=meses_disp,
                    y=[r.get("aporte_total_hm3",0) for r in _m_data],
                    marker_color="#38bdf8", text=[f"{r.get('aporte_total_hm3',0):.1f}" for r in _m_data],
                    textposition="auto"))
                fig_sem_ap.add_trace(go.Bar(
                    name="Aporte Neto", x=meses_disp,
                    y=[r.get("aporte_neto_hm3",0) for r in _m_data],
                    marker_color="#0072b8", text=[f"{r.get('aporte_neto_hm3',0):.1f}" for r in _m_data],
                    textposition="auto"))
                fig_sem_ap.update_layout(barmode="group", height=280, title="Aportes por mes",
                    plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                    margin=dict(t=40,b=10,l=10,r=10), legend=dict(orientation="h",y=1.15),
                    yaxis=dict(title="hm³", gridcolor="rgba(148,163,184,.2)"))
                st.plotly_chart(fig_sem_ap, use_container_width=True)


# ═══════════════════════════════════════════════════════════════
#  TAB 1 — BALANCE HÍDRICO
# ═══════════════════════════════════════════════════════════════
with tabs[5]:
    st.markdown('<div class="sec-hdr">♻️ Balance Hídrico</div>', unsafe_allow_html=True)
    if not data_ready:
        st.info("Carga el archivo DAILY_AND_ACUMMEN para ver el balance hídrico.")
    else:
        # Sankey / Waterfall del balance
        col_bh1, col_bh2 = st.columns([2,1])
        with col_bh1:
            st.markdown("#### Diagrama de Flujo Hídrico (Sankey)")
            # Sankey con todos los usos del agua (incluyendo ZZFlush auto)
            _zzf = D("zzflush_auto_hm3", 0.0)
            label_s = ["Aportes<br>CHCP", "Aporte<br>Neto", "Evaporación",
                       "Esclusajes", "Potabilización", "Concesiones",
                       "Hidrogeneración", "ZZFlush"]
            source_s = [0, 0, 1, 1, 1, 1, 1]
            target_s = [1, 2, 3, 4, 5, 6, 7]
            value_s  = [max(D("aporte_neto_hm3"), 0.001),
                        max(D("evaporacion_hm3"), 0.001),
                        max(D("esclusaje_hm3"), 0.001),
                        max(D("potabilizacion_hm3"), 0.001),
                        max(D("concesiones_hm3"), 0.001),
                        max(D("hidro_gatun_hm3"), 0.001),
                        max(_zzf, 0.001)]
            fig_sk = go.Figure(go.Sankey(
                node=dict(label=label_s,
                    color=["#003a6e","#0072b8","#fbbf24","#e55c00",
                           "#059669","#8b5cf6","#f97316","#0ea5e9"],
                    pad=18, thickness=18),
                link=dict(source=source_s, target=target_s, value=value_s,
                    color=["rgba(0,114,184,.35)","rgba(251,191,36,.35)",
                           "rgba(229,92,0,.35)","rgba(5,150,105,.35)",
                           "rgba(139,92,246,.35)","rgba(249,115,22,.35)",
                           "rgba(14,165,233,.35)"])))
            fig_sk.update_layout(plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                height=400, margin=dict(t=10,b=10,l=10,r=10))
            st.plotly_chart(fig_sk, use_container_width=True)

        with col_bh2:
            st.markdown("#### Resumen de Entradas y Salidas")
            # Tabla resumen
            deficit = D("total_salidas_hm3") - D("aporte_neto_hm3")
            tabla = [
                ("📥 ENTRADAS","",""),
                ("Aporte Total CHCP",    f"{D('aporte_total_hm3'):.2f} hm³",  f"{D('aporte_total_m3s'):.2f} m³/s"),
                ("  Subcuenca GL",       f"{D('gatun_hm3'):.2f} hm³",         ""),
                ("  Subcuenca MD/Alh.",  f"{D('alhajuela_hm3'):.2f} hm³",     ""),
                ("(–) Evaporación",      f"{D('evaporacion_hm3'):.2f} hm³",   f"{D('evaporacion_m3s'):.2f} m³/s"),
                ("= Aporte Neto",        f"{D('aporte_neto_hm3'):.2f} hm³",   f"{D('aporte_neto_m3s'):.2f} m³/s"),
                ("","",""),
                ("📤 SALIDAS (Usos)","",""),
                ("  Esclusajes PM",      f"{D('panamax_hm3'):.2f} hm³",       ""),
                ("  Esclusajes NEO",     f"{D('neopanamax_hm3'):.2f} hm³",    ""),
                ("  Potabilización Gat.",f"{D('potabilizacion_gat_hm3'):.2f} hm³",""),
                ("  Potabilización Alh.",f"{D('potabilizacion_alh_hm3'):.2f} hm³",""),
                ("  Evaporación",        f"{D('evaporacion_hm3'):.2f} hm³",   f"{D('evaporacion_m3s'):.2f} m³/s"),
                ("  Concesiones/Fugas",  f"{D('concesiones_hm3'):.2f} hm³",   f"{D('concesiones_m3s'):.2f} m³/s"),
                ("  Hidrogen. Gatún",    f"{D('hidro_gatun_hm3'):.4f} hm³",   ""),
                ("  ZZFlush",            f"{D('zzflush_auto_hm3'):.4f} hm³",  ""),
                ("TOTAL USOS CHCP",      f"{D('total_salidas_hm3'):.2f} hm³", ""),
                ("  Mov. internos",      f"{D('movimientos_operativos_hm3'):.2f} hm³", ""),
                ("","",""),
                ("⚖️ Balance", f"{'−' if (D('total_salidas_hm3')-D('aporte_neto_hm3'))>0 else '+'}{abs(D('total_salidas_hm3')-D('aporte_neto_hm3')):.2f} hm³",""),
            ]
            rows_html = ""
            for r in tabla:
                if r[0] in ("📥 ENTRADAS","📤 SALIDAS","⚖️ Balance") or r[0] in ("= Aporte Neto","Total Salidas"):
                    cls = "subtotal"
                elif r[0] == "":
                    cls = ""; 
                else:
                    cls = ""
                if r[0].startswith("📥") or r[0].startswith("📤"):
                    rows_html += f"<tr class='subtotal'><td colspan='3'><b>{r[0]}</b></td></tr>"
                elif r[0] == "":
                    rows_html += "<tr><td colspan='3' style='padding:2px'></td></tr>"
                else:
                    b1 = "<b>" if r[0] in ("= Aporte Neto","TOTAL SALIDAS","⚖️ Balance") else ""
                    b2 = "</b>" if b1 else ""
                    rows_html += f"<tr><td>{b1}{r[0]}{b2}</td><td>{b1}{r[1]}{b2}</td><td>{r[2]}</td></tr>"
            st.markdown(f"""
            <table class='styled-table'>
            <thead><tr><th>Concepto</th><th>hm³</th><th>m³/s</th></tr></thead>
            <tbody>{rows_html}</tbody>
            </table>""", unsafe_allow_html=True)


        # ── Balance final del período (formato del informe) ──────────────────
        st.markdown("---")
        st.markdown('<div class="sec-hdr">⚖️ Balance Final del Período — Resumen del Informe</div>', unsafe_allow_html=True)
        if data_ready:
            _dias_b = n_dias
            def _eed_b(h): return round(h / (_dias_b * 0.0864) / 28.317, 3) if _dias_b > 0 else 0.0
            _at_b   = D("aporte_total_hm3");  _atm_b = D("aporte_total_m3s")
            _an_b   = D("aporte_neto_hm3");   _anm_b = D("aporte_neto_m3s")
            _ev_b   = D("evaporacion_hm3");   _evm_b = D("evaporacion_m3s")
            _esc_b  = D("esclusaje_hm3");     _escm_b= D("esclusaje_m3s")
            _pot_b  = D("potabilizacion_hm3");_potm_b= D("potabilizacion_m3s")
            _zz_b   = D("zzflush_auto_hm3")
            _conc_b = D("concesiones_gat_hm3")
            _hidg_b = D("hidro_gatun_hm3")
            _vgat_b = D("vertidos_gat_hm3")
            _tot_b  = D("total_salidas_hm3")
            _exc_b  = D("excedente_hm3")
            _trav_b = D("trasvase_hm3");      _travm_b= D("trasvase_m3s")
            _pnx_b  = D("panamax_hm3");       _npx_b  = D("neopanamax_hm3")

            _tot_sal_pct = lambda h: round(h/_tot_b*100,1) if _tot_b>0 else 0.0

            brows = [
                ("ENTRADAS", None, None, None, None, True),
                ("Aporte Total CHCP", _at_b, _atm_b, _eed_b(_at_b), None, False),
                ("  Subcuenca Alhajuela", D("alhajuela_hm3"), None, None, f"{D('alhajuela_pct'):.1f}%", False),
                ("  Subcuenca Gatún",     D("gatun_hm3"),     None, None, f"{D('gatun_pct'):.1f}%",     False),
                ("(−) Evaporación", _ev_b, _evm_b, _eed_b(_ev_b), None, False),
                ("= Aporte Neto CHCP", _an_b, _anm_b, _eed_b(_an_b), None, True),
                ("", None, None, None, None, False),
                ("SALIDAS / USOS", None, None, None, None, True),
                ("Esclusajes total", _esc_b, _escm_b, _eed_b(_esc_b), f"{_tot_sal_pct(_esc_b):.1f}%", False),
                ("  Panamax (GAT+PM)", _pnx_b, None, None, f"{D('panamax_pct'):.1f}%", False),
                ("  Neopanamax (ACL+CCL)", _npx_b, None, None, f"{D('neopanamax_pct'):.1f}%", False),
                ("Potabilización", _pot_b, _potm_b, _eed_b(_pot_b), f"{_tot_sal_pct(_pot_b):.1f}%", False),
                ("ZZFlush / CCA", _zz_b, None, _eed_b(_zz_b), f"{_tot_sal_pct(_zz_b):.1f}%", False),
                ("Concesiones / Fugas Gatún", _conc_b, None, None, f"{_tot_sal_pct(_conc_b):.1f}%", False),
                ("Evaporación", _ev_b, _evm_b, _eed_b(_ev_b), f"{_tot_sal_pct(_ev_b):.1f}%", False),
                ("Hidrogeneración Gatún", _hidg_b, None, _eed_b(_hidg_b), f"{_tot_sal_pct(_hidg_b):.1f}%", False),
                ("Vertidos Gatún (GATSPILL)", _vgat_b, None, None, f"{_tot_sal_pct(_vgat_b):.1f}%", False),
                ("TOTAL SALIDAS", _tot_b, D("total_salidas_m3s"), _eed_b(_tot_b), "100.0%", True),
                ("", None, None, None, None, False),
                ("Trasvase Alh.→Gatún", _trav_b, _travm_b, None, None, False),
                ("  Por hidrogeneración", D("trasvase_hidro_hm3"), None, None, None, False),
                ("  Por vertidos", D("trasvase_vertidos_hm3"), None, None, None, False),
                ("  Por fugas", D("trasvase_fugas_hm3"), None, None, None, False),
                ("", None, None, None, None, False),
                ("⚖️ BALANCE (Neto − Salidas)", _exc_b, None, None,
                 "✅ Superávit" if _exc_b >= 0 else "⚠️ Déficit", True),
            ]
            rows_html_b = ""
            for br in brows:
                label, hm3, m3s, eed_v, pct, bold = br
                if label == "":
                    rows_html_b += "<tr><td colspan='5' style='padding:2px'></td></tr>"
                    continue
                if bold and hm3 is None:
                    rows_html_b += f"<tr class='subtotal'><td colspan='5'><b>{label}</b></td></tr>"
                    continue
                hm3_s  = f"<b>{hm3:.2f}</b>" if (hm3 is not None and bold) else (f"{hm3:.2f}" if hm3 is not None else "—")
                m3s_s  = f"{m3s:.2f}" if m3s is not None else "—"
                eed_s  = f"{eed_v:.3f}" if eed_v is not None else "—"
                pct_s  = pct if pct else "—"
                lbl_s  = f"<b>{label}</b>" if bold else label
                # Color balance row
                if label.startswith("⚖️"):
                    c = "#16a34a" if (_exc_b >= 0) else "#dc2626"
                    rows_html_b += f"<tr style='background:{c}22'><td><b style='color:{c}'>{label}</b></td><td><b style='color:{c}'>{hm3:.2f} hm³</b></td><td>—</td><td>—</td><td><b style='color:{c}'>{pct_s}</b></td></tr>"
                else:
                    rows_html_b += f"<tr><td>{lbl_s}</td><td>{hm3_s} hm³</td><td>{m3s_s} m³/s</td><td>{eed_s} EED</td><td>{pct_s}</td></tr>"
            st.markdown(f"""
            <table class='styled-table' style='width:100%'>
            <thead><tr><th>Concepto</th><th>hm³</th><th>m³/s</th><th>EED</th><th>%</th></tr></thead>
            <tbody>{rows_html_b}</tbody>
            </table>""", unsafe_allow_html=True)
        else:
            st.info("Carga el DAILY_AND_ACUMMEN para ver el balance final.")

        # Nivel embalse timeline si hay datos
        if data_ready and f_daily:
            st.markdown(f"<br>#### Nivel Embalse Gatún — Evolución del período ({label_per})")
            try:
                if tipo == "Mensual":
                    mes_n = MESES_NUM.get(mes_sel, 3)
                    yr    = anio_sel if mes_n <= 9 else anio_sel - 1
                    df_daily2 = load_daily(_source_bytes(f_daily))
                    periodo_data = df_daily2[(df_daily2["ACTDATE"].dt.year==yr) &
                                             (df_daily2["ACTDATE"].dt.month==mes_n)]
                else:
                    # Período multi-mes: filtrar todos los meses del período
                    df_daily2 = load_daily(_source_bytes(f_daily))
                    meses_per = _get_meses_periodo(tipo, int(anio_sel), trimestre, semestre)
                    mask = pd.Series(False, index=df_daily2.index)
                    for mn, yr in meses_per:
                        mask |= ((df_daily2["ACTDATE"].dt.year==yr) &
                                 (df_daily2["ACTDATE"].dt.month==mn))
                    periodo_data = df_daily2[mask]

                fig_niv = go.Figure()
                fig_niv.add_trace(go.Scatter(
                    x=periodo_data["ACTDATE"], y=periodo_data["ACTGATEL"],
                    mode="lines", name="Nivel Gatún (pies)",
                    line=dict(color="#0072b8", width=2.5)))
                fig_niv.add_trace(go.Scatter(
                    x=periodo_data["ACTDATE"],
                    y=periodo_data["ACTMADEL"].apply(lambda x: x/3.0 if x else 0),
                    mode="lines", name="Nivel Alhajuela ÷3 (pies)",
                    line=dict(color="#7dd3fc", width=1.5, dash="dot"),
                    yaxis="y"))
                # Referencia nivel óptimo
                fig_niv.add_hline(y=87.0, line_dash="dash", line_color="#16a34a",
                    annotation_text="Nivel óptimo Gatún (87 pies)")
                fig_niv.update_layout(height=300, plot_bgcolor="rgba(0,0,0,0)",
                    paper_bgcolor="rgba(0,0,0,0)", margin=dict(t=10,b=10,l=10,r=10),
                    yaxis=dict(title="Pies PLD", gridcolor="rgba(148,163,184,.2)"),
                    xaxis=dict(title="Fecha"),
                    legend=dict(orientation="h", y=1.12))
                st.plotly_chart(fig_niv, use_container_width=True)
            except Exception:
                pass


# ═══════════════════════════════════════════════════════════════
#  TAB 2 — FUENTES DE AGUA
# ═══════════════════════════════════════════════════════════════
with tabs[1]:
    st.markdown('<div class="sec-hdr">🏞️ Fuentes de Agua por Tributario</div>', unsafe_allow_html=True)

    if not data_ready:
        st.info("Carga el archivo DAILY_AND_ACUMMEN para continuar.")
    else:
        # Usa los datos cargados del archivo Fuentes_Agua o los defaults
        rios_alh = fuentes_rios_alh
        rios_gat = fuentes_rios_gat
        _aporte_src = datos.get("_aporte_source", "missing") if data_ready else "—"
        if _aporte_src == "fuentes_agua":
            fuente_origen = "📂 Fuentes_Agua.xlsx — ✅ Aporte Total CHCP calculado desde tributarios medidos"
            _src_cls = "ok-box"
        else:
            fuente_origen = "⚠️ Fuentes_Agua no cargada — Aportes = 0. Carga el archivo para obtener los aportes hidrológicos reales."
            _src_cls = "warn-box"
        st.markdown(f"<div class='{_src_cls}'>🏞️ <b>{fuente_origen}</b></div>",
                    unsafe_allow_html=True)

        aporte_total_fuentes = D('aporte_total_hm3', 0)
        df_fuentes_excel, df_alh_disp, df_gat_disp = _build_fuentes_excel_section(
            rios_alh, rios_gat, fuentes_sub, aporte_total_fuentes, n_dias
        )

        col_fa1, col_fa2 = st.columns(2)
        with col_fa1:
            st.markdown("##### 🔵 Subcuenca Embalse Alhajuela")
            st.dataframe(df_alh_disp.style.format({"m³/s":"{:.3f}","hm³":"{:.2f}","%":"{:.1f}"}),
                         use_container_width=True, hide_index=True)
            alh_total = float(df_fuentes_excel.loc[df_fuentes_excel['Tipo']=='subtotal'].query("Grupo == 'Alhajuela'")['hm³'].sum()) if not df_fuentes_excel.empty else 0.0
            st.metric("Subtotal Alhajuela", f"{D('alhajuela_hm3', alh_total):.2f} hm³",
                      f"{D('alhajuela_pct', alh_total/D('aporte_total_hm3',1)*100 if D('aporte_total_hm3',0)>0 else 0):.1f}% del total CHCP")

        with col_fa2:
            st.markdown("##### 🟦 Subcuenca Embalse Gatún")
            st.dataframe(df_gat_disp.style.format({"m³/s":"{:.3f}","hm³":"{:.2f}","%":"{:.1f}"}),
                         use_container_width=True, hide_index=True)
            gat_total = float(df_fuentes_excel.loc[df_fuentes_excel['Tipo']=='subtotal'].query("Grupo == 'Gatún'")['hm³'].sum()) if not df_fuentes_excel.empty else 0.0
            st.metric("Subtotal Gatún", f"{D('gatun_hm3', gat_total):.2f} hm³",
                      f"{D('gatun_pct', gat_total/D('aporte_total_hm3',1)*100 if D('aporte_total_hm3',0)>0 else 0):.1f}% del total CHCP")

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("##### 📋 Fuentes_Agua — formato tipo Excel")
        st.markdown("<div class='info-box'>Esta sección replica la lógica del Excel de <b>Fuentes_Agua</b> y completa automáticamente <b>m³/s</b>, <b>hm³</b> y <b>%</b> del total cuando alguno de esos valores no venga explícito en la hoja.</div>", unsafe_allow_html=True)
        df_fuentes_show = df_fuentes_excel[["Fuente","m³/s","hm³","%","Tipo"]].copy() if not df_fuentes_excel.empty else pd.DataFrame(columns=["Fuente","m³/s","hm³","%","Tipo"])
        df_fuentes_show = df_fuentes_show.rename(columns={"Fuente": "Fuente (tributario)"})
        st.dataframe(
            df_fuentes_show.style.format({"m³/s":"{:.3f}","hm³":"{:.2f}","%":"{:.1f}"}),
            use_container_width=True,
            hide_index=True
        )

        # Gráfico unificado
        df_all = df_fuentes_excel[df_fuentes_excel['Tipo']=='tributario'][['Fuente','hm³','Grupo']].rename(columns={'Grupo':'Subcuenca'}) if not df_fuentes_excel.empty else pd.DataFrame(columns=['Fuente','hm³','Subcuenca'])
        if df_all.empty:
            st.markdown("<div class='warn-box'>No hay tributarios disponibles para el período seleccionado. La hoja exacta en Fuentes_Agua no se encontró o no contiene filas válidas; por eso los subtotales se muestran en 0.</div>", unsafe_allow_html=True)
        else:
            fig_f = px.bar(df_all, x="Fuente", y="hm³", color="Subcuenca",
                color_discrete_map={"Alhajuela":"#003a6e","Gatún":"#0072b8"},
                text="hm³", labels={"hm³":"Volumen (hm³)"},
                title=f"Fuentes de Agua — {label_per}")
            fig_f.update_traces(texttemplate="%{text:.1f}", textposition="outside",
                textfont=dict(size=10))
            fig_f.update_layout(height=420, plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                margin=dict(t=40,b=90,l=10,r=10),
                xaxis_tickangle=-35, legend=dict(orientation="h", y=1.05),
                yaxis=dict(gridcolor="rgba(148,163,184,.2)"))
            st.plotly_chart(fig_f, use_container_width=True)

        # Donut subcuencas
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            if df_alh_disp.empty or float(df_alh_disp["hm³"].sum()) <= 0:
                st.markdown("<div class='info-box'>Sin tributarios válidos de Alhajuela para graficar en este período.</div>", unsafe_allow_html=True)
            else:
                fig_da = px.pie(df_alh_disp, names="Fuente", values="hm³", hole=0.45,
                    title="Distribución Subcuenca Alhajuela",
                    color_discrete_sequence=px.colors.sequential.Blues_r)
                fig_da.update_traces(textinfo="percent+label", textfont_size=10)
                fig_da.update_layout(plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)", height=320, margin=dict(t=40,b=10,l=10,r=10), showlegend=False)
                st.plotly_chart(fig_da, use_container_width=True)
        with col_d2:
            if df_gat_disp.empty or float(df_gat_disp["hm³"].sum()) <= 0:
                st.markdown("<div class='info-box'>Sin tributarios válidos de Gatún para graficar en este período.</div>", unsafe_allow_html=True)
            else:
                fig_dg = px.pie(df_gat_disp, names="Fuente", values="hm³", hole=0.45,
                    title="Distribución Subcuenca Gatún",
                    color_discrete_sequence=px.colors.sequential.Blues_r)
                fig_dg.update_traces(textinfo="percent+label", textfont_size=10)
                fig_dg.update_layout(plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)", height=320, margin=dict(t=40,b=10,l=10,r=10), showlegend=False)
                st.plotly_chart(fig_dg, use_container_width=True)

    # ── Sección narrativa de Fuentes de Agua ──────────────────────────────
    if data_ready:
        st.markdown("---")
        st.markdown("#### 📝 Narrativa auto-generada — Fuentes de Agua")
        st.markdown("<div class='info-box' style='font-size:.8rem'>Los valores <b>en negro</b> provienen del DAILY. Los <b style='background:#fed7aa;padding:1px 5px;border-radius:3px'>en naranja</b> requieren ingreso manual (Tab Meteorología / Comparativa Histórica).</div>", unsafe_allow_html=True)

        # Valores complementarios — leídos de session_state (se ingresan en 📈 Comparativa y 🚢 Esclusajes)
        fn_hist_chcp_hm3 = float(st.session_state.get("x_hist_hm3", hist.get("hist_chcp_hm3", 0.0) or 0.0))
        fn_hist_chcp_m3s = float(st.session_state.get("x_hist_m3s", hist.get("hist_chcp_m3s", 0.0) or 0.0))
        fn_hist_neto_hm3 = round(float(hist.get("hist_neto_hm3", 0.0) or 0.0), 2)
        fn_hist_neto_m3s = round(float(hist.get("hist_neto_m3s", 0.0) or 0.0), 3)
        fn_calado_pan    = float(st.session_state.get("x_cpan", 39.50))
        fn_calado_neo    = float(st.session_state.get("x_cneo", 50.0))
        fn_pos_hist      = st.session_state.get("fn_pos_hist_txt", "")
        fn_periodo_hist  = st.session_state.get("fn_periodo_hist_txt", "")
        fn_pos_hist_neto = str(datos.get("hist_neto_scale_pos", "") or "")
        fn_cls_hist_neto = str(datos.get("hist_neto_scale_class", "") or "")
        _years_netfl = hist.get("years_netfl") if isinstance(hist, dict) else None
        fn_periodo_hist_neto = (f"{_years_netfl[0]}-{_years_netfl[1]}" if _years_netfl else "")

        st.markdown("""
        <div class='info-box' style='font-size:.82rem'>
        ℹ️ Los valores históricos (prom. CHCP, calados, posición) se configuran en las pestañas
        <b>📈 Comparativa Histórica</b> y <b>🚢 Esclusajes</b>.<br>
        Posición histórica y período de comparación se pueden ajustar abajo:
        </div>""", unsafe_allow_html=True)
        fn_col1, fn_col2 = st.columns(2)
        with fn_col1:
            fn_pos_hist   = st.text_input("Posición histórica (ej: 4 de 128)", value=fn_pos_hist, key="fn_pos_hist_txt", help="Formato 'N de M': ej '4 de 128'. El app convierte pos. 1 a 'el mes más húmedo' automáticamente.")
        with fn_col2:
            fn_periodo_hist = st.text_input("Período comparación (ej: 1898-2025)", value=fn_periodo_hist, key="fn_periodo_hist_txt")
        # Métricas de referencia
        _h_ref1, _h_ref2, _h_ref3, _h_ref4 = st.columns(4)
        _h_ref1.metric("Hist. prom. CHCP (hm³)", f"{fn_hist_chcp_hm3:.1f}" if fn_hist_chcp_hm3 else "—")
        _h_ref2.metric("Hist. prom. CHCP (m³/s)", f"{fn_hist_chcp_m3s:.2f}" if fn_hist_chcp_m3s else "—")
        _h_ref3.metric("Calado Panamax (pies)", f"{fn_calado_pan:.2f}" if fn_calado_pan else "—")
        _h_ref4.metric("Calado Neopanamax (pies)", f"{fn_calado_neo:.1f}" if fn_calado_neo else "—")

        def _fv(val, empty_cond, fmt=".1f"):
            bg = "background:#fed7aa;padding:1px 5px;border-radius:3px;font-weight:700" if empty_cond else "font-weight:700"
            txt = format(val, fmt) if isinstance(val, (int,float)) else str(val)
            return f"<span style='{bg}'>{txt}</span>"

        _at   = D("aporte_total_hm3"); _atm  = D("aporte_total_m3s")
        _an   = D("aporte_neto_hm3");  _anm  = D("aporte_neto_m3s")
        _ev   = D("evaporacion_hm3");  _evm  = D("evaporacion_m3s")
        _alh  = D("alhajuela_hm3");    _alh_p= D("alhajuela_pct")
        _gat  = D("gatun_hm3");        _gat_p= D("gatun_pct")
        _ngl  = D("nivel_gatun_pies"); _nglm = D("nivel_gatun_m")
        _pct_at  = ((_at - fn_hist_chcp_hm3) / fn_hist_chcp_hm3 * 100) if fn_hist_chcp_hm3 > 0 else 0
        _pct_an  = ((_an - fn_hist_neto_hm3) / fn_hist_neto_hm3 * 100) if fn_hist_neto_hm3 > 0 else 0
        _pct_at_txt, _pct_at_dir = _pct_rel_text(_pct_at, positive_text="por encima al", negative_text="por debajo del", zero_text="igual al", decimals=0)
        _pct_an_txt, _pct_an_dir = _pct_rel_text(_pct_an, positive_text="por arriba del", negative_text="por debajo del", zero_text="igual al", decimals=0)
        _fn_pos_hist_txt = _posicion_historica_texto(fn_pos_hist, tipo, mes_sel)
        _fn_hist_years = _anios_registro_incluyendo_actual(
            fn_periodo_hist, anio_sel, st.session_state.get("h_n_years", hist.get("n_years_chcp", 0) if hist else 0)
        )

        bullets_fa = [
            (f"Los aportes totales a la Cuenca Hidrográfica del Canal de Panamá (CHCP) durante "
             f"<b>{mes_sel.lower()} {anio_sel}</b> fueron "
             f"{_fv(_at, False)} hm³ ({_fv(_atm, False)} m³/s), desglosados en "
             f"{_fv(_an, False)} hm³ ({_fv(_anm, False)} m³/s) de aportes netos y "
             f"{_fv(_ev, False)} hm³ ({_fv(_evm, False, '.2f')} m³/s) de evaporación directa en los embalses."),

            (f"Los aportes totales a la CHCP ({_fv(_at, False, '.0f')} hm³) de {mes_sel.lower()} {anio_sel}, "
             f"se encuentra como {_fv(_fn_pos_hist_txt if _fn_pos_hist_txt else '[posición]', not _fn_pos_hist_txt)} "
             f"en {_fv(_fn_hist_years, _fn_hist_years==0, '.0f')} años de registro continuo, "
             f"en una escala de húmedo a seco para el periodo {_fv(fn_periodo_hist if fn_periodo_hist else '[período]', not fn_periodo_hist)}. "
             f"Los cuales se encuentran en un {_fv(_pct_at_txt, fn_hist_chcp_hm3==0)} % {_pct_at_dir} promedio histórico de "
             f"{_fv(fn_hist_chcp_hm3, fn_hist_chcp_hm3==0, '.0f')} hm³ ({_fv(fn_hist_chcp_m3s, fn_hist_chcp_m3s==0, '.1f')} m³/s)."),

            (f"Los aportes netos de {mes_sel.lower()} {anio_sel} fueron igual a "
             f"{_fv(_an, False)} hm³ ({_fv(_anm, False)} m³/s), se ubican en la posición "
             f"{_fv(fn_pos_hist_neto if fn_pos_hist_neto else '[posición]', not fn_pos_hist_neto)} dentro de la escala histórica de húmedo a seco"
             f"{(' (' + fn_cls_hist_neto + ')') if fn_cls_hist_neto else ''} para el periodo "
             f"{_fv(fn_periodo_hist_neto if fn_periodo_hist_neto else '[período]', not fn_periodo_hist_neto)}, y se encuentran "
             f"{_fv(_pct_an_txt, fn_hist_neto_hm3==0)}% {_pct_an_dir} promedio histórico "
             f"{_fv(fn_hist_neto_hm3, fn_hist_neto_hm3==0)} hm³ ({_fv(fn_hist_neto_m3s, fn_hist_neto_m3s==0, '.2f')} m³/s)."),

            (f"En {mes_sel.lower()} {anio_sel} la distribución de los aportes hídricos totales a la CHCP fue: "
             f"Subcuenca embalse Alhajuela el {_fv(_alh_p, False, '.1f')}% ({_fv(_alh, False)} hm³), "
             f"Subcuenca del embalse Gatún el {_fv(_gat_p, False, '.1f')}% ({_fv(_gat, False)} hm³)."),

            (f"Los niveles de los embalses Gatún y Alhajuela presentan suficiente agua disponible. "
             f"El nivel de agua promedio en el embalse Gatún durante este mes fue de "
             f"{_fv(_ngl, False, '.2f')} pies ({_fv(_nglm, False, '.2f')} m) PLD. "
             f"Esta condición ha permitido mantener el calado máximo permisible de "
             f"{_fv(fn_calado_pan, fn_calado_pan==0, '.2f')} pies para los buques Panama y "
             f"{_fv(fn_calado_neo, fn_calado_neo==0, '.0f')} pies para los buques Neopanamax."),
        ]

        for bi, btext in enumerate(bullets_fa, 1):
            st.markdown(f"<div class='info-box' style='font-size:.85rem;line-height:1.7;margin-bottom:6px'><b>•</b> {btext}</div>",
                        unsafe_allow_html=True)

        # Auditoría Fuentes
        st.markdown("---")
        st.markdown("#### 🧮 Auditoría — Fuentes de Agua")
        _fn_missing = []
        if not fn_pos_hist:       _fn_missing.append("Posición histórica")
        if not fn_periodo_hist:   _fn_missing.append("Período comparación")
        if fn_hist_chcp_hm3 == 0: _fn_missing.append("Hist. prom. CHCP (hm³)")
        if fn_hist_neto_hm3 == 0: _fn_missing.append("Hist. prom. Aporte Neto (hm³)")
        if fn_calado_pan == 0:    _fn_missing.append("Calado Panamax")
        if fn_calado_neo == 0:    _fn_missing.append("Calado Neopanamax")

        if _fn_missing:
            st.markdown(f"<div class='warn-box'>⚠️ Revisión del informe: <b>{', '.join(_fn_missing)}</b></div>", unsafe_allow_html=True)
        else:
            st.markdown("<div class='ok-box'>✅ Todos los campos de Fuentes de Agua están completos.</div>", unsafe_allow_html=True)

        df_audit_fa = pd.DataFrame([
            ("Aporte total CHCP (hm³)",         round(_at, 2),  "DAILY ✅", "Auto"),
            ("Aporte total CHCP (m³/s)",         round(_atm, 2), "DAILY ✅", "Auto"),
            ("Aporte neto (hm³)",                round(_an, 2),  "DAILY ✅", "Auto"),
            ("Evaporación (hm³)",                round(_ev, 2),  "DAILY ✅", "Auto"),
            ("Subcuenca Alhajuela (hm³ / %)",    f"{_alh:.1f} / {_alh_p:.1f}%", "DAILY ✅", "Auto"),
            ("Subcuenca Gatún (hm³ / %)",        f"{_gat:.1f} / {_gat_p:.1f}%", "DAILY ✅", "Auto"),
            ("Nivel Gatún (pies / m)",           f"{_ngl:.2f} / {_nglm:.2f}", "DAILY ✅", "Auto"),
            ("Hist. prom. CHCP (hm³)",           fn_hist_chcp_hm3 or "—", "—", "Manual ✅" if fn_hist_chcp_hm3>0 else "⚠️ Falta"),
            ("Hist. prom. Aporte Neto (hm³)",    fn_hist_neto_hm3 or "—", "—", "Manual ✅" if fn_hist_neto_hm3>0 else "⚠️ Falta"),
            ("Posición histórica",               fn_pos_hist or "—", "—", "Manual ✅" if fn_pos_hist else "⚠️ Falta"),
            ("Calado Panamax (pies)",            fn_calado_pan or "—", "—", "Manual ✅" if fn_calado_pan>0 else "⚠️ Falta"),
            ("Calado Neopanamax (pies)",         fn_calado_neo or "—", "—", "Manual ✅" if fn_calado_neo>0 else "⚠️ Falta"),
        ], columns=["Variable", "Valor", "Fuente DAILY", "Estado"])
        st.dataframe(df_audit_fa, use_container_width=True, hide_index=True)


with tabs[6]:
    st.markdown('<div class="sec-hdr">🚢 Esclusajes y Ahorros</div>', unsafe_allow_html=True)
    if not data_ready:
        st.info("Carga el archivo DAILY_AND_ACUMMEN para continuar.")
    else:
        ce1, ce2, ce3 = st.columns(3)
        with ce1:
            st.markdown("**Panamax**")
            st.metric("Volumen", f"{D('panamax_hm3'):.2f} hm³", f"{D('panamax_pct'):.1f}%")
            st.metric("Tránsitos", f"{D('transitos_panamax'):.0f}")
            st.metric("Agua / tránsito", f"{D('agua_panamax_trans'):.4f} hm³")
        with ce2:
            st.markdown("**Neopanamax**")
            st.metric("Volumen", f"{D('neopanamax_hm3'):.2f} hm³", f"{D('neopanamax_pct'):.1f}%")
            st.metric("Tránsitos", f"{D('transitos_neopanamax'):.0f}")
            st.metric("Agua / tránsito", f"{D('agua_neo_trans'):.4f} hm³")
        with ce3:
            st.markdown("**Uso Promedio**")
            st.metric("Total esclusajes", f"{D('esclusaje_hm3'):.2f} hm³")
            st.metric("Uso diario promedio", f"{D('uso_prom_diario_hm3'):.4f} hm³")
            st.metric("Uso diario (m³/s)", f"{D('uso_prom_diario_m3s'):.2f}")

        st.markdown("---")
        st.markdown("##### 🔎 Diagnóstico ahorros — columnas W y X del DAILY")
        if f_daily_src and data_ready:
            try:
                _df_wx = load_daily(_source_bytes(f_daily_src))
                import openpyxl as _opx_wx
                def _sum_col_wx(df, col_letter):
                    if tipo == "Mensual":
                        _mn = MESES_NUM.get(mes_sel, 1)
                        _yr = anio_sel - 1 if _mn in [10, 11, 12] else anio_sel
                        _mf = df[(df["ACTDATE"].dt.year == _yr) & (df["ACTDATE"].dt.month == _mn)]
                    else:
                        _meses_p = _get_meses_periodo(tipo, int(anio_sel), trimestre, semestre)
                        _mask_wx = pd.Series(False, index=df.index)
                        for _mn2, _yr2 in _meses_p:
                            _mask_wx |= ((df["ACTDATE"].dt.year == _yr2) & (df["ACTDATE"].dt.month == _mn2))
                        _mf = df[_mask_wx]
                    if _mf.empty:
                        return 0.0, 0, "(sin datos del período)"
                    try:
                        _idx = _opx_wx.utils.column_index_from_string(col_letter) - 1
                    except Exception:
                        return 0.0, 0, "(error col)"
                    if _idx >= len(_mf.columns):
                        return 0.0, 0, "(fuera de rango)"
                    _col_name = str(_mf.columns[_idx])
                    _vals = pd.to_numeric(_mf.iloc[:, _idx], errors="coerce")
                    _n_llenos = int(_vals.notna().sum())
                    _suma = float(_vals.sum(min_count=1) or 0.0)
                    return _suma, _n_llenos, _col_name
                _w_ee, _w_n, _w_col = _sum_col_wx(_df_wx, "W")
                _x_ee, _x_n, _x_col = _sum_col_wx(_df_wx, "X")
                _w_hm3 = _w_ee * EE_HM3
                _x_hm3 = _x_ee * EE_HM3
                _wxc1, _wxc2, _wxc3, _wxc4 = st.columns(4)
                _wxc1.metric(f"Col W  — {_w_col}", f"{_w_hm3:.3f} hm³", f"{_w_n} días con dato · {_w_ee:.1f} EE")
                _wxc2.metric(f"Col X  — {_x_col}", f"{_x_hm3:.3f} hm³", f"{_x_n} días con dato · {_x_ee:.1f} EE")
                _wxc3.metric("Total DAILY (W+X)", f"{_w_hm3+_x_hm3:.3f} hm³", "× 0.2081976 hm³/EE")
                _wxc4.metric("Total en uso (ajustes finos)", f"{D('ahorro_total_hm3'):.3f} hm³",
                             "⚠️ manual" if _w_n == 0 and _x_n == 0 else "✅ del DAILY")
                if _w_n == 0 and _x_n == 0:
                    st.markdown("<div class='warn-box'>⚠️ Las columnas W y X están <b>vacías</b> para el período seleccionado. Los ahorros deben ingresarse manualmente en la sección <b>⚙️ Ajustes finos del informe</b> (pestaña 📝 Exportar).</div>", unsafe_allow_html=True)
                elif abs(_w_hm3 + _x_hm3 - D("ahorro_total_hm3")) > 0.05:
                    st.markdown(f"<div class='info-box'>ℹ️ El DAILY (W+X) da <b>{_w_hm3+_x_hm3:.3f} hm³</b>; los Ajustes Finos usan <b>{D('ahorro_total_hm3'):.3f} hm³</b>. Si difieren, los Ajustes Finos fueron editados manualmente.</div>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<div class='ok-box'>✅ W y X del DAILY consistentes con los ahorros en uso ({_w_hm3+_x_hm3:.3f} hm³).</div>", unsafe_allow_html=True)
            except Exception as _e_wx:
                st.markdown(f"<div class='warn-box'>⚠️ No se pudo leer W/X del DAILY: {_e_wx}</div>", unsafe_allow_html=True)
        else:
            st.markdown("<div class='info-box'>ℹ️ Carga el DAILY para verificar las columnas W y X de ahorros.</div>", unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("##### CCA y Ahorros")
        col_cca, col_aho = st.columns(2)
        with col_cca:
            st.markdown("""
            <div class='info-box'>ℹ️ CCA y ahorros no son equivalentes. Usa este campo para el CCA reportado en el informe.</div>
            """, unsafe_allow_html=True)
            _cca_esc = float(st.session_state.get("rep_cca_hm3", 0.0) or 0.0)
            transitos_neo = D("transitos_neopanamax", 304)
            cca_unit = _cca_esc / transitos_neo if transitos_neo > 0 else 0
            st.metric("CCA total (hm³)", f"{_cca_esc:.2f}" if _cca_esc else "⚠️ Sin dato",
                      "✏️ Configura en pestaña 🌊 Salinidad")
            st.metric("CCA unitario", f"{cca_unit:.4f} hm³/tránsito",
                      f"{transitos_neo:.0f} tránsitos Neopanamax")

        with col_aho:
            st.markdown("**Ahorros implementados · DAILY W/X × 0.2081976**")
            st.metric("NeoPanamax · columna X", f"{D('ahorro_neopanamax_hm3'):.2f} hm³")
            st.metric("Panamax · columna W", f"{D('ahorro_panamax_hm3'):.2f} hm³")
            st.caption(f"EE base: Panamax {D('ahorro_panamax_ee'):.2f} · NeoPanamax {D('ahorro_neopanamax_ee'):.2f}")
            st.metric("Total ahorrado", f"{D('ahorro_total_hm3'):.2f} hm³",
                      f"{D('ahorro_lamina_pies'):.3f} pies en Gatún")

        # Gauge de uso del calado
        st.markdown("---")
        st.markdown("##### Calado disponible vs nivel Gatún")
        fig_gauge = make_subplots(rows=1, cols=2, specs=[[{"type":"indicator"},{"type":"indicator"}]])
        fig_gauge.add_trace(go.Indicator(
            mode="gauge+number",
            value=D("nivel_gatun_pies"),
            title={"text": "Nivel Gatún (pies)"},
            gauge={"axis":{"range":[75,95]},
                   "bar":{"color":"#0072b8"},
                   "steps":[{"range":[75,82],"color":"#fee2e2"},
                             {"range":[82,87],"color":"#fef9c3"},
                             {"range":[87,95],"color":"#dcfce7"}],
                   "threshold":{"line":{"color":"#003a6e","width":3},
                                "thickness":0.8,"value":87}},
            number={"suffix":" pies", "font":{"size":24}}), row=1, col=1)
        fig_gauge.add_trace(go.Indicator(
            mode="gauge+number",
            value=D("nivel_alh_pies"),
            title={"text": "Nivel Alhajuela (pies)"},
            gauge={"axis":{"range":[200,270]},
                   "bar":{"color":"#003a6e"},
                   "steps":[{"range":[200,220],"color":"#fee2e2"},
                             {"range":[220,240],"color":"#fef9c3"},
                             {"range":[240,270],"color":"#dcfce7"}]},
            number={"suffix":" pies", "font":{"size":24}}), row=1, col=2)
        fig_gauge.update_layout(plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
            height=280, margin=dict(t=20,b=10,l=10,r=10))
        st.plotly_chart(fig_gauge, use_container_width=True)


# ═══════════════════════════════════════════════════════════════
#  TAB 5 — ENERGÍA
# ═══════════════════════════════════════════════════════════════
with tabs[7]:
    st.markdown('<div class="sec-hdr">📤 Usos del Agua (Salidas)</div>', unsafe_allow_html=True)
    if not data_ready:
        st.info("Carga el archivo DAILY_AND_ACUMMEN para continuar.")
    else:
        col_u1, col_u2 = st.columns([1,1])

        usos_data = {
            "Esclusajes PNX":          (D("panamax_hm3"), D("panamax_m3s"), D("panamax_cfs")),
            "Esclusajes NPX":          (D("neopanamax_hm3"), D("neopanamax_m3s"), D("neopanamax_cfs")),
            "Evaporación Gatún":       (D("evap_gatun_hm3"), hm3_to_m3s_period(D("evap_gatun_hm3"), n_dias) if n_dias else 0, hm3_to_cfs_period(D("evap_gatun_hm3"), n_dias) if n_dias else 0),
            "Evaporación Alhajuela":   (D("evap_alh_hm3"), hm3_to_m3s_period(D("evap_alh_hm3"), n_dias) if n_dias else 0, hm3_to_cfs_period(D("evap_alh_hm3"), n_dias) if n_dias else 0),
            "Potabilización Gatún":    (D("potabilizacion_gat_hm3"), D("potabilizacion_gat_m3s"), D("potabilizacion_gat_cfs")),
            "Potabilización Alh.":     (D("potabilizacion_alh_hm3"), D("potabilizacion_alh_m3s"), D("potabilizacion_alh_cfs")),
            "Concesiones/Fugas Gatún": (D("concesiones_gat_hm3"), D("concesiones_gat_m3s"), D("concesiones_gat_cfs")),
            "Concesiones/Fugas Alh.":  (D("concesiones_alh_hm3"), D("concesiones_alh_m3s"), D("concesiones_alh_cfs")),
            "Hidrogeneración Gatún":   (D("hidro_gatun_hm3"), hm3_to_m3s_period(D("hidro_gatun_hm3"), n_dias) if n_dias else 0, hm3_to_cfs_period(D("hidro_gatun_hm3"), n_dias) if n_dias else 0),
            "ZZ Flush (manual)":       (D("zzflush_auto_hm3"), D("zzflush_auto_m3s"), D("zzflush_auto_cfs")),
            "Vertidos Gatún (Q)":      (D("vertidos_gat_hm3"), D("vertidos_gat_m3s"), D("vertidos_gat_cfs")),
        }
        total_sal = D("total_salidas_hm3", 0)

        with col_u1:
            st.markdown("##### Tabla de Usos")
            rows_html = ""
            for uso, (hm3, m3s, cfs) in usos_data.items():
                pct = hm3 / total_sal * 100 if total_sal > 0 else 0
                rows_html += f"<tr><td>{uso}</td><td>{hm3:.3f}</td><td>{m3s:.3f}</td><td>{cfs:.2f}</td><td>{pct:.1f}%</td></tr>"
            rows_html += f"<tr class='total'><td><b>TOTAL SALIDAS</b></td><td><b>{total_sal:.2f}</b></td><td><b>{D('total_salidas_m3s'):.3f}</b></td><td><b>{D('total_salidas_cfs'):.2f}</b></td><td><b>100%</b></td></tr>"
            st.markdown(f"""
            <table class='styled-table'>
            <thead><tr><th>Uso</th><th>hm³</th><th>m³/s</th><th>pies³/s</th><th>%</th></tr></thead>
            <tbody>{rows_html}</tbody>
            </table>""", unsafe_allow_html=True)

            st.markdown("<br>")
            st.markdown("##### 🔄 Trasvase Alhajuela → Gatún")
            periodo_lbl = "mensual" if tipo=="Mensual" else "semestral"
            trasvase_pct_esc = D("trasvase_pct_esc")
            st.markdown(f"""
            <div class='success-box'>
            🔄 <b>Volumen trasvasado:</b> {D("trasvase_hm3"):.2f} hm³ ({D("trasvase_m3s"):.2f} m³/s)<br>
            Representa el <b>{trasvase_pct_esc:.1f}%</b> del volumen {periodo_lbl} de esclusajes
            </div>""", unsafe_allow_html=True)
            if tipo == "Mensual":
                st.markdown(f"""
                <div class='ok-box'>
                🧂 <b>ZZ Flush manual:</b> {D('zzflush_auto_hm3'):.4f} hm³ ({D('zzflush_auto_m3s'):.3f} m³/s).<br>
                🌊 <b>Vertidos Gatún (columna Q / GATSPILL):</b> {D('vertidos_gat_hm3'):.3f} hm³ ({D('vertidos_gat_m3s'):.3f} m³/s).<br>
                ℹ️ <b>Descarga operativa Madden (MADSPILL):</b> {D('vertidos_mad_ops_hm3'):.3f} hm³ — movimiento operacional interno.
                </div>""", unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class='info-box'>
                🧂 <b>ZZ Flush manual del período:</b> {D('zzflush_auto_hm3'):.3f} hm³<br>
                🌊 <b>Vertidos Gatún:</b> {D('vertidos_gat_hm3'):.3f} hm³<br>
                💾 <b>Excedente almacenado:</b> {D('excedente_hm3',0):.2f} hm³
                </div>""", unsafe_allow_html=True)
        with col_u2:
            st.markdown("##### Distribución porcentual")
            u_names = list(usos_data.keys())
            u_vals  = [v[0] for v in usos_data.values()]
            fig_u = px.pie(names=u_names, values=u_vals, hole=0.42,
                color_discrete_sequence=["#003a6e","#0072b8","#ef4444","#f87171",
                                          "#38bdf8","#7dd3fc","#fbbf24","#34d399","#0ea5e9"])
            fig_u.update_traces(textinfo="percent+label", textfont_size=10,
                pull=[0.06 if "Esclusaje" in n else 0 for n in u_names])
            fig_u.update_layout(plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                height=400, margin=dict(t=10,b=10,l=10,r=10), showlegend=True,
                legend=dict(orientation="v", font=dict(size=10)))
            st.plotly_chart(fig_u, use_container_width=True)


# ═══════════════════════════════════════════════════════════════
#  TAB 4 — ESCLUSAJES
# ═══════════════════════════════════════════════════════════════
with tabs[8]:
    st.markdown('<div class="sec-hdr">⚡ Producción Hidroeléctrica</div>', unsafe_allow_html=True)
    if not data_ready:
        st.info("Carga el archivo DAILY_AND_ACUMMEN para continuar.")
    else:
        col_e1, col_e2 = st.columns(2)
        with col_e1:
            st.markdown("**🏭 Hidroeléctrica Madden**")
            st.metric("Producción promedio", f"{D('energia_madden_mw'):.2f} MW")
            st.metric("Producción diaria", f"{D('energia_madden_mwh_dia'):.1f} MWh/día")
            demanda_acp = st.number_input("Demanda interna ACP (MW)", value=19.0, step=0.1)
            diferencia = D("energia_madden_mw") - demanda_acp
            if diferencia >= 0:
                st.markdown(f"""<div class='success-box'>
                ✅ Madden cubre la demanda interna. Excedente: <b>{diferencia:.2f} MW</b></div>""",
                unsafe_allow_html=True)
            else:
                st.markdown(f"""<div class='warn-box'>
                ⚠️ Madden NO cubre la demanda interna. Déficit: <b>{abs(diferencia):.2f} MW</b>
                (Se requiere termoeléctrica o importación)</div>""", unsafe_allow_html=True)

        with col_e2:
            st.markdown("**🏭 Hidroeléctrica Gatún**")
            st.metric("Producción promedio",   f"{D('energia_gatun_mw'):.4f} MW")
            st.metric("Producción diaria",     f"{D('energia_gatun_mwh_dia'):.2f} MWh/día")
            st.metric("Vol. turbinado Gatún",  f"{D('hidro_gatun_hm3'):.4f} hm³",
                      help="GATMCF × 0.028317 — verificado contra informe Hidroestadística")
            st.metric("Vol. turbinado Madden", f"{D('hidro_madden_hm3'):.2f} hm³",
                      help="MADMCF × 0.028316846592 — millones de pies cúbicos a hm³; flujo operacional Madden (hidrogeneración)")

        # Energy timeline
        if f_daily:
            st.markdown("---")
            st.markdown("##### Producción diaria — valores directos de columnas AS/AT")
            try:
                df_d2 = load_daily(_source_bytes(f_daily))
                if tipo == "Mensual":
                    mes_n = MESES_NUM.get(mes_sel, 3)
                    yr = anio_sel if mes_n <= 9 else anio_sel - 1
                    mes_e = df_d2[(df_d2["ACTDATE"].dt.year==yr) & (df_d2["ACTDATE"].dt.month==mes_n)]
                else:
                    meses_e = _get_meses_periodo(tipo, int(anio_sel), trimestre, semestre)
                    mask_e = pd.Series(False, index=df_d2.index)
                    for mn, yr in meses_e:
                        mask_e |= ((df_d2["ACTDATE"].dt.year==yr) & (df_d2["ACTDATE"].dt.month==mn))
                    mes_e = df_d2[mask_e]
                fig_en = go.Figure()
                # Producción diaria auditada para el tablero:
                #   Madden/Alhajuela = valores directos de AT (MADMW)
                #   Gatún            = valores directos de AS (GATMW)
                # No multiplicar por 24: estas columnas ya vienen ajustadas para el criterio del informe.
                y_madden_mwh_d = pd.to_numeric(mes_e["MADMW"], errors="coerce")
                y_gatun_mwh_d  = pd.to_numeric(mes_e["GATMW"], errors="coerce")
                fig_en.add_trace(go.Scatter(x=mes_e["ACTDATE"], y=y_madden_mwh_d,
                    mode="lines+markers", name="Madden — AT/MADMW",
                    line=dict(color="#f59e0b", width=2.5), marker=dict(size=5)))
                fig_en.add_trace(go.Scatter(x=mes_e["ACTDATE"], y=y_gatun_mwh_d,
                    mode="lines+markers", name="Gatún — AS/GATMW",
                    line=dict(color="#0072b8", width=2), marker=dict(size=4)))
                fig_en.add_hline(y=demanda_acp, line_dash="dash",
                    line_color="#dc2626", annotation_text="Demanda ACP (MW)")
                fig_en.update_layout(height=300, plot_bgcolor="rgba(0,0,0,0)",
                    paper_bgcolor="rgba(0,0,0,0)", margin=dict(t=10,b=10,l=10,r=10),
                    yaxis=dict(title="Valor directo columna MW", gridcolor="rgba(148,163,184,.2)"),
                    legend=dict(orientation="h", y=1.05))
                st.plotly_chart(fig_en, use_container_width=True)
            except Exception:
                pass



# ═══════════════════════════════════════════════════════════════
#  TAB 6 — METEOROLOGÍA
# ═══════════════════════════════════════════════════════════════
# Estado de fuente salinidad (compartido con tab 7)
sal_mode_widget = st.session_state.get("_sal_mode_select", "Manual / Calculado del Daily")

with tabs[2]:
    st.markdown('<div class="sec-hdr">🌡️ Meteorología — Precipitación CHCP</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class='info-box'>
    ℹ️ La precipitación <b>no está en el DAILY</b>. Ingresa los valores manualmente.<br>
    Los promedios históricos provienen principalmente de <b>03_QAVGMCHCP</b> y, si está disponible, de <b>04_GATNETFL</b>.
    Los campos en cero o con incongruencias aparecerán en <b style='background:#fed7aa;padding:1px 6px;border-radius:3px'>naranja</b> en el informe exportado.
    </div>""", unsafe_allow_html=True)

    st.markdown("#### 🌧️ Precipitación — CHCP")

    _met_modo_tx = st.radio(
        "Modo de ingreso — precipitación",
        ["Campos individuales", "Texto oficial (pegar íntegramente)"],
        index=1,
        horizontal=True, key="met_texto_mode",
        help="'Texto oficial': pega la viñeta completa del informe fuente; reemplazará la auto-generada en el DOCX."
    )
    if _met_modo_tx == "Texto oficial (pegar íntegramente)":
        _met_libre_v = st.text_area(
            "📋 Viñeta oficial de precipitación",
            value=str(st.session_state.get("met_texto_libre", "") or ""),
            height=120, key="met_texto_libre",
            help="Pega aquí el párrafo oficial tal cual. Reemplazará la primera viñeta en el DOCX exportado."
        )
        # Auto-poblar campos numéricos si están en 0
        import re as _re_met
        if str(_met_libre_v or "").strip():
            _m1 = _re_met.search(r"(\d+[\.,]?\d*)\s*mm", _met_libre_v)
            _m2 = _re_met.search(r"promedio histórico de\s+([\d.,]+)\s*mm", _met_libre_v, _re_met.I)
            if _m1:
                try:
                    _vv = float(_m1.group(1).replace(",", "."))
                    if float(st.session_state.get("x_prec", 0) or 0) < 0.001:
                        st.session_state["x_prec"] = _vv
                except: pass
            if _m2:
                try:
                    _vv = float(_m2.group(1).replace(",", "."))
                    if float(st.session_state.get("x_phist", 0) or 0) < 0.001:
                        st.session_state["x_phist"] = _vv
                except: pass
        st.markdown("<div class='info-box' style='font-size:.79rem'>ℹ️ El texto pegado reemplazará la primera viñeta auto-generada en el DOCX. Los campos numéricos abajo siguen activos para estadísticas y comparativas.</div>", unsafe_allow_html=True)
        st.markdown("")

    mc1, mc2, mc3 = st.columns(3)
    with mc1:
        st.markdown("**Período actual**")
        met_precip_mm   = st.number_input("Precipitación total del período (mm)", value=float(st.session_state.get("x_prec", 0.0)), min_value=0.0, step=0.1, format="%.1f", key="x_prec")
        met_precip_hist = st.number_input("Promedio histórico del período (mm)",  value=float(st.session_state.get("x_phist", 0.0)), min_value=0.0, step=0.1, format="%.1f", key="x_phist")
        met_pct_informe = st.number_input("% sobre/bajo promedio histórico (informe)",
            value=float(st.session_state.get("x_met_pct_informe",
                round(abs((float(st.session_state.get("x_prec", 0.0)) - float(st.session_state.get("x_phist", 0.0)))
                      / float(st.session_state.get("x_phist", 1.0)) * 100), 1)
                if float(st.session_state.get("x_phist", 0.0)) > 0 else 0.0)),
            min_value=0.0, step=0.1, format="%.1f",
            key="x_met_pct_informe",
            help="Se auto-calcula desde precipitación/histórico; edítalo si el valor oficial difiere (ej: 514.0%).")
    with mc2:
        st.markdown("**Registro y ranking histórico**")
        _periodo_default = st.session_state.get("x_pper", "")
        met_periodo = st.text_input("Período registro (ej: 1950-2026)", value=_periodo_default, key="x_pper")
        _n_anios_auto = _periodo_registro_n_anios(met_periodo)
        _n_anios_default = int(st.session_state.get("h_n_years", _n_anios_auto) or _n_anios_auto or 0)
        met_n_anios = st.number_input("Años de registro", value=_n_anios_default, min_value=0, step=1, key="met_n_anios")
        met_posicion = st.number_input("Posición histórica de humedad", value=int(st.session_state.get("x_ppos", 0) or 0), min_value=0, step=1, key="x_ppos",
            help="Ejemplo: 25 para redactar 'posición 25 de los marzos más húmedos'.")
        met_rank_label = st.text_input("Comparación ranking", value=st.session_state.get("x_prank_label", _meteo_rank_label_default(tipo, mes_sel)), key="x_prank_label",
            help="Ejemplos: marzos más húmedos, abriles más húmedos, trimestres más húmedos.")
    with mc3:
        st.markdown("**Días/eventos de mayor lluvia**")
        met_dias_lluvia = st.text_area("Descripción días máx. (texto libre)",
            value=st.session_state.get("x_prec_dias", ""),
            height=100, key="x_prec_dias",
            help="Opcional. Ejemplo: principalmente los días 20 y 21.")

    # Auto-cálculos precipitación
    met_superavit  = met_precip_mm - met_precip_hist
    met_pct_sobre  = (met_superavit / met_precip_hist * 100) if met_precip_hist > 0 else 0.0
    met_dir        = "por encima" if met_precip_mm >= met_precip_hist else "por debajo"

    ka, kb, kc, kd, ke = st.columns(5)
    ka.metric("Superávit / Déficit (mm)", f"{met_superavit:+.1f}")
    kb.metric("% auto-calculado",         f"{met_pct_sobre:+.1f}%")
    kc.metric("% en informe",             f"{met_pct_informe:.1f}%")
    kd.metric("Dirección",                met_dir.capitalize())
    ke.metric("N° años registro",         met_n_anios)

    _met_missing = []
    if met_precip_mm == 0:    _met_missing.append("Precipitación total (mm)")
    if met_precip_hist == 0:  _met_missing.append("Promedio histórico (mm)")
    if not met_periodo:       _met_missing.append("Período de registro")
    if met_n_anios == 0:      _met_missing.append("Años de registro")
    if met_posicion == 0:     _met_missing.append("Posición histórica de humedad")
    if not met_rank_label:    _met_missing.append("Comparación ranking")

    if _met_missing:
        st.markdown(f"<div class='warn-box'>⚠️ <b>Pendiente de ingreso</b> (naranja en informe): {', '.join(_met_missing)}</div>", unsafe_allow_html=True)
    else:
        st.markdown("<div class='ok-box'>✅ Todos los campos de meteorología están completos y listos para el escrito.</div>", unsafe_allow_html=True)

    # Viñeta narrativa meteorología
    st.markdown("---")
    st.markdown("#### 📝 Viñeta auto-generada — Precipitación")

    def _met_val(val, empty_cond, fmt=".1f"):
        bg = "background:#fed7aa;padding:1px 5px;border-radius:3px;font-weight:700" if empty_cond else "font-weight:700"
        return f"<span style='{bg}'>{format(val, fmt) if isinstance(val, (int,float)) else val}</span>"

    # Usar % editable del informe si fue modificado por el usuario
    _pct_informe_display = float(st.session_state.get("x_met_pct_informe", abs(met_pct_sobre)))
    _periodo_lbl_met = f"{mes_sel.lower()} de {anio_sel}" if tipo == "Mensual" else label_per.lower()

    # ── Oración 2: superávit / déficit ──────────────────────────────────────
    if met_superavit >= 0:
        _oracion2 = (
            f"El superávit acumulado en este periodo fue de "
            f"{_met_val(round(met_superavit, 0), met_precip_mm==0 or met_precip_hist==0, '.0f')} mm."
        )
    else:
        _oracion2 = (
            f"El déficit acumulado en este periodo fue de "
            f"{_met_val(round(abs(met_superavit), 0), met_precip_mm==0 or met_precip_hist==0, '.0f')} mm."
        )

    # ── Oración 3: posición histórica ──────────────────────────────────────
    _rank_lbl_lower = (met_rank_label or "").lower()
    if met_posicion == 1:
        if "húmedo" in _rank_lbl_lower:
            _pos_frase = "el más húmedo"
        elif "seco" in _rank_lbl_lower:
            _pos_frase = "el más seco"
        else:
            _pos_frase = "el primero"
    else:
        _pos_frase = f"el {_met_val(met_posicion, met_posicion==0, 'd')}° lugar"

    _oracion3 = (
        f"Estos registros sitúan a este {'mes' if tipo == 'Mensual' else 'período'} como "
        f"{_met_val(_pos_frase, met_posicion==0)} en "
        f"{_met_val(met_n_anios, met_n_anios==0, 'd')} años de registros "
        f"({_met_val(met_periodo if met_periodo else '[período]', not met_periodo)})."
    )

    bullet_met = (
        f"La precipitación total sobre la Cuenca Hidrográfica del Canal de Panamá (CHCP) durante "
        f"<b>{_periodo_lbl_met}</b> fue de "
        f"{_met_val(met_precip_mm, met_precip_mm==0, '.0f')} mm, "
        f"{_met_val(_pct_informe_display, met_precip_hist==0, '.1f')}% {met_dir} del promedio histórico de "
        f"{_met_val(met_precip_hist, met_precip_hist==0, '.0f')} mm. "
        f"{_oracion2} "
        f"{_oracion3}"
    )
    if met_dias_lluvia:
        bullet_met += f" La mayor parte de la lluvia del período se concentró {met_dias_lluvia}."

    st.markdown(f"<div class='info-box' style='font-size:.88rem;line-height:1.75'>• {bullet_met}</div>",
                unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════
#  TAB 7 — SALINIDAD
# ═══════════════════════════════════════════════════════════════
with tabs[3]:
    st.markdown('<div class="sec-hdr">🌊 Salinidad — Embalse Gatún</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class='info-box'>
    ℹ️ La salinidad ya no se toma del <b>DAILY</b>.<br>
    Ingresa SPC, SPV y Decremento manualmente o carga un archivo externo.
    </div>""", unsafe_allow_html=True)
    # ── Selector de fuente de salinidad ────────────────────────────────────
    sal_mode_widget = st.radio(
        "Fuente de datos de salinidad",
        ["Manual", "Subir archivo Excel"],
        index=0 if sal_mode.startswith("Manual") else 1,
        key="_sal_mode_select",
        horizontal=True,
        help="Selecciona 'Manual' para ingresar SPC/SPV/Decremento manualmente. 'Subir archivo Excel' para usar un archivo externo."
    )
    st.markdown("#### 🌊 Salinidad — Embalse Gatún")

    _sal_modo_tx = st.radio(
        "Modo de ingreso — salinidad",
        ["Campos individuales", "Texto oficial (pegar íntegramente)"],
        index=1,
        horizontal=True, key="sal_texto_mode",
        help="'Texto oficial': pega la viñeta completa de salinidad del informe fuente."
    )
    if _sal_modo_tx == "Texto oficial (pegar íntegramente)":
        st.text_area(
            "📋 Viñeta oficial de salinidad",
            value=str(st.session_state.get("met_sal_texto_libre", "") or ""),
            height=120, key="met_sal_texto_libre",
            help="Reemplazará la viñeta de salinidad en el DOCX exportado. SPC, SPV y Decremento siguen activos para el dashboard."
        )
        import re as _re_sal
        _sal_tl_v = str(st.session_state.get("met_sal_texto_libre", "") or "")
        if _sal_tl_v.strip():
            _sp1 = _re_sal.search(r"SPC[^\d]*([\.\d,]+)\s*unidades", _sal_tl_v, _re_sal.I)
            _sp2 = _re_sal.search(r"SPV[^\d]*([\.\d,]+)\s*ups", _sal_tl_v, _re_sal.I)
            _sp3 = _re_sal.search(r"decremento del\s+([\d.,]+)\s*%", _sal_tl_v, _re_sal.I)
            for _pat, _key in [(_sp1, "x_sspc"), (_sp2, "x_sspv"), (_sp3, "x_sdec")]:
                if _pat:
                    try:
                        _vv = float(_pat.group(1).replace(",", "."))
                        if float(st.session_state.get(_key, 0) or 0) < 0.0001:
                            st.session_state[_key] = _vv
                    except: pass
        st.markdown("<div class='info-box' style='font-size:.79rem'>ℹ️ El texto pegado reemplazará la viñeta de salinidad. SPC, SPV y Decremento abajo siguen activos para el dashboard y la auditoría.</div>", unsafe_allow_html=True)
        st.markdown("")

    sal_source_lbl = datos.get("salinidad_source","manual") if data_ready else "manual"
    sal_spc_val    = D("salinidad_spc", 0.0)

    col_s1, col_s2 = st.columns([1,1])
    with col_s1:
        st.markdown(f"""
        <div class='info-box'>
        📡 <b>Fuente actual de salinidad:</b> {sal_source_lbl.upper()}<br>
        {'✅ SPC leído desde archivo externo.' if sal_source_lbl == 'archivo' else '✍️ SPC/SPV/Decremento definidos manualmente en esta pestaña.'}
        </div>""", unsafe_allow_html=True)

        if sal_mode_widget == "Manual" or not f_sal:
            sal_spc_inp = st.number_input("Salinidad SPC (ups)",
                value=float(sal_spc_val) if sal_spc_val else 0.0,
                step=0.0001, format="%.4f", key="x_sspc")
            sal_spv_inp = st.number_input("Salinidad SPV (ups)",
                value=float(datos.get("salinidad_spv", 0.0) or 0.0) if data_ready else 0.0,
                step=0.0001, format="%.4f", key="x_sspv")
            sal_dec_pct = st.number_input("Decremento vs máx. histórico (%)",
                value=float(datos.get("salinidad_dec_pct", 0.0) or 0.0) if data_ready else 0.0,
                step=0.01, key="x_sdec")
            if data_ready:
                datos["salinidad_spc"] = sal_spc_inp
                datos["salinidad_spv"] = sal_spv_inp
                datos["salinidad_dec_pct"] = sal_dec_pct
                datos["salinidad_source"] = "manual"
                datos["sal_series"] = pd.Series(dtype=float)
        else:
            if sal_manual:
                st.success(f"✅ Salinidad del archivo: **{sal_manual:.4f} ups**")
            sal_spc_inp = float(sal_manual or 0.0)
            st.metric("Salinidad SPC del archivo (ups)", f"{sal_spc_inp:.4f}")
            sal_spv_inp = st.number_input("Salinidad SPV (ups) — manual",
                value=float(st.session_state.get("_sal_spv_file", sal_manual*0.97 if sal_manual else 0.19)),
                step=0.0001, format="%.4f", key="_sal_spv_file")
            sal_dec_pct = st.number_input("Decremento vs máx. histórico (%)", 
                value=float(st.session_state.get("_sal_dec_file", float(datos.get("salinidad_dec_pct", 0.0) or 0.0) if data_ready else 0.0)),
                step=0.01, key="_sal_dec_file")
            if data_ready:
                datos["salinidad_spc"] = sal_spc_inp
                datos["salinidad_spv"] = sal_spv_inp
                datos["salinidad_dec_pct"] = sal_dec_pct
                datos["salinidad_source"] = "archivo"
                datos["sal_series"] = sal_series_ext if sal_series_ext is not None else pd.Series(dtype=float)

        st.markdown("""
        <div class='info-box' style='border-radius:8px;padding:12px;font-size:.8rem;margin-top:12px'>
        <b>Clasificación IHO:</b><br>
        🟢 &lt; 0.5 ups — Agua dulce<br>
        🟡 0.5–3.0 ups — Ligeramente salobre<br>
        🟠 3.0–10 ups — Salobre<br>
        🔴 &gt; 10 ups — Salino
        </div>""", unsafe_allow_html=True)

    with col_s2:
        spc_plot = sal_spc_inp if isinstance(sal_spc_inp, float) else (sal_spc_val or 0.0)
        fig_sal = go.Figure(go.Indicator(
            mode="gauge+number+delta",
            value=spc_plot,
            number={"suffix":" ups", "font":{"size":28, "color":"#002a4d"}},
            delta={"reference":0.5, "decreasing":{"color":"#15803d"},
                   "increasing":{"color":"#b91c1c"}, "suffix":" vs límite"},
            title={"text":"Salinidad SPC<br><span style='font-size:.75em'>Cauce de Navegación</span>",
                   "font":{"size":15}},
            gauge={
                "axis":{"range":[0, 1.5], "tickwidth":1, "tickcolor":"#94a3b8"},
                "bar":{"color":"#0072b8","thickness":0.25},
                "bgcolor":"white",
                "steps":[
                    {"range":[0, 0.5],  "color":"#dcfce7"},
                    {"range":[0.5, 3.0],"color":"#fef9c3"},
                ],
                "threshold":{"line":{"color":"#dc2626","width":3},
                             "thickness":0.8,"value":0.5}
            }))
        fig_sal.update_layout(plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
            height=320, margin=dict(t=40,b=20,l=30,r=30))
        st.plotly_chart(fig_sal, use_container_width=True)

    # Serie temporal de salinidad (solo si proviene de archivo externo)
    sal_series = datos.get("sal_series") if data_ready else None
    if sal_series is not None and len(sal_series) > 0:
        st.markdown("---")
        st.markdown("##### Evolución diaria de salinidad en el período")
        fig_sal_ts = go.Figure(go.Scatter(
            y=sal_series.values, mode="lines+markers",
            line=dict(color="#0072b8", width=2),
            marker=dict(size=6, color="#003a6e")))
        fig_sal_ts.add_hline(y=0.5, line_dash="dash", line_color="#dc2626",
            annotation_text="Límite agua dulce (0.5 ups)")
        fig_sal_ts.update_layout(height=260, plot_bgcolor="rgba(0,0,0,0)",
            paper_bgcolor="rgba(0,0,0,0)", margin=dict(t=10,b=10,l=10,r=10),
            yaxis=dict(title="ups", gridcolor="rgba(148,163,184,.2)"),
            xaxis=dict(title="Día del período"))
        st.plotly_chart(fig_sal_ts, use_container_width=True)
    else:
        st.markdown("""
        <div class='info-box'>
        ✍️ No hay serie temporal cargada. La salinidad del período se está manejando manualmente en esta pestaña.
        </div>""", unsafe_allow_html=True)

    # ── CCA — Conservación de Calidad del Agua ──────────────────────────────
    st.markdown("---")
    st.markdown('<div class="sec-hdr">💧 CCA — Conservación de Calidad del Agua</div>', unsafe_allow_html=True)
    _cca_default7 = float(st.session_state.get("rep_cca_hm3", 0.0) or 0.0)
    _tneo_tab7    = D("transitos_neopanamax", 0)
    cc1, cc2, cc3 = st.columns(3)
    with cc1:
        _cca_tab7 = st.number_input(
            "CCA / ZZFlush (hm³)",
            value=_cca_default7,
            min_value=0.0, step=0.01, format="%.2f",
            key="rep_cca_hm3",
            help="Conservación de Calidad del Agua = ZZFlush — volumen invertido en esclusas NPX para mitigar intrusión salina.")
    _cca_unit7 = (_cca_tab7 / _tneo_tab7) if _tneo_tab7 > 0 else 0.0
    cc2.metric("Tránsitos Neopanamax", f"{int(_tneo_tab7)}")
    cc3.metric("CCA unitario",
        f"{_cca_unit7:.4f} hm³/tráns." if _tneo_tab7 else "—")

    # ── Tabla de alerta operacional ──────────────────────────────────────────
    st.markdown("---")
    st.markdown('<div class="sec-hdr">⚠️ Tabla de Alerta Operacional — Proyección de Salinidad</div>', unsafe_allow_html=True)

    _spc_alert7 = sal_spc_inp if isinstance(sal_spc_inp, float) else (sal_spc_val or 0.0)

    def _sal_alert_level(spc):
        if   spc < 0.44: return "Normal",                       "#16a34a", "🟢"
        elif spc < 0.52: return "Vigilancia",                   "#d97706", "🟡"
        elif spc < 0.54: return "Pre-alerta",                   "#ea580c", "🟠"
        elif spc < 0.56: return "Preparar lavado / suspensión", "#dc2626", "🔴"
        elif spc < 0.58: return "Alerta — lavado urgente",      "#b91c1c", "🔴"
        else:            return "Crítico",                      "#7f1d1d", "🚨"

    _alrt_label, _alrt_color, _alrt_emoji = _sal_alert_level(_spc_alert7)

    st.markdown(f"""
    <div style='background:{_alrt_color}22;border:2px solid {_alrt_color};border-radius:12px;
    padding:14px 20px;margin-bottom:14px;text-align:center'>
    <span style='font-size:1.5rem'>{_alrt_emoji}</span>
    <span style='font-size:1.1rem;font-weight:800;color:{_alrt_color};margin-left:10px'>
    Estado actual: {_alrt_label}</span>
    <span style='font-size:.85rem;color:var(--c-text-muted);margin-left:12px'>
    SPC = {_spc_alert7:.4f} ups</span>
    </div>""", unsafe_allow_html=True)

    st.markdown("""
    <table class='styled-table' style='width:100%'>
    <thead><tr>
      <th>Mediana diaria</th><th>P(MAX &ge; 0.60)</th>
      <th>P(MAX &ge; 0.55)</th><th>Acción recomendada</th>
    </tr></thead>
    <tbody>
      <tr><td>&lt; 0.44 ups</td><td>0%</td><td>0%</td><td>Normal</td></tr>
      <tr><td>0.44&ndash;0.52 ups</td><td>0&ndash;9%</td><td>0&ndash;72%</td><td>Vigilancia</td></tr>
      <tr style='background:rgba(234,88,12,.13);font-weight:700;font-style:italic'>
        <td>0.52 ups</td><td>9%</td><td>72%</td>
        <td style='color:#ea580c'>Pre-alerta</td></tr>
      <tr style='background:rgba(220,38,38,.13);font-weight:700;font-style:italic'>
        <td>0.54 ups</td><td>20%</td><td>93%</td>
        <td style='color:#dc2626'><u>Preparar lavado</u> o <u>suspensi&oacute;n</u></td></tr>
      <tr style='background:rgba(185,28,28,.15);font-weight:700;font-style:italic'>
        <td>0.56 ups</td><td>56%</td><td>100%</td>
        <td style='color:#b91c1c'><u>Alerta</u> &mdash; <u>lavado urgente</u></td></tr>
      <tr style='background:rgba(127,29,29,.17);font-weight:700'>
        <td>&ge; 0.58 ups</td><td>92%</td><td>100%</td>
        <td style='color:#7f1d1d'><b>Cr&iacute;tico</b></td></tr>
    </tbody>
    </table>
    <p style='font-size:.72rem;font-style:italic;color:var(--c-text-muted);margin-top:6px'>
    El umbral de alerta operacional en la proyecci&oacute;n (promedio/mediana diaria) es <b>0.52 ups</b>.
    </p>""", unsafe_allow_html=True)

    # Viñeta narrativa salinidad
    # Viñeta narrativa salinidad
    st.markdown("---")
    st.markdown("#### 📝 Viñeta auto-generada — Salinidad")
    _spc_used = sal_spc_inp if isinstance(sal_spc_inp, float) else (sal_spc_val or 0.0)
    _spv_used = sal_spv_inp if isinstance(sal_spv_inp, (int,float)) else 0.0
    _dec_used = sal_dec_pct if isinstance(sal_dec_pct, (int,float)) else 0.0
    _spc_from_daily = False

    def _sal_val(val, empty_cond, fmt=".4f"):
        bg = "background:#fed7aa;padding:1px 5px;border-radius:3px;font-weight:700" if empty_cond else "font-weight:700"
        return f"<span style='{bg}'>{format(val, fmt) if isinstance(val, (int,float)) else val}</span>"

    bullet_sal = (
        f"Durante {mes_sel.lower()} de {anio_sel}, la salinidad promedio diaria en el cauce de navegación "
        f"(SPC) del embalse Gatún fue en promedio "
        f"{_sal_val(_spc_used, _spc_used == 0.0)} ups. "
        f"Al cierre del mes, la salinidad promedio ponderada (SPV) se registró en "
        f"{_sal_val(_spv_used, _spv_used == 0.0)} ups, representando un decremento del "
        f"{_sal_val(_dec_used, _dec_used == 0.0, '.1f')}% respecto al máximo histórico alcanzado."
    )
    st.markdown(f"<div class='info-box' style='font-size:.88rem;line-height:1.75'>• {bullet_sal}</div>",
                unsafe_allow_html=True)

    # Auditoría combinada
    st.markdown("---")
    st.markdown("#### 🧮 Estado de campos — Salinidad")
    _sal_missing_list = []
    if _spv_used == 0.0: _sal_missing_list.append("SPV al cierre")
    if _dec_used == 0.0: _sal_missing_list.append("Decremento %")
    if _spc_used == 0.0: _sal_missing_list.append("SPC promedio")
    all_missing = _met_missing + _sal_missing_list
    if all_missing:
        st.markdown(f"<div class='warn-box'>⚠️ Pendiente: <b>{', '.join(all_missing)}</b> — en naranja en el informe.</div>", unsafe_allow_html=True)
    else:
        st.markdown("<div class='ok-box'>✅ Todos los campos de meteorología y salinidad están completos.</div>", unsafe_allow_html=True)

    def _est(val, auto=False):
        if auto:       return "Auto ✅"
        if val and str(val) not in ("0", "0.0", "—", ""): return "Manual ✅"
        return "⚠️ Falta"

    df_audit_ms = pd.DataFrame([
        ("Precipitación total (mm)",     met_precip_mm,              _est(met_precip_mm)),
        ("Promedio histórico (mm)",       met_precip_hist,            _est(met_precip_hist)),
        ("Superávit/Déficit (mm)",        round(met_superavit, 1),    _est(met_superavit, auto=True)),
        ("% vs promedio histórico",       round(met_pct_sobre, 1),    _est(met_pct_sobre, auto=True)),
        ("Período histórico",             met_periodo or "—",         _est(met_periodo)),
        ("Años de registro",              met_n_anios or "—",         _est(met_n_anios)),
        ("Posición ranking",              int(met_posicion) if met_posicion else "—", _est(met_posicion)),
        ("SPC promedio (ups)",            round(_spc_used, 4),        _est(_spc_used)),
        ("SPV al cierre (ups)",           round(_spv_used, 4),        _est(_spv_used)),
        ("Decremento vs máx. hist. (%)",  round(_dec_used, 1),        _est(_dec_used)),
    ], columns=["Variable", "Valor", "Estado"])
    st.dataframe(df_audit_ms, use_container_width=True, hide_index=True)


# ═══════════════════════════════════════════════════════════════
#  TAB 8 — COMPARATIVA HISTÓRICA
# ═══════════════════════════════════════════════════════════════
with tabs[9]:
    st.markdown('<div class="sec-hdr">💧 Evaporación</div>', unsafe_allow_html=True)
    if not data_ready:
        st.info("Carga el archivo DAILY_AND_ACUMMEN para ver la evaporación del período.")
    else:
        _d = n_dias if n_dias and n_dias > 0 else 31
        ev_total_hm3 = float(D("evaporacion_hm3", 0.0) or 0.0)
        ev_total_m3s = float(D("evaporacion_m3s", 0.0) or 0.0)
        ev_total_cfs = hm3_to_cfs_period(ev_total_hm3, _d) if _d else 0.0
        ev_total_mm  = float(D("evaporacion_diaria_mm", 0.0) or 0.0)

        ev_g_hm3 = float(D("evap_gatun_hm3", 0.0) or 0.0)
        ev_g_m3s = hm3_to_m3s_period(ev_g_hm3, _d) if _d else 0.0
        ev_g_cfs = hm3_to_cfs_period(ev_g_hm3, _d) if _d else 0.0
        ev_g_hm3d = float(D("evap_gatun_hm3_dia", 0.0) or 0.0)
        ev_g_mmd = float(D("evap_gatun_mm_dia", 0.0) or 0.0)
        sup_g = float(D("sup_prom_gatun_km2", 0.0) or 0.0)

        ev_a_hm3 = float(D("evap_alh_hm3", 0.0) or 0.0)
        ev_a_m3s = hm3_to_m3s_period(ev_a_hm3, _d) if _d else 0.0
        ev_a_cfs = hm3_to_cfs_period(ev_a_hm3, _d) if _d else 0.0
        ev_a_hm3d = float(D("evap_alh_hm3_dia", 0.0) or 0.0)
        ev_a_mmd = float(D("evap_alh_mm_dia", 0.0) or 0.0)
        sup_a = float(D("sup_prom_alh_km2", 0.0) or 0.0)

        k1, k2, k3, k4 = st.columns(4)
        k1.metric("Evaporación total", f"{ev_total_hm3:.2f} hm³", f"{ev_total_m3s:.2f} m³/s")
        k2.metric("Evaporación diaria total", f"{ev_total_mm:.2f} mm/día", f"{ev_total_cfs:.1f} pies³/s")
        k3.metric("Sup. promedio Gatún", f"{sup_g:.2f} km²", f"{ev_g_hm3d:.2f} hm³/día")
        k4.metric("Sup. promedio Alhajuela", f"{sup_a:.2f} km²", f"{ev_a_hm3d:.2f} hm³/día")

        st.markdown("#### Detalle de evaporación por embalse")
        df_ev = pd.DataFrame([
            {"Concepto": "Evaporación total", "hm³": ev_total_hm3, "m³/s": ev_total_m3s, "pies³/s": ev_total_cfs, "mm/día": ev_total_mm, "Fuente": "DAILY · volúmenes + áreas promedio"},
            {"Concepto": "Evaporación Gatún", "hm³": ev_g_hm3, "m³/s": ev_g_m3s, "pies³/s": ev_g_cfs, "mm/día": ev_g_mmd, "Fuente": "DAILY · Volumen evaporado embalse Gatún, hm3"},
            {"Concepto": "Evaporación Alhajuela", "hm³": ev_a_hm3, "m³/s": ev_a_m3s, "pies³/s": ev_a_cfs, "mm/día": ev_a_mmd, "Fuente": "DAILY · Volumen evaporado embalse Alhajuela, hm3"},
            {"Concepto": "Evap. Gatún por día", "hm³": ev_g_hm3d, "m³/s": hm3_to_m3s_period(ev_g_hm3d, 1), "pies³/s": hm3_to_cfs_period(ev_g_hm3d, 1), "mm/día": ev_g_mmd, "Fuente": "hm³/día = hm³ del período / días"},
            {"Concepto": "Evap. Alhajuela por día", "hm³": ev_a_hm3d, "m³/s": hm3_to_m3s_period(ev_a_hm3d, 1), "pies³/s": hm3_to_cfs_period(ev_a_hm3d, 1), "mm/día": ev_a_mmd, "Fuente": "hm³/día = hm³ del período / días"},
        ])
        st.dataframe(df_ev.style.format({"hm³":"{:.3f}","m³/s":"{:.3f}","pies³/s":"{:.1f}","mm/día":"{:.3f}"}), use_container_width=True, hide_index=True)

        st.markdown("#### Fórmulas usadas")
        st.markdown(f"""
        <div class='info-box'>
        <b>Evaporación mensual (hm³)</b> = Evap. Gatún + Evap. Alhajuela<br>
        <b>Evaporación diaria total (mm/día)</b> = Evaporación mensual × 1000 / ((Sup. promedio Gatún + Sup. promedio Alhajuela) × días)<br>
        <b>Evaporación Gatún (hm³/día)</b> = Evap. Gatún / días<br>
        <b>Evaporación Gatún (mm/día)</b> = Evap. Gatún × 1000 / (Sup. promedio Gatún × días)<br>
        <b>Evaporación Alhajuela (hm³/día)</b> = Evap. Alhajuela / días<br>
        <b>Evaporación Alhajuela (mm/día)</b> = Evap. Alhajuela × 1000 / (Sup. promedio Alhajuela × días)
        </div>
        """, unsafe_allow_html=True)

        if hidro_file_rows1:
            _ev_file = pd.DataFrame([{
                "Fila": r["row"], "Variable": r["label"], "Valor 1": r["value1"], "Unidad 1": r["unit1"],
                "Valor 2": r["value2"], "Unidad 2": r["unit2"], "Valor 3": r["value3"], "Unidad 3": r["unit3"],
            } for r in hidro_file_rows1 if "Evaporación" in str(r.get("label","")) or "Sup. Promedio" in str(r.get("label","")) or "Elev. Prom." in str(r.get("label",""))])
            if not _ev_file.empty:
                st.markdown("#### Referencia desde archivo de Hidroestadística detectado")
                st.dataframe(_ev_file, use_container_width=True, hide_index=True)

# ═══════════════════════════════════════════════════════════════
#  TAB 10 — HIDROESTADÍSTICA
# ═══════════════════════════════════════════════════════════════
with tabs[10]:
    st.markdown('<div class="sec-hdr">📈 Comparativa Histórica</div>', unsafe_allow_html=True)

    # ── Diagnóstico de archivos históricos ─────────────────────────────────
    if (f_hist_chcp_src or f_hist_netfl_src or f_hist_alh_src or f_hist_gat_src) and not hist:
        with st.expander("🔍 Diagnóstico — archivos históricos cargados pero sin datos", expanded=True):
            st.markdown("""
            <div class='warn-box'>
            ⚠️ Se detectaron archivos históricos pero no se pudieron calcular promedios.<br>
            Posibles causas:
            <ul>
            <li>El archivo no tiene fila de encabezado con columnas <b>AÑO, OCT, NOV, DIC, ENE, FEB, MAR, ABR, MAY, JUN, JUL, AGO, SEP</b></li>
            <li>La columna <b>AÑO</b> tiene valores no numéricos o está en otra posición</li>
            <li>Los valores de caudal están en una unidad o formato distinto al esperado (m³/s)</li>
            </ul>
            Usa la sección de <b>Valores manuales</b> abajo para ingresar los promedios históricos directamente.
            </div>""", unsafe_allow_html=True)
            # Mostrar preview del archivo para diagnóstico
            if f_hist_chcp_src:
                try:
                    for h in [0, 1, 2, 3]:
                        df_preview = pd.read_excel(f_hist_chcp_src, header=h, nrows=5)
                        st.caption(f"Preview CHCP (header={h}): {list(df_preview.columns[:8])}")
                        st.dataframe(df_preview.head(3), use_container_width=True)
                        break
                except Exception as e:
                    st.error(f"Error al previsualizar: {e}")

    n_d = n_dias if data_ready and n_dias > 0 else 31

    # ── Valores históricos de referencia ───────────────────────────────────
    _hist_from_file = bool(hist)
    _h_chcp_hm3_def = round(float(hist.get("hist_chcp_hm3") or hist.get("hist_chcp_m3s", 0) * n_d * HM3_PER_M3S_DAY), 2) if hist else 0.0
    _h_chcp_m3s_def = float(hist.get("hist_chcp_m3s", 0.0)) if hist else 0.0
    _h_alh_hm3_def  = round(float(hist.get("hist_alh_hm3")  or hist.get("hist_alh_m3s",  0) * n_d * HM3_PER_M3S_DAY), 2) if hist else 0.0
    _h_gat_hm3_def  = round(float(hist.get("hist_gat_hm3")  or hist.get("hist_gat_m3s",  0) * n_d * HM3_PER_M3S_DAY), 2) if hist else 0.0
    _n_years_def    = int(hist.get("n_years_chcp", 0)) if hist else 0
    _years_str_def  = (f"{hist['years_chcp'][0]}-{hist['years_chcp'][1]}" if hist and hist.get("years_chcp") else "")

    # CRÍTICO: Streamlit ignora value= si el key ya existe en session_state (aunque sea 0).
    # Forzar session_state antes de renderizar widgets para que la tabla siempre
    # muestre los valores del archivo histórico.
    if _hist_from_file:
        if _h_chcp_hm3_def > 0: st.session_state["h_chcp_hm3"] = _h_chcp_hm3_def
        if _h_chcp_m3s_def > 0: st.session_state["h_chcp_m3s"] = _h_chcp_m3s_def
        if _h_alh_hm3_def  > 0: st.session_state["h_alh_hm3"]  = _h_alh_hm3_def
        if _h_gat_hm3_def  > 0: st.session_state["h_gat_hm3"]  = _h_gat_hm3_def
        if _n_years_def    > 0: st.session_state["h_n_years"]   = _n_years_def
        if _years_str_def:       st.session_state["h_years_str"] = _years_str_def

    _src_label = "📂 Desde archivos históricos" if _hist_from_file else "✏️ Ingreso manual — se marcará con ★ en el informe"
    _src_color = "ok-box" if _hist_from_file else "warn-box"
    st.markdown(f"<div class='{_src_color}'>Fuente de promedios históricos: <b>{_src_label}</b></div>", unsafe_allow_html=True)

    with st.expander("📊 Valores históricos de referencia" + ("  ← edita aquí si no tienes archivos" if not _hist_from_file else ""), expanded=not _hist_from_file):
        hv_c1, hv_c2, hv_c3, hv_c4 = st.columns(4)
        with hv_c1:
            h_chcp_hm3 = st.number_input("Prom. hist. CHCP (hm³)", value=_h_chcp_hm3_def, step=0.1, format="%.2f", key="h_chcp_hm3",
                                          help="Promedio histórico del aporte total CHCP para este período (en hm³)")
            h_chcp_m3s = st.number_input("Prom. hist. CHCP (m³/s)", value=_h_chcp_m3s_def, step=0.1, format="%.2f", key="h_chcp_m3s")
        with hv_c2:
            h_alh_hm3  = st.number_input("Prom. hist. Alhajuela (hm³)", value=_h_alh_hm3_def, step=0.1, format="%.2f", key="h_alh_hm3")
            h_gat_hm3  = st.number_input("Prom. hist. Gatún (hm³)",     value=_h_gat_hm3_def, step=0.1, format="%.2f", key="h_gat_hm3")
        with hv_c3:
            h_n_years  = st.number_input("N° años registro", value=_n_years_def, step=1, key="h_n_years")
            h_years_str= st.text_input("Período histórico (ej: 1950-2026)", value=_years_str_def, key="h_years_str")
        with hv_c4:
            h_precip_hist = st.number_input("Precip. hist. (mm)", value=float(st.session_state.get("x_phist", 0.0)), step=0.1, key="h_precip_hist_tab7",
                                             help="Precipitación promedio histórica para este mes/período")
            h_pos_ranking = st.number_input("Posición en ranking (húmedo→seco)", value=int(st.session_state.get("x_ppos", 0)), step=1, key="h_pos_ranking_tab7")

        # Nota: x_phist, x_pper, x_hist_hm3, x_hist_m3s son widgets en Tab Exportar —
        # no se asignan aquí para evitar conflicto con Streamlit session_state.
        # Los valores de hist.get() ya se propagan al dict 'hist' que usa Tab 9.
        pass

        _raw_cls = hist.get("series_chcp_raw") if hist else None
        if isinstance(_raw_cls, pd.DataFrame) and not _raw_cls.empty and "Clas. Aporte" in _raw_cls.columns:
            with st.expander("🗂️ Clasificación anual del 03_QAVGMCHCP (columna R)", expanded=False):
                _show_cols = [c for c in ["AÑO", "ANUAL", "hm3", "Dif%", "%ACUM.", "Clas. Aporte"] if c in _raw_cls.columns]
                st.caption("Tabla filtrada con el rango histórico activo. La clasificación proviene de la columna R del archivo histórico.")
                st.dataframe(_raw_cls[_show_cols], use_container_width=True, hide_index=True)

    # Indicador de si los valores son manuales (para DOCX con alertas)
    _hist_is_manual = not _hist_from_file or (h_chcp_hm3 != _h_chcp_hm3_def) or (h_chcp_m3s != _h_chcp_m3s_def)
    if _hist_is_manual and not _hist_from_file:
        st.markdown("<div class='warn-box'>✏️ Los valores históricos son ingresados manualmente — se solo se resaltarán en <b style='background:#fed7aa;padding:2px 5px;border-radius:3px'>naranja</b> cuando el valor sea 0 o haya incongruencias en el informe exportado.</div>", unsafe_allow_html=True)

    if not data_ready:
        st.info("Carga el archivo DAILY_AND_ACUMMEN para ver el análisis comparativo.")
    else:
        pct_vs_hist = ((D("aporte_total_hm3") - h_chcp_hm3) / h_chcp_hm3 * 100) if h_chcp_hm3 > 0 else 0

        # ── KPIs comparativos del DAILY ──────────────────────────────────
        st.markdown("---")
        st.markdown(f"#### 📊 {label_per} — Datos del DAILY vs promedio histórico")
        kc1, kc2, kc3, kc4 = st.columns(4)
        def _kpi_c(col, lbl, val, hist_val, unidad="hm³"):
            if hist_val and hist_val > 0:
                delta = (val - hist_val) / hist_val * 100
                arrow = "▲" if delta >= 0 else "▼"
                color = "#16a34a" if delta >= 0 else "#dc2626"
                sub = f"<div style='color:{color};font-size:.74rem;font-weight:700'>{arrow} {abs(delta):.1f}% vs hist. ({hist_val:,.1f} {unidad})</div>"
            else:
                sub = "<div style='color:#64748b;font-size:.72rem'>— sin histórico (ingresa en panel abajo)</div>"
            col.markdown(f"""
            <div class='kpi'>
              <div class='kpi-label'>{lbl}</div>
              <div class='kpi-val'>{val:,.2f} {unidad}</div>
              {sub}
            </div>""", unsafe_allow_html=True)
        _kpi_c(kc1, "Aporte Total CHCP",   D("aporte_total_hm3"), h_chcp_hm3)
        _kpi_c(kc2, "Subcuenca Alhajuela",  D("alhajuela_hm3"),    h_alh_hm3)
        _kpi_c(kc3, "Subcuenca Gatún",      D("gatun_hm3"),        h_gat_hm3)
        _kpi_c(kc4, "Aporte Neto",          D("aporte_neto_hm3"),  max(h_chcp_hm3 - D("evaporacion_hm3"), 0))

        # ── Sección GATNETFL — Aportes Netos históricos ──────────────────────
        _has_netfl = bool(hist.get("hist_neto_m3s", 0) or hist.get("hist_neto_hm3", 0))
        _netfl_years  = hist.get("years_netfl")
        _netfl_n      = int(hist.get("n_years_netfl", 0) or 0)
        _netfl_yr_str = f"{_netfl_years[0]}–{_netfl_years[1]}" if _netfl_years else ""
        _s_netfl      = hist.get("series_netfl") if hist else None
        _h_neto_hm3   = fn_hist_neto_hm3
        _h_neto_m3s   = fn_hist_neto_m3s

        st.markdown("---")
        if _has_netfl:
            st.markdown(f"#### 📉 GATNETFL — Aportes Netos históricos CHCP"
                        f"{'  ·  ' + _netfl_yr_str + f' ({_netfl_n} años)' if _netfl_yr_str else ''}")
            nc1, nc2, nc3, nc4 = st.columns(4)
            _pct_neto = ((D('aporte_neto_hm3') - _h_neto_hm3) / _h_neto_hm3 * 100) if _h_neto_hm3 > 0 else 0
            _pct_neto_txt, _pct_neto_dir = _pct_rel_text(_pct_neto, positive_text="por arriba del", negative_text="por debajo del", zero_text="igual al", decimals=0)
            _net_pos = str(datos.get('hist_neto_scale_pos', '') or '')
            _net_cls = str(datos.get('hist_neto_scale_class', '') or '')
            nc1.markdown(f"""<div class='kpi'>
              <div class='kpi-label'>Aporte Neto Actual</div>
              <div class='kpi-val'>{D('aporte_neto_hm3'):,.2f} hm³</div>
              <div style='font-size:.74rem'>{D('aporte_neto_m3s'):.2f} m³/s</div>
            </div>""", unsafe_allow_html=True)
            nc2.markdown(f"""<div class='kpi'>
              <div class='kpi-label'>Prom. Hist. GATNETFL</div>
              <div class='kpi-val'>{_h_neto_hm3:,.2f} hm³</div>
              <div style='font-size:.74rem'>{_h_neto_m3s:.2f} m³/s</div>
            </div>""", unsafe_allow_html=True)
            _col_neto = "#16a34a" if _pct_neto >= 0 else "#dc2626"
            nc3.markdown(f"""<div class='kpi'>
              <div class='kpi-label'>% vs histórico neto</div>
              <div class='kpi-val' style='color:{_col_neto}'>{abs(_pct_neto):.0f}%</div>
              <div style='font-size:.74rem'>{_pct_neto_dir}</div>
            </div>""", unsafe_allow_html=True)
            nc4.markdown(f"""<div class='kpi'>
              <div class='kpi-label'>Escala húmedo→seco</div>
              <div class='kpi-val'>{_net_pos if _net_pos else '—'}</div>
              <div style='font-size:.74rem'>{_net_cls if _net_cls else 'Sin clasificación'}</div>
            </div>""", unsafe_allow_html=True)

            st.markdown(f"<div class='info-box'>• Los aportes netos del período actual se ubican en la posición <b>{_net_pos if _net_pos else '—'}</b> dentro de la escala histórica de húmedo a seco{(' (' + _net_cls + ')') if _net_cls else ''}, y están <b>{_pct_neto_txt}% {_pct_neto_dir}</b> promedio histórico de <b>{_h_neto_hm3:.1f} hm³ ({_h_neto_m3s:.2f} m³/s)</b>.</div>", unsafe_allow_html=True)

            if _s_netfl is not None and not _s_netfl.empty and "valor" in _s_netfl.columns:
                fig_netfl = go.Figure()
                fig_netfl.add_trace(go.Scatter(
                    x=_s_netfl["AÑO"].astype(int), y=_s_netfl["valor"].astype(float),
                    mode="lines+markers", name="Aporte Neto histórico",
                    line=dict(color="#0072b8", width=1.5), marker=dict(size=4)))
                if _h_neto_m3s > 0:
                    fig_netfl.add_hline(y=_h_neto_m3s, line_dash="dash",
                        line_color="#64748b", annotation_text=f"Prom={_h_neto_m3s:.1f} m³/s")
                if data_ready:
                    fig_netfl.add_scatter(
                        x=[anio_sel], y=[D("aporte_neto_m3s")],
                        mode="markers", name="Período actual",
                        marker=dict(color="#e55c00", size=12, symbol="star"))
                fig_netfl.update_layout(
                    height=300, title="Serie histórica — Aportes Netos CHCP (GATNETFL)",
                    plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                    margin=dict(t=30,b=10,l=10,r=10),
                    yaxis=dict(title="m³/s", gridcolor="rgba(148,163,184,.2)"),
                    xaxis=dict(title="Año"), legend=dict(orientation="h", y=1.15))
                st.plotly_chart(fig_netfl, use_container_width=True)
        else:
            st.markdown("""<div class='info-box'>
            📉 <b>GATNETFL no cargado.</b> Sube el archivo <code>04_GATNETFL.xls</code> en el panel lateral
            para ver la comparativa histórica de aportes netos y calcular automáticamente el promedio histórico neto.
            </div>""", unsafe_allow_html=True)
            # Entrada manual de neto histórico
            with st.expander("✏️ Ingresar promedio histórico neto manualmente"):
                _h_neto_hm3_inp = st.number_input("Prom. hist. Aporte Neto (hm³)", value=_h_neto_hm3, step=0.1, format="%.2f", key="h_neto_hm3_manual")
                _h_neto_m3s_inp = st.number_input("Prom. hist. Aporte Neto (m³/s)", value=_h_neto_m3s, step=0.01, format="%.3f", key="h_neto_m3s_manual")
                if _h_neto_hm3_inp != _h_neto_hm3:
                    fn_hist_neto_hm3 = _h_neto_hm3_inp
                if _h_neto_m3s_inp != _h_neto_m3s:
                    fn_hist_neto_m3s = _h_neto_m3s_inp

        # ── Gráfico comparativo ───────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        col_h1, col_h2 = st.columns(2)

        with col_h1:
            st.markdown(f"##### Actual vs Promedio Histórico")
            cats_c = ["CHCP Total", "Sub. Alhajuela", "Sub. Gatún"]
            act_c  = [D("aporte_total_hm3"), D("alhajuela_hm3"), D("gatun_hm3")]
            his_c  = [h_chcp_hm3, h_alh_hm3, h_gat_hm3]
            _hist_lbl = f"Prom. Hist. ({h_years_str})" if h_years_str else "Prom. Histórico"
            if _hist_is_manual:
                _hist_lbl += " ✏️"
            fig_comp = go.Figure()
            fig_comp.add_trace(go.Bar(name="Actual (DAILY)", x=cats_c, y=act_c,
                marker_color="#0072b8", text=[f"{v:.1f}" for v in act_c], textposition="auto",
                textfont=dict(color="white", size=11)))
            fig_comp.add_trace(go.Bar(name=_hist_lbl, x=cats_c, y=his_c,
                marker_color="#94a3b8" if _hist_from_file else "#fbbf24",
                text=[f"{v:.1f}" for v in his_c], textposition="auto",
                textfont=dict(size=11)))
            fig_comp.update_layout(barmode="group", height=320,
                plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)",
                margin=dict(t=10, b=10, l=10, r=10),
                legend=dict(orientation="h", y=1.1),
                yaxis=dict(title="hm³", gridcolor="rgba(148,163,184,.2)"))
            st.plotly_chart(fig_comp, use_container_width=True)

        with col_h2:
            # Serie histórica si hay datos de archivos
            s_chcp = hist.get("series_chcp") if hist else None
            if s_chcp is not None and not s_chcp.empty and "valor" in s_chcp.columns:
                st.markdown("##### Serie histórica — CHCP")
                fig_hist = go.Figure()
                fig_hist.add_trace(go.Scatter(
                    x=s_chcp["AÑO"].astype(int), y=s_chcp["valor"].astype(float),
                    mode="lines", name="CHCP histórico",
                    line=dict(color="#94a3b8", width=1.5)))
                if h_chcp_m3s > 0:
                    fig_hist.add_hline(y=h_chcp_m3s, line_dash="dash",
                        line_color="#0072b8", annotation_text=f"Prom={h_chcp_m3s:.1f} m³/s")
                if data_ready:
                    fig_hist.add_scatter(
                        x=[anio_sel], y=[D("aporte_total_m3s")],
                        mode="markers", name="Período actual",
                        marker=dict(color="#e55c00", size=12, symbol="star"))
                fig_hist.update_layout(height=320, plot_bgcolor="rgba(0,0,0,0)",
                    paper_bgcolor="rgba(0,0,0,0)", margin=dict(t=10, b=10, l=10, r=10),
                    yaxis=dict(title="m³/s", gridcolor="rgba(148,163,184,.2)"),
                    xaxis=dict(title="Año"), legend=dict(orientation="h", y=1.1))
                st.plotly_chart(fig_hist, use_container_width=True)
            else:
                # Sin serie histórica: mostrar gauge vs histórico
                st.markdown("##### % vs promedio histórico")
                gauges_data = [
                    ("CHCP Total",   D("aporte_total_hm3"), h_chcp_hm3),
                    ("Alhajuela",    D("alhajuela_hm3"),    h_alh_hm3),
                    ("Gatún",        D("gatun_hm3"),        h_gat_hm3),
                ]
                for _lbl, _act, _hist_v in gauges_data:
                    if _hist_v > 0:
                        _pct = _act / _hist_v * 100
                        _color = "#16a34a" if _pct >= 100 else "#f59e0b" if _pct >= 80 else "#dc2626"
                        st.markdown(f"""
                        <div style='margin:6px 0;padding:8px 12px;border-radius:8px;border:1px solid rgba(148,163,184,.3);background:var(--c-surface)'>
                          <div style='display:flex;justify-content:space-between;align-items:center;margin-bottom:4px'>
                            <span style='font-size:.78rem;font-weight:600;color:var(--c-text)'>{_lbl}</span>
                            <span style='font-size:.78rem;color:{_color};font-weight:700'>{_pct:.0f}% del hist.</span>
                          </div>
                          <div style='background:rgba(148,163,184,.2);border-radius:4px;height:8px;overflow:hidden'>
                            <div style='background:{_color};height:100%;width:{min(_pct,150):.0f}%;border-radius:4px;transition:width .3s'></div>
                          </div>
                          <div style='font-size:.72rem;color:var(--c-text-muted);margin-top:3px'>Actual: {_act:.1f} hm³ · Histórico: {_hist_v:.1f} hm³</div>
                        </div>""", unsafe_allow_html=True)

        # ── Tabla resumen comparativa ─────────────────────────────────────
        st.markdown("---")
        st.markdown("##### Tabla resumen comparativa")
        if _hist_is_manual:
            st.markdown("<div class='warn-box' style='font-size:.78rem'>✏️ Los valores históricos son manuales</div>", unsafe_allow_html=True)
        pct_alh = ((D("alhajuela_hm3") - h_alh_hm3) / h_alh_hm3 * 100) if h_alh_hm3 > 0 else 0
        pct_gat = ((D("gatun_hm3")    - h_gat_hm3)  / h_gat_hm3  * 100) if h_gat_hm3  > 0 else 0
        df_cmp = pd.DataFrame({
            "Subcuenca":            ["CHCP Total", "Alhajuela", "Gatún"],
            "Actual DAILY (hm³)":   [round(D("aporte_total_hm3"),2), round(D("alhajuela_hm3"),2), round(D("gatun_hm3"),2)],
            "Actual DAILY (m³/s)":  [round(D("aporte_total_m3s"),2), "—", "—"],
            f"Prom. Hist. (hm³){' ✏️' if _hist_is_manual else ''}": [h_chcp_hm3, h_alh_hm3, h_gat_hm3],
            "% vs histórico":       [f"{pct_vs_hist:+.1f}%", f"{pct_alh:+.1f}%", f"{pct_gat:+.1f}%"],
            "Años registro":        [h_n_years, "—", "—"],
        })
        st.dataframe(df_cmp, use_container_width=True, hide_index=True)

        # Propagar a hist para el resto del app
        if not hist:
            hist["hist_chcp_hm3"]  = h_chcp_hm3
            hist["hist_chcp_m3s"]  = h_chcp_m3s
            hist["hist_alh_hm3"]   = h_alh_hm3
            hist["hist_gat_hm3"]   = h_gat_hm3
            hist["n_years_chcp"]   = h_n_years
            hist["years_chcp"]     = tuple(int(x) for x in h_years_str.split("-")) if "-" in h_years_str else None
        hist["_hist_is_manual"] = _hist_is_manual


# ═══════════════════════════════════════════════════════════════
#  TAB 9 — EVAPORACIÓN
# ═══════════════════════════════════════════════════════════════
with tabs[4]:
    st.markdown('<div class="sec-hdr">📋 ROCC — Texto final del informe</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class='info-box'>
    ℹ️ El texto ingresado aquí se añadirá al <b>final del DOCX</b> exportado, después del balance hídrico y las tablas.
    Úsalo para la sección ROCC, recomendaciones operacionales u observaciones adicionales.
    </div>""", unsafe_allow_html=True)
    _rocc_val = st.text_area(
        "📝 Texto ROCC (se añade al final del DOCX)",
        value=str(st.session_state.get("rocc_texto", "") or ""),
        height=420, key="rocc_texto",
        help="Aparecerá al final del DOCX. Usa saltos de línea para separar párrafos."
    )
    if str(st.session_state.get("rocc_texto", "") or "").strip():
        st.markdown("<div class='ok-box'>✅ Texto ROCC listo — se incluirá al final del informe exportado.</div>", unsafe_allow_html=True)
    else:
        st.markdown("<div class='info-box'>ℹ️ Sin texto ROCC — la sección no se añadirá al DOCX.</div>", unsafe_allow_html=True)

with tabs[11]:
    st.markdown('<div class="sec-hdr">📝 Exportar Informe</div>', unsafe_allow_html=True)

    # ── Salinidad consolidada para el informe ─────────────────────────────
    sal_final_spc = float(datos.get("salinidad_spc", 0.0) or 0.0) if data_ready else 0.0
    sal_final_spv = float(datos.get("salinidad_spv", 0.0) or 0.0) if data_ready else 0.0
    sal_final_dec = float(datos.get("salinidad_dec_pct", 0.0) or 0.0) if data_ready else 0.0

    st.markdown(f"#### Informe: **{label_per}** · {tipo}")
    st.markdown("#### Fuente auditable del informe")
    fuente_informe = st.radio(
        "Fuente base para números del escrito",
        ["Hydro Interna", "Manual / Ajustado"],
        horizontal=True,
        index=0,
        help="Hydro Interna usa la Hidroestadística generada internamente desde DAILY + históricos; Manual permite afinar el escrito."
    )
    if hidro_rows1:
        st.caption("Modo recomendado: Hydro Interna, generada desde DAILY + históricos y usada como base del escrito.")
    if oficial_doc:
        st.info(f"Documento oficial detectado para este período: **{getattr(f_doc_oficial_src, 'name', None) or Path(str(f_doc_oficial_src)).name}**. Puedes usarlo como narrativa y como tabla de referencia para consolidar el reporte.")

    base_dias = n_dias if n_dias else (
        calendar.monthrange(
            anio_sel if MESES_NUM.get(mes_sel,1) not in [10,11,12] else anio_sel-1,
            MESES_NUM.get(mes_sel,1))[1] if tipo=="Mensual"
        else int(PERIODO_DIAS.get(_periodo_key(tipo, trimestre, semestre), 182)))
    base_aporte_total_hm3 = H("Aporte Total","value1","hm3", default=D("aporte_total_hm3")) if fuente_informe == "Hydro Interna" else D("aporte_total_hm3")
    base_aporte_total_m3s = H("Aporte Total","value2","m3/s", default=D("aporte_total_m3s")) if fuente_informe == "Hydro Interna" else D("aporte_total_m3s")
    base_evap_hm3 = H("Evaporación mensual","value1","hm3", default=D("evaporacion_hm3")) if fuente_informe == "Hydro Interna" else D("evaporacion_hm3")
    base_evap_m3s = H("Evaporación mensual","value2","m3/s", default=D("evaporacion_m3s")) if fuente_informe == "Hydro Interna" else D("evaporacion_m3s")
    base_aporte_neto_hm3 = H("Aporte Neto","value1","hm3", default=D("aporte_neto_hm3")) if fuente_informe == "Hydro Interna" else D("aporte_neto_hm3")
    base_aporte_neto_m3s = H("Aporte Neto","value2","m3/s", default=D("aporte_neto_m3s")) if fuente_informe == "Hydro Interna" else D("aporte_neto_m3s")
    base_pan_hm3 = H("Esclusajes PNX","value1","hm3", default=D("panamax_hm3")) if fuente_informe == "Hydro Interna" else D("panamax_hm3")
    base_neo_hm3 = H("Esclusajes NPX","value1","hm3", default=D("neopanamax_hm3")) if fuente_informe == "Hydro Interna" else D("neopanamax_hm3")
    base_esclusaje_hm3 = (base_pan_hm3 or 0) + (base_neo_hm3 or 0) if fuente_informe == "Hydro Interna" else D("esclusaje_hm3")
    base_esclusaje_m3s = hm3_to_m3s_period(base_esclusaje_hm3, base_dias)
    base_tpan = H("Tránsitos Panamax","value1","tránsitos", row_min=26, default=D("transitos_panamax")) if fuente_informe == "Hydro Interna" else D("transitos_panamax")
    base_tneo = H("Tránsitos NeoPanamax","value1","tránsitos", row_min=30, default=D("transitos_neopanamax")) if fuente_informe == "Hydro Interna" else D("transitos_neopanamax")
    base_pot_gat_hm3 = H("Potabilización Gatún","value1","hm3", default=D("potabilizacion_gat_hm3")) if fuente_informe == "Hydro Interna" else D("potabilizacion_gat_hm3")
    base_pot_alh_hm3 = H("Potabilización Alhajuela","value1","hm3", default=D("potabilizacion_alh_hm3")) if fuente_informe == "Hydro Interna" else D("potabilizacion_alh_hm3")
    base_pot_hm3 = (base_pot_gat_hm3 or 0) + (base_pot_alh_hm3 or 0)
    base_pot_m3s = hm3_to_m3s_period(base_pot_hm3, base_dias)
    base_conc_hm3 = H("Concesiones y Misceláneos","value1","hm3", default=D("concesiones_hm3")) if fuente_informe == "Hydro Interna" else D("concesiones_hm3")
    base_conc_m3s = H("Concesiones y Misceláneos","value2","m3/s", default=D("concesiones_m3s")) if fuente_informe == "Hydro Interna" else D("concesiones_m3s")
    base_hidg_hm3 = H("Hidrogeneración Gatún","value1","hm3", default=D("hidro_gatun_hm3")) if fuente_informe == "Hydro Interna" else D("hidro_gatun_hm3")
    base_hidg_m3s = H("Hidrogeneración Gatún","value2","m3/s", default=hm3_to_m3s_period(D("hidro_gatun_hm3"), base_dias)) if fuente_informe == "Hydro Interna" else hm3_to_m3s_period(D("hidro_gatun_hm3"), base_dias)
    base_zz_hm3 = D("zzflush_auto_hm3", 0.0)
    # Trasvase: siempre se toma del cálculo auditado DAILY G+R+U.
    # No se sustituye por Hidro Interna ni por narrativa DOCX para evitar acumulaciones viejas.
    base_trav_hm3 = D("trasvase_hm3")
    base_trav_m3s = D("trasvase_m3s")
    # Ahorros: siempre se toman del cálculo auditado DAILY W/X × 0.2081976.
    # No se sustituyen por Hidro Interna ni por narrativa DOCX para evitar valores antiguos o de otra fuente.
    base_ahp_hm3 = D("ahorro_panamax_hm3")
    base_ahn_hm3 = D("ahorro_neopanamax_hm3")
    base_ahorro_total_hm3 = D("ahorro_total_hm3")
    base_ahorro_lamina = D("ahorro_lamina_pies")
    # Energía: siempre desde DAILY AS/AT.
    # MW = promedio de la columna; MWh/día/valor del informe = suma directa de la misma columna.
    # No se sustituye por Hidro Interna ni por narrativa DOCX para evitar valores antiguos.
    base_emw = D("energia_madden_mw")
    base_emwh = D("energia_madden_mwh_dia")
    base_egw = D("energia_gatun_mw")
    base_egwh = D("energia_gatun_mwh_dia")
    base_total_salidas_hm3 = H("Usos Totales","value1","hm3", default=D("total_salidas_hm3")) if fuente_informe == "Hydro Interna" else D("total_salidas_hm3")
    base_total_salidas_m3s = H("Usos Totales","value2","m3/s", default=hm3_to_m3s_period(base_total_salidas_hm3, base_dias)) if fuente_informe == "Hydro Interna" else hm3_to_m3s_period(base_total_salidas_hm3, base_dias)
    base_deficit_hm3 = H("Deficit","value1","hm3", default=base_aporte_neto_hm3 - base_total_salidas_hm3) if fuente_informe == "Hydro Interna" else base_aporte_neto_hm3 - base_total_salidas_hm3

    # Perfil auditado del informe mensual:
    # 1) Prioriza Hidroestadística para el escrito
    # 2) Usa Fuentes_Agua para subtotales de subcuenca
    # 3) Solo usa DAILY como respaldo si no hay hoja oficial cargada
    if fuente_informe == "Hydro Interna":
        if base_aporte_total_hm3 is None:
            base_aporte_total_hm3 = D("aporte_total_hm3")
        if base_aporte_neto_hm3 is None:
            base_aporte_neto_hm3 = D("aporte_neto_hm3")
        if base_total_salidas_hm3 is None:
            base_total_salidas_hm3 = D("total_salidas_hm3")

    # ── Datos adicionales manuales ────────────────────────────────────────
    st.markdown("#### Datos complementarios para el informe")

    # ── Imagen para el informe ─────────────────────────────────────────────
    st.markdown("##### 🖼️ Gráfica / Imagen para el informe")
    st.markdown("""<div class='info-box' style='font-size:.82rem'>
    Sube aquí la gráfica o imagen que debe aparecer en el informe exportado (PNG, JPG o JPEG).<br>
    Se insertará al final del DOCX, centrada, después del balance hídrico.
    </div>""", unsafe_allow_html=True)
    report_image_file = st.file_uploader(
        "📎 Imagen para el informe (PNG / JPG)",
        type=["png", "jpg", "jpeg"],
        key="f_report_image",
        help="Gráfica generada externamente que se incluirá en el DOCX exportado."
    )
    _img_bytes = None
    _img_ext   = "png"
    if report_image_file is not None:
        _img_bytes = report_image_file.read()
        _img_ext   = report_image_file.name.rsplit(".", 1)[-1].lower()
        if _img_ext == "jpg":
            _img_ext = "jpeg"
        st.image(_img_bytes, caption=f"Vista previa: {report_image_file.name}", use_container_width=True)
        st.success(f"✅ Imagen lista para incluir en el informe ({len(_img_bytes)//1024} KB)")
    else:
        st.caption("Sin imagen cargada — el informe se exportará sin gráfica.")

    st.markdown("---")
    col_x1, col_x2, col_x3 = st.columns(3)
    with col_x1:
        st.markdown("**🌧 Precipitación**")
        st.markdown("""<div class='info-box' style='font-size:.79rem'>
        Configura estos valores en la pestaña <b>🌡️ Meteorología</b>.</div>""",
        unsafe_allow_html=True)
        precip_mm   = float(st.session_state.get("x_prec", 0.0))
        precip_hist = float(st.session_state.get("x_phist", 0.0))
        precip_pos  = int(st.session_state.get("x_ppos", 0) or 0)
        precip_rank_label = str(st.session_state.get("x_prank_label", _meteo_rank_label_default(tipo, mes_sel)) or _meteo_rank_label_default(tipo, mes_sel))
        _pper_ss    = st.session_state.get("x_pper", "")
        _pper_default = (f"{hist.get('years_chcp', (0,0))[0]}-{hist.get('years_chcp', (0,0))[1]}"
                          if hist.get("years_chcp") else "")
        precip_per  = str(_pper_ss) if _pper_ss else _pper_default
        precip_dias = str(st.session_state.get("x_prec_dias", ""))
        _e1, _e2 = st.columns(2)
        _e1.metric("Precipitación (mm)", f"{precip_mm:.0f}" if precip_mm else "—")
        _e2.metric("Hist. promedio (mm)", f"{precip_hist:.0f}" if precip_hist else "—")
        _e3, _e4 = st.columns(2)
        _e3.metric("Posición ranking", str(precip_pos) if precip_pos else "—")
        _e4.metric("Período", precip_per or "—")
        st.caption(f"Ranking usado en el escrito: {precip_rank_label or '—'}")

    with col_x2:
        st.markdown("**🏊 Embalses, Calado e Histórico**")
        _x_cpan_default = float(D("calado_panamax_pies", 39.50) or 39.50) if data_ready else 39.50
        _x_cneo_default = float(D("calado_neopanamax_pies", 50.0) or 50.0) if data_ready else 50.0
        _hist_rank_auto = int(hidro_audit.get("hist_rank", hist.get("hist_rank", 0)) or 0)
        _hist_rank_total_auto = int(hidro_audit.get("hist_rank_total", hist.get("hist_rank_total", hist.get("n_years_chcp", 0))) or 0) + 1  # +1 incluye el año en curso
        _hist_scale_class_auto = str(hidro_audit.get("hist_scale_class", hist.get("hist_scale_class", "")) or "")
        _x_pesc_default = (f"{_hist_rank_auto} de {_hist_rank_total_auto}" if _hist_rank_auto and _hist_rank_total_auto else (str(_hist_rank_auto) if _hist_rank_auto else ""))
        # Forzar actualización del widget si el total cambió (1898-2025 → +año actual)
        _cur_pesc = str(st.session_state.get("x_pesc", "") or "")
        if _hist_rank_auto and _hist_rank_total_auto and _cur_pesc:
            import re as _re_pesc
            _old_tot = _re_pesc.search(r"de\s+(\d+)$", _cur_pesc.strip())
            if _old_tot and int(_old_tot.group(1)) != _hist_rank_total_auto:
                st.session_state["x_pesc"] = _x_pesc_default
        _x_pct_default = (round(((D("aporte_total_hm3") - hist.get("hist_chcp_hm3",0))/hist.get("hist_chcp_hm3",1)*100),1) if hist.get("hist_chcp_hm3",0) else 0.0)
        _x_hist_hm3_default = float(hist.get("hist_chcp_hm3", 0.0) or 0.0)
        _x_hist_m3s_default = float(hist.get("hist_chcp_m3s", 0.0) or 0.0)
        _x_demanda_acp_default = 19.0

        _x_cpan_val = _ensure_numeric_state("x_cpan", _x_cpan_default, replace_if_zero=True)
        _x_cneo_val = _ensure_numeric_state("x_cneo", _x_cneo_default, replace_if_zero=True)
        _x_pesc_val = _ensure_text_state("x_pesc", _x_pesc_default, replace_if_blank=True)
        _x_pct_val = _ensure_numeric_state("x_pct", _x_pct_default, replace_if_zero=True)
        _x_hist_hm3_val = _ensure_numeric_state("x_hist_hm3", _x_hist_hm3_default, replace_if_zero=True)
        _x_hist_m3s_val = _ensure_numeric_state("x_hist_m3s", _x_hist_m3s_default, replace_if_zero=True)
        _x_demanda_acp_val = _ensure_numeric_state("x_demanda_acp", _x_demanda_acp_default, replace_if_zero=True)

        st.caption("Valores por defecto: 39.50 pies (Panamax) y 50 pies (Neopanamax).")
        calado_pan     = st.number_input("Calado Panamax (pies)",         value=float(_x_cpan_val), step=0.01, key="x_cpan")
        calado_neo     = st.number_input("Calado Neopanamax (pies)",      value=float(_x_cneo_val), step=0.01, key="x_cneo")
        pos_escala     = st.text_input("Posición en escala húmedo→seco (ej: 4 de 128)",
            value=str(_x_pesc_val),
            key="x_pesc",
            help="Se llena automáticamente con el ranking histórico del período activo. Puedes editarlo solo si necesitas alinear el texto final.")
        if _hist_scale_class_auto:
            st.caption(f"Clasificación automática: {_hist_scale_class_auto}")
        pct_sobre_prom = st.number_input("% sobre promedio histórico",    value=float(_x_pct_val), step=0.1,  key="x_pct")
        hist_hm3       = st.number_input("Aporte histórico prom (hm³)",   value=float(_x_hist_hm3_val), step=0.1,  key="x_hist_hm3")
        hist_m3s       = st.number_input("Aporte histórico prom (m³/s)",  value=float(_x_hist_m3s_val), step=0.01, key="x_hist_m3s")
        demanda_acp_exp = st.number_input("Demanda interna ACP (MW)",     value=float(_x_demanda_acp_val), step=0.1,  key="x_demanda_acp")

    with col_x3:
        st.markdown("**🌊 Salinidad**")
        st.markdown("""<div class='info-box' style='font-size:.79rem'>
        Configura estos valores en la pestaña <b>🌡️ Meteorología</b>.</div>""",
        unsafe_allow_html=True)
        sal_spc_x = float(st.session_state.get("x_sspc", st.session_state.get("_sal_spc_file", sal_final_spc)))
        sal_spv_x = float(st.session_state.get("x_sspv", st.session_state.get("_sal_spv_file", sal_final_spv)))
        sal_dec_x = float(st.session_state.get("x_sdec", st.session_state.get("_sal_dec_file", sal_final_dec)))
        _s1, _s2, _s3 = st.columns(3)
        _s1.metric("SPC (ups)", f"{sal_spc_x:.4f}" if sal_spc_x else "—")
        _s2.metric("SPV (ups)", f"{sal_spv_x:.4f}" if sal_spv_x else "—")
        _s3.metric("Decremento (%)", f"{sal_dec_x:.1f}" if sal_dec_x else "—")

    st.markdown("---")

    # ── Ajustes finos del informe (para alinear el reporte con la versión oficial) ──
    fa_alh = float(fuentes_sub.get("Alhajuela", {}).get("hm3", D("alhajuela_hm3", 0.0)) or 0.0)
    fa_gat = float(fuentes_sub.get("Gatún", {}).get("hm3", D("gatun_hm3", 0.0)) or 0.0)
    fa_total = fa_alh + fa_gat if (fa_alh > 0 or fa_gat > 0) else float(D("aporte_total_hm3", 0.0) or 0.0)

    zzflush_rep_default = float(D("zzflush_auto_hm3", st.session_state.get("zzflush_hm3", 0.0)) or 0.0)
    rpt_aporte_total_default = float(oficial_metrics.get("aporte_total_hm3") if oficial_metrics.get("aporte_total_hm3") is not None else (base_aporte_total_hm3 if base_aporte_total_hm3 else (fa_total if fa_total > 0 else D("aporte_total_hm3", 0.0))))
    rpt_aporte_neto_default = float(oficial_metrics.get("aporte_neto_hm3") if oficial_metrics.get("aporte_neto_hm3") is not None else (base_aporte_neto_hm3 if base_aporte_neto_hm3 else max(rpt_aporte_total_default - base_evap_hm3, 0.0)))
    rpt_alh_default = float(oficial_metrics.get("alhajuela_hm3") if oficial_metrics.get("alhajuela_hm3") is not None else (fa_alh if fa_alh > 0 else D("alhajuela_hm3", 0.0)))
    rpt_gat_default = float(oficial_metrics.get("gatun_hm3") if oficial_metrics.get("gatun_hm3") is not None else (fa_gat if fa_gat > 0 else D("gatun_hm3", 0.0)))

    # ── Sembrar TODOS los widgets rep_* cuando están en 0 y hay datos reales ────
    # Streamlit conserva session_state entre reruns; value= solo aplica la primera vez.
    # Esta función actualiza el estado ANTES de renderizar el widget, sin pisar
    # ediciones manuales del usuario (solo actúa si el widget vale 0).
    def _seed_if_zero(key, default_val):
        if default_val and float(default_val) > 0.0001:
            _ensure_numeric_state(key, default_val, replace_if_zero=True)

    _seed_if_zero("rep_aporte_total",  rpt_aporte_total_default)
    _seed_if_zero("rep_aporte_neto",   rpt_aporte_neto_default)
    _seed_if_zero("rep_alh_hm3",       rpt_alh_default)
    _seed_if_zero("rep_gat_hm3",       rpt_gat_default)
    rpt_esclusaje_default = float(oficial_metrics.get("esclusaje_hm3") if oficial_metrics.get("esclusaje_hm3") is not None else (base_esclusaje_hm3 or 0.0))
    rpt_pan_default = float(oficial_metrics.get("panamax_hm3") if oficial_metrics.get("panamax_hm3") is not None else (base_pan_hm3 or 0.0))
    rpt_neo_default = float(oficial_metrics.get("neopanamax_hm3") if oficial_metrics.get("neopanamax_hm3") is not None else (base_neo_hm3 or 0.0))
    rpt_tpan_default = float(base_tpan or 0.0)
    rpt_tneo_default = float(base_tneo or 0.0)
    # Ahorros del informe: quedan amarrados a DAILY W/X × 0.2081976.
    rpt_ahp_default = float(base_ahp_hm3 or 0.0)
    rpt_ahn_default = float(base_ahn_hm3 or 0.0)
    rpt_cca_default = float(oficial_metrics.get("cca_hm3") if oficial_metrics.get("cca_hm3") is not None else (base_zz_hm3 if base_zz_hm3 else D("zzflush_auto_hm3", zzflush_rep_default) or 0.0))
    # Trasvase del informe queda amarrado a DAILY G+R+U; se ignoran narrativas previas.
    rpt_trav_default = float(base_trav_hm3 or 0.0)

    # Sembrar los widgets restantes ahora que todos los defaults están definidos
    _seed_if_zero("rep_esclusaje_hm3", rpt_esclusaje_default)
    # PNX/NPX: forzar siempre el valor del DAILY para evitar que session_state
    # conserve valores invertidos de versiones anteriores del app.
    if rpt_pan_default > 0: st.session_state["rep_pan_hm3"] = float(rpt_pan_default)
    if rpt_neo_default > 0: st.session_state["rep_neo_hm3"] = float(rpt_neo_default)
    # Forzar trasvase auditado antes de instanciar el widget; evita que Streamlit conserve
    # un valor viejo (p.ej. 199 hm³) de una corrida anterior.
    _ensure_numeric_state("rep_trasvase_hm3", rpt_trav_default, replace_if_zero=True)
    st.session_state["rep_trasvase_hm3"] = float(rpt_trav_default or 0.0)
    _seed_if_zero("rep_tpan",          rpt_tpan_default)
    _seed_if_zero("rep_tneo",          rpt_tneo_default)
    # Estos dos campos se fuerzan antes de instanciar los widgets para evitar
    # que Streamlit conserve valores viejos en session_state.
    _ensure_numeric_state("rep_ahorro_pan", rpt_ahp_default, replace_if_zero=True)
    _ensure_numeric_state("rep_ahorro_neo", rpt_ahn_default, replace_if_zero=True)
    st.session_state["rep_ahorro_pan"] = float(rpt_ahp_default or 0.0)
    st.session_state["rep_ahorro_neo"] = float(rpt_ahn_default or 0.0)
    _seed_if_zero("rep_cca_hm3",       rpt_cca_default)
    _seed_if_zero("rep_zzflush_hm3",   zzflush_rep_default)

    with st.expander("⚙️ Ajustes finos del informe", expanded=True):
        st.markdown("""
        <div class='warn-box'>
        Sin archivos cargados, este bloque inicia en 0. Cuando cargues archivos para el período seleccionado,
        los valores base se recalculan automáticamente y luego puedes ajustarlos manualmente sin afectar el resto del dashboard.<br>
        <b>Nota:</b> los <i>Usos Totales CHCP</i> del balance excluyen trasvase Madden→Gatún, hidrogeneración Madden y MADSPILL
        porque corresponden a movimientos/operaciones internas y no a consumos/pérdidas netas del sistema.
        </div>""", unsafe_allow_html=True)

        a1, a2, a3, a4 = st.columns(4)
        with a1:
            rpt_aporte_total = st.number_input("Aporte total informe (hm³)", value=float(st.session_state.get("rep_aporte_total", rpt_aporte_total_default)), step=0.1, key="rep_aporte_total")
            rpt_aporte_neto  = st.number_input("Aporte neto informe (hm³)", value=float(st.session_state.get("rep_aporte_neto", rpt_aporte_neto_default)), step=0.1, key="rep_aporte_neto")
            rpt_alh_hm3      = st.number_input("Subtotal Alhajuela informe (hm³)", value=float(st.session_state.get("rep_alh_hm3", rpt_alh_default)), step=0.1, key="rep_alh_hm3")
            rpt_gat_hm3      = st.number_input("Subtotal Gatún informe (hm³)", value=float(st.session_state.get("rep_gat_hm3", rpt_gat_default)), step=0.1, key="rep_gat_hm3")
        with a2:
            rpt_esclusaje_hm3 = st.number_input("Esclusaje informe (hm³)", value=float(st.session_state.get("rep_esclusaje_hm3", rpt_esclusaje_default)), step=0.01, key="rep_esclusaje_hm3")
            rpt_pan_hm3       = st.number_input("PNX informe (hm³)", value=float(st.session_state.get("rep_pan_hm3", rpt_pan_default)), step=0.01, key="rep_pan_hm3")
            rpt_neo_hm3       = st.number_input("NPX informe (hm³)", value=float(st.session_state.get("rep_neo_hm3", rpt_neo_default)), step=0.01, key="rep_neo_hm3")
            rpt_trasvase_hm3  = st.number_input("Trasvase informe (hm³)", value=float(st.session_state.get("rep_trasvase_hm3", rpt_trav_default)), step=0.1, key="rep_trasvase_hm3")
        with a3:
            rpt_tpan        = st.number_input("Tránsitos PNX informe", value=float(st.session_state.get("rep_tpan", rpt_tpan_default)), step=0.5, key="rep_tpan")
            rpt_tneo        = st.number_input("Tránsitos NPX informe", value=float(st.session_state.get("rep_tneo", rpt_tneo_default)), step=0.5, key="rep_tneo")
            rpt_ahorro_pan  = st.number_input("Ahorro PNX informe (hm³)", value=float(st.session_state.get("rep_ahorro_pan", rpt_ahp_default)), step=0.01, key="rep_ahorro_pan")
            rpt_ahorro_neo  = st.number_input("Ahorro NeoPNX informe (hm³)", value=float(st.session_state.get("rep_ahorro_neo", rpt_ahn_default)), step=0.01, key="rep_ahorro_neo")
        with a4:
            rpt_cca_hm3     = float(st.session_state.get("rep_cca_hm3", rpt_cca_default))
            st.metric("CCA / ZZFlush (hm³)", f"{rpt_cca_hm3:.2f}",
                      "✏️ Configura en Tab 🌊 Salinidad")
            rpt_zzflush_hm3 = rpt_cca_hm3  # ZZFlush = CCA, ingresado en pestaña Salinidad
            st.metric("Evaporación base", f"{base_evap_hm3:.2f} hm³")
            st.metric("Potabilización base", f"{base_pot_hm3:.2f} hm³")
            st.metric("Mov. internos base", f"{D('movimientos_operativos_hm3', 0.0):.2f} hm³")

    # ── Derivados del informe ──────────────────────────────────────────────
    mes_lbl = mes_sel if tipo=="Mensual" else label_per
    prec_pct_auto = abs(precip_mm - precip_hist) / precip_hist * 100 if precip_hist else 0
    prec_pct  = float(st.session_state.get("x_met_pct_informe", prec_pct_auto))
    prec_dir  = "por encima" if precip_mm >= precip_hist else "por debajo"
    precip_rank_label = str(st.session_state.get("x_prank_label", _meteo_rank_label_default(tipo, mes_sel)) or _meteo_rank_label_default(tipo, mes_sel))
    precip_rank_txt = (f", colocándola en la posición {precip_pos} de los {precip_rank_label} ({precip_per})"
                       if precip_pos and precip_per and precip_rank_label else "")
    precip_dias_txt = (f" La mayor parte de la lluvia del período se concentró {precip_dias}."
                       if str(precip_dias).strip() else "")

    # ── Fallback: si algún rpt_* aún es 0 después del seed, usar D() ──────────
    # Protege el primer rerun tras cargar archivos cuando session_state
    # no había sido sembrado previamente.
    def _rpt(widget_val, datos_key, scale=1.0):
        """Devuelve widget_val si >0, si no D(datos_key)*scale."""
        v = float(widget_val or 0)
        if abs(v) > 0.0001:
            return v
        return float(D(datos_key, 0.0) or 0.0) * scale

    rpt_aporte_total   = _rpt(rpt_aporte_total,   "aporte_total_hm3")
    rpt_aporte_neto    = _rpt(rpt_aporte_neto,    "aporte_neto_hm3")
    rpt_alh_hm3        = _rpt(rpt_alh_hm3,        "alhajuela_hm3")
    rpt_gat_hm3        = _rpt(rpt_gat_hm3,        "gatun_hm3")
    rpt_esclusaje_hm3  = _rpt(rpt_esclusaje_hm3,  "esclusaje_hm3")
    rpt_pan_hm3        = _rpt(rpt_pan_hm3,        "panamax_hm3")
    rpt_neo_hm3        = _rpt(rpt_neo_hm3,        "neopanamax_hm3")
    rpt_trasvase_hm3   = float(D("trasvase_hm3", rpt_trasvase_hm3) or 0.0)
    rpt_tpan           = _rpt(rpt_tpan,           "transitos_panamax")
    rpt_tneo           = _rpt(rpt_tneo,           "transitos_neopanamax")
    rpt_ahorro_pan     = _rpt(rpt_ahorro_pan,     "ahorro_panamax_hm3")
    rpt_ahorro_neo     = _rpt(rpt_ahorro_neo,     "ahorro_neopanamax_hm3")
    rpt_cca_hm3        = _rpt(rpt_cca_hm3,        "zzflush_auto_hm3")
    rpt_zzflush_hm3    = _rpt(rpt_zzflush_hm3,    "zzflush_auto_hm3")

    rpt_alh_pct = (rpt_alh_hm3 / rpt_aporte_total * 100) if rpt_aporte_total > 0 else 0.0
    rpt_gat_pct = (rpt_gat_hm3 / rpt_aporte_total * 100) if rpt_aporte_total > 0 else 0.0
    rpt_pan_pct = (rpt_pan_hm3 / rpt_esclusaje_hm3 * 100) if rpt_esclusaje_hm3 > 0 else 0.0
    rpt_neo_pct = (rpt_neo_hm3 / rpt_esclusaje_hm3 * 100) if rpt_esclusaje_hm3 > 0 else 0.0
    rpt_aporte_total_m3s = H("Aporte Total","value2","m3/s", default=hm3_to_m3s_period(rpt_aporte_total, n_dias)) if fuente_informe == "Hydro Interna" else hm3_to_m3s_period(rpt_aporte_total, n_dias)
    rpt_aporte_neto_m3s = H("Aporte Neto","value2","m3/s", default=hm3_to_m3s_period(rpt_aporte_neto, n_dias)) if fuente_informe == "Hydro Interna" else hm3_to_m3s_period(rpt_aporte_neto, n_dias)
    rpt_uso_prom_diario_hm3 = rpt_esclusaje_hm3 / n_dias if n_dias > 0 else 0.0
    rpt_uso_prom_diario_m3s = rpt_esclusaje_hm3 / (n_dias * HM3_PER_M3S_DAY) if n_dias > 0 else 0.0
    rpt_agua_pan_trans = (rpt_pan_hm3 / rpt_tpan) if rpt_tpan > 0 else 0.0
    rpt_agua_neo_trans = (rpt_neo_hm3 / rpt_tneo) if rpt_tneo > 0 else 0.0
    rpt_ahorro_total = rpt_ahorro_pan + rpt_ahorro_neo
    rpt_ahorro_lamina = rpt_ahorro_total / 148.0 if rpt_ahorro_total > 0 else 0.0
    rpt_total_salidas = (
        rpt_esclusaje_hm3 + base_pot_hm3 + rpt_zzflush_hm3 +
        float(base_conc_hm3 or 0) + float(base_evap_hm3 or 0) + float(base_hidg_hm3 or 0) + float(D("vertidos_gat_hm3") or 0)
    )
    rpt_trasvase_m3s = rpt_trasvase_hm3 / (n_dias * HM3_PER_M3S_DAY) if n_dias > 0 else 0.0
    # Trasvase total = MADMCF (G) + MADSPILL (R) + LEAK MAD (U).
    # rpt_trasvase_hm3 ya contiene ese total; no se vuelve a sumar.
    rpt_trasvase_total_hm3 = rpt_trasvase_hm3
    rpt_trasvase_total_m3s = rpt_trasvase_total_hm3 / (n_dias * HM3_PER_M3S_DAY) if n_dias > 0 else 0.0
    rpt_trasvase_pct = (rpt_trasvase_total_hm3 / rpt_esclusaje_hm3 * 100) if rpt_esclusaje_hm3 > 0 else 0.0
    rpt_trasvase_hidro_hm3 = float(D("hidro_madden_hm3", 0) or 0)
    rpt_trasvase_vert_hm3 = float(D("vertidos_mad_ops_hm3", 0) or 0)
    rpt_trasvase_fugas_hm3 = float(D("trasvase_fugas_hm3", D("concesiones_alh_hm3", 0)) or 0)
    base_demanda_acp = float(st.session_state.get("x_demanda_acp", 19.0) or 19.0)
    rpt_cca_unitario = (rpt_cca_hm3 / rpt_tneo) if rpt_tneo > 0 else 0.0

    # ── Auditoría automática del escrito ───────────────────────────────────
    # Esta tabla verifica los valores que alimentan directamente las viñetas del
    # informe. Si alguno queda en cero, el app avisa antes de exportar.
    def _audit_num(label, value, source, required=True, allow_zero=False, fmt="{:.2f}", note=""):
        try:
            v = float(value if value is not None else 0.0)
        except Exception:
            v = 0.0
        missing = bool(required and (not allow_zero) and abs(v) <= 0.000001)
        try:
            val_txt = fmt.format(v)
        except Exception:
            val_txt = str(v)
        return {
            "Campo del escrito": label,
            "Valor": val_txt,
            "Fuente / cálculo": source,
            "Estado": "⚠️ Quedará en 0" if missing else "✅ OK",
            "Observación": note,
            "_missing": missing,
        }

    # Modo "Texto oficial" activo → los datos están en el párrafo pegado
    _met_oficial = (
        st.session_state.get("met_texto_mode") == "Texto oficial (pegar íntegramente)"
    )
    _sal_oficial = (
        st.session_state.get("sal_texto_mode") == "Texto oficial (pegar íntegramente)"
    )
    _nota_met = "Dato en texto oficial pegado." if _met_oficial else "Requerido para la primera viñeta."
    _nota_sal = "Dato en texto oficial pegado." if _sal_oficial else ""

    audit_rows = [
        _audit_num("Precipitación del período", precip_mm, "Manual · pestaña Meteorología", required=not _met_oficial, fmt="{:.0f}", note=_nota_met),
        _audit_num("Precipitación histórica", precip_hist, "Manual / histórico · pestaña Meteorología", required=not _met_oficial, fmt="{:.0f}", note=_nota_met),
        _audit_num("Ranking precipitación", precip_pos, "Manual · pestaña Meteorología", required=not _met_oficial, allow_zero=_met_oficial, fmt="{:.0f}", note="Posición histórica de humedad." if not _met_oficial else _nota_met),
        _audit_num("Período histórico precipitación", 1 if str(precip_per).strip() else 0, "Manual · pestaña Meteorología", required=not _met_oficial, fmt="{:.0f}", note=_nota_met),
        _audit_num("Aporte total CHCP", rpt_aporte_total, "Fuentes_Agua / Hidro Interna", fmt="{:.1f}"),
        _audit_num("Aporte neto CHCP", rpt_aporte_neto, "Fuentes_Agua - evaporación / GATNETFL", fmt="{:.1f}"),
        _audit_num("Aporte subcuenca Alhajuela", rpt_alh_hm3, "Fuentes_Agua", fmt="{:.1f}"),
        _audit_num("Aporte subcuenca Gatún", rpt_gat_hm3, "Fuentes_Agua", fmt="{:.1f}"),
        _audit_num("Esclusaje total", rpt_esclusaje_hm3, "DAILY · J+L+N+P / Hidro Interna", fmt="{:.2f}"),
        _audit_num("Esclusajes Panamax", rpt_pan_hm3, "DAILY · columnas J+L", fmt="{:.2f}"),
        _audit_num("Esclusajes NeoPanamax", rpt_neo_hm3, "DAILY · columnas N+P", fmt="{:.2f}"),
        _audit_num("Tránsitos Panamax", rpt_tpan, "DAILY", fmt="{:.0f}"),
        _audit_num("Tránsitos NeoPanamax", rpt_tneo, "DAILY", fmt="{:.0f}", note="También se usa para el CCA unitario."),
        _audit_num("Ahorro Panamax", rpt_ahorro_pan, "DAILY · columna W × 0.2081976", fmt="{:.2f}", note="Si realmente no hubo ahorro, puede quedar en 0."),
        _audit_num("Ahorro NeoPanamax", rpt_ahorro_neo, "DAILY · columna X × 0.2081976", fmt="{:.2f}", note="Si realmente no hubo ahorro, puede quedar en 0."),
        _audit_num("CCA total", rpt_cca_hm3, "Manual · pestaña Salinidad", fmt="{:.2f}", allow_zero=True, note="Si no hubo ZZFlush, 0 es el valor correcto."),
        _audit_num("CCA unitario", rpt_cca_unitario, "CCA total / tránsitos NeoPanamax", fmt="{:.4f}"),
        _audit_num("Salinidad SPC", sal_spc_x, "Manual o archivo externo · pestaña Salinidad", required=not _sal_oficial, fmt="{:.4f}", note=_nota_sal),
        _audit_num("Salinidad SPV", sal_spv_x, "Manual o archivo externo · pestaña Salinidad", required=not _sal_oficial, fmt="{:.4f}", note=_nota_sal),
        _audit_num("Decremento salinidad", sal_dec_x, "Manual o archivo externo · pestaña Salinidad", required=not _sal_oficial, fmt="{:.1f}", note=_nota_sal),
        _audit_num("Madden MW promedio", base_emw, "DAILY · columna AT = MADMW", fmt="{:.2f}"),
        _audit_num("Madden MWh/día", base_emwh, "Suma directa DAILY · columna AT = MADMW", fmt="{:.1f}"),
        _audit_num("Gatún MW promedio", base_egw, "DAILY · columna AS = GATMW", fmt="{:.4f}", note="Si Gatún no generó, validar que el cero sea real."),
        _audit_num("Gatún MWh/día", base_egwh, "Suma directa DAILY · columna AS = GATMW", fmt="{:.2f}", note="Si Gatún no generó, validar que el cero sea real."),
        _audit_num("Demanda interna ACP", base_demanda_acp, "Manual · Exportar Informe", fmt="{:.1f}"),
        _audit_num("Trasvase total", rpt_trasvase_total_hm3, "DAILY · G + R + U × 0.028316846592", fmt="{:.2f}"),
        _audit_num("Trasvase por hidrogeneración", rpt_trasvase_hidro_hm3, "DAILY · columna G × 0.028316846592", fmt="{:.2f}"),
        _audit_num("Trasvase por vertidos", rpt_trasvase_vert_hm3, "DAILY · columna R × 0.028316846592", fmt="{:.2f}", note="Puede ser 0 si no hubo vertidos."),
        _audit_num("Trasvase por fugas", rpt_trasvase_fugas_hm3, "DAILY · columna U × 0.028316846592", fmt="{:.2f}", note="Puede ser 0 si no hubo fugas registradas."),
    ]
    audit_missing = [r for r in audit_rows if r.get("_missing")]
    if audit_missing:
        campos = ", ".join(r["Campo del escrito"] for r in audit_missing[:8])
        if len(audit_missing) > 8:
            campos += f" y {len(audit_missing)-8} más"
        st.markdown(f"""<div class='warn-box'>
        ⚠️ <b>Auditoría del escrito:</b> hay valores que alimentarían el informe con <b>0</b>.
        Revisa antes de exportar: {campos}.
        </div>""", unsafe_allow_html=True)
    else:
        st.markdown("<div class='ok-box'>✅ Auditoría del escrito completa: los campos principales tienen valores distintos de 0.</div>", unsafe_allow_html=True)

    with st.expander("🔎 Ver auditoría de valores incluidos automáticamente en el escrito", expanded=bool(audit_missing)):
        audit_df = pd.DataFrame([{k: v for k, v in r.items() if k != "_missing"} for r in audit_rows])
        st.dataframe(audit_df, use_container_width=True, hide_index=True)
        st.caption("Los campos marcados como 0 pueden ser reales en ciertos casos operativos; la alerta existe para evitar omisiones por falta de carga o configuración.")

    use_official_doc = st.checkbox(
        "Usar narrativa y tablas del documento oficial detectado",
        value=False,
        help="Para evitar errores por números antiguos, el informe usa por defecto el escrito automático del app. Actívalo solo si deseas conservar literalmente la narrativa/tablas del DOCX oficial."
    ) if oficial_doc else False
    if use_official_doc:
        st.markdown("""<div class='warn-box'>
        ⚠️ <b>Advertencia de auditoría:</b> al usar narrativa del DOCX oficial, algunos números pueden venir del documento cargado y no del cálculo automático del app.
        Para el informe dinámico auditado, deja esta opción desactivada.
        </div>""", unsafe_allow_html=True)

    # El período termina en el año evaluado aunque el archivo histórico llegue a N-1
    _hist_per_raw = str(st.session_state.get("h_years_str", "") or
                        (f"{hist['years_chcp'][0]}-{hist['years_chcp'][1]}" if hist.get("years_chcp") else f"1898-{anio_sel}"))
    _hist_per_parts = re.findall(r"\d{4}", _hist_per_raw)
    if len(_hist_per_parts) >= 2 and int(_hist_per_parts[-1]) < int(anio_sel):
        _hist_periodo_chcp = f"{_hist_per_parts[0]}-{anio_sel}"
    else:
        _hist_periodo_chcp = _hist_per_raw
    # Incluir el año en curso en el conteo total de años (serie histórica + año evaluado)
    _hist_years_chcp = _anios_registro_incluyendo_actual(
        _hist_periodo_chcp, anio_sel,
        fallback=int(st.session_state.get("h_n_years", hist.get("n_years_chcp", 0) if hist else 0) or 0)
    )
    if _hist_years_chcp == 0:
        _nums_yr = [int(x) for x in __import__("re").findall(r"\d{4}", _hist_periodo_chcp)]
        if len(_nums_yr) >= 2:
            _hist_years_chcp = _nums_yr[-1] - _nums_yr[0] + 1
    _hist_pos_chcp_txt = _posicion_historica_texto(pos_escala, tipo, mes_lbl)
    _hist_pct_dir = "por encima al" if float(pct_sobre_prom) > 0 else ("por debajo del" if float(pct_sobre_prom) < 0 else "igual al")
    # Fragmento "N de M años" sin repetir el total cuando ya viene en "N de M"
    _pos_txt_chcp = _hist_pos_chcp_txt or pos_escala or "—"
    _pos_anos_chcp = (f"{_pos_txt_chcp} años"
                      if re.search(r"\d+\s+de\s+\d+", _pos_txt_chcp)
                      else f"{_pos_txt_chcp} en {_hist_years_chcp} años")

    bullets = []
    if use_official_doc and oficial_doc.get("bullets"):
        bullets = oficial_doc.get("bullets", [])
    elif tipo == "Mensual":
        _bullets_raw = [
            f"La precipitación total sobre la Cuenca Hidrográfica del Canal de Panamá (CHCP) durante {mes_lbl.lower()} de {anio_sel} fue de {sig3(precip_mm)} mm, {sig3(prec_pct)}% {prec_dir} del promedio histórico de {sig3(precip_hist)} mm{precip_rank_txt}.{precip_dias_txt}",
            f"Los aportes totales a la Cuenca Hidrográfica del Canal de Panamá (CHCP) durante {mes_lbl.lower()} {anio_sel} fueron {sig3(rpt_aporte_total)} hm³ ({sig3(rpt_aporte_total_m3s)} m³/s), desglosados en {sig3(rpt_aporte_neto)} hm³ ({sig3(rpt_aporte_neto_m3s)} m³/s) de aportes netos y {sig3(base_evap_hm3)} hm³ ({sig3(base_evap_m3s)} m³/s) de evaporación directa en los embalses.",
            f"Los aportes totales a la CHCP de {sig3(rpt_aporte_total)} hm³ ({sig3(rpt_aporte_total_m3s)} m³/s) de {mes_lbl.lower()} {anio_sel}, "
            f"se encuentra como {_pos_anos_chcp} de registro continuo, en una escala de húmedo a seco "
            f"para el periodo {_hist_periodo_chcp}. Los cuales se encuentran en un {sig3(abs(float(pct_sobre_prom)))} % "
            f"{_hist_pct_dir} promedio histórico de {sig3(hist_hm3)} hm³ ({sig3(hist_m3s)} m³/s).",
            # Bullet aportes netos — solo si GATNETFL tiene datos históricos netos
            (lambda _hn=float(hist.get('hist_neto_hm3', 0) or 0.0),
                    _hm=float(hist.get('hist_neto_m3s', 0) or 0.0),
                    _pos=str(datos.get('hist_neto_scale_pos', '') or ''),
                    _cls=str(datos.get('hist_neto_scale_class','') or ''),
                    _per=f"{hist['years_netfl'][0]}-{anio_sel}" if hist.get('years_netfl') else f'1914-{anio_sel}',
                    _nyr=int(hist.get('n_years_netfl', 0) or 0) + 1:
                (f"Los aportes netos de {mes_lbl.lower()} {anio_sel} fueron {sig3(rpt_aporte_neto)} hm³ "
                 f"({sig3(rpt_aporte_neto_m3s)} m³/s). "
                 + (f"Se ubican en la posición {(_pos.split()[0] + ' de ' + str(_nyr)) if _pos and _pos.split()[0].isdigit() else (str(_pos) + ' de ' + str(_nyr) if _pos else '')} "
                    f"en la escala histórica de húmedo a seco"
                    f"{(' para el periodo ' + _per) if _per else ''}. "
                    f"Los cuales se encuentran en un "
                    f"{sig3(abs(round(((rpt_aporte_neto - _hn) / _hn * 100), 0)))}% "
                    f"{'por arriba del' if rpt_aporte_neto > _hn else ('por debajo del' if rpt_aporte_neto < _hn else 'igual al')} "
                    f"promedio histórico de {sig3(_hn)} hm³ ({sig3(_hm)} m³/s)."
                    if _hn > 0 else "")
                ) if _hn > 0 else ""
            )(),
            f"En {mes_lbl.lower()} {anio_sel} la distribución de los aportes hídricos totales a la CHCP fue: Subcuenca embalse Alhajuela el {sig3(rpt_alh_pct)}% ({sig3(rpt_alh_hm3)} hm³), Subcuenca del embalse Gatún el {sig3(rpt_gat_pct)}% ({sig3(rpt_gat_hm3)} hm³).",
            # ── ROCC: se inserta aquí, después de subcuencas ──────────────
            *[_rl.strip() for _rl in str(st.session_state.get("rocc_texto", "") or "").strip().split("\n") if _rl.strip()],
            f"Los niveles de los embalses Gatún y Alhajuela presentan suficiente agua disponible para el consumo humano e industrial. El nivel de agua promedio en el embalse Gatún durante este mes fue de {sig3(D('nivel_gatun_pies'))} pies ({sig3(D('nivel_gatun_m'))} m) PLD. Esta condición ha permitido mantener el calado máximo permisible de {calado_pan} pies para los buques Panama y {calado_neo} pies para los buques Neopanamax.",
            f"El uso de agua en las esclusas del Canal de Panamá ({mes_lbl.lower()} {anio_sel}) ha sido de {sig3(rpt_esclusaje_hm3)} hm³. De este volumen, {sig3(rpt_pan_pct)}% ({sig3(rpt_pan_hm3)} hm³) corresponde a las esclusas Panamax y {sig3(rpt_neo_pct)}% ({sig3(rpt_neo_hm3)} hm³) a las Neopanamax.",
            f"El uso promedio diario de agua en las esclusas fue de {sig3(rpt_uso_prom_diario_hm3)} hm³ ({sig3(rpt_uso_prom_diario_m3s)} m³/s).",
            f"Durante el mes de {mes_lbl.lower()}, se implementaron ahorros en las esclusas NeoPanamax representando un volumen de {sig3(rpt_ahorro_neo)} hm³ y Panamax representando un volumen de {sig3(rpt_ahorro_pan)} hm³, para un total de {sig3(rpt_ahorro_total)} hm³, lo cual en términos de lámina del embalse Gatún representa {sig3(rpt_ahorro_lamina)} pies.",
            f"El uso de agua promedio por tránsito de un buque panamax fue de {sig3(rpt_agua_pan_trans)} hm³.",
            f"Durante {mes_lbl.lower()} de {anio_sel}, la salinidad promedio diaria en el cauce de navegación (SPC) del embalse Gatún fue en promedio {sal_spc_x:.4f} unidades prácticas de salinidad (ups). Al cierre del mes, la salinidad promedio ponderada (SPV) se registró en {sal_spv_x:.4f} ups, representando un decremento del {sal_dec_x:.1f}% respecto al máximo histórico alcanzado.",
            f"El volumen de agua destinado para la conservación de la calidad de agua (CCA) fue de {sig3(rpt_cca_hm3)} hm³, para {rpt_tneo:.1f} tránsitos totales NeoPanamax durante {mes_lbl.lower()} de {anio_sel}, equivalente a un CCA unitario de {rpt_cca_unitario:.4f} hm³/tránsito.",
            f"La producción de energía eléctrica en promedio en la hidroeléctrica Madden fue de {sig3(base_emw)} MW (equivalente a {sig3(base_emwh)} MWh por día en Madden). La producción en la hidroeléctrica Gatún fue de {sig3(base_egw)} MW (equivalente a {sig3(base_egwh)} MWh por día). {'Se logró cubrir la demanda interna de energía de la ACP (' + f'{base_demanda_acp:.1f}' + ' MW promedio), sin utilizar la planta termoeléctrica o importar energía del mercado.' if base_emw >= base_demanda_acp else 'No se logró cubrir la demanda interna de energía de la ACP (' + f'{base_demanda_acp:.1f}' + ' MW promedio), requiriendo uso de termoeléctrica o importación de energía.'}",
            f"Durante este mes, se trasvasaron {sig3(rpt_trasvase_total_hm3)} hm³ ({sig3(rpt_trasvase_total_m3s)} m³/s) de agua desde Alhajuela hacia Gatún ({sig3(rpt_trasvase_hidro_hm3)} hm³ por hidrogeneración, {sig3(rpt_trasvase_vert_hm3)} hm³ por vertidos y {sig3(rpt_trasvase_fugas_hm3)} hm³ por fugas), lo que representó un {sig3(rpt_trasvase_pct)}% del volumen mensual utilizado por las esclusas ({sig3(rpt_esclusaje_hm3)} hm³). Esta estrategia optimiza la gestión del recurso hídrico al mejorar la calidad del agua en el embalse Alhajuela, reduciendo su tiempo de residencia hidráulica (TRH) y potenciando la eficiencia operativa.",
        ]
        # Filtrar bullets del bloque Mensual que estén vacíos o sin datos de hist
        bullets_mensual = [b for b in _bullets_raw
                           if b and str(b).strip() and len(str(b).strip()) > 10
                           and "None" not in str(b)
                           and not (str(b).strip().startswith("Los aportes netos")
                                    and not (hist and hist.get("hist_neto_m3s", 0)))]
        bullets = bullets_mensual
        # ── Reemplazar viñeta de precipitación si el usuario pegó texto oficial ──
        _met_tl = str(st.session_state.get("met_texto_libre", "") or "").strip()
        if st.session_state.get("met_texto_mode") == "Texto oficial (pegar íntegramente)" and _met_tl:
            if bullets:
                bullets[0] = _met_tl
            else:
                bullets.insert(0, _met_tl)
        # ── Reemplazar viñeta de salinidad si el usuario pegó texto oficial ──
        _sal_tl = str(st.session_state.get("met_sal_texto_libre", "") or "").strip()
        if st.session_state.get("sal_texto_mode") == "Texto oficial (pegar íntegramente)" and _sal_tl:
            for _bi, _bb in enumerate(bullets):
                if "salinidad" in str(_bb).lower() and ("spc" in str(_bb).lower() or "ups" in str(_bb).lower()):
                    bullets[_bi] = _sal_tl
                    break
    else:
        # Genérico para Trimestral, Semestral S2 y Anual
        _per_lbl  = label_per.lower()
        _excedente = base_aporte_total_hm3 - base_total_salidas_hm3  # balance correcto: total − salidas
        _cond      = "hubo excedentes" if _excedente > 0 else "no hubo excedentes"
        _accion    = "almacenamiento y reserva de agua en ambos embalses" if _excedente > 0 else "uso de agua almacenada en ambos embalses"
        # Etiqueta del período para CCA/ahorros (s = semestre/trimestre/año)
        _per_noun = {
            "Trimestral": "trimestre",
            "Semestral":  "semestre",
            "Anual":      "año fiscal",
        }.get(tipo, "período")
        bullets = [
            f"La precipitación total sobre la Cuenca Hidrográfica del Canal de Panamá (CHCP) durante {_per_lbl} fue de {precip_mm:.0f} mm, {prec_pct:.1f}% {prec_dir} del promedio histórico de {precip_hist:.0f} mm{precip_rank_txt}.{precip_dias_txt}",
            f"Durante {_per_lbl}, los aportes totales a la CHCP resultaron en {rpt_aporte_total:.1f} hm³ ({rpt_aporte_total_m3s:.1f} m³/s), con un aporte neto de {rpt_aporte_neto:.1f} hm³ ({rpt_aporte_neto_m3s:.1f} m³/s) y una evaporación directa de {base_evap_hm3:.1f} hm³ ({base_evap_m3s:.2f} m³/s).",
            f"La distribución de los aportes hídricos totales a la CHCP fue: subcuenca embalse Alhajuela {rpt_alh_pct:.1f}% ({rpt_alh_hm3:.1f} hm³) y subcuenca del embalse Gatún {rpt_gat_pct:.1f}% ({rpt_gat_hm3:.1f} hm³).",
            f"Los usos de agua del balance CHCP (incluida la evaporación directa de los embalses) durante {_per_lbl} fueron de {base_total_salidas_hm3:.1f} hm³ ({base_total_salidas_m3s:.1f} m³/s), sin incluir trasvase Madden→Gatún, hidrogeneración Madden ni MADSPILL por tratarse de movimientos internos/operativos.",
            f"El uso de agua en las esclusas del Canal de Panamá durante {_per_lbl} fue de {rpt_esclusaje_hm3:.1f} hm³. De este volumen, {rpt_pan_pct:.1f}% ({rpt_pan_hm3:.1f} hm³) correspondió a las esclusas Panamax con {rpt_tpan:.0f} tránsitos, y {rpt_neo_pct:.1f}% ({rpt_neo_hm3:.1f} hm³) a las Neopanamax con {rpt_tneo:.0f} tránsitos. El uso promedio diario de agua en las esclusas fue de {rpt_uso_prom_diario_hm3:.4f} hm³ ({rpt_uso_prom_diario_m3s:.2f} m³/s).",
            f"El uso de agua promedio por tránsito fue de {rpt_agua_pan_trans:.4f} hm³ (Panamax) y {rpt_agua_neo_trans:.4f} hm³ (Neopanamax).",
            f"Los niveles promedio de los embalses durante {_per_lbl}: Gatún {D('nivel_gatun_pies'):.2f} pies ({D('nivel_gatun_m'):.2f} m PLD) y Alhajuela {D('nivel_alh_pies'):.2f} pies, con calado máximo de {calado_neo} pies (Neopanamax) y {calado_pan:.2f} pies (Panamax).",
            f"En {_per_lbl}, se implementaron ahorros en las esclusas NeoPanamax ({rpt_ahorro_neo:.2f} hm³) y Panamax ({rpt_ahorro_pan:.2f} hm³), totalizando {rpt_ahorro_total:.2f} hm³ de ahorro, equivalente a {rpt_ahorro_lamina:.3f} pies en la lámina del embalse Gatún.",
            f"La salinidad promedio diaria en el cauce de navegación (SPC) del embalse Gatún durante {_per_lbl} fue de {sal_spc_x:.4f} ups.",
            f"El volumen destinado para la conservación de la calidad de agua (CCA) fue de {rpt_cca_hm3:.2f} hm³ en {rpt_tneo:.0f} tránsitos NeoPanamax, equivalente a {rpt_cca_unitario:.4f} hm³/tránsito.",
            f"La producción de energía eléctrica promedio durante {_per_lbl}: Madden {base_emw:.2f} MW ({base_emwh:.0f} MWh/día) y Gatún {base_egw:.4f} MW ({base_egwh:.2f} MWh/día), comparada con una demanda interna ACP de {base_demanda_acp:.1f} MW.",
            f"Durante {_per_lbl}, se trasvasaron {rpt_trasvase_total_hm3:.1f} hm³ ({rpt_trasvase_total_m3s:.2f} m³/s) de Alhajuela hacia Gatún, representando el {rpt_trasvase_pct:.1f}% del volumen de esclusajes.",
            f"En el {_per_noun}, el Canal de Panamá alcanzó un ahorro total de {rpt_ahorro_total:.2f} hm³ en sus esclusas, equivalente a {rpt_ahorro_lamina:.3f} pies en el embalse Gatún.",
            f"Durante el período, {_cond}; los aportes netos frente a las salidas requirieron {_accion} para sostener la operación y el abastecimiento.",
        ]

    # Filtrar bullets vacíos (evita viñetas huérfanas en la vista previa)
    bullets = [b for b in bullets if b and str(b).strip() and len(str(b).strip()) > 10]

    col_prev, col_exp = st.columns([3,2])
    with col_prev:
        st.markdown("#### Vista previa del informe")
        preview_title = oficial_doc.get("title", "Agua y Sostenibilidad") if use_official_doc else "Agua y Sostenibilidad"
        preview_subtitle = oficial_doc.get("subtitle", label_per) if use_official_doc else label_per
        preview_txt = f"**{preview_title}**\n**{preview_subtitle}**\n\n"
        for b in bullets:
            preview_txt += f"• {b}\n\n"
        st.text_area("", preview_txt, height=500, label_visibility="collapsed")

    with col_exp:
        st.markdown("#### Opciones de exportación")

        # ── DOCX (con python-docx si disponible, o XML nativo) ─────────────
        def build_docx_xml(title_str, subtitle_str, bullets_list, usos_rows,
                            fuentes_rows, extra_data=None):
            """
            Genera un DOCX fiel a la plantilla oficial Agua y Sostenibilidad.
            Compatible con Word 2010+, con formato homologado al informe oficial.
            """
            ed = extra_data or {}

            def _normalize_table_rows(_rows, kind="usos"):
                out = []
                for _row in (_rows or []):
                    if _row is None:
                        continue
                    try:
                        _vals = list(_row)
                    except Exception:
                        _vals = [_row]
                    if not _vals:
                        continue
                    _name = str(_vals[0]).strip() if len(_vals) > 0 else ""
                    _val_raw = _vals[1] if len(_vals) > 1 else None
                    _pct_raw = _vals[2] if len(_vals) > 2 else ""
                    _flag = bool(_vals[3]) if len(_vals) > 3 else False
                    _num = _to_float(_val_raw)
                    _val = float(_num) if _num is not None else (str(_val_raw) if _val_raw is not None else "")
                    _pct = str(_pct_raw).strip() if _pct_raw is not None else ""
                    out.append((_name, _val, _pct, _flag))
                return out

            fuentes_rows = _normalize_table_rows(fuentes_rows, "fuentes")
            usos_rows = _normalize_table_rows(usos_rows, "usos")

            # ── Infraestructura OOXML ─────────────────────────────────────────
            content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
<Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>
</Types>"""
            rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
            # Image relationship inserted dynamically if image present
            _rpt_img_bytes = ed.get("report_image_bytes")
            _rpt_img_ext   = ed.get("report_image_ext", "png")
            _img_mime      = "image/jpeg" if _rpt_img_ext in ("jpeg","jpg") else "image/png"
            _img_media     = f"media/rpt_image.{_rpt_img_ext}"
            word_rels_img_rel = (
                f'\n<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="{_img_media}"/>'
                if _rpt_img_bytes else ""
            )
            word_rels = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" Target="numbering.xml"/>{word_rels_img_rel}
</Relationships>"""

            numbering = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="&#8226;"/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>
      <w:rPr><w:rFonts w:ascii="Courier New" w:hAnsi="Courier New" w:cs="Courier New"/>
             <w:sz w:val="22"/></w:rPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
</w:numbering>"""

            styles = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
          xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">
  <w:docDefaults>
    <w:rPrDefault><w:rPr>
      <w:rFonts w:ascii="Calibri Light" w:hAnsi="Calibri Light" w:eastAsia="Calibri Light" w:cs="Calibri Light"/>
      <w:sz w:val="20"/><w:szCs w:val="20"/>
      <w:color w:val="003E69"/>
      <w:lang w:val="es-PA" w:eastAsia="es-PA" w:bidi="ar-SA"/>
    </w:rPr></w:rPrDefault>
  </w:docDefaults>
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:qFormat/>
    <w:pPr><w:spacing w:after="60"/></w:pPr>
    <w:rPr><w:rFonts w:ascii="Calibri Light" w:hAnsi="Calibri Light"/>
           <w:color w:val="003E69"/>
           <w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="ListParagraph">
    <w:name w:val="List Paragraph"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>
    <w:rPr><w:rFonts w:ascii="Calibri Light" w:hAnsi="Calibri Light"/>
           <w:color w:val="003E69"/>
           <w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading1">
    <w:name w:val="heading 1"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr><w:spacing w:before="280" w:after="80"/></w:pPr>
    <w:rPr><w:b/><w:bCs/>
      <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
      <w:sz w:val="28"/><w:szCs w:val="28"/>
      <w:color w:val="003E69"/>
    </w:rPr>
  </w:style>
  <w:style w:type="table" w:styleId="TableGrid">
    <w:name w:val="Table Grid"/>
    <w:tblPr>
      <w:tblBorders>
        <w:top    w:val="single" w:sz="4" w:space="0" w:color="C0CFDF"/>
        <w:left   w:val="single" w:sz="4" w:space="0" w:color="C0CFDF"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="C0CFDF"/>
        <w:right  w:val="single" w:sz="4" w:space="0" w:color="C0CFDF"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="C0CFDF"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="C0CFDF"/>
      </w:tblBorders>
    </w:tblPr>
    <w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
           <w:sz w:val="18"/><w:szCs w:val="18"/>
           <w:color w:val="1F3864"/></w:rPr>
  </w:style>
</w:styles>"""

            # ── Helpers XML ───────────────────────────────────────────────────
            NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

            def esc(s):
                return (str(s)
                    .replace("&","&amp;").replace("<","&lt;")
                    .replace(">","&gt;").replace('"',"&quot;")
                    .replace("³","&#179;").replace("²","&#178;")
                    .replace("·","&#183;").replace("→","&#8594;"))

            def rpr(bold=False, italic=False, size=20, color="003E69",
                    font="Calibri Light", highlight=None, underline=False):
                t = "<w:rPr>"
                t += f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}"/>'
                t += f'<w:sz w:val="{size}"/><w:szCs w:val="{size}"/>'
                if color and color not in ("003E69",):
                    t += f'<w:color w:val="{color}"/>'
                else:
                    t += f'<w:color w:val="003E69"/>'
                if bold:
                    t += "<w:b/><w:bCs/>"
                if italic:
                    t += "<w:i/><w:iCs/>"
                if underline:
                    t += '<w:u w:val="single"/>'
                if highlight:
                    t += f'<w:highlight w:val="{highlight}"/>'
                t += "</w:rPr>"
                return t

            def run(text, **kwargs):
                """Single text run."""
                return f'<w:r>{rpr(**kwargs)}<w:t xml:space="preserve">{esc(text)}</w:t></w:r>'

            def run_bold(text, **kwargs):
                return run(text, bold=True, **kwargs)

            def para_empty(spacing_after=80):
                return f'<w:p><w:pPr><w:spacing w:after="{spacing_after}"/></w:pPr></w:p>'

            def para_simple(text, bold=False, size=20, color="003E69",
                            align="left", spacing_after=60, spacing_before=0,
                            indent_left=0, keep_next=False, font="Calibri Light"):
                """Paragraph with a single run."""
                ppr = "<w:pPr>"
                if align != "left":
                    ppr += f'<w:jc w:val="{align}"/>'
                if spacing_before or spacing_after != 120:
                    ppr += f'<w:spacing w:before="{spacing_before}" w:after="{spacing_after}"/>'
                if indent_left:
                    ppr += f'<w:ind w:left="{indent_left}"/>'
                if keep_next:
                    ppr += '<w:keepNext/>'
                ppr += "</w:pPr>"
                r = run(text, bold=bold, size=size, color=color, font=font)
                return f"<w:p>{ppr}{r}</w:p>"

            def para_mixed(runs_list, align="both", spacing_after=60,
                           spacing_before=0, indent_left=0, num_id=None,
                           ilvl=0, keep_with_next=False):
                """
                Paragraph with multiple runs.
                runs_list = [(text, bold, size, color, highlight), ...]
                """
                ppr = "<w:pPr>"
                if num_id:
                    ppr += f'<w:numPr><w:ilvl w:val="{ilvl}"/><w:numId w:val="{num_id}"/></w:numPr>'
                if align != "left":
                    ppr += f'<w:jc w:val="{align}"/>'
                ppr += f'<w:spacing w:before="{spacing_before}" w:after="{spacing_after}"/>'
                if indent_left:
                    ppr += f'<w:ind w:left="{indent_left}"/>'
                if keep_with_next:
                    ppr += '<w:keepNext/>'
                ppr += "</w:pPr>"
                body_runs = ""
                for item in runs_list:
                    if isinstance(item, str):
                        body_runs += run(item)
                    elif len(item) == 2:
                        body_runs += run(item[0], bold=item[1])
                    elif len(item) == 4:
                        body_runs += run(item[0], bold=item[1], size=item[2], color=item[3])
                    elif len(item) == 5:
                        body_runs += run(item[0], bold=item[1], size=item[2],
                                         color=item[3], highlight=item[4])
                    else:
                        body_runs += run(str(item[0]))
                return f"<w:p>{ppr}{body_runs}</w:p>"

            def bullet(runs_list, spacing_after=100):
                """Bullet paragraph matching the template style."""
                return para_mixed(runs_list, num_id="1", spacing_after=spacing_after,
                                  spacing_before=0, align="both")

            def bullet_text(text, spacing_after=100):
                """Simple text bullet."""
                return bullet([(text, False, 20, "003E69", None)])

            # Table helpers
            def tbl_cell(text, bold=False, color="1F3864", bg=None,
                         size=18, align="center", width_pct=None, alt_row=False):
                tcpr = "<w:tcPr>"
                if width_pct:
                    w_twips = int(width_pct * 97.2 * 60)
                    tcpr += f'<w:tcW w:w="{w_twips}" w:type="dxa"/>'
                _fill = bg if bg else ("E8F0F8" if alt_row else "FFFFFF")
                tcpr += f'<w:shd w:val="clear" w:color="auto" w:fill="{_fill}"/>'
                tcpr += (
                    '<w:tcBorders>'
                    '<w:top w:val="single" w:sz="4" w:space="0" w:color="C0CFDF"/>'
                    '<w:left w:val="single" w:sz="4" w:space="0" w:color="C0CFDF"/>'
                    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="C0CFDF"/>'
                    '<w:right w:val="single" w:sz="4" w:space="0" w:color="C0CFDF"/>'
                    '</w:tcBorders>'
                    '<w:tcMar>'
                    '<w:top w:w="80" w:type="dxa"/><w:left w:w="108" w:type="dxa"/>'
                    '<w:bottom w:w="80" w:type="dxa"/><w:right w:w="108" w:type="dxa"/>'
                    '</w:tcMar>'
                    '<w:vAlign w:val="center"/>'
                )
                tcpr += "</w:tcPr>"
                txt_color = color
                ppr = f'<w:pPr><w:jc w:val="{align}"/><w:spacing w:after="40"/></w:pPr>'
                r_xml = f'<w:r>{rpr(bold=bold, size=size, color=txt_color, font="Calibri")}<w:t xml:space="preserve">{esc(str(text))}</w:t></w:r>'
                return f"<w:tc>{tcpr}<w:p>{ppr}{r_xml}</w:p></w:tc>"

            def tbl_row(cells_data, is_header=False, is_subtotal=False, is_total=False, row_index=0):
                bg       = "003E69" if (is_header or is_total) else ("D6E8F7" if is_subtotal else None)
                txt_color= "FFFFFF" if (is_header or is_total) else ("003E69" if is_subtotal else "1F3864")
                bold_row = is_header or is_subtotal or is_total
                alt      = (row_index % 2 == 1) and not is_header and not is_subtotal and not is_total
                cells_xml = ""
                for cell in cells_data:
                    if isinstance(cell, tuple):
                        txt   = cell[0]
                        bold_f= cell[1] if len(cell) > 1 else bold_row
                        cells_xml += tbl_cell(txt, bold=bold_f, color=txt_color, bg=bg, alt_row=alt)
                    else:
                        cells_xml += tbl_cell(cell, bold=bold_row, color=txt_color, bg=bg, alt_row=alt)
                return f"<w:tr>{cells_xml}</w:tr>"

            def make_table(rows, col_widths=None):
                """rows = list of dicts: {cells, header, subtotal, total}"""
                tbl_xml = '<w:tbl><w:tblPr><w:tblStyle w:val="TableGrid"/>'
                tbl_xml += '<w:tblW w:w="0" w:type="auto"/>'
                tbl_xml += '<w:jc w:val="center"/>'
                tbl_xml += '<w:tblLayout w:type="autofit"/>'
                tbl_xml += '<w:tblCellMar>'
                tbl_xml += '<w:top w:w="80" w:type="dxa"/><w:left w:w="108" w:type="dxa"/>'
                tbl_xml += '<w:bottom w:w="80" w:type="dxa"/><w:right w:w="108" w:type="dxa"/>'
                tbl_xml += '</w:tblCellMar></w:tblPr>'
                data_idx = 0
                for row in rows:
                    is_h = row.get("header", False)
                    is_s = row.get("subtotal", False)
                    is_t = row.get("total", False)
                    tbl_xml += tbl_row(
                        row["cells"],
                        is_header=is_h,
                        is_subtotal=is_s,
                        is_total=is_t,
                        row_index=data_idx if not is_h else 0,
                    )
                    if not is_h:
                        data_idx += 1
                tbl_xml += "</w:tbl>"
                return tbl_xml

            def section_heading(text):
                return para_simple(text, bold=True, size=28, color="003E69",
                                   spacing_before=280, spacing_after=80, keep_next=True, font="Calibri")

            def sub_heading(text):
                return para_simple(text, bold=True, size=22, color="003E69",
                                   spacing_before=180, spacing_after=60, keep_next=True, font="Calibri")

            def body_para(runs_list, spacing_after=120):
                return para_mixed(runs_list, spacing_after=spacing_after, align="both")

            # ── EED helper ────────────────────────────────────────────────────
            def hm3_to_eed(hm3, dias=31):
                return round(hm3_to_eed_period(hm3, dias), 2)

            def m3s_to_eed_local(m3s):
                return round(m3s_to_eed(m3s), 2)

            # Short aliases
            # Días del período: debe definirse antes de cualquier conversión hm³ → m³/s.
            # Antes se asignaba más abajo, pero se usaba primero en trav_total_m3s,
            # provocando: UnboundLocalError: cannot access local variable 'dias'.
            try:
                dias = float(ed.get("dias_mes", ed.get("_dias_periodo", 31)) or 31)
            except Exception:
                dias = 31.0
            if dias <= 0:
                dias = 31.0

            at   = ed.get("aporte_total_hm3", D("aporte_total_hm3"))
            atm  = ed.get("aporte_total_m3s", D("aporte_total_m3s"))
            an   = ed.get("aporte_neto_hm3",  D("aporte_neto_hm3"))
            anm  = ed.get("aporte_neto_m3s",  D("aporte_neto_m3s"))
            ev   = ed.get("evaporacion_hm3",  D("evaporacion_hm3"))
            evm  = ed.get("evaporacion_m3s",  D("evaporacion_m3s"))
            esc_hm3 = ed.get("esclusaje_hm3",    D("esclusaje_hm3"))
            escm = ed.get("esclusaje_m3s",    D("esclusaje_m3s"))
            pan  = ed.get("panamax_hm3",      D("panamax_hm3"))
            panp = ed.get("panamax_pct",      D("panamax_pct"))
            neo  = ed.get("neopanamax_hm3",   D("neopanamax_hm3"))
            neop = ed.get("neopanamax_pct",   D("neopanamax_pct"))
            pot  = ed.get("potabilizacion_hm3", D("potabilizacion_hm3"))
            potm = ed.get("potabilizacion_m3s", D("potabilizacion_m3s"))
            conc = ed.get("concesiones_hm3",  D("concesiones_hm3"))
            concm= ed.get("concesiones_m3s",  D("concesiones_m3s"))
            hidg = ed.get("hidro_gatun_hm3",  D("hidro_gatun_hm3"))
            hidgm= ed.get("hidro_gatun_m3s",  0.0)
            trav = ed.get("trasvase_hm3",     D("trasvase_hm3"))
            travm= ed.get("trasvase_m3s",     D("trasvase_m3s"))
            travp= ed.get("trasvase_pct_esc", D("trasvase_pct_esc"))
            # Trasvase total = MADMCF (G) + MADSPILL (R) + LEAK MAD (U).
            # trav ya contiene el total; no se debe volver a sumar componentes.
            trav_hidro = ed.get("trasvase_hidro_hm3", D("hidro_madden_hm3", 0))
            trav_vert  = ed.get("trasvase_vertidos_hm3", D("vertidos_mad_ops_hm3", 0))
            trav_fugas = ed.get("trasvase_fugas_hm3", D("trasvase_fugas_hm3", 0))
            trav_total = ed.get("trasvase_total_hm3", trav)
            trav_total_m3s = ed.get("trasvase_total_m3s", (round(trav_total / (dias * 0.0864), 2) if dias else 0))
            ngl  = ed.get("nivel_gatun_pies", D("nivel_gatun_pies"))
            nglm = ed.get("nivel_gatun_m",    D("nivel_gatun_m"))
            nalp = ed.get("nivel_alh_pies",   D("nivel_alh_pies"))
            tpan = ed.get("transitos_panamax",    D("transitos_panamax"))
            tneo = ed.get("transitos_neopanamax", D("transitos_neopanamax"))
            upd  = ed.get("uso_prom_diario_hm3",  D("uso_prom_diario_hm3"))
            updm = ed.get("uso_prom_diario_m3s",  D("uso_prom_diario_m3s"))
            wpt  = ed.get("agua_panamax_trans",   D("agua_panamax_trans"))
            aho  = ed.get("ahorro_total_hm3",     D("ahorro_total_hm3"))
            ahop = ed.get("ahorro_panamax_hm3",   D("ahorro_panamax_hm3"))
            ahon = ed.get("ahorro_neopanamax_hm3",D("ahorro_neopanamax_hm3"))
            ahol = ed.get("ahorro_lamina_pies",   D("ahorro_lamina_pies"))
            sal_spc = ed.get("sal_spc", 0.0)
            sal_spv = ed.get("sal_spv", 0.0)
            sal_dec = ed.get("sal_dec", 0.0)
            cca_hm3 = ed.get("cca_hm3", D("zzflush_auto_hm3", st.session_state.get("rep_cca_hm3", 0.0)))
            cca_u   = ed.get("cca_unitario", cca_hm3 / tneo if tneo > 0 else 0)
            emw     = ed.get("energia_madden_mw",      D("energia_madden_mw"))
            emwh    = ed.get("energia_madden_mwh_dia",  D("energia_madden_mwh_dia"))
            egw     = ed.get("energia_gatun_mw",       D("energia_gatun_mw"))
            egwh    = ed.get("energia_gatun_mwh_dia",  D("energia_gatun_mwh_dia"))
            dem_acp = ed.get("demanda_acp_mw", 19.0)
            precip  = ed.get("precipitacion_mm", 0.0)
            prec_h  = ed.get("precip_hist_mm", 0.0)
            prec_pct= ed.get("precip_pct_sobre",
                             abs(precip - prec_h) / prec_h * 100 if prec_h else 0)
            prec_dir= "por encima" if precip >= prec_h else "por debajo"
            prec_pos= ed.get("precip_posicion", 25)
            prec_rank_label_doc = ed.get("precip_rank_label", _meteo_rank_label_default(tipo, mes_sel))
            prec_per= ed.get("precip_periodo", "1950-2026")
            prec_dias= ed.get("precip_dias_max", "los días 20 y 21, cuando se acumularon 26 mm y 9 mm, respectivamente")
            # Procedencia de datos. El informe solo resaltará en naranja valores en 0 o incongruentes.
            _hist_manual = ed.get("hist_is_manual", False)
            _prec_manual = ed.get("precip_is_manual", False) or (precip == 0.0)
            def _needs_orange(text, flag=False):
                s = str(text).strip().lower()
                if not s:
                    return True
                if any(tok in s for tok in ["[", "]", "n/d", "nd", "nan", "none", "—"]):
                    return True
                nums = re.findall(r"-?\d+(?:\.\d+)?", s.replace(",", ""))
                if nums:
                    try:
                        vals = [float(x) for x in nums]
                        return all(abs(v) < 1e-12 for v in vals)
                    except Exception:
                        return bool(flag)
                return False
            def _y(text, flag=False):
                """Run normal o con alerta naranja si el valor está en 0 o luce inconsistente."""
                if _needs_orange(text, flag):
                    return (str(text), True, 20, "9A5A00", "darkYellow")
                return (str(text), False, 20, "003E69", None)
            alh_hm3 = ed.get("alhajuela_hm3", D("alhajuela_hm3"))
            alh_pct = ed.get("alhajuela_pct", D("alhajuela_pct"))
            gat_hm3 = ed.get("gatun_hm3",     D("gatun_hm3"))
            gat_pct = ed.get("gatun_pct",     D("gatun_pct"))
            cal_pan = ed.get("calado_panamax",    39.50)
            cal_neo = ed.get("calado_neopanamax", 50.0)
            pos_esc = ed.get("posicion_escala",   "")
            cls_esc = ed.get("clasificacion_escala", "")
            pct_sob = ed.get("pct_sobre_prom",    98.0)
            def _sum_usos_rows_hm3(_rows):
                total = 0.0
                for _r in (_rows or []):
                    try:
                        _name = str(_r[0]).strip().lower() if len(_r) > 0 else ""
                    except Exception:
                        _name = ""
                    if "total" in _name and "salida" in _name:
                        continue
                    _val = _to_float(_r[1] if len(_r) > 1 else None)
                    if _val is not None:
                        total += _val
                return total
            _tot_sal_ed = _to_float(ed.get("total_salidas_hm3", None))
            if _tot_sal_ed is not None:
                tot_sal = float(_tot_sal_ed)
            elif usos_rows:
                tot_sal = _sum_usos_rows_hm3(usos_rows)
            else:
                tot_sal = float(_to_float(D("total_salidas_hm3")) or 0.0)
            dias    = ed.get("dias_mes", 31)
            vertg   = ed.get("vertidos_gat_hm3", D("vertidos_gat_hm3", 0))
            excedente = at - tot_sal  # balance correcto: aporte_total − total_salidas (no neto, que doblaría evap)

            # EED values
            at_eed  = round(hm3_to_eed(at, dias), 1)
            ev_eed  = round(hm3_to_eed(ev, dias), 1)
            an_eed  = round(hm3_to_eed(an, dias), 1)
            esc_eed = round(hm3_to_eed(esc_hm3, dias), 1)
            pot_eed = round(hm3_to_eed(pot, dias), 2)
            conc_eed= round(hm3_to_eed(conc, dias), 2)
            hidg_eed= round(hm3_to_eed(hidg, dias), 3)
            ev_pct  = round(ev / tot_sal * 100, 1) if tot_sal > 0 else 0
            esc_pct = round(esc_hm3/ tot_sal * 100, 1) if tot_sal > 0 else 0
            pot_pct = round(pot/ tot_sal * 100, 1) if tot_sal > 0 else 0
            conc_pct= round(conc/tot_sal * 100, 1) if tot_sal > 0 else 0
            hidg_pct= round(hidg/tot_sal * 100, 1) if tot_sal > 0 else 0
            cca_pct = round(cca_hm3/tot_sal*100,1) if tot_sal > 0 else 0
            cca_m3s = round(cca_hm3/(dias*0.0864),3) if dias>0 else 0
            cca_eed = round(hm3_to_eed(cca_hm3, dias),3)

            # ── DOCUMENT BODY ─────────────────────────────────────────────────
            body = ""

            # ── TÍTULO ────────────────────────────────────────────────────────
            body += para_simple(title_str, bold=True, size=32,
                                color="003E69", align="left", spacing_after=60, font="Calibri")
            body += para_simple(subtitle_str, bold=True, size=18,
                                color="003E69", align="left", spacing_after=180, font="Calibri")

            # ── VIÑETAS PRINCIPALES ───────────────────────────────────────────
            if bullets_list:
                for _b in bullets_list:
                    if str(_b).strip():
                        body += bullet_text(str(_b))
            else:
                # 1. Precipitación — valores siempre manuales (no están en DAILY)
                body += bullet([
                    ("La precipitación total sobre la Cuenca Hidrográfica del Canal de Panamá (CHCP) "
                     f"durante {subtitle_str.lower()} fue de ", False, 20, "003E69", None),
                    _y(f"{precip:.0f} mm", _prec_manual),
                    (f", {prec_pct:.1f}% {prec_dir} del promedio histórico de ", False, 20, "003E69", None),
                    _y(f"{prec_h:.0f} mm", _hist_manual or _prec_manual),
                    (", colocándola en la posición ", False, 20, "003E69", None),
                    _y(str(prec_pos), _hist_manual),
                    (" de los ", False, 20, "003E69", None),
                    _y(prec_rank_label_doc if prec_rank_label_doc else "meses más húmedos", _hist_manual),
                    (" (", False, 20, "003E69", None),
                    _y(prec_per, _hist_manual),
                    (f")." + (f" La mayor parte de la lluvia del período se concentró {prec_dias}." if str(prec_dias).strip() else ""), False, 20, "003E69", None),
                ])
                # 2. Aportes totales (datos del DAILY — no amarillo)
                body += bullet_text(
                    f"Los aportes totales a la Cuenca Hidrográfica del Canal de Panamá (CHCP) durante "
                    f"{subtitle_str.lower()} fueron {at:.1f} hm³ ({atm:.1f} m³/s), desglosados en "
                    f"{an:.1f} hm³ ({anm:.1f} m³/s) de aportes netos y {ev:.1f} hm³ "
                    f"({evm:.1f} m³/s) de evaporación directa en los embalses."
                )
                # 3. Posición histórica — valores históricos pueden ser manuales
                _pct_sob_txt, _pct_sob_dir = _pct_rel_text(pct_sob, positive_text="por encima al", negative_text="por debajo del", zero_text="igual al", decimals=0)
                _pct_neto_doc = ((an - float(ed.get('aporte_neto_hist_hm3', 0.0) or 0.0)) / float(ed.get('aporte_neto_hist_hm3', 0.0) or 1.0) * 100) if float(ed.get('aporte_neto_hist_hm3', 0.0) or 0.0) > 0 else 0
                _pct_neto_txt, _pct_neto_dir = _pct_rel_text(_pct_neto_doc, positive_text="por arriba del", negative_text="por debajo del", zero_text="igual al", decimals=0)
                _hist_periodo_doc = str(ed.get('hist_years_str', '') or '1898-2025')
                _hist_years_doc = _anios_registro_incluyendo_actual(_hist_periodo_doc, None, ed.get('h_n_years', 129))
                _pos_esc_doc = _posicion_historica_texto(pos_esc, 'Mensual' if ' de ' in subtitle_str.lower() else None, subtitle_str.split()[0] if subtitle_str else '')
                body += bullet([
                    (f"Los aportes totales a la CHCP ({at:.0f} hm³) de {subtitle_str.lower()}, se encuentra como ", False, 20, "003E69", None),
                    _y(_pos_esc_doc if _pos_esc_doc else "—", _hist_manual),
                    (" en ", False, 20, "003E69", None),
                    _y(f"{_hist_years_doc} años", _hist_manual),
                    (" de registro continuo, en una escala de húmedo a seco para el periodo ", False, 20, "003E69", None),
                    _y(_hist_periodo_doc, _hist_manual),
                    (f". Los cuales se encuentran en un ", False, 20, "003E69", None),
                    _y(f"{_pct_sob_txt} %", _hist_manual),
                    (f" {_pct_sob_dir} promedio histórico de ", False, 20, "003E69", None),
                    _y(f"{ed.get('aporte_hist_hm3', 0.0):.0f} hm³ ({ed.get('aporte_hist_m3s', 0.0):.1f} m³/s)", _hist_manual),
                    (".", False, 20, "003E69", None),
                ])
                body += bullet([
                    (f"Los aportes netos de {subtitle_str.lower()} fueron igual a {an:.1f} hm³ ({anm:.1f} m³/s), se ubican en la posición ", False, 20, "003E69", None),
                    _y(str(ed.get('posicion_escala_neto', '—') or '—'), _hist_manual),
                    (" dentro de la escala histórica de húmedo a seco", False, 20, "003E69", None),
                    ((f" ({ed.get('clasificacion_escala_neto', '')})" if ed.get('clasificacion_escala_neto', '') else ""), False, 20, "003E69", None),
                    (" para el periodo ", False, 20, "003E69", None),
                    _y(str(ed.get('hist_neto_years_str', '') or '1914-2025'), _hist_manual),
                    (", y se encuentran en ", False, 20, "003E69", None),
                    _y(f"{_pct_neto_txt}%", _hist_manual),
                    (f" {_pct_neto_dir} promedio histórico ", False, 20, "003E69", None),
                    _y(f"{ed.get('aporte_neto_hist_hm3', 0.0):.1f} hm³ ({ed.get('aporte_neto_hist_m3s', 0.0):.2f} m³/s)", _hist_manual),
                    (".", False, 20, "003E69", None),
                ])
                # 4. Distribución subcuencas
                body += bullet_text(
                    f"En {subtitle_str.lower()} la distribución de los aportes hídricos totales a la CHCP fue "
                    f"de la siguiente manera: Subcuenca embalse Alhajuela el {alh_pct:.1f}% ({alh_hm3:.0f} hm³), "
                    f"Subcuenca del embalse Gatún el {gat_pct:.1f}% ({gat_hm3:.0f} hm³)."
                )
                # 5. Niveles embalses
                body += bullet_text(
                    f"Los niveles de los embalses Gatún y Alhajuela presentan suficiente agua disponible para "
                    f"el consumo humano e industrial. El nivel de agua promedio en el embalse Gatún durante este "
                    f"mes fue de {ngl:.2f} pies ({nglm:.2f} m) PLD. Esta condición ha permitido mantener el "
                    f"calado máximo permisible de {cal_pan:.2f} pies para los buques Panama y {cal_neo:.0f} pies "
                    f"para los buques Neopanamax."
                )
                # 6. Uso esclusas
                body += bullet_text(
                    f"El uso de agua en las esclusas del Canal de Panamá ({subtitle_str.lower()}) ha sido de "
                    f"{esc_hm3:.2f} hm³. De este volumen, {panp:.2f}% ({pan:.2f} hm³) corresponde a las "
                    f"esclusas Panamax y {neop:.2f}% ({neo:.2f} hm³) a las Neopanamax."
                )
                # 7. Uso diario promedio
                body += bullet_text(
                    f"El uso promedio diario de agua en las esclusas fue de {upd:.4f} hm³ ({updm:.2f} m³/s)."
                )
                # 8. Ahorros esclusas
                body += bullet_text(
                    f"Durante el mes, se implementaron ahorros en las esclusas NeoPanamax (cambio de dirección) "
                    f"representando un volumen de {ahon:.2f} hm³ y Panamax (cámara corta en Pedro Miguel) "
                    f"representando un volumen de {ahop:.2f} hm³, para un total de {aho:.2f} hm³, "
                    f"lo cual en términos de lámina del embalse Gatún representa {ahol:.3f} pies."
                )
                # 9. Agua por tránsito Panamax
                body += bullet_text(
                    f"El uso de agua promedio por tránsito de un buque panamax fue de {wpt:.4f} hm³."
                )
                # 10. Salinidad
                if sal_spc > 0:
                    body += bullet_text(
                        f"Durante {subtitle_str.lower()}, la salinidad promedio diaria en el cauce de navegación "
                        f"(SPC) del embalse Gatún fue en promedio {sal_spc:.4f} unidades prácticas de salinidad "
                        f"(ups). Al cierre del mes, la salinidad promedio ponderada (SPV) se registró en "
                        f"{sal_spv:.4f} ups, representando un decremento del {sal_dec:.1f}% respecto al máximo "
                        f"histórico alcanzado."
                    )
                # 11. CCA
                if tneo > 0:
                    body += bullet_text(
                        f"El volumen de agua destinado para la conservación de la calidad de agua (CCA) fue de "
                        f"{cca_hm3:.2f} hm³, para {tneo:.0f} tránsitos totales NeoPanamax durante el mes, "
                        f"equivalente a un CCA unitario de {cca_u:.4f} hm³/tránsito."
                    )
                # 12. Energía
                body += bullet_text(
                    f"La producción de energía eléctrica en promedio en la hidroeléctrica Madden fue de "
                    f"{emw:.2f} MW (equivalente a {emwh:.0f} MWh por día). La producción en la hidroeléctrica "
                    f"Gatún fue de {egw:.4f} MW (equivalente a {egwh:.2f} MWh por día). "
                    f"{'Se logró cubrir la demanda interna de energía de la ACP (' + str(dem_acp) + ' MW promedio), sin utilizar la planta termoeléctrica o importar energía del mercado.' if emw >= dem_acp else 'No se logró cubrir la demanda interna de energía de la ACP (' + str(dem_acp) + ' MW promedio), requiriendo uso de termoeléctrica o importación de energía.'}"
                )
                # 13. Trasvase
                body += bullet_text(
                    f"Durante este mes, se trasvasaron {trav_total:.2f} hm³ ({trav_total_m3s:.2f} m³/s) de agua desde "
                    f"Alhajuela hacia Gatún ({float(trav_hidro or 0):.2f} hm³ por hidrogeneración, {float(trav_vert or 0):.2f} hm³ "
                    f"por vertidos y {float(trav_fugas or 0):.2f} hm³ por fugas), lo que representó un {(trav_total / esc_hm3 * 100) if esc_hm3 > 0 else 0:.1f}% "
                    f"del volumen mensual utilizado por las esclusas ({esc_hm3:.2f} hm³). Esta estrategia optimiza la gestión del recurso hídrico "
                    f"al mejorar la calidad del agua en el embalse Alhajuela, reduciendo su tiempo de residencia hidráulica "
                    f"(TRH) y potenciando la eficiencia operativa."
                )

            # ── BALANCE HÍDRICO ───────────────────────────────────────────────
            # Imagen del Tablero de Mando — se inserta ANTES del balance, igual que en el informe oficial
            if _rpt_img_bytes:
                _ct_ext  = "jpeg" if _rpt_img_ext in ("jpeg","jpg") else "png"
                _ct_mime = f"image/{_ct_ext}"
                content_types = content_types.replace(
                    "</Types>",
                    f'<Default Extension="{_ct_ext}" ContentType="{_ct_mime}"/>\n</Types>'
                )
                _img_cx = 5486400   # 6 pulgadas de ancho en EMU
                _img_cy = 3200400
                try:
                    import struct as _s
                    if _rpt_img_ext == "png":
                        _iw, _ih = _s.unpack(">II", _rpt_img_bytes[16:24])
                    else:
                        _ii = 2
                        while _ii < len(_rpt_img_bytes):
                            if _rpt_img_bytes[_ii] != 0xFF: break
                            _mm = _rpt_img_bytes[_ii+1]
                            if _mm in (0xC0, 0xC2):
                                _ih, _iw = _s.unpack(">HH", _rpt_img_bytes[_ii+5:_ii+9])
                                break
                            _ll = _s.unpack(">H", _rpt_img_bytes[_ii+2:_ii+4])[0]
                            _ii += 2 + _ll
                    if _iw > 0:
                        _img_cy = int(_img_cx * (_ih / _iw))
                except Exception:
                    pass
                body += para_empty(160)
                body += f"""<w:p>
  <w:pPr><w:jc w:val="center"/><w:spacing w:before="160" w:after="160"/></w:pPr>
  <w:r><w:drawing>
    <wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
      <wp:extent cx="{_img_cx}" cy="{_img_cy}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:docPr id="1" name="Grafica_informe"/>
      <a:graphic>
        <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <pic:pic>
            <pic:nvPicPr>
              <pic:cNvPr id="0" name="Grafica_informe"/>
              <pic:cNvPicPr/>
            </pic:nvPicPr>
            <pic:blipFill>
              <a:blip r:embed="rId3"/>
              <a:stretch><a:fillRect/></a:stretch>
            </pic:blipFill>
            <pic:spPr>
              <a:xfrm><a:off x="0" y="0"/><a:ext cx="{_img_cx}" cy="{_img_cy}"/></a:xfrm>
              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
            </pic:spPr>
          </pic:pic>
        </a:graphicData>
      </a:graphic>
    </wp:inline>
  </w:drawing></w:r>
</w:p>"""
            body += para_empty(160)
            body += section_heading(f"Balance hídrico {subtitle_str}")

            # Texto balance — párrafos como en la plantilla
            at_eed_s  = sig3(at_eed)
            ev_eed_s  = sig3(ev_eed)
            an_eed_s  = sig3(an_eed)
            body += body_para([
                (f"El aporte total a la Cuenca Hidrográfica del Canal de Panamá (CHCP) en "
                 f"{subtitle_str.lower()}, fue de {sig3(at)} hm³ ({sig3(atm)} m³/s = "
                 f"{at_eed_s} EED). Se evaporó directamente desde la superficie de los embalses "
                 f"Gatún y Alhajuela un volumen de {sig3(ev)} hm³ ({sig3(evm)} m³/s = "
                 f"{ev_eed_s} EED) de agua resultando un aporte neto de {sig3(an)} hm³ "
                 f"({sig3(anm)} m³/s = {an_eed_s} EED).",
                 False, 22, "003E69", None)
            ], spacing_after=120)

            body += body_para([
                ("Los usos o salidas de agua, incluida la evaporación, fueron de ", False, 20, "003E69", None),
                (f"{sig3(tot_sal)} hm³", True, 20, "1F3864", None),
                (f" ({sig3(round(tot_sal/(dias*0.0864),3))} m³/s = {sig3(round(hm3_to_eed(tot_sal, dias),3))} EED), "
                 "desglosados de la siguiente manera:", False, 20, "003E69", None),
            ], spacing_after=100)

            # Viñetas balance
            body += bullet([
                ("Las operaciones de esclusajes consumieron: ", False, 20, "003E69", None),
                (f"{sig3(esc_hm3)} hm³", True, 20, "1F3864", None),
                (f" ({sig3(escm)} m³/s = {sig3(esc_eed)} EED), lo que equivale al "
                 f"{sig3(esc_pct)}%", False, 20, "003E69", None),
                (" del total de las salidas de agua.", False, 20, "003E69", None),
            ])
            body += bullet([
                ("El agua extraída para potabilización (IDAAN, ACP y APSA) fue ", False, 20, "003E69", None),
                (f"{sig3(pot)} hm³", True, 20, "1F3864", None),
                (f" ({sig3(potm)} m³/s = {sig3(pot_eed)} EED), lo que equivale al ", False, 20, "003E69", None),
                (f"{sig3(pot_pct)}%", False, 20, "003E69", None),
                (" del total de las salidas.", False, 20, "003E69", None),
            ])
            body += bullet([
                ("El volumen de agua invertido en las esclusas Neopanamax para mitigar la intrusión "
                 "salina (ZZFlush), fue de ", False, 20, "003E69", None),
                (f"{sig3(cca_hm3)} hm³", True, 20, "1F3864", None),
                (f" ({sig3(cca_m3s)} m³/s = {cca_eed:.3f} EED), el {sig3(cca_pct)}% de las salidas.",
                 False, 22, "003E69", None),
            ])
            body += bullet([
                ("Las concesiones, misceláneos, fugas y filtraciones fueron de ", False, 20, "003E69", None),
                (f"{sig3(conc)} hm³", True, 20, "1F3864", None),
                (f" ({sig3(concm)} m³/s = {sig3(conc_eed)} EED), el ", False, 20, "003E69", None),
                (f"{sig3(conc_pct)}%", False, 20, "003E69", None),
                (" de las salidas.", False, 20, "003E69", None),
            ])
            body += bullet([
                ("El volumen de agua evaporada de los embalses Gatún y Alhajuela, fue ", False, 20, "003E69", None),
                (f"{sig3(ev)} hm³", True, 20, "1F3864", None),
                (f" ({sig3(evm)} m³/s = {sig3(ev_eed)} EED), el ", False, 20, "003E69", None),
                (f"{sig3(ev_pct)}%", False, 20, "003E69", None),
                (" de las salidas.", False, 20, "003E69", None),
            ])
            body += bullet([
                ("Producción de energía en la hidroeléctrica Gatún fue ", False, 20, "003E69", None),
                (f"{sig3(hidg)} hm³", True, 20, "1F3864", None),
                (f" ({hidg_eed:.3f} EED), el {sig3(hidg_pct)}% de las salidas.", False, 20, "003E69", None),
            ])

            # Párrafo balance final
            if excedente < 0:
                balance_txt = (
                    f"Durante el mes de {subtitle_str.lower()}, no hubo excedentes, los aportes netos "
                    f"fueron menores que las salidas ({sig3(abs(excedente))} hm³ de déficit), "
                    f"requiriendo el uso del agua almacenada en ambos embalses (Gatún y Alhajuela) "
                    f"y abatimiento de niveles de agua."
                )
            else:
                balance_txt = (
                    f"Durante el mes de {subtitle_str.lower()}, los aportes netos superaron las salidas "
                    f"en {sig3(excedente)} hm³, permitiendo el almacenamiento de agua en los embalses "
                    f"Gatún y Alhajuela."
                )
            body += para_empty(80)
            body += body_para([(balance_txt, False, 20, "003E69", None)], spacing_after=200)

            # ── FUENTES DE AGUA Y USOS ────────────────────────────────────────
            body += section_heading("Fuentes de Agua y Usos")

            # Tabla fuentes
            fuentes_tbl_rows = [{"cells": ["Fuente (tributario)",
                                           "Cantidad de agua (hm³)",
                                           "% del total"], "header": True}]
            for row in fuentes_rows:
                is_sub   = len(row) > 3 and row[3] is True
                is_total = str(row[0]).upper().startswith("APORTE TOTAL")
                fuentes_tbl_rows.append({
                    "cells": [(row[0], is_sub or is_total),
                              (sig3(row[1]) if isinstance(row[1], float) else str(row[1]), is_sub or is_total),
                              (str(row[2]), is_sub or is_total)],
                    "subtotal": is_sub,
                    "total": is_total,
                })
            body += make_table(fuentes_tbl_rows)

            body += para_empty(60)
            body += body_para([
                ("*Embalse Alhajuela y embalse Gatún como área tributaria se refiere a la lluvia "
                 "directa sobre el espejo de agua de estos embalses.", False, 18, "5A7A9A", None)
            ], spacing_after=160)

            # Tabla usos
            body += sub_heading(f"Usos de agua {subtitle_str}")
            usos_tbl_rows = [{"cells": ["Tipo de uso",
                                        "Cantidad de agua (hm³)",
                                        "% del total"], "header": True}]
            for row in usos_rows:
                is_total = str(row[0]).upper().startswith("TOTAL")
                # Highlight values for esclusaje, potabilización, etc.
                usos_tbl_rows.append({
                    "cells": [(row[0], is_total),
                              (sig3(row[1]) if isinstance(row[1], float) else str(row[1]), is_total),
                              (str(row[2]), is_total)],
                    "total": is_total,
                })
            body += make_table(usos_tbl_rows)

            # ── OOXML ASSEMBLY ────────────────────────────────────────────────
            doc_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
  xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">
<w:body>
{body}
<w:sectPr>
  <w:pgSz w:w="12240" w:h="15840"/>
  <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"
           w:header="720" w:footer="720" w:gutter="0"/>
</w:sectPr>
</w:body></w:document>"""

            buf = io.BytesIO()
            with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
                zf.writestr("[Content_Types].xml", content_types)
                zf.writestr("_rels/.rels", rels)
                zf.writestr("word/_rels/document.xml.rels", word_rels)
                zf.writestr("word/document.xml", doc_xml)
                zf.writestr("word/styles.xml",   styles)
                zf.writestr("word/numbering.xml", numbering)
                if _rpt_img_bytes:
                    zf.writestr(f"word/{_img_media}", _rpt_img_bytes)
            buf.seek(0)
            return buf.read()
        # Construir datos para tablas
        # Build fuentes rows from actual loaded data
        fuentes_rows_exp = []
        for r in fuentes_rios_alh:
            fuentes_rows_exp.append((r["nombre"], round(r["hm3"],2), f"{r['pct']:.1f} %", False))
        fuentes_rows_exp.append(("Subtotal Alhajuela", round(rpt_alh_hm3,2),
                                  f"{rpt_alh_pct:.1f} %", True))
        for r in fuentes_rios_gat:
            fuentes_rows_exp.append((r["nombre"], round(r["hm3"],2), f"{r['pct']:.1f} %", False))
        fuentes_rows_exp.append(("Subtotal Gatún", round(rpt_gat_hm3,2),
                                  f"{rpt_gat_pct:.1f} %", True))
        fuentes_rows_exp.append(("APORTE TOTAL CHCP", round(rpt_aporte_total,2), "100.0 %", True))

        # Tabla final de usos: consolidada en las mismas categorías del informe oficial
        _zzflush = float(rpt_zzflush_hm3 or 0)
        _vertidos = float(D("vertidos_gat_hm3") or 0)
        _esclusaje = float(rpt_esclusaje_hm3 or 0)
        _pot = float(base_pot_hm3 or 0)
        _conc = float(base_conc_hm3 or 0)
        _evap = float(base_evap_hm3 or 0)
        _hidg = float(base_hidg_hm3 or 0)
        _total_s = float(rpt_total_salidas or 0)

        def _fmt_hm3_uso(x, total=False):
            return sig3(x)

        def _fmt_pct_uso(x, total=False):
            if total:
                return "100%"
            x = float(x or 0)
            if x == 0:
                return "0.00%"
            return f"{sig3(x)}%"

        usos_rows_exp = [
            ("Esclusaje", _fmt_hm3_uso(_esclusaje),
             _fmt_pct_uso((_esclusaje / _total_s * 100) if _total_s > 0 else 0)),
            ("Potabilización", _fmt_hm3_uso(_pot),
             _fmt_pct_uso((_pot / _total_s * 100) if _total_s > 0 else 0)),
            ("Mitigación de salinidad en esclusas neopanamax (ZZFlush)", _fmt_hm3_uso(_zzflush),
             _fmt_pct_uso((_zzflush / _total_s * 100) if _total_s > 0 else 0)),
            ("Concesiones, fugas, filtraciones, misceláneos", _fmt_hm3_uso(_conc),
             _fmt_pct_uso((_conc / _total_s * 100) if _total_s > 0 else 0)),
            ("Evaporación", _fmt_hm3_uso(_evap),
             _fmt_pct_uso((_evap / _total_s * 100) if _total_s > 0 else 0)),
            ("Hidrogeneración Gatún", _fmt_hm3_uso(_hidg),
             _fmt_pct_uso((_hidg / _total_s * 100) if _total_s > 0 else 0)),
            ("Vertidos", _fmt_hm3_uso(_vertidos),
             _fmt_pct_uso((_vertidos / _total_s * 100) if _total_s > 0 else 0)),
            ("Total de salidas", _fmt_hm3_uso(_total_s, total=True), "100%"),
        ]
        if use_official_doc and oficial_doc.get("fuentes_rows"):
            fuentes_rows_exp = [tuple(r) for r in oficial_doc.get("fuentes_rows", fuentes_rows_exp)]
        if use_official_doc and oficial_doc.get("usos_rows"):
            usos_rows_exp = [tuple(r[:3]) for r in oficial_doc.get("usos_rows", [])]

        label_file = _safe_label_file(tipo, mes_sel, int(anio_sel), trimestre, semestre)

        st.markdown("##### 📄 Informe DOCX")
        st.markdown("""
        <div class='success-box'>
        ✅ Exportación DOCX disponible <b>sin necesidad de python-docx</b> — 
        XML nativo compatible con Word 2010+
        </div>""", unsafe_allow_html=True)

        docx_bytes = build_docx_xml(
            (oficial_doc.get("title") if use_official_doc and oficial_doc.get("title") else "Agua y sostenibilidad"),
            (oficial_doc.get("subtitle") if use_official_doc and oficial_doc.get("subtitle") else label_per),
            bullets,
            usos_rows_exp,
            fuentes_rows_exp,
            extra_data={
                # Precipitación
                "precipitacion_mm":   precip_mm,
                "precip_hist_mm":     precip_hist,
                "precip_pct_sobre":   float(st.session_state.get("x_met_pct_informe", abs(float(precip_mm - precip_hist) / precip_hist * 100) if precip_hist > 0 else 0.0)),
                "precip_posicion":    int(precip_pos),
                "precip_rank_label":   precip_rank_label,
                "precip_periodo":     precip_per,
                "precip_dias_max":    st.session_state.get("x_prec_dias",
                    "los días 20 y 21, cuando se acumularon 26 mm y 9 mm, respectivamente"),
                # Comparativa histórica
                "posicion_escala":    pos_escala,
                "clasificacion_escala": _hist_scale_class_auto,
                "pct_sobre_prom":     pct_sobre_prom,
                "aporte_hist_hm3":    st.session_state.get("x_hist_hm3", 0.0),
                "aporte_hist_m3s":    st.session_state.get("x_hist_m3s", 0.0),
                "posicion_escala_neto": str(datos.get("hist_neto_scale_pos", "") or ""),
                "clasificacion_escala_neto": str(datos.get("hist_neto_scale_class", "") or ""),
                "aporte_neto_hist_hm3": float(hist.get("hist_neto_hm3", 0.0) or 0.0),
                "aporte_neto_hist_m3s": float(hist.get("hist_neto_m3s", 0.0) or 0.0),
                "hist_neto_n_years":  int(hist.get("n_years_netfl", 0) or 0),
                "hist_neto_years_str": (f"{hist['years_netfl'][0]}-{hist['years_netfl'][1]}"
                                        if hist.get("years_netfl") else ""),
                # Calados
                "calado_panamax":     calado_pan,
                "calado_neopanamax":  calado_neo,
                # Salinidad
                "sal_spc":            sal_spc_x,
                "sal_spv":            sal_spv_x,
                "sal_dec":            sal_dec_x,
                # CCA (ZZFlush) — ahora auto desde GATSPILL
                "cca_hm3":            rpt_cca_hm3,
                # Energía
                "demanda_acp_mw":     base_demanda_acp,
                # Días del mes
                "dias_mes":           n_dias if n_dias > 0 else 31,
                # Ajustes finos del informe
                "aporte_total_hm3":   rpt_aporte_total,
                "aporte_total_m3s":   rpt_aporte_total_m3s,
                "aporte_neto_hm3":    rpt_aporte_neto,
                "aporte_neto_m3s":    rpt_aporte_neto_m3s,
                "alhajuela_hm3":      rpt_alh_hm3,
                "alhajuela_pct":      rpt_alh_pct,
                "gatun_hm3":          rpt_gat_hm3,
                "gatun_pct":          rpt_gat_pct,
                "esclusaje_hm3":      rpt_esclusaje_hm3,
                "esclusaje_m3s":      rpt_uso_prom_diario_m3s,
                "panamax_hm3":        rpt_pan_hm3,
                "panamax_pct":        rpt_pan_pct,
                "neopanamax_hm3":     rpt_neo_hm3,
                "neopanamax_pct":     rpt_neo_pct,
                "transitos_panamax":  rpt_tpan,
                "transitos_neopanamax": rpt_tneo,
                "uso_prom_diario_hm3": rpt_uso_prom_diario_hm3,
                "uso_prom_diario_m3s": rpt_uso_prom_diario_m3s,
                "agua_panamax_trans": rpt_agua_pan_trans,
                "agua_neo_trans":     rpt_agua_neo_trans,
                "potabilizacion_hm3": base_pot_hm3,
                "potabilizacion_m3s": base_pot_m3s,
                "potabilizacion_gat_hm3": base_pot_gat_hm3,
                "potabilizacion_alh_hm3": base_pot_alh_hm3,
                "concesiones_hm3": base_conc_hm3,
                "concesiones_m3s": base_conc_m3s,
                "evaporacion_hm3": base_evap_hm3,
                "evaporacion_m3s": base_evap_m3s,
                "hidro_gatun_hm3": base_hidg_hm3,
                "hidro_gatun_m3s": base_hidg_m3s,
                "energia_madden_mw": base_emw,
                "energia_madden_mwh_dia": base_emwh,
                "energia_gatun_mw": base_egw,
                "energia_gatun_mwh_dia": base_egwh,
                "total_salidas_hm3": rpt_total_salidas,
                "ahorro_panamax_hm3": rpt_ahorro_pan,
                "ahorro_neopanamax_hm3": rpt_ahorro_neo,
                "ahorro_total_hm3":   rpt_ahorro_total,
                "ahorro_lamina_pies": rpt_ahorro_lamina,
                "trasvase_hm3":       rpt_trasvase_hm3,
                "trasvase_m3s":       rpt_trasvase_total_m3s,
                "trasvase_pct_esc":   rpt_trasvase_pct,
                "trasvase_hidro_hm3": rpt_trasvase_hidro_hm3,
                "trasvase_vertidos_hm3": rpt_trasvase_vert_hm3,
                "trasvase_fugas_hm3": rpt_trasvase_fugas_hm3,
                "trasvase_total_hm3": rpt_trasvase_total_hm3,
                "trasvase_total_m3s": rpt_trasvase_total_m3s,
                "cca_hm3":            rpt_cca_hm3,
                # Flag informativo de procedencia; el DOCX solo marca en naranja ceros/incongruencias
                "hist_is_manual":     hist.get("_hist_is_manual", True),
                "precip_is_manual":   (precip_mm == 0.0),
                "h_n_years":          int(_hist_years_chcp or st.session_state.get("h_n_years", 0) or 0),
                "hist_years_str":     _hist_periodo_chcp,
                # Imagen del informe
                "report_image_bytes": _img_bytes,
                "report_image_ext":   _img_ext,
                "rocc_texto":         str(st.session_state.get("rocc_texto", "") or ""),
            }
        )
        st.download_button(
            "📄 Descargar Informe DOCX",
            data=docx_bytes,
            file_name=f"Agua y sostenibilidad {label_file.replace('_', ' ')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True)

        # Excel export
        st.markdown("<br>##### 📊 Datos en Excel")
        def gen_excel():
            buf2 = io.BytesIO()
            with pd.ExcelWriter(buf2, engine="openpyxl") as wr:
                periodo_txt = label_per
                df_res = pd.DataFrame({
                    "Variable":["Período","Precipitación (mm)","Aporte Total (hm³)","Aporte Neto (hm³)",
                                "Evaporación (hm³)","Nivel Gatún (pies)","Nivel Alhajuela (pies)",
                                "Esclusaje Total (hm³)","Panamax (hm³)","Panamax %",
                                "Neopanamax (hm³)","Neopanamax %",
                                "Tránsitos Panamax","Tránsitos Neopanamax",
                                "Agua/tránsito Panamax (hm³)","Agua/tránsito Neopanamax (hm³)",
                                "Uso diario prom. esclusas (hm³)",
                                "Potabilización (hm³)","Potabilización Gatún (hm³)","Potabilización Alh. (hm³)",
                                "Concesiones (hm³)","ZZFlush auto GATSPILL (hm³)",
                                "Hidrogeneración Madden (MW)","Hidrogeneración Madden (MWh/día)",
                                "Hidrogeneración Gatún (MW)","Hidrogeneración Gatún (hm³)",
                                "Hidrogeneración Madden (hm³)",
                                "Salinidad SPC (ups)","Trasvase (hm³)","Trasvase (m³/s)",
                                "Ahorro Panamax (hm³)","Ahorro Neopanamax (hm³)","Ahorro Total (hm³)",
                                "Ahorro lámina Gatún (pies)","Desc. Ops. Madden MADSPILL (hm³)"],
                    "Valor":[periodo_txt,
                             precip_mm, D("aporte_total_hm3"), D("aporte_neto_hm3"),
                             D("evaporacion_hm3"), D("nivel_gatun_pies"), D("nivel_alh_pies"),
                             D("esclusaje_hm3"), D("panamax_hm3"), D("panamax_pct"),
                             D("neopanamax_hm3"), D("neopanamax_pct"),
                             D("transitos_panamax"), D("transitos_neopanamax"),
                             D("agua_panamax_trans"), D("agua_neo_trans"),
                             D("uso_prom_diario_hm3"),
                             D("potabilizacion_hm3"), D("potabilizacion_gat_hm3"), D("potabilizacion_alh_hm3"),
                             D("concesiones_hm3"), D("zzflush_auto_hm3"),
                             D("energia_madden_mw"), D("energia_madden_mwh_dia"),
                             D("energia_gatun_mw"), D("hidro_gatun_hm3"),
                             D("hidro_madden_hm3"),
                             sal_spc_x, D("trasvase_hm3"), D("trasvase_m3s"),
                             D("ahorro_panamax_hm3"), D("ahorro_neopanamax_hm3"), D("ahorro_total_hm3"),
                             D("ahorro_lamina_pies"), D("vertidos_mad_ops_hm3")],
                })
                df_res.to_excel(wr, sheet_name="Resumen", index=False)
                pd.DataFrame(fuentes_rows_exp, columns=["Fuente (tributario)","hm³","%","Subtotal"]).to_excel(
                    wr, sheet_name="Fuentes de Agua", index=False)
                pd.DataFrame(usos_rows_exp, columns=["Uso","hm³","%"]).to_excel(
                    wr, sheet_name="Usos de Agua", index=False)
            buf2.seek(0)
            return buf2.read()

        st.download_button(
            "📊 Descargar Datos Excel",
            data=gen_excel(),
            file_name=f"Datos_Agua_{label_file}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True)

# ── Footer ─────────────────────────────────────────────────────────────────
st.markdown("""
<div class="footer">
♻️ Agua y Sostenibilidad · Autoridad del Canal de Panamá &nbsp;|&nbsp;
HIMH - Hidrología &nbsp;|&nbsp; app_AyS v3.0 &nbsp;|&nbsp; JFRodriguez
</div>""", unsafe_allow_html=True)